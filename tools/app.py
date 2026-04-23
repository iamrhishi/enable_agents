from flask import Flask,request, jsonify, redirect
import requests
import sqlite3
import shutil
import psutil
import tempfile
from datetime import datetime
from uuid import uuid4
from time import time
import json
import os
import bs4
import fitz 
import faiss
from typing import Dict, List
from typing_extensions import TypedDict
from dotenv import load_dotenv
from rake_nltk import Rake
import numpy as np
from scipy.spatial.distance import cosine
import openai
import nltk
import pandas as pd
import pickle
import hashlib
import http.client
import json
import glob
from google_auth_oauthlib.flow import Flow
from google.oauth2.credentials import Credentials
import googleapiclient.discovery
from email.message import EmailMessage
import base64
from flask_cors import CORS  # Import Flask-CORS
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash
from werkzeug.security import check_password_hash
from werkzeug.utils import secure_filename
import openpyxl
from flask_cors import cross_origin
from urllib.parse import urlencode

# LangChain imports with fallbacks for version compatibility
try:
    from langchain_openai import OpenAIEmbeddings, ChatOpenAI
except (ImportError, Exception) as e:
    try:
        from langchain.embeddings.openai import OpenAIEmbeddings
        from langchain.chat_models import ChatOpenAI
    except (ImportError, Exception):
        # Set placeholders for testing without full dependencies
        OpenAIEmbeddings = None
        ChatOpenAI = None
        print(f"Warning: LangChain OpenAI modules not available: {e}")

try:
    from langchain_community.vectorstores import FAISS
except (ImportError, Exception):
    try:
        from langchain.vectorstores import FAISS
    except (ImportError, Exception):
        FAISS = None
        print("Warning: FAISS not available")

try:
    from langchain_community.document_loaders import WebBaseLoader, PyPDFLoader
except (ImportError, Exception):
    try:
        from langchain.document_loaders import WebBaseLoader, PyPDFLoader
    except (ImportError, Exception):
        WebBaseLoader = None
        PyPDFLoader = None
        print("Warning: Document loaders not available")

try:
    from langchain_core.documents import Document
except (ImportError, Exception):
    try:
        from langchain.schema import Document
    except (ImportError, Exception):
        Document = None
        print("Warning: Document class not available")

try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except (ImportError, Exception):
    try:
        from langchain.text_splitter import RecursiveCharacterTextSplitter
    except (ImportError, Exception):
        RecursiveCharacterTextSplitter = None
        print("Warning: RecursiveCharacterTextSplitter not available")

try:
    from langchain_core.prompts import ChatPromptTemplate
except (ImportError, Exception):
    try:
        from langchain.prompts import ChatPromptTemplate
    except (ImportError, Exception):
        ChatPromptTemplate = None
        print("Warning: ChatPromptTemplate not available")
from bs4 import BeautifulSoup
from urllib.parse import urlparse
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.common.exceptions import TimeoutException, NoSuchElementException
from webdriver_manager.chrome import ChromeDriverManager
import time
from datetime import datetime, timedelta
import re
from docx import Document as DocxDocument
import networkx as nx
from google_business_helper import GoogleBusinessHelper, GoogleBusinessAnalyzer, GoogleBusinessSearcher

ENV_FILE = os.path.join(os.path.dirname(__file__), '.env')
load_dotenv(ENV_FILE, override=True)
LINKEDIN_CLIENT_ID = os.getenv('LINKEDIN_CLIENT_ID')
LINKEDIN_CLIENT_SECRET = os.getenv('LINKEDIN_CLIENT_SECRET')
LINKEDIN_REDIRECT_URI = os.getenv('LINKEDIN_REDIRECT_URI', 'http://localhost:5000/linkedin/callback')


nltk.download('stopwords')
nltk.download('punkt_tab')

PROMPTS_FILE = "/Users/rhishikeshthakur/Enable/Software Development/enable_agents/data/prompts.json"


app = Flask(__name__)
CORS(app)

GOOGLE_CLIENT_ID = os.getenv('GOOGLE_CLIENT_ID')
GOOGLE_CLIENT_SECRET = os.getenv('GOOGLE_CLIENT_SECRET')
GOOGLE_REDIRECT_URI = os.getenv('GOOGLE_REDIRECT_URI', 'http://localhost:5000/auth/google/callback')
os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1' # allow HTTP for local dev

# Database config (env override + local fallback)
app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv(
    'DATABASE_URI',
    'sqlite:///enable_agents.db'
)
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)

ALLOWED_EXTENSIONS = {'csv', 'xlsx', 'xls'}
MAX_FILE_SIZE = 16 * 1024 * 1024  # 16MB
app.config['MAX_CONTENT_LENGTH'] = MAX_FILE_SIZE

# Content Marketing Agent Configuration
CONTENT_MARKETING_UPLOAD_FOLDER = os.environ.get('CONTENT_MARKETING_UPLOAD_FOLDER', os.path.join(os.path.dirname(__file__), 'data', 'content_marketing_uploads'))
CONTENT_MARKETING_DB_PATH = os.environ.get('CONTENT_MARKETING_DB_PATH', os.path.join(os.path.dirname(__file__), 'data', 'content_marketing.db'))
CONTENT_MARKETING_ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt', 'xlsx', 'html', 'md'}
os.makedirs(CONTENT_MARKETING_UPLOAD_FOLDER, exist_ok=True)
os.makedirs(os.path.dirname(CONTENT_MARKETING_DB_PATH), exist_ok=True)

# Initialize Content Marketing Database
def init_content_marketing_db():
    """Initialize SQLite database for content marketing agent"""
    conn = sqlite3.connect(CONTENT_MARKETING_DB_PATH)
    cursor = conn.cursor()
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS projects (
            project_id TEXT PRIMARY KEY,
            user_id TEXT NOT NULL,
            project_name TEXT NOT NULL,
            description TEXT,
            industry TEXT,
            sector TEXT,
            function TEXT,
            role TEXT,
            created_at TIMESTAMP,
            updated_at TIMESTAMP,
            metadata JSON
        )
    ''')
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS documents (
            doc_id TEXT PRIMARY KEY,
            project_id TEXT NOT NULL,
            file_name TEXT NOT NULL,
            file_type TEXT,
            file_path TEXT,
            file_size INTEGER,
            upload_date TIMESTAMP,
            document_type TEXT,
            extracted_content TEXT,
            FOREIGN KEY(project_id) REFERENCES projects(project_id)
        )
    ''')
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS knowledge_graphs (
            kg_id TEXT PRIMARY KEY,
            project_id TEXT NOT NULL,
            kg_data JSON,
            entities INT,
            relationships INT,
            created_at TIMESTAMP,
            FOREIGN KEY(project_id) REFERENCES projects(project_id)
        )
    ''')
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS generated_content (
            content_id TEXT PRIMARY KEY,
            project_id TEXT NOT NULL,
            channel TEXT,
            content_type TEXT,
            content TEXT,
            source_docs JSON,
            domain_context JSON,
            created_at TIMESTAMP,
            modified_at TIMESTAMP,
            FOREIGN KEY(project_id) REFERENCES projects(project_id)
        )
    ''')
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS conversation_history (
            msg_id TEXT PRIMARY KEY,
            project_id TEXT NOT NULL,
            user_message TEXT,
            agent_response TEXT,
            context JSON,
            timestamp TIMESTAMP,
            FOREIGN KEY(project_id) REFERENCES projects(project_id)
        )
    ''')
    
    conn.commit()
    conn.close()

init_content_marketing_db()

class User(db.Model):
    __tablename__ = 'users'
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password = db.Column(db.String(512), nullable=False)
    first_name = db.Column(db.String(80))
    last_name = db.Column(db.String(80))
    email = db.Column(db.String(120))
    company = db.Column(db.String(120))
    linkedin = db.Column(db.String(256))
    short_intro = db.Column(db.String(256))
    company_intro = db.Column(db.String(256))

class GoogleOAuthToken(db.Model):
    __tablename__ = 'google_oauth_tokens'
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), db.ForeignKey('users.username'), nullable=False)
    token = db.Column(db.Text, nullable=False)
    refresh_token = db.Column(db.Text)
    token_uri = db.Column(db.String(512))
    client_id = db.Column(db.String(512))
    client_secret = db.Column(db.String(512))
    scopes = db.Column(db.Text)


EMAIL_EXTRACTION_UNIT_COST = float(os.getenv('EMAIL_EXTRACTION_UNIT_COST', '0.20'))
DEFAULT_EMAIL_EXTRACTION_LIMIT = int(os.getenv('EMAIL_EXTRACTION_DEFAULT_LIMIT', '500'))


class EmailExtractionQuota(db.Model):
    __tablename__ = 'email_extraction_quotas'
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(120), unique=True, nullable=False, index=True)
    total_allowed = db.Column(db.Integer, nullable=False, default=DEFAULT_EMAIL_EXTRACTION_LIMIT)
    used_count = db.Column(db.Integer, nullable=False, default=0)
    created_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow, onupdate=datetime.utcnow)


class EmailExtractionUsageLog(db.Model):
    __tablename__ = 'email_extraction_usage_logs'
    id = db.Column(db.Integer, primary_key=True)
    request_id = db.Column(db.String(64), unique=True, nullable=False, index=True)
    username = db.Column(db.String(120), nullable=False, index=True)
    processed_count = db.Column(db.Integer, nullable=False, default=0)
    billable_count = db.Column(db.Integer, nullable=False, default=0)
    charged_count = db.Column(db.Integer, nullable=False, default=0)
    cost_this_request = db.Column(db.Float, nullable=False, default=0.0)
    total_cost_after = db.Column(db.Float, nullable=False, default=0.0)
    created_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)

class EmailCampaign(db.Model):
    __tablename__ = 'email_campaigns'
    id = db.Column(db.String(36), primary_key=True)
    name = db.Column(db.String(255), nullable=False)
    subject = db.Column(db.String(255), nullable=False)
    username = db.Column(db.String(120), index=True)
    created_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)

class EmailCampaignRecipient(db.Model):
    __tablename__ = 'email_campaign_recipients'
    id = db.Column(db.Integer, primary_key=True)
    campaign_id = db.Column(db.String(36), db.ForeignKey('email_campaigns.id'), nullable=False)
    receiver_email = db.Column(db.String(255), nullable=False, index=True)
    receiver_name = db.Column(db.String(255), nullable=True)
    status = db.Column(db.String(50), default='Sent')
    reply_status = db.Column(db.String(50), default='No Reply')
    sent_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)
    replied_at = db.Column(db.DateTime, nullable=True)

class SavedProject(db.Model):
    __tablename__ = 'saved_projects'
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(120), nullable=False, index=True)
    name = db.Column(db.String(255), nullable=False)
    query_used = db.Column(db.String(512), nullable=True)
    created_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)

class SavedLead(db.Model):
    __tablename__ = 'saved_leads'
    id = db.Column(db.Integer, primary_key=True)
    project_id = db.Column(db.Integer, db.ForeignKey('saved_projects.id', ondelete='CASCADE'), nullable=False)
    
    name = db.Column(db.String(255))
    website = db.Column(db.String(512))
    phone = db.Column(db.String(100))
    address = db.Column(db.String(512))
    
    emails = db.Column(db.Text) 
    linkedin_links = db.Column(db.Text)
    social_links = db.Column(db.Text)
    
    has_extracted = db.Column(db.Boolean, default=False)
    
    raw_data = db.Column(db.Text) 
    
    created_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)



def _ensure_email_usage_tables():
    """Create usage tracking tables if they don't exist yet."""
    db.create_all()


def _normalize_username(value):
    candidate = (value or '').strip()
    return candidate if candidate else 'anonymous'


def _is_billable_email(value):
    if not value:
        return False

    email_str = str(value).strip()
    lowered = email_str.lower()
    if lowered in {'n/a', 'na', 'none', 'error'}:
        return False
    if lowered.startswith('phone:'):
        return False

    return re.fullmatch(r'[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}', email_str) is not None


def _get_or_create_quota(username):
    quota = EmailExtractionQuota.query.filter_by(username=username).first()
    if quota:
        return quota

    quota = EmailExtractionQuota(
        username=username,
        total_allowed=DEFAULT_EMAIL_EXTRACTION_LIMIT,
        used_count=0
    )
    db.session.add(quota)
    db.session.commit()
    return quota


def _build_usage_summary(username, quota=None):
    quota = quota or _get_or_create_quota(username)
    used = max(quota.used_count, 0)
    remaining = max(quota.total_allowed - used, 0)

    return {
        'username': username,
        'totalAllowed': quota.total_allowed,
        'usedCount': used,
        'remainingCount': remaining,
        'unitCost': EMAIL_EXTRACTION_UNIT_COST,
        'totalCost': round(used * EMAIL_EXTRACTION_UNIT_COST, 2)
    }

# 1. Load: First we need to load our data. This is done with Document Loaders.
# 2. Split: Text splitters break large Documents into smaller chunks. This is useful both for indexing data and passing it into a model, as large chunks are harder to search over and won't fit in a model's finite context window.
# 3. Store: We need somewhere to store and index our splits, so that they can be searched over later. This is often done using a VectorStore and Embeddings model.
# 4. Retrieve: Given a user input, relevant splits are retrieved from storage using a Retriever.
# 5. Generate: A ChatModel / LLM produces an answer using a prompt that includes both the question with the retrieved data

cache = {}

# Cache for KG+RAG to avoid recreating embeddings and graphs for same documents
kg_rag_cache = {
    'embeddings': {},      # Store embeddings by document hash
    'faiss_indices': {},   # Store FAISS indices by document hash
    'chunks': {},          # Store text chunks by document hash
    'knowledge_graphs': {} # Store knowledge graphs by nodes/edges hash
}

# Define state for application
class State(TypedDict):
    question: str
    context: 'List'  # List of Document objects
    answer: str


# ====== CONTENT MARKETING AGENT CLASSES & HELPERS ======

class DomainSpecializationAnalyzer:
    """Analyzes documents to extract domain specialization information"""
    
    def __init__(self):
        self.llm = ChatOpenAI(model="gpt-4", temperature=0)
        self.industry_keywords = self._load_industry_keywords()
    
    def _load_industry_keywords(self) -> Dict[str, List[str]]:
        """Load industry-specific keywords"""
        return {
            'Technology': ['software', 'cloud', 'api', 'infrastructure', 'devops', 'saas'],
            'Healthcare': ['medical', 'patient', 'pharmaceutical', 'clinical', 'health', 'disease'],
            'Finance': ['banking', 'investment', 'portfolio', 'trading', 'compliance', 'regulatory'],
            'Retail': ['ecommerce', 'inventory', 'customer', 'sales', 'purchase', 'product'],
            'Manufacturing': ['production', 'supply chain', 'logistics', 'quality', 'automation'],
            'Real Estate': ['property', 'tenant', 'lease', 'valuation', 'construction'],
            'Education': ['student', 'curriculum', 'learning', 'course', 'assessment'],
        }
    
    def analyze_documents(self, documents: List[str]) -> Dict:
        """
        Analyze documents to extract domain specialization
        
        Args:
            documents: List of document texts
            
        Returns:
            Dictionary with industry, sector, function, role analysis
        """
        combined_text = ' '.join(documents[:3]) if documents else ''
        
        prompt = ChatPromptTemplate.from_template("""
        Analyze the following business documents and extract domain specialization information.
        
        Documents:
        {documents}
        
        Provide a JSON response with:
        {{
            "industry": "identified industry",
            "sector": "business sector",
            "function": "primary business function",
            "role": "primary role/persona",
            "target_audience": "target customer/audience",
            "value_proposition": "key value proposition",
            "tone": "recommended tone (professional/casual/formal)",
            "key_themes": ["theme1", "theme2", ...]
        }}
        """)
        
        try:
            chain = prompt | self.llm
            response = chain.invoke({"documents": combined_text[:2000]})
            
            import re
            json_match = re.search(r'\{.*\}', response.content, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
        except:
            pass
        
        return {
            "industry": "General",
            "sector": "Unknown",
            "function": "Marketing",
            "role": "Marketing Manager",
            "target_audience": "Business Professionals",
            "value_proposition": "Enhanced marketing through AI",
            "tone": "professional",
            "key_themes": ["innovation", "value", "efficiency"]
        }


def extract_text_from_file_content_marketing(file_path: str, file_type: str) -> str:
    """Extract text content from various file formats"""
    try:
        if file_type == 'pdf':
            text = []
            pdf_document = fitz.open(file_path)
            for page in pdf_document:
                text.append(page.get_text())
            pdf_document.close()
            return '\n'.join(text)
        
        elif file_type == 'docx':
            doc = DocxDocument(file_path)
            return '\n'.join([para.text for para in doc.paragraphs])
        
        elif file_type == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        elif file_type == 'html':
            with open(file_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
                return soup.get_text()
        
        elif file_type == 'md':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        else:
            return ''
    
    except Exception as e:
        print(f"Error extracting text from {file_path}: {str(e)}")
        return ''


def setup_driver(headless=True):
    """Setup Chrome WebDriver for interactive scraping with popup handling"""
    chrome_options = Options()
    
    if headless:
        chrome_options.add_argument('--headless')
    
    chrome_options.add_argument('--no-sandbox')
    chrome_options.add_argument('--disable-dev-shm-usage')
    chrome_options.add_argument('--disable-gpu')
    chrome_options.add_argument('--window-size=1920,1080')
    chrome_options.add_argument('--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36')
    
    # Disable notifications and popups
    prefs = {
        "profile.default_content_setting_values.notifications": 2,
        "profile.default_content_settings.popups": 0,
        "profile.managed_default_content_settings.images": 2  # Block images for faster loading
    }
    chrome_options.add_experimental_option("prefs", prefs)
    
    # Additional options to handle consent and cookies
    chrome_options.add_argument('--disable-features=VizDisplayCompositor')
    chrome_options.add_argument('--disable-extensions')
    chrome_options.add_argument('--disable-plugins')
    
    service = Service(ChromeDriverManager().install())
    return webdriver.Chrome(service=service, options=chrome_options)

def get_chrome_history_path():
    """Get Chrome history file path based on OS"""
    if os.name == 'nt':  # Windows
        return os.path.expanduser('~\\AppData\\Local\\Google\\Chrome\\User Data\\Default\\History')
    elif 'darwin' in os.sys.platform.lower():  # macOS
        return os.path.expanduser('~/Library/Application Support/Google/Chrome/Default/History')
    else:  # Linux
        return os.path.expanduser('~/.config/google-chrome/Default/History')

def identify_saas_tools_with_openai(history_data):
    """Use OpenAI to identify which URLs are web applications, tools, SaaS, PaaS, or productivity platforms"""
    try:
        # Extract URLs for analysis (limit to avoid token limits)
        urls_to_analyze = [item['url'] for item in history_data[:50]]  # Analyze top 50 URLs
        urls_text = "\n".join([f"{i+1}. {url}" for i, url in enumerate(urls_to_analyze)])
        
        prompt = f"""Analyze the following URLs and identify which ones are web applications, tools, websites, or platforms. This includes:

- Web Applications
- Websites
- Tools
- Platforms
- PaaS (Platform as a Service) platforms of any kind
- Web-based productivity tools 
- Cloud platforms and services
- Development tools and platforms
- Business applications and tools
- Design and creative tools
- Communication and collaboration platforms
- Analytics and monitoring tools
- Project management tools
- CRM and business software
- Educational and learning platforms
- Entertainment and media platforms (if they're tools/apps)
- Social media platforms (if used as business tools)
- Storage and file-sharing platforms
- Any other web-based tools or platforms not mentioned above

For each URL, return a JSON array with this structure:
[
  {{"url_index": 1, "is_tool": true, "tool_name": "Google Docs", "category": "Productivity", "type": "SaaS", "description": "Document creation and collaboration"}},
  {{"url_index": 2, "is_tool": true, "tool_name": "AWS Console", "category": "Cloud Platform", "type": "PaaS", "description": "Cloud computing services"}},
  {{"url_index": 3, "is_tool": false, "tool_name": null, "category": null, "type": null, "description": "Regular website"}}
]

URLs to analyze:
{urls_text}

Rules:
- Identify ANY website, web-based tool, application, or platform
- Include Google Workspace (Docs, Sheets, Drive, Gmail), Microsoft 365, Slack, Zoom, etc.
- Include development platforms (GitHub, GitLab, Heroku, Vercel)
- Include cloud platforms (AWS, Azure, GCP)  
- Include design tools (Figma, Canva, Adobe Creative Cloud)
- Include business tools (Salesforce, HubSpot, Trello, Asana)
- Include social media platforms if used as business tools
- Categories: Development, Communication, Productivity, Design, Analytics, Cloud Platform, CRM, Project Management, Storage, Entertainment, Education, Social Media, Other
- Types: PaaS, Web App, Platform, Tool, Website
- Differentiate between a tool, WebApp, Website, and Platform
- Keep descriptions short (under 60 characters)
- Only mark as "false" if it's clearly a regular informational website
- Return valid JSON only"""

        client = openai.OpenAI()
        client.api_key = os.environ['OPENAI_API_KEY']
        
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a web application and tool identifier. You identify ALL types of web-based tools, applications, and platforms. Return only valid JSON arrays."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=3000,  # Increased for more detailed analysis
            temperature=0.1
        )
        
        response_text = response.choices[0].message.content.strip()
        response_text = response_text.replace('```json', '').replace('```', '').strip()
        
        print(response_text)

        try:
            tools_analysis = json.loads(response_text)

            
    
            # Create a mapping from URL index to tool info
            tools_mapping = {}
            for item in tools_analysis:
                if 'url_index' in item:
                    tools_mapping[item['url_index'] - 1] = {  # Convert to 0-based index
                        'is_tool': item.get('is_tool', False),
                        'tool_name': item.get('tool_name'),
                        'category': item.get('category'),
                        'type': item.get('type'),  # SaaS, PaaS, Web App, etc.
                        'description': item.get('description')
                    }
            
            return {
                'success': True,
                'mapping': tools_mapping
            }
            
        except json.JSONDecodeError as e:
            print(f"JSON parsing failed for tools analysis: {str(e)}")
            return {
                'success': False,
                'error': f'JSON parsing failed: {str(e)}'
            }
            
    except Exception as e:
        print(f"OpenAI tools analysis error: {str(e)}")
        return {
            'success': False,
            'error': f'OpenAI error: {str(e)}'
        }



def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def clean_dataframe_strict(df, required_field_groups=None):
    """Clean dataframe with strict field group requirements"""
    # ... existing cleaning code ...
    
    if required_field_groups:
        condition = pd.Series([True] * len(df))
        
        for group_name, field_list in required_field_groups.items():
            # Find which fields from this group exist in the dataframe
            existing_fields = [field for field in field_list if field in df.columns]
            
            if existing_fields:
                # For this group, at least one field must have data
                group_condition = pd.Series([False] * len(df))
                
                for field in existing_fields:
                    field_condition = (df[field] != '') & (df[field].str.lower() != 'n/a') & (df[field].str.lower() != 'na') & (df[field].str.lower() != 'none') & (df[field].str.lower() != 'null')
                    group_condition = group_condition | field_condition
                
                # All groups must have at least one field with data
                condition = condition & group_condition
        
        df = df[condition]
        print(f"After strict filtering: {len(df)} rows remaining")
    
    return df


def clean_dataframe(df, required_columns=None):
    """Clean dataframe by removing rows with missing required columns"""
    if required_columns is None:
        return df
    
    # Filter rows where at least one of the required columns has a value
    if required_columns:
        condition = pd.Series([False] * len(df))
        for col in required_columns:
            if col in df.columns:
                condition = condition | (df[col].notna() & (df[col] != ''))
        df = df[condition]
    
    return df


def clean_csv(df, required_columns=None):
    """Remove rows where any specified columns have blank/missing values"""
    if required_columns is None:
        required_columns = df.columns
    
    # Replace empty strings and common null representations with NaN
    df[required_columns] = df[required_columns].replace(['', ' ', 'N/A', 'n/a', 'NA', 'na', 'null', 'NULL'], pd.NA)
    
    # Drop rows with any NaN in required columns
    return df.dropna(subset=required_columns, how='any')


def csv_to_json(file_path):
    """Convert CSV file to JSON object with row filtering"""
    try:
        # Read CSV with error handling for different encodings
        try:
            df = pd.read_csv(file_path, encoding='utf-8')
        except UnicodeDecodeError:
            try:
                df = pd.read_csv(file_path, encoding='latin-1')
            except UnicodeDecodeError:
                df = pd.read_csv(file_path, encoding='cp1252')
        
        original_count = len(df)
        
        # Clean the dataframe - specify required columns or clean all
        required_columns = ['company', 'Company', 'name', 'Name', 'title', 'Title']
        existing_required = [col for col in required_columns if col in df.columns]
        
        if existing_required:
            df = clean_csv(df, existing_required)
        else:
            df = clean_csv(df)  # Clean all columns if no specific ones found
        
        filtered_count = len(df)
        
        # Convert to JSON
        json_data = df.to_dict('records')
        
        return {
            'success': True,
            'data': json_data,
            'total_records': len(json_data),
            'original_records': original_count,
            'filtered_records': filtered_count,
            'rows_removed': original_count - filtered_count,
            'columns': list(df.columns),
            'message': f'Successfully converted {len(json_data)} records (removed {original_count - filtered_count} empty rows)'
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': f'Error processing CSV file: {str(e)}',
            'data': []
        }

def xlsx_to_json(file_path):
    """Convert XLSX file to JSON object with row filtering"""
    try:
        # Read Excel file (first sheet by default)
        df = pd.read_excel(file_path, engine='openpyxl')
        
        original_count = len(df)
        
        # Define important columns that should not be empty
        required_columns = [
            'company', 'Company', 'organization', 'Organization', 'employer', 'Employer',
            'title', 'Title', 'position', 'Position', 'job_title', 'Job Title', 'role', 'Role',
            'name', 'Name', 'full_name', 'Full Name', 'first_name', 'First Name', 'last_name', 'Last Name'
        ]
        
        # Clean the dataframe with filtering
        df = clean_dataframe(df, required_columns)
        
        filtered_count = len(df)
        
        # Convert to JSON
        json_data = df.to_dict('records')
        
        return {
            'success': True,
            'data': json_data,
            'total_records': len(json_data),
            'original_records': original_count,
            'filtered_records': filtered_count,
            'rows_removed': original_count - filtered_count,
            'columns': list(df.columns),
            'message': f'Successfully converted {len(json_data)} records (removed {original_count - filtered_count} empty rows)'
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': f'Error processing XLSX file: {str(e)}',
            'data': []
        }

def xlsx_to_json_multiple_sheets(file_path):
    """Convert XLSX file with multiple sheets to JSON object with row filtering"""
    try:
        # Read all sheets
        excel_file = pd.ExcelFile(file_path)
        sheets_data = {}
        total_original = 0
        total_filtered = 0
        
        # Define important columns that should not be empty
        required_columns = [
            'company', 'Company', 'organization', 'Organization', 'employer', 'Employer',
            'title', 'Title', 'position', 'Position', 'job_title', 'Job Title', 'role', 'Role',
            'name', 'Name', 'full_name', 'Full Name', 'first_name', 'First Name', 'last_name', 'Last Name'
        ]
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name, engine='openpyxl')
            original_count = len(df)
            total_original += original_count
            
            df = clean_dataframe(df, required_columns)
            filtered_count = len(df)
            total_filtered += filtered_count
            
            sheets_data[sheet_name] = df.to_dict('records')
        
        return {
            'success': True,
            'data': sheets_data,
            'sheets': list(excel_file.sheet_names),
            'total_records': total_filtered,
            'original_records': total_original,
            'filtered_records': total_filtered,
            'rows_removed': total_original - total_filtered,
            'message': f'Successfully converted {len(excel_file.sheet_names)} sheets with {total_filtered} total records (removed {total_original - total_filtered} empty rows)'
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': f'Error processing XLSX file: {str(e)}',
            'data': {}
        }
    
def extract_unique_companies(json_data):
    """Extract unique company names from JSON data"""
    companies = set()
    
    for record in json_data:
        # Check various possible company field names
        company_fields = ['company', 'Company', 'organization', 'Organization', 'employer', 'Employer']
        
        for field in company_fields:
            if field in record and record[field]:
                company_name = str(record[field]).strip()
                if company_name and company_name.lower() not in ['', 'n/a', 'na', 'none', 'null']:
                    companies.add(company_name)
                break
    
    return list(companies)

def get_company_skills_from_openai(company_list):
    """Send company list to OpenAI and get required skills for each company"""
    try:
        companies_text = ", ".join(company_list)
        
        # Simplified prompt with clearer instructions
        prompt = f"""For the following companies, provide required skills in JSON format:

Companies: {companies_text}

Return a JSON array with this exact structure:
[
  {{"company": "Company1", "required_skills": ["skill1", "skill2", "skill3", "skill4", "skill5", "skill6"]}},
  {{"company": "Company2", "required_skills": ["skill1", "skill2", "skill3", "skill4", "skill5", "skill6"]}}
]

Rules:
- Exactly 6 skills per company
- Mix of technical and soft skills
- Valid JSON only, no explanations
- Keep skills concise (1-3 words each)"""

        client = openai.OpenAI()
        client.api_key=os.environ['OPENAI_API_KEY']
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a JSON generator. Return only valid JSON arrays. Keep responses concise."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1500,  # Reduced token limit to prevent truncation
            temperature=0.1   # Lower temperature for more consistent output
        )
        
        response_text = response.choices[0].message.content.strip()
        
        # More robust cleaning
        response_text = response_text.replace('```json', '').replace('```', '').strip()
        
        # Remove any trailing incomplete content
        if response_text.endswith(','):
            response_text = response_text[:-1]
        if response_text.endswith('}'):
            response_text += ']'
        elif not response_text.endswith(']'):
            # Find the last complete object and truncate there
            last_complete = response_text.rfind('}')
            if last_complete > 0:
                response_text = response_text[:last_complete + 1] + ']'
        
        # Try to fix common JSON issues
        response_text = fix_json_issues(response_text)
        
        try:
            company_skills_data = json.loads(response_text)
            
            # Validate the structure
            if not isinstance(company_skills_data, list):
                raise ValueError("Response is not a list")
            
            for item in company_skills_data:
                if not isinstance(item, dict) or 'company' not in item or 'required_skills' not in item:
                    raise ValueError("Invalid item structure")
                if not isinstance(item['required_skills'], list):
                    raise ValueError("required_skills is not a list")
            
            return {
                'success': True,
                'data': company_skills_data
            }
            
        except (json.JSONDecodeError, ValueError) as e:
            print(f"JSON parsing failed: {str(e)}")
            print(f"Raw response: {response_text[:500]}...")
            
            # Fallback: try to extract data manually
            fallback_result = extract_skills_manually(companies_text, response_text)
            if fallback_result['success']:
                return fallback_result
            
            return {
                'success': False,
                'error': f'Invalid JSON response: {str(e)}',
                'raw_response': response_text[:300]
            }
        
    except Exception as e:
        print(f"OpenAI API error: {str(e)}")
        return {
            'success': False,
            'error': f'OpenAI API error: {str(e)}'
        }
    

def fix_json_issues(json_text):
    """Fix common JSON formatting issues"""
    # Remove any text before the first [
    start_idx = json_text.find('[')
    if start_idx > 0:
        json_text = json_text[start_idx:]
    
    # Remove any text after the last ]
    end_idx = json_text.rfind(']')
    if end_idx > 0:
        json_text = json_text[:end_idx + 1]
    
    # Fix common escape issues
    json_text = json_text.replace('\n', ' ').replace('\r', ' ')
    json_text = json_text.replace('"', '"').replace('"', '"')  # Fix smart quotes
    
    return json_text

def extract_skills_manually(companies_text, response_text):
    """Fallback method to extract skills manually if JSON parsing fails"""
    try:
        import re
        
        company_list = [c.strip() for c in companies_text.split(',')]
        result_data = []
        
        # Default skills for different company types
        default_skills = {
            'tech': ['Python', 'JavaScript', 'Problem Solving', 'Communication', 'Teamwork', 'Leadership'],
            'finance': ['Excel', 'Financial Analysis', 'Risk Management', 'Communication', 'Attention to Detail', 'Leadership'],
            'healthcare': ['Patient Care', 'Medical Knowledge', 'Communication', 'Empathy', 'Attention to Detail', 'Teamwork'],
            'default': ['Communication', 'Problem Solving', 'Leadership', 'Teamwork', 'Analytical Thinking', 'Adaptability']
        }
        
        for company in company_list:
            # Try to determine company type and assign appropriate skills
            company_lower = company.lower()
            if any(tech_word in company_lower for tech_word in ['google', 'microsoft', 'apple', 'facebook', 'amazon', 'tech', 'software']):
                skills = default_skills['tech']
            elif any(fin_word in company_lower for fin_word in ['bank', 'finance', 'capital', 'investment', 'goldman', 'morgan']):
                skills = default_skills['finance']
            elif any(health_word in company_lower for health_word in ['hospital', 'medical', 'health', 'pharma', 'clinic']):
                skills = default_skills['healthcare']
            else:
                skills = default_skills['default']
            
            result_data.append({
                'company': company,
                'required_skills': skills
            })
        
        return {
            'success': True,
            'data': result_data
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': f'Manual extraction failed: {str(e)}'
        }

def enrich_json_with_openai(json_data):
    """Enhanced function: Extract companies -> Get skills from OpenAI -> Enrich original data"""
    try:
        # Step 1: Extract unique company names from JSON data
        unique_companies = extract_unique_companies(json_data)
        
        if not unique_companies:
            return {
                'success': False,
                'error': 'No company names found in the data'
            }
        
        print(f"Found {len(unique_companies)} unique companies: {unique_companies}")
        
        # Step 2: Send company list to OpenAI to get required skills
        openai_result = get_company_skills_from_openai(unique_companies)
        
        if not openai_result['success']:
            return {
                'success': False,
                'error': f'Failed to get skills from OpenAI: {openai_result.get("error", "Unknown error")}'
            }
        
        # Step 3: Create company-skills mapping from OpenAI response
        company_skills_map = {}
        for item in openai_result['data']:
            if 'company' in item and 'required_skills' in item:
                company_skills_map[item['company']] = item['required_skills']
        
        print(f"Created skills mapping for {len(company_skills_map)} companies")
        
        # Step 4: Enrich original JSON data by matching company names
        enriched_data = []
        skills_added_count = 0
        
        for record in json_data:
            # Create a copy of the original record
            enriched_record = record.copy()
            
            # Find company name in this record
            company_name = None
            company_fields = ['company', 'Company', 'organization', 'Organization', 'employer', 'Employer']
            
            for field in company_fields:
                if field in record and record[field]:
                    company_name = str(record[field]).strip()
                    break
            
            # Add required_skills based on company match
            if company_name and company_name in company_skills_map:
                enriched_record['required_skills'] = company_skills_map[company_name]
                skills_added_count += 1
            else:
                enriched_record['required_skills'] = []
            
            enriched_data.append(enriched_record)
        
        return {
            'success': True,
            'data': enriched_data,
            'message': f'Successfully enriched {len(enriched_data)} profiles. Added skills to {skills_added_count} records based on {len(unique_companies)} companies.',
            'companies_processed': len(unique_companies),
            'records_enriched': skills_added_count
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': f'Enrichment process error: {str(e)}'
        }

def get_credentials():
    load_dotenv()
    return os.getenv("OPENAI_API_KEY")

def get_file_hash(file_path):
    hasher = hashlib.md5()
    with open(file_path, 'rb') as f:
        buf = f.read()
        hasher.update(buf)
    return hasher.hexdigest()

def save_embeddings(file_hash, index, phrase_embeddings, page_chunks):
    embeddings_folder = "/Users/rhishikeshthakur/Enable/Software Development/enable_agents/data/embeddings"
    os.makedirs(embeddings_folder, exist_ok=True)
    with open(os.path.join(embeddings_folder, f"{file_hash}_index.pkl"), "wb") as f:
        pickle.dump(index, f)
    with open(os.path.join(embeddings_folder, f"{file_hash}_phrase_embeddings.pkl"), "wb") as f:
        pickle.dump(phrase_embeddings, f)
    with open(os.path.join(embeddings_folder, f"{file_hash}_page_chunks.pkl"), "wb") as f:
        pickle.dump(page_chunks, f)

def load_embeddings(file_hash):
    embeddings_folder = "/Users/rhishikeshthakur/Enable/Software Development/enable_agents/data/embeddings"
    with open(os.path.join(embeddings_folder, f"{file_hash}_index.pkl"), "rb") as f:
        index = pickle.load(f)
    with open(os.path.join(embeddings_folder, f"{file_hash}_phrase_embeddings.pkl"), "rb") as f:
        phrase_embeddings = pickle.load(f)
    with open(os.path.join(embeddings_folder, f"{file_hash}_page_chunks.pkl"), "rb") as f:
        page_chunks = pickle.load(f)
    print("Embeddings loaded successfully.")
    return index, phrase_embeddings, page_chunks

def init_llm():
    llm = ChatOpenAI(model="gpt-4o-mini", temperature=0)
    return llm

def init_embeddings():
    embeddings = OpenAIEmbeddings(model="text-embedding-3-large")
    return embeddings

def init_vector_store(embeddings):
    return None

def pdf_loader(file_path):
    pdf_document = fitz.open(file_path)
    pdf_text ={}
    for page_number in range(pdf_document.page_count):
        page = pdf_document.load_page(page_number)
        pdf_text[page_number + 1] = page.get_text()
    pdf_document.close()
    return pdf_text

def web_loader():
    loader = WebBaseLoader(
        web_paths=("https://lilianweng.github.io/posts/2023-06-23-agent/",),
        bs_kwargs=dict(parse_only=bs4.SoupStrainer(class_=("post-content", "post-title", "post-header"))),
    )
    docs = loader.load()
    return docs

def pdf_splitter(pdf_text):
    number_of_characters = sum(len(text) for text in pdf_text.values())
    print(f"Total number of characters in PDF: {number_of_characters}")  # Debug print for total characters
    # Set chunk size and overlap based on the total number of characters
    # For example, if the total number of characters is less than 100,000, use smaller chunks
    # Adjust chunk size and overlap based on the total number of characters
    if number_of_characters < 100000:
        chunk_size = 1000
        chunk_overlap = 200
    elif number_of_characters < 500000:
        chunk_size = 800
        chunk_overlap = 200
    elif number_of_characters < 1000000:
        chunk_size = 600
        chunk_overlap = 200
    elif number_of_characters < 2000000:
        chunk_size = 500
        chunk_overlap = 200
    elif number_of_characters < 5000000:
        chunk_size = 400
        chunk_overlap = 200
    elif number_of_characters < 10000000:
        chunk_size = 300
        chunk_overlap = 200
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    page_chunks = {}
    for page, text in pdf_text.items():
        # print(f"Page {page} length: {len(text)}")  # Debug print for text length
        chunks = text_splitter.split_text(text)
        # print(f"Page {page} chunks: {len(chunks)}")  # Debug print for number of chunks
        page_chunks[page] = chunks
    return page_chunks

def extract_keywords_from_pdf(pdf_text):
    rake = Rake()
    page_phrases = {}
    for page, text in pdf_text.items():
        rake.extract_keywords_from_text(text)
        phrases = rake.get_ranked_phrases()[:5]
        page_phrases[page] = phrases
    return page_phrases

def extract_keywords_from_chunks(page_chunks):
    rake = Rake()
    chunk_phrases = {}
    for page, chunks in page_chunks.items():
        for chunk_number, chunk in enumerate(chunks, start=1):
            rake.extract_keywords_from_text(chunk)
            phrases = rake.get_ranked_phrases()[:5]
            chunk_phrases[(page, chunk_number)] = phrases
    return chunk_phrases

def get_embeddings(phrase):
    client = openai.OpenAI()
    client.api_key = get_credentials()
    response = client.embeddings.create(model="text-embedding-ada-002", input=phrase)
    return response.data[0].embedding

def get_embeddings_batch(phrases):
    client = openai.OpenAI()
    client.api_key = get_credentials()
    response = client.embeddings.create(model="text-embedding-ada-002", input=phrases)
    return [data.embedding for data in response.data]

def store_embeddings(page_phrases, chunk_phrases):
    """Store embeddings with better error handling"""
    print("Processing embeddings...")
    print(f"Page phrases: {page_phrases}")
    print(f"Chunk phrases keys: {list(chunk_phrases.keys())}")
    
    phrase_embeddings = {}
    all_embeddings = []  # Store all embeddings for FAISS index
    
    # Process chunk phrases and get embeddings
    for (page, chunk_number), phrases in chunk_phrases.items():
        if phrases:  # Only process if there are phrases
            try:
                embeddings = get_embeddings_batch(phrases)
                phrase_embeddings[(page, chunk_number)] = list(zip(phrases, embeddings))
                # Collect all embeddings for FAISS index
                all_embeddings.extend(embeddings)
            except Exception as e:
                print(f"Error getting embeddings for page {page}, chunk {chunk_number}: {e}")
                phrase_embeddings[(page, chunk_number)] = []
        else:
            phrase_embeddings[(page, chunk_number)] = []
    
    # Check if we have any embeddings
    if not all_embeddings:
        print("Warning: No embeddings were created. Using dummy embeddings.")
        # Create a dummy embedding for testing
        dummy_embedding = [0.0] * 1536  # OpenAI ada-002 embedding dimension
        all_embeddings = [dummy_embedding]
        # Add dummy phrase_embeddings entry
        phrase_embeddings[(1, 1)] = [("no content found", dummy_embedding)]
    
    # Initialize FAISS index
    dimension = len(all_embeddings[0])
    print(f"Creating FAISS index with dimension: {dimension}")
    index = faiss.IndexFlatIP(dimension)
    
    # Add all embeddings to the index
    if all_embeddings:
        embeddings_array = np.array(all_embeddings, dtype=np.float32)
        index.add(embeddings_array)
        print(f"Added {len(all_embeddings)} embeddings to FAISS index")
    
    return index, phrase_embeddings

def extract_phrases_from_query(query):
    rake = Rake()
    rake.extract_keywords_from_text(query)
    return rake.get_ranked_phrases()

def get_embeddings_for_query(phrases):
    client = openai.OpenAI()
    client.api_key = get_credentials()
    return [client.embeddings.create(model="text-embedding-ada-002", input=phrase).data[0].embedding for phrase in phrases]

def get_cosine_similarity(embedding1, embedding2):
    return 1 - cosine(embedding1, embedding2)

def store_cosine_similarities(query_embeddings, phrase_embeddings, page_chunks):
    chunk_similarities = {}
    for (page, chunk_number), phrases in phrase_embeddings.items():
        similarities = []
        for phrase, embedding in phrases:
            phrase_similarities = [get_cosine_similarity(embedding, query_embedding) for query_embedding in query_embeddings] 
        similarities.append(max(phrase_similarities)) 
        # Choose the highest similarity for each phrase 
        average_similarity = np.mean(similarities) 
        # Average similarity for the chunk 
        chunk_similarities[(page, chunk_number)] = average_similarity 
    # Get top 5 chunks by similarity 
    top_chunks = sorted(chunk_similarities.items(), key=lambda x: x[1], reverse=True)[:5] 
    # Output top 5 chunks 
    print("Top 5 most relatable chunks:") 
    selected_chunks = []
    for (page, chunk_number), similarity in top_chunks: 
        print(f"Page: {page}, Chunk: {chunk_number}, Similarity: {similarity}") 
        print(f"Chunk text:\n{page_chunks[page][chunk_number-1]}\n")
        selected_chunks.append(page_chunks[page][chunk_number-1])
    return selected_chunks

def retrieve_similar_chunks(query_embeddings, index, phrase_embeddings, page_chunks):
    """Retrieve similar chunks with better error handling"""
    try:
        if not query_embeddings or not phrase_embeddings:
            print("Warning: No query embeddings or phrase embeddings available")
            return ["No relevant content found."]
        
        query_embeddings_np = np.array(query_embeddings, dtype=np.float32)
        
        # Ensure we don't search for more chunks than available
        available_chunks = index.ntotal
        k = min(5, available_chunks) if available_chunks > 0 else 1
        
        print(f"Searching for {k} similar chunks from {available_chunks} total")
        
        if available_chunks == 0:
            return ["No indexed content available for search."]
        
        D, I = index.search(query_embeddings_np, k=k)
        
        selected_chunks = []
        processed_indices = set()
        
        # Map FAISS indices back to chunks
        embedding_index = 0
        index_to_chunk_map = {}
        
        # Create mapping from FAISS index to chunk content
        for (page, chunk_number), phrases in phrase_embeddings.items():
            for phrase, embedding in phrases:
                index_to_chunk_map[embedding_index] = (page, chunk_number)
                embedding_index += 1
        
        # Retrieve chunks based on similarity
        for i in range(len(I)):
            for j in range(len(I[i])):
                chunk_idx = int(I[i][j])
                
                if chunk_idx in index_to_chunk_map and chunk_idx not in processed_indices:
                    page, chunk_number = index_to_chunk_map[chunk_idx]
                    
                    # Get the actual chunk content
                    if page in page_chunks and chunk_number - 1 < len(page_chunks[page]):
                        chunk_content = page_chunks[page][chunk_number - 1]
                        selected_chunks.append(chunk_content)
                        processed_indices.add(chunk_idx)
                        
                        if len(selected_chunks) >= 5:  # Limit to 5 chunks
                            break
            
            if len(selected_chunks) >= 5:
                break
        
        # Fallback: if no chunks found, return some content from page_chunks
        if not selected_chunks and page_chunks:
            print("Using fallback: returning first available chunks")
            for page, chunks in page_chunks.items():
                for chunk in chunks[:2]:  # Take first 2 chunks from each page
                    selected_chunks.append(chunk)
                    if len(selected_chunks) >= 3:
                        break
                if len(selected_chunks) >= 3:
                    break
        
        if not selected_chunks:
            selected_chunks = ["Unable to find relevant content in the document."]
        
        print(f"Retrieved {len(selected_chunks)} chunks for context")
        return selected_chunks
        
    except Exception as e:
        print(f"Error in retrieve_similar_chunks: {e}")
        return [f"Error retrieving content: {str(e)}"]

def extract_keywords_from_chunks(page_chunks):
    """Extract keywords from chunks with better error handling"""
    rake = Rake()
    chunk_phrases = {}
    
    for page, chunks in page_chunks.items():
        for chunk_number, chunk in enumerate(chunks, start=1):
            try:
                if chunk and chunk.strip():  # Only process non-empty chunks
                    rake.extract_keywords_from_text(chunk)
                    phrases = rake.get_ranked_phrases()[:5]
                    chunk_phrases[(page, chunk_number)] = phrases
                else:
                    chunk_phrases[(page, chunk_number)] = []
            except Exception as e:
                print(f"Error extracting keywords from page {page}, chunk {chunk_number}: {e}")
                chunk_phrases[(page, chunk_number)] = []
    
    print(f"Extracted keywords from {len(chunk_phrases)} chunks")
    return chunk_phrases


def parse_simple_query_enhanced(user_query):
    """Enhanced keyword extraction using OpenAI function calling with special commands support"""
    try:
        # Define the function schema for OpenAI with special commands
        extract_function = {
            "name": "extract_search_criteria",
            "description": "Extract search criteria from natural language query, including special commands",
            "parameters": {
                "type": "object",
                "properties": {
                    "company": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "Company names, organizations, or employers mentioned"
                    },
                    "title": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "Job titles, positions, or roles mentioned"
                    },
                    "name": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "Person names mentioned"
                    },
                    "skills": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "Technical skills, technologies, or expertise mentioned"
                    },
                    "location": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "Cities, countries, or geographic locations mentioned"
                    },
                    "phrases": {
                        "type": "array",
                        "items": {
                            "type": "string",
                            "enum": [
                                "show_all_companies",
                                "show_all_titles", 
                                "show_all_locations",
                                "show_all_skills",
                                "show_favorites",
                                "list_companies",
                                "list_titles",
                                "list_locations", 
                                "list_skills",
                                "my_favorites",
                                "saved_profiles",
                                "all_companies",
                                "all_titles",
                                "all_locations",
                                "all_skills"
                            ]
                        },
                        "description": "Special command phrases detected in the query like 'show all companies', 'show my favorites', 'list all titles', etc."
                    }
                },
                "required": []
            }
        }

        client = openai.OpenAI()
        client.api_key = os.environ['OPENAI_API_KEY']
        
        # Enhanced system prompt to detect special commands while maintaining your existing logic
        system_prompt = """You are a search query parser that extracts search criteria accurately from user queries. Extract search criteria accurately from user queries. Include information that is explicitly mentioned but also matching terms that are close to what is mentioned. They could be singular or plural forms, sub-strings or variations, etc.

SPECIAL COMMANDS TO DETECT:
- "show all companies" / "list companies" / "all companies" â†’ show_all_companies
- "show all titles" / "list titles" / "all titles" / "job titles" â†’ show_all_titles  
- "show all locations" / "list locations" / "all locations" â†’ show_all_locations
- "show all skills" / "list skills" / "all skills" â†’ show_all_skills
- "show favorites" / "my favorites" / "show saved" / "saved profiles" â†’ show_favorites

Extract both specific search terms AND any special commands detected. If the query contains both search terms and special commands, include both in the response."""

        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {
                    "role": "system", 
                    "content": system_prompt
                },
                {
                    "role": "user", 
                    "content": f"Extract search criteria and special commands from this query: '{user_query}'"
                }
            ],
            functions=[extract_function],
            function_call={"name": "extract_search_criteria"},
            temperature=0.1
        )
        
        # Extract function call result
        function_call = response.choices[0].message.function_call
        keywords = json.loads(function_call.arguments)
        
        # Clean up empty arrays
        keywords = {k: v for k, v in keywords.items() if v and len(v) > 0}
        
        # Check if this is a special command
        special_phrases = keywords.get('phrases', [])
        is_special_command = len(special_phrases) > 0
        
        # Determine command type from phrases
        command_type = None
        if is_special_command:
            phrase = special_phrases[0]  # Use first detected phrase
            if phrase in ['show_all_companies', 'list_companies', 'all_companies']:
                command_type = 'show_companies'
            elif phrase in ['show_all_titles', 'list_titles', 'all_titles']:
                command_type = 'show_titles'
            elif phrase in ['show_all_locations', 'list_locations', 'all_locations']:
                command_type = 'show_locations'
            elif phrase in ['show_all_skills', 'list_skills', 'all_skills']:
                command_type = 'show_skills'
            elif phrase in ['show_favorites', 'my_favorites', 'saved_profiles']:
                command_type = 'show_favorites'
        
        print(f"OpenAI extracted keywords: {keywords}")
        print(f"Special command detected: {is_special_command}, Type: {command_type}")
        
        return {
            'success': True, 
            'keywords': keywords,
            'special_command': is_special_command,
            'command_type': command_type,
            'phrases': special_phrases
        }
        
    except Exception as e:
        print(f"OpenAI parsing error: {e}")
        return {
            'success': False, 
            'keywords': {}, 
            'special_command': False,
            'command_type': None,
            'phrases': [],
            'error': str(e)
        }


def simple_search_json(json_data, keywords):
    """Simple search function with substring matching for both keys and values"""
    results = []
    print(f"Searching with keywords: {keywords}")
    
    for record in json_data:
        match = True
        
        # Check each keyword type
        for field, values in keywords.items():
            if not values:  # Skip empty fields
                continue
                
            field_match = False
            
            # Map search fields to possible record fields with substring matching
            if field == 'company':
                record_fields = ['company', 'Company', 'organization', 'Organization', 'employer', 'Employer']
            elif field == 'title':
                record_fields = ['title', 'Title', 'position', 'Position', 'job_title', 'Job Title', 'role', 'Role']
            elif field == 'name':
                record_fields = ['name', 'Name', 'full_name', 'Full Name', 'first_name', 'First Name', 'last_name', 'Last Name']
            elif field == 'location':
                record_fields = ['location', 'Location', 'city', 'City', 'country', 'Country']
            elif field == 'skills':
                record_fields = ['skills', 'Skills', 'required_skills', 'technologies', 'Technologies']
            else:
                record_fields = [field]
            
            # First try exact field matching
            for record_field in record_fields:
                if record_field in record and record[record_field]:
                    record_value = str(record[record_field]).lower()
                    
                    for search_value in values:
                        search_value_lower = search_value.lower()
                        
                        # Check if search value is substring of record value
                        if search_value_lower in record_value:
                            field_match = True
                            print(f"Exact match: '{search_value}' found in '{record_value}' (field: {record_field})")
                            break
                        
                        # Check if record value is substring of search value (reverse match)
                        if record_value in search_value_lower:
                            field_match = True
                            print(f"Reverse match: '{record_value}' found in '{search_value}' (field: {record_field})")
                            break
                    
                    if field_match:
                        break
            
            # If no exact field match, try substring matching on field names
            if not field_match:
                for record_field in record.keys():
                    # Check if search field is substring of record field or vice versa
                    field_lower = field.lower()
                    record_field_lower = record_field.lower()
                    
                    # Match if field names have substring relationship
                    if (field_lower in record_field_lower or record_field_lower in field_lower) and record[record_field]:
                        record_value = str(record[record_field]).lower()
                        
                        for search_value in values:
                            search_value_lower = search_value.lower()
                            
                            # Check substring matches in both directions
                            if search_value_lower in record_value:
                                field_match = True
                                print(f"Field substring match: '{search_value}' found in '{record_value}' (field: {record_field})")
                                break
                            
                            if record_value in search_value_lower:
                                field_match = True
                                print(f"Field reverse substring match: '{record_value}' found in '{search_value}' (field: {record_field})")
                                break
                        
                        if field_match:
                            break
            
            # If still no match, try word-level substring matching
            if not field_match:
                for record_field in record_fields:
                    if record_field in record and record[record_field]:
                        record_value = str(record[record_field]).lower()
                        
                        for search_value in values:
                            search_value_lower = search_value.lower()
                            
                            # Split search value into words and check each word
                            search_words = [word.strip() for word in search_value_lower.replace('-', ' ').split() if len(word.strip()) > 2]
                            
                            for word in search_words:
                                if word in record_value:
                                    field_match = True
                                    print(f"Word substring match: '{word}' from '{search_value}' found in '{record_value}' (field: {record_field})")
                                    break
                            
                            # Also split record value and check against search value
                            if not field_match:
                                record_words = [word.strip() for word in record_value.replace('-', ' ').split() if len(word.strip()) > 2]
                                
                                for word in record_words:
                                    if word in search_value_lower:
                                        field_match = True
                                        print(f"Record word substring match: '{word}' from '{record_value}' found in '{search_value}' (field: {record_field})")
                                        break
                            
                            if field_match:
                                break
                    
                    if field_match:
                        break
            
            # If this field didn't match, exclude the record
            if not field_match:
                match = False
                break
        
        
        if match:
            results.append(record)
    print(match)
    return results

def make_request(url, config):
    """Centralized HTTP request handler"""
    headers = config.get('headers', {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
    })
    
    timeout = config.get('timeout', 30)
    
    response = requests.get(url, headers=headers, timeout=timeout)
    response.raise_for_status()
    
    return response

def parse_html(response):
    """Centralized HTML parsing"""
    return BeautifulSoup(response.text, 'html.parser')

def get_page_metadata(soup, url):
    """Extract basic page metadata"""
    title = None
    title_tag = soup.find('title')
    if title_tag:
        title = title_tag.get_text(strip=True)
    
    meta_description = None
    meta_desc = soup.find('meta', attrs={'name': 'description'})
    if meta_desc:
        meta_description = meta_desc.get('content', '')
    
    return {
        'page_title': title,
        'meta_description': meta_description,
        'url': url
    }

def scrape_basic_content(url, config):
    """Basic content scraping - common elements"""
    try:
        response = make_request(url, config)
        soup = parse_html(response)
        
        # Get page metadata
        result = get_page_metadata(soup, url)
        result['status_code'] = response.status_code
        
        # Get selectors from config or use defaults
        selectors = config.get('selectors', {})
        
        if selectors:
            # Use custom selectors
            for name, selector_config in selectors.items():
                result[name] = extract_with_selector(soup, selector_config)
        else:
            # Default extraction
            result.update(extract_common_elements(soup))
        
        return {
            'success': True,
            'type': 'basic',
            'data': result,
            'scraped_at': datetime.now().isoformat()
        }
        
    except requests.exceptions.RequestException as e:
        return {
            'success': False,
            'error': f'Request failed: {str(e)}',
            'url': url
        }
    except Exception as e:
        return {
            'success': False,
            'error': f'Scraping error: {str(e)}'
        }

def scrape_text_content(url, config):
    """Text-only content scraping"""
    try:
        response = make_request(url, config)
        soup = parse_html(response)
        
        selector = config.get('selector')
        max_length = config.get('max_length')
        clean_text = config.get('clean_text', True)
        
        if selector:
            # Extract text from specific selector
            element = soup.select_one(selector)
            text = element.get_text(strip=clean_text) if element else None
        else:
            # Extract all text
            text = soup.get_text()
            if clean_text:
                text = ' '.join(text.split())
        
        # Apply length limit
        if text and max_length and len(text) > max_length:
            text = text[:max_length] + '...'
        
        return {
            'success': True,
            'type': 'text',
            'url': url,
            'selector': selector if selector else 'entire_page',
            'text': text,
            'text_length': len(text) if text else 0,
            'scraped_at': datetime.now().isoformat()
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }

def scrape_json_ld_content(url, config):
    """JSON-LD structured data scraping"""
    try:
        response = make_request(url, config)
        soup = parse_html(response)
        
        # Extract JSON-LD scripts
        json_ld_data = []
        scripts = soup.find_all("script", type="application/ld+json")
        
        filters = config.get('filters', {})
        schema_type = filters.get('schema_type')
        keywords = filters.get('keywords', [])
        
        for script in scripts:
            try:
                if script.string:
                    data_obj = json.loads(script.string)
                    
                    # Apply filters
                    if schema_type and data_obj.get('@type', '').lower() != schema_type.lower():
                        continue
                    
                    if keywords:
                        data_str = str(data_obj).lower()
                        if not any(keyword.lower() in data_str for keyword in keywords):
                            continue
                    
                    json_ld_data.append(data_obj)
            except json.JSONDecodeError:
                continue
        
        return {
            'success': True,
            'type': 'json_ld',
            'url': url,
            'json_ld': json_ld_data,
            'count': len(json_ld_data),
            'filters_applied': filters,
            'scraped_at': datetime.now().isoformat()
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }

def scrape_product_content(url, config):
    """Enhanced product data scraping"""
    try:
        response = make_request(url, config)
        soup = parse_html(response)
        
        result = get_page_metadata(soup, url)
        
        # Product-specific selectors
        product_selectors = {
            'name': [
                'h1.product-title', 'h1.product-name', '.product-title', '.pdp-product-name',
                'h1[data-testid="product-title"]', '.product-info h1'
            ],
            'price': [
                '.price', '.current-price', '.product-price', '[data-testid="price"]',
                '.price-current', '.final-price'
            ],
            'description': [
                '.product-description', '.pdp-description', '.product-details-description',
                '[data-testid="product-description"]', '.product-info-description'
            ],
            'brand': [
                '.brand', '.product-brand', '[data-testid="brand"]', '.manufacturer'
            ],
            'availability': [
                '.availability', '.stock-status', '[data-testid="availability"]', '.in-stock'
            ],
            'images': [
                '.product-image img', '.product-gallery img', '.pdp-images img'
            ]
        }
        
        # Override with custom selectors if provided
        custom_selectors = config.get('selectors', {})
        if custom_selectors:
            product_selectors.update(custom_selectors)
        
        # Extract product data
        product_data = {}
        for field, selector_list in product_selectors.items():
            if field == 'images':
                # Handle images specially
                images = []
                for selector in selector_list:
                    elements = soup.select(selector)
                    for img in elements[:5]:  # Limit to 5 images
                        src = img.get('src') or img.get('data-src')
                        if src:
                            images.append({
                                'src': src,
                                'alt': img.get('alt', '')
                            })
                    if images:
                        break
                product_data[field] = images
            else:
                # Handle text fields
                for selector in selector_list:
                    element = soup.select_one(selector)
                    if element:
                        product_data[field] = element.get_text(strip=True)
                        break
        
        result['product_data'] = product_data
        
        return {
            'success': True,
            'type': 'product',
            'data': result,
            'scraped_at': datetime.now().isoformat()
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }

def scrape_table_content(url, config):
    """Table data scraping"""
    try:
        response = make_request(url, config)
        soup = parse_html(response)
        
        table_selector = config.get('table_selector', 'table')
        include_headers = config.get('include_headers', True)
        max_rows = config.get('max_rows')
        
        tables = soup.select(table_selector)
        tables_data = []
        
        for i, table in enumerate(tables):
            table_data = {
                'table_index': i + 1,
                'headers': [],
                'rows': []
            }
            
            # Extract headers
            if include_headers:
                header_row = table.find('thead') or table.find('tr')
                if header_row:
                    headers = header_row.find_all(['th', 'td'])
                    table_data['headers'] = [th.get_text(strip=True) for th in headers]
            
            # Extract rows
            body = table.find('tbody') or table
            rows = body.find_all('tr')
            
            if include_headers and table_data['headers']:
                rows = rows[1:]  # Skip header row
            
            for row_idx, row in enumerate(rows):
                if max_rows and row_idx >= max_rows:
                    break
                    
                cells = row.find_all(['td', 'th'])
                if cells:  # Skip empty rows
                    row_data = [cell.get_text(strip=True) for cell in cells]
                    table_data['rows'].append(row_data)
            
            if table_data['headers'] or table_data['rows']:
                tables_data.append(table_data)
        
        return {
            'success': True,
            'type': 'tables',
            'url': url,
            'tables': tables_data,
            'total_tables': len(tables_data),
            'scraped_at': datetime.now().isoformat()
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }

def scrape_custom_selectors(url, config):
    """Custom selector-based scraping"""
    try:
        response = make_request(url, config)
        soup = parse_html(response)
        
        selectors = config.get('selectors', {})
        if not selectors:
            return {
                'success': False,
                'error': 'No selectors provided for custom scraping'
            }
        
        result = get_page_metadata(soup, url)
        
        for name, selector_config in selectors.items():
            result[name] = extract_with_selector(soup, selector_config)
        
        return {
            'success': True,
            'type': 'custom',
            'data': result,
            'scraped_at': datetime.now().isoformat()
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }

def extract_with_selector(soup, selector_config):
    """Universal selector extraction handler"""
    try:
        if isinstance(selector_config, str):
            # Simple string selector
            element = soup.select_one(selector_config)
            return element.get_text(strip=True) if element else None
        
        elif isinstance(selector_config, dict):
            # Advanced selector with options
            selector = selector_config.get('selector')
            attribute = selector_config.get('attribute', 'text')
            multiple = selector_config.get('multiple', False)
            
            if multiple:
                elements = soup.select(selector)
                if attribute == 'text':
                    return [el.get_text(strip=True) for el in elements]
                else:
                    return [el.get(attribute) for el in elements if el.get(attribute)]
            else:
                element = soup.select_one(selector)
                if element:
                    if attribute == 'text':
                        return element.get_text(strip=True)
                    else:
                        return element.get(attribute)
                return None
        
        return None
        
    except Exception as e:
        return f"Error: {str(e)}"

def extract_common_elements(soup):
    """Extract common page elements when no specific selectors provided"""
    common_data = {}
    
    try:
        # Headings (limit to first 10)
        headings = []
        for i in range(1, 7):  # h1 to h6
            heading_elements = soup.find_all(f'h{i}')
            for h in heading_elements[:3]:  # Max 3 per level
                text = h.get_text(strip=True)
                if text:
                    headings.append({
                        'level': i,
                        'text': text
                    })
        common_data['headings'] = headings[:10]
        
        # Paragraphs (first 5 meaningful ones)
        paragraphs = []
        p_elements = soup.find_all('p')
        for p in p_elements:
            text = p.get_text(strip=True)
            if text and len(text) > 30:  # Only meaningful paragraphs
                paragraphs.append(text)
                if len(paragraphs) >= 5:
                    break
        common_data['paragraphs'] = paragraphs
        
        # Links (first 10)
        links = []
        a_elements = soup.find_all('a', href=True)
        for a in a_elements[:10]:
            href = a.get('href')
            text = a.get_text(strip=True)
            if href and text:
                links.append({
                    'url': href,
                    'text': text
                })
        common_data['links'] = links
        
        # Images (first 5)
        images = []
        img_elements = soup.find_all('img')
        for img in img_elements[:5]:
            src = img.get('src') or img.get('data-src')
            if src:
                images.append({
                    'src': src,
                    'alt': img.get('alt', '')
                })
        common_data['images'] = images
        
    except Exception as e:
        common_data['error'] = f"Error extracting common elements: {str(e)}"
    
    return common_data

def handle_cookie_consent_and_popups(driver):
    """Enhanced cookie consent handler specifically for ZARA and similar sites"""
    try:
        # Wait a moment for popups to appear
        time.sleep(3)
        
        # Specific selectors for ZARA and other fashion sites
        zara_selectors = [
            # ZARA specific
            'button[data-qa-action="accept-all"]',
            'button[data-qa-action="accept"]',
            '#onetrust-accept-btn-handler',
            '.ot-pc-refuse-all-handler',
            
            # Common cookie consent buttons
            'button:contains("Accept All")',
            'button:contains("Alle akzeptieren")',  # German
            'button:contains("Aceptar todo")',      # Spanish
            'button:contains("Accepter tout")',     # French
            'button:contains("Accept")',
            'button:contains("Akzeptieren")',
            'button:contains("I Accept")',
            'button:contains("Agree")',
            'button:contains("Einverstanden")',
            'button:contains("OK")',
            'button:contains("Continue")',
            'button:contains("Weiter")',
            'button:contains("Got it")',
            
            # Common class names and IDs
            '.accept-all',
            '.accept-cookies',
            '.accept-btn',
            '.cookie-accept',
            '.consent-accept',
            '.terms-accept',
            '#accept-all',
            '#accept-cookies',
            '#cookie-accept',
            '.onetrust-close-btn-handler',
            '.optanon-allow-all',
            
            # Data attributes
            '[data-accept="all"]',
            '[data-cookie-accept]',
            '[data-consent="accept"]',
            '[data-testid="accept"]',
            '[data-testid="accept-all"]',
            '[data-qa-action="accept-all"]',
            
            # OneTrust specific (used by ZARA)
            '#onetrust-accept-btn-handler',
            '.onetrust-close-btn-handler',
            '.ot-pc-refuse-all-handler',
            
            # ARIA labels
            '[aria-label*="Accept"]',
            '[aria-label*="Agree"]',
            '[aria-label*="akzeptieren"]',
            '[role="button"][aria-label*="Accept"]'
        ]
        
        print("Looking for cookie consent buttons...")
        
        # Try each selector
        for selector in zara_selectors:
            try:
                if ':contains(' in selector:
                    # Handle text-based selectors with XPath
                    text = selector.split(':contains("')[1].split('")')[0]
                    tag = selector.split(':contains(')[0]
                    xpath = f"//{tag}[contains(translate(text(), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), '{text.lower()}')]"
                    elements = driver.find_elements(By.XPATH, xpath)
                else:
                    elements = driver.find_elements(By.CSS_SELECTOR, selector)
                
                for element in elements:
                    if element.is_displayed() and element.is_enabled():
                        print(f"Found and clicking accept button: {selector}")
                        # Scroll to element first
                        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", element)
                        time.sleep(1)
                        
                        # Try multiple click methods
                        try:
                            element.click()
                        except Exception:
                            try:
                                driver.execute_script("arguments[0].click();", element)
                            except Exception:
                                # Force click using coordinates
                                from selenium.webdriver.common.action_chains import ActionChains
                                ActionChains(driver).move_to_element(element).click().perform()
                        
                        time.sleep(3)  # Wait for dialog to close
                        print(f"Successfully clicked consent button")
                        return True
                        
            except Exception as e:
                print(f"Error trying selector {selector}: {e}")
                continue
        
        print("No cookie consent button found or clicked")
        return False
        
    except Exception as e:
        print(f"Error handling popups: {e}")
        return False

def setup_driver(headless=True):
    """Setup Chrome WebDriver for interactive scraping with popup handling"""
    chrome_options = Options()
    
    if headless:
        chrome_options.add_argument('--headless')
    
    chrome_options.add_argument('--no-sandbox')
    chrome_options.add_argument('--disable-dev-shm-usage')
    chrome_options.add_argument('--disable-gpu')
    chrome_options.add_argument('--window-size=1920,1080')
    chrome_options.add_argument('--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36')
    
    # Disable notifications and popups
    prefs = {
        "profile.default_content_setting_values.notifications": 2,
        "profile.default_content_settings.popups": 0,
        "profile.managed_default_content_settings.images": 2  # Block images for faster loading
    }
    chrome_options.add_experimental_option("prefs", prefs)
    
    # Additional options to handle consent and cookies
    chrome_options.add_argument('--disable-features=VizDisplayCompositor')
    chrome_options.add_argument('--disable-extensions')
    chrome_options.add_argument('--disable-plugins')
    
    service = Service(ChromeDriverManager().install())
    return webdriver.Chrome(service=service, options=chrome_options)

def scrape_with_interaction(url, headless, config):
    """Scrape product information with interactive capabilities and popup handling"""
    driver = None
    try:
        driver = setup_driver(headless)
        
        # Set page load timeout
        driver.set_page_load_timeout(30)
        
        print(f"Loading page: {url}")
        driver.get(url)
        
        # Handle cookie consent and terms acceptance first
        print("Handling cookie consent and popups...")
        popup_handled = handle_cookie_consent_and_popups(driver)
        if popup_handled:
            print("Successfully handled popup/consent dialog")
        else:
            print("No popup found or couldn't handle it")
        
        # Wait for page to load after handling popups
        wait = WebDriverWait(driver, 15)
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        
        # Additional wait for dynamic content
        time.sleep(3)
        
        result = {
            'success': True,
            'url': url,
            'scraped_at': datetime.now().isoformat(),
            'method': 'interactive',
            'popup_handled': popup_handled,
            'product_info': {}
        }
        
        # Get basic page info
        page_title = driver.title
        result['page_title'] = page_title
        print(f"Page title: {page_title}")
        
        # Extract product information
        product_info = extract_product_information(driver, wait, config)
        result['product_info'] = product_info
        
        return result
        
    except Exception as e:
        error_msg = f'Interactive scraping error: {str(e)}'
        print(error_msg)
        return {
            'success': False,
            'error': error_msg,
            'method': 'interactive'
        }
    finally:
        if driver:
            try:
                driver.quit()
            except:
                pass

def extract_product_information(driver, wait, config):
    """Extract product description and size information with enhanced interaction"""
    product_info = {
        'description': None,
        'size_info': None,
        'size_chart': None,
        'size_measurements': None,  # Add this new field
        'materials': None,
        'care_instructions': None,
        'specifications': None,
        'interactive_content_found': False,
        'extraction_methods_used': []
    }
    
    try:
        # Step 1: Try to extract visible content first
        print("Extracting visible content...")
        soup = BeautifulSoup(driver.page_source, 'html.parser')
        visible_content = extract_static_product_info(soup)
        product_info.update(visible_content)
        product_info['extraction_methods_used'].append('static_html')
        
        # Step 2: Look for and interact with clickable elements
        print("Looking for interactive elements...")
        interactive_elements = find_interactive_elements(driver)
        
        if interactive_elements:
            product_info['interactive_content_found'] = True
            product_info['extraction_methods_used'].append('interactive_clicks')
            
            for element_info in interactive_elements:
                try:
                    print(f"Interacting with {element_info['type']}: {element_info['text']}")
                    
                    # Click the element
                    element = element_info['element']
                    element_type = element_info['type']
                    
                    # Scroll to element and click
                    driver.execute_script("arguments[0].scrollIntoView(true);", element)
                    time.sleep(1)
                    
                    # Get page source before click for comparison
                    before_click_content = driver.page_source
                    
                    # Try different click methods
                    try:
                        element.click()
                    except:
                        # If regular click fails, try JavaScript click
                        driver.execute_script("arguments[0].click();", element)
                    
                    # Wait longer for content to load
                    time.sleep(4)  # Increased wait time
                    
                    # Get page source after click
                    after_click_content = driver.page_source
                    
                    # Check if content changed
                    if len(after_click_content) != len(before_click_content):
                        print(f"Page content changed after clicking {element_type}")
                    
                    # Extract content after interaction
                    new_soup = BeautifulSoup(driver.page_source, 'html.parser')
                    new_content = extract_dynamic_content(new_soup, element_type)
                    
                    # Merge new content
                    for key, value in new_content.items():
                        if value and (not product_info.get(key) or product_info.get(key) == element_info['text']):
                            product_info[key] = value
                            print(f"Updated {key} with new content from {element_type}")
                    
                    # Close popup/modal if it opened
                    close_popup(driver)
                    time.sleep(1)  # Wait after closing
                    
                except Exception as e:
                    print(f"Error interacting with element {element_info['type']}: {e}")
                    continue
        
        # Step 3: Try scrolling to load more content
        print("Scrolling to load additional content...")
        driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
        time.sleep(2)
        driver.execute_script("window.scrollTo(0, 0);")
        time.sleep(1)
        
        # Extract content after scrolling
        final_soup = BeautifulSoup(driver.page_source, 'html.parser')
        scroll_content = extract_static_product_info(final_soup)
        
        # Merge scroll content
        for key, value in scroll_content.items():
            if value and not product_info.get(key):
                product_info[key] = value
        
        if scroll_content:
            product_info['extraction_methods_used'].append('scroll_content')
        
        # Step 4: Clean and format the extracted data
        product_info = clean_product_info(product_info)
        
        print(f"Extraction complete. Methods used: {product_info['extraction_methods_used']}")
        print(f"Final product info keys: {list(product_info.keys())}")
        
        return product_info
        
    except Exception as e:
        print(f"Error in extract_product_information: {e}")
        product_info['error'] = str(e)
        return product_info



def find_interactive_elements(driver):
    """Enhanced function to find clickable elements for multiple brands"""
    interactive_elements = []
    
    # Enhanced selectors for different brands and languages
    element_selectors = {
        'size_guide': [
            # English selectors
            'a:contains("Size Guide")', 'button:contains("Size Guide")',
            'a:contains("Size Chart")', 'button:contains("Size Chart")',
            'a:contains("Sizing")', 'button:contains("Sizing")',
            'a:contains("Measurements")', 'button:contains("Measurements")',
            'a:contains("Product Measurements")', 'button:contains("Product Measurements")',
            
            # German selectors
            'a:contains("GrÃ¶ÃŸentabelle")', 'button:contains("GrÃ¶ÃŸentabelle")',
            'a:contains("GrÃ¶ÃŸenfÃ¼hrung")', 'button:contains("GrÃ¶ÃŸenfÃ¼hrung")',
            'a:contains("MaÃŸe")', 'button:contains("MaÃŸe")',
            
            # French selectors
            'a:contains("Guide des tailles")', 'button:contains("Guide des tailles")',
            'a:contains("Tableau des tailles")', 'button:contains("Tableau des tailles")',
            
            # Spanish selectors
            'a:contains("GuÃ­a de tallas")', 'button:contains("GuÃ­a de tallas")',
            'a:contains("Tabla de tallas")', 'button:contains("Tabla de tallas")',
            
            # Generic attribute selectors
            'a[href*="size"]', 'a[href*="sizing"]', 'a[href*="measurement"]',
            'button[class*="size"]', 'a[class*="size"]',
            '[data-size-guide]', '[data-testid*="size"]',
            '[data-qa-action*="size"]', '[data-qa*="size"]',
            '[data-modal*="size"]', '[data-popup*="size"]',
            
            # Brand specific selectors
            'button[data-qa-action="size-guide"]',
            'a[data-qa-action="size-guide"]',
            '.product-size-guide-link',
            '.size-guide-trigger',
            'button.size-chart-btn',
            'a.size-chart-link',
            '.size-guide', '.size-chart', '.sizing-info',
            
            # Common class patterns
            '.size-guide-link', '.sizing-link', '.measurement-link',
            '[class*="size-guide"]', '[class*="sizing"]', '[class*="measurement"]'
        ],
        'description': [
            # English
            'a:contains("Description")', 'button:contains("Description")',
            'a:contains("More Details")', 'button:contains("More Details")',
            'a:contains("Product Details")', 'button:contains("Product Details")',
            'a:contains("Details")', 'button:contains("Details")',
            
            # German
            'a:contains("Beschreibung")', 'button:contains("Beschreibung")',
            'a:contains("Mehr Details")', 'button:contains("Mehr Details")',
            'a:contains("Produktdetails")', 'button:contains("Produktdetails")',
            
            # Generic
            '.description-toggle', '.product-description-toggle',
            '.more-info', '.view-more', '.expand-description',
            '[data-qa-action*="description"]', '[data-testid*="description"]'
        ],
        'specifications': [
            'a:contains("Specifications")', 'button:contains("Specifications")',
            'a:contains("Specs")', 'button:contains("Specs")',
            'a:contains("Technical Details")', 'button:contains("Technical Details")',
            'a:contains("Spezifikationen")', 'button:contains("Spezifikationen")',
            '.specs-toggle', '.specifications-toggle', '.product-specs',
            '[data-qa-action*="specs"]', '[data-qa-action*="details"]'
        ],
        'care': [
            'a:contains("Care")', 'button:contains("Care Instructions")',
            'a:contains("Care Instructions")', 'a:contains("Washing Instructions")',
            'button:contains("Pflege")', 'button:contains("Pflegeanleitung")',
            '.care-instructions', '.washing-instructions', '.garment-care',
            '[data-qa-action*="care"]'
        ]
    }
    
    print("Searching for interactive elements...")
    
    for element_type, selectors in element_selectors.items():
        print(f"Searching for {element_type} elements...")
        
        for selector in selectors:
            try:
                if ':contains(' in selector:
                    # Handle text-based selectors with XPath
                    text = selector.split(':contains("')[1].split('")')[0]
                    tag = selector.split(':contains(')[0]
                    xpath = f"//{tag}[contains(translate(text(), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), '{text.lower()}')]"
                    elements = driver.find_elements(By.XPATH, xpath)
                else:
                    elements = driver.find_elements(By.CSS_SELECTOR, selector)
                
                for element in elements:
                    if element.is_displayed() and element.is_enabled():
                        element_text = element.text.strip()
                        element_tag = element.tag_name
                        
                        print(f"Found interactive element: {element_type} - '{element_text}' ({element_tag}) (selector: {selector})")
                        
                        interactive_elements.append({
                            'element': element,
                            'type': element_type,
                            'selector': selector,
                            'text': element_text,
                            'tag': element_tag
                        })
                        break  # Only take first working element of each type
                
                if any(e['type'] == element_type for e in interactive_elements):
                    break  # Found element for this type, move to next
                    
            except Exception as e:
                print(f"Error trying selector {selector}: {e}")
                continue
    
    print(f"Found {len(interactive_elements)} interactive elements total")
    return interactive_elements

def extract_static_product_info(soup):
    """Enhanced static product extraction for multiple brands"""
    info = {}
    
    # Enhanced description selectors for various brands
    description_selectors = [
        # ZARA specific
        '.expandable-text__inner-content p',
        '.expandable-text__inner-content',
        '.product-detail-info__description',
        
        # H&M specific
        '.product-description-text',
        '.pdp-product-description',
        '.product-description-content',
        
        # Uniqlo specific
        '.product-description',
        '.pdp-description',
        
        # Generic brand selectors
        '.product-detail-view__description',
        '.product-info__description',
        '[data-testid="product-description"]',
        '.product-info-description',
        '.product-long-description',
        '.description-content',
        '.product-content',
        '.item-description',
        '.product-summary',
        '.product-details-content',
        '.product-information',
        '#product-description',
        '.description',
        '.details',
        '.overview',
        '.about-product',
        '.product-info',
        '.product-detail',
        '.product-overview',
        '.item-details',
        '.product-specs',
        
        # Common patterns
        '[class*="description"]',
        '[class*="product-info"]',
        '[id*="description"]'
    ]
    
    # Extract description with enhanced search
    for selector in description_selectors:
        try:
            element = soup.select_one(selector)
            if element:
                text = element.get_text(strip=True)
                if text and len(text) > 20:
                    info['description'] = text
                    print(f"Found description using selector: {selector}")
                    break
        except Exception as e:
            continue
    
    # Enhanced fallback for description
    if not info.get('description'):
        # Look for expandable content patterns
        expandable_patterns = [
            '[class*="expandable"]',
            '[class*="collapsible"]',
            '[class*="toggle"]'
        ]
        
        for pattern in expandable_patterns:
            elements = soup.select(pattern)
            for element in elements:
                text = element.get_text(strip=True)
                if (text and len(text) > 50 and 
                    not any(exclude in text.lower() for exclude in ['cookie', 'privacy', 'terms', 'size', 'shipping'])):
                    info['description'] = text
                    print(f"Found description in expandable content: {pattern}")
                    break
            if info.get('description'):
                break
    
    # Enhanced size information extraction
    size_selectors = [
        # Direct size guide links/buttons
        '.size-guide', '.size-chart', '.sizing-info', '.size-information',
        '.product-sizing', '.fit-guide', '.measurements', '.dimensions',
        '[data-testid="size-info"]', '.size-details', '.sizing-chart',
        
        # Brand specific
        '.product-detail-size-info', '.size-fit-info', '.size-guide-trigger',
        '.size-chart-link', '.product-size-info', '.sizing-information',
        
        # Interactive elements that might contain size info
        'button[class*="size"]', 'a[class*="size"]',
        '[data-qa-action*="size"]', '[data-testid*="size"]',
        '[data-modal*="size"]', '[data-popup*="size"]'
    ]
    
    for selector in size_selectors:
        try:
            element = soup.select_one(selector)
            if element:
                text = element.get_text(strip=True)
                if text and len(text) > 5:
                    info['size_info'] = text
                    print(f"Found size info using selector: {selector}")
                    break
        except Exception as e:
            continue
    
    # Enhanced materials extraction
    material_selectors = [
        '.materials', '.fabric', '.composition', '.material-composition',
        '.product-materials', '.fabric-composition', '[data-testid="materials"]',
        '.product-detail-composition', '.composition-info',
        '[class*="composition"]', '[class*="material"]', '[class*="fabric"]'
    ]
    
    for selector in material_selectors:
        element = soup.select_one(selector)
        if element:
            text = element.get_text(strip=True)
            if text and any(keyword in text.lower() for keyword in ['cotton', 'polyester', 'material', 'fabric', '%', 'composition']):
                info['materials'] = text
                print(f"Found materials using selector: {selector}")
                break
    
    print(f"Static extraction completed. Found: {list(info.keys())}")
    return info

def extract_dynamic_content(soup, content_type):
    """Extract content that appeared after interaction - Enhanced for multiple brands"""
    content = {}
    
    # Enhanced modal/popup selectors for different brands and platforms
    modal_selectors = [
        # Generic modal selectors
        '.modal-body', '.popup-content', '.modal-content', '.modal-dialog',
        '.dialog-content', '.lightbox-content', '.overlay-content',
        '.modal', '.popup', '.dialog', '.lightbox', '.overlay',
        
        # Size guide specific
        '.size-guide-modal', '.size-modal', '.sizing-modal',
        '.measurements-modal', '.product-measurements', '.size-chart-modal',
        
        # Brand specific selectors
        '.zara-modal', '.h-and-m-modal', '.uniqlo-modal',
        
        # Role-based selectors
        '[role="dialog"]', '[role="modal"]', '[role="alertdialog"]',
        
        # Data attribute selectors
        '[data-modal]', '[data-popup]', '[data-size-guide]',
        
        # Class patterns
        '[class*="modal"]', '[class*="popup"]', '[class*="dialog"]',
        '[class*="size-guide"]', '[class*="measurement"]',
        
        # ID patterns
        '#modal', '#popup', '#size-guide', '#measurements',
        
        # Recently appeared content (might not be in modal)
        '.size-guide-content', '.measurements-content', '.sizing-content',
        '.product-measurements-content', '.size-chart-content'
    ]
    
    modal_content = None
    modal_selector_used = None
    
    # Try each selector to find modal content
    for selector in modal_selectors:
        try:
            element = soup.select_one(selector)
            if element:
                text = element.get_text(strip=True)
                if text and len(text) > 20:  # Must have meaningful content
                    modal_content = element
                    modal_selector_used = selector
                    print(f"Found modal content using selector: {selector}")
                    break
        except Exception as e:
            continue
    
    # If no modal found, look for any recently appeared content with size keywords
    if not modal_content and content_type == 'size_guide':
        print("No modal found, searching for size-related content...")
        
        # Look for any visible element containing size information
        size_keywords = ['size', 'measurement', 'dimension', 'length', 'width', 'chest', 'waist', 'hip']
        
        # Search through all elements for size-related content
        all_elements = soup.find_all(['div', 'section', 'article', 'table', 'ul', 'ol'])
        
        for element in all_elements:
            text = element.get_text(strip=True).lower()
            
            # Check if element contains size-related keywords
            if (len(text) > 50 and 
                any(keyword in text for keyword in size_keywords) and
                # Exclude navigation and header content
                not any(exclude in text for exclude in ['cookie', 'privacy', 'terms', 'navigation', 'menu'])):
                
                modal_content = element
                modal_selector_used = f"size_keyword_search_{element.name}"
                print(f"Found size content in {element.name} element: {text[:100]}...")
                break
    
    if modal_content:
        modal_text = modal_content.get_text(strip=True)
        print(f"Modal content found ({len(modal_text)} chars): {modal_text[:200]}...")
        
        if content_type == 'size_guide':
            # Extract size chart table from modal
            tables = modal_content.select('table')
            if tables:
                print(f"Found {len(tables)} tables in modal")
                for i, table in enumerate(tables):
                    table_data = extract_table_data(table)
                    if table_data and (table_data.get('headers') or table_data.get('rows')):
                        content['size_chart'] = table_data
                        print(f"Extracted size chart table {i}")
                        break
            
            # Look for structured measurement data
            measurements = extract_measurements_from_text(modal_text)
            if measurements:
                content['size_measurements'] = measurements
                print(f"Extracted measurements: {list(measurements.keys())}")
            
            # Always include the full text if it's substantial
            if modal_text and len(modal_text) > 20:
                content['size_info'] = modal_text
                print("Added full modal text as size_info")
            
            # Look for specific measurement patterns
            size_patterns = extract_size_patterns(modal_text)
            if size_patterns:
                content['size_patterns'] = size_patterns
                print(f"Extracted size patterns: {size_patterns}")
        
        elif content_type == 'description':
            if modal_text and len(modal_text) > 20:
                content['description'] = modal_text
                print("Added modal text as description")
        
        elif content_type == 'specifications':
            if modal_text and len(modal_text) > 20:
                content['specifications'] = modal_text
        
        elif content_type == 'care':
            if modal_text and len(modal_text) > 10:
                content['care_instructions'] = modal_text
        
        # Add metadata about extraction
        content['extraction_method'] = 'modal_extraction'
        content['modal_selector'] = modal_selector_used
        content['modal_text_length'] = len(modal_text)
    
    # If still no content found for size guide, try alternative approaches
    if not content and content_type == 'size_guide':
        print("Trying alternative size extraction methods...")
        
        # Method 1: Look for any new tables that might have appeared
        all_tables = soup.select('table')
        for table in all_tables:
            table_text = table.get_text(strip=True).lower()
            if any(keyword in table_text for keyword in ['size', 'measurement', 'cm', 'inch', 'xs', 'sm', 'md', 'lg', 'xl']):
                table_data = extract_table_data(table)
                if table_data:
                    content['size_chart'] = table_data
                    content['extraction_method'] = 'table_scan'
                    print("Found size table through table scan")
                    break
        
        # Method 2: Look for lists with size information
        all_lists = soup.select('ul, ol')
        for list_elem in all_lists:
            list_text = list_elem.get_text(strip=True)
            if (len(list_text) > 50 and 
                any(keyword in list_text.lower() for keyword in ['size', 'measurement', 'dimension'])):
                content['size_info'] = list_text
                content['extraction_method'] = 'list_extraction'
                print("Found size info in list element")
                break
        
        # Method 3: Check page for any content changes
        page_text = soup.get_text()
        if any(keyword in page_text.lower() for keyword in ['size guide', 'measurements', 'sizing chart']):
            measurements = extract_measurements_from_text(page_text)
            if measurements:
                content['size_measurements'] = measurements
                content['size_info'] = "Measurements found in page content"
                content['extraction_method'] = 'page_scan'
                print("Extracted measurements from full page")
    
    print(f"Dynamic content extraction result: {list(content.keys())}")
    return content

def extract_size_patterns(text):
    """Extract common size patterns from text"""
    import re
    
    patterns = {}
    
    # Pattern for size ranges (e.g., "XS: 32-34", "Size S: 36-38")
    size_range_pattern = r'(?:Size\s+)?([XS]{1,2}|[SML]{1,2}|\d{2})\s*:?\s*(\d{1,3}(?:\.\d)?)\s*[-â€“]\s*(\d{1,3}(?:\.\d)?)'
    matches = re.finditer(size_range_pattern, text, re.IGNORECASE)
    
    for match in matches:
        size, min_val, max_val = match.groups()
        patterns[f"size_{size.upper()}_range"] = f"{min_val}-{max_val}"
    
    # Pattern for measurements with units
    measurement_pattern = r'(\w+(?:\s+\w+)?)\s*:?\s*(\d{1,3}(?:\.\d)?)\s*(cm|inches?|in)'
    matches = re.finditer(measurement_pattern, text, re.IGNORECASE)
    
    for match in matches:
        measurement_type, value, unit = match.groups()
        clean_type = measurement_type.strip().lower().replace(' ', '_')
        patterns[f"{clean_type}_{unit.lower()}"] = f"{value} {unit}"
    
    return patterns if patterns else None

def extract_measurements_from_text(text):
    """Extract structured measurements from text content"""
    import re
    
    measurements = {}
    
    # Common measurement patterns
    measurement_patterns = [
        # Pattern: "Size XS: Length 95cm, Waist 68cm"
        r'Size\s+(\w+):\s*([^,]+(?:,\s*[^,]+)*)',
        # Pattern: "XS - Length: 95cm, Waist: 68cm"
        r'(\w+)\s*-\s*([^-]+(?=\w+\s*-|$))',
        # Pattern: "Length: 95cm" 
        r'(\w+):\s*(\d+(?:\.\d+)?\s*cm|\d+(?:\.\d+)?\s*inches?)',
        # Pattern: "95cm length"
        r'(\d+(?:\.\d+)?\s*cm|\d+(?:\.\d+)?\s*inches?)\s+(\w+)',
    ]
    
    # Size categories to look for
    size_categories = ['XS', 'S', 'M', 'L', 'XL', 'XXL', '34', '36', '38', '40', '42', '44', '46']
    measurement_types = ['length', 'width', 'chest', 'waist', 'hips', 'shoulder', 'sleeve', 'inseam', 'rise']
    
    for pattern in measurement_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            groups = match.groups()
            if len(groups) >= 2:
                key, value = groups[0], groups[1]
                
                # Clean and structure the data
                key = key.strip()
                value = value.strip()
                
                if key.upper() in size_categories:
                    # This is a size with measurements
                    measurements[f"size_{key.upper()}"] = value
                elif any(mtype in key.lower() for mtype in measurement_types):
                    # This is a measurement type
                    measurements[key.lower()] = value
    
    return measurements if measurements else None

def extract_table_data(table):
    """Extract data from size chart tables"""
    try:
        table_data = {
            'headers': [],
            'rows': []
        }
        
        # Extract headers
        headers = table.select('thead th, tr:first-child th, tr:first-child td')
        if headers:
            table_data['headers'] = [th.get_text(strip=True) for th in headers]
        
        # Extract rows
        rows = table.select('tbody tr, tr')
        for i, row in enumerate(rows):
            if i == 0 and table_data['headers']:
                continue  # Skip header row
            
            cells = row.select('td, th')
            row_data = [cell.get_text(strip=True) for cell in cells]
            if row_data and any(row_data):  # Skip empty rows
                table_data['rows'].append(row_data)
        
        return table_data
        
    except Exception as e:
        return {'error': f'Error extracting table: {str(e)}'}

def close_popup(driver):
    """Enhanced popup closing with more selectors"""
    close_selectors = [
        '.modal-close', '.close', '.popup-close', '[aria-label="Close"]',
        '.close-btn', '.modal-dismiss', 'button[data-dismiss="modal"]',
        '.overlay-close', '.lightbox-close',
        # ZARA specific close buttons
        '.zara-modal-close', '.modal-backdrop',
        # Generic close patterns
        'button:contains("Close")', 'button:contains("Ã—")', 
        'a:contains("Close")', '[title="Close"]',
        # ESC key simulation
        '.modal.show', '.modal.in'  # For backdrop click
    ]
    
    for selector in close_selectors:
        try:
            if ':contains(' in selector:
                # Handle text-based selectors with XPath
                text = selector.split(':contains("')[1].split('")')[0]
                tag = selector.split(':contains(')[0]
                xpath = f"//{tag}[contains(text(), '{text}')]"
                close_elements = driver.find_elements(By.XPATH, xpath)
            else:
                close_elements = driver.find_elements(By.CSS_SELECTOR, selector)
            
            for close_btn in close_elements:
                if close_btn.is_displayed():
                    try:
                        close_btn.click()
                        time.sleep(1)
                        print(f"Closed popup using selector: {selector}")
                        return True
                    except:
                        # Try JavaScript click
                        try:
                            driver.execute_script("arguments[0].click();", close_btn)
                            time.sleep(1)
                            print(f"Closed popup with JS using selector: {selector}")
                            return True
                        except:
                            continue
                            
        except Exception as e:
            continue
    
    # Try pressing ESC key as fallback
    try:
        from selenium.webdriver.common.keys import Keys
        from selenium.webdriver.common.action_chains import ActionChains
        ActionChains(driver).send_keys(Keys.ESCAPE).perform()
        time.sleep(1)
        print("Closed popup using ESC key")
        return True
    except:
        pass
    
    return False

def scrape_static_product(url, config):
    """Fallback: scrape product info using regular HTTP requests"""
    try:
        response = make_request(url, config)
        soup = parse_html(response)
        
        result = {
            'success': True,
            'url': url,
            'scraped_at': datetime.now().isoformat(),
            'method': 'static',
            'page_title': get_page_metadata(soup, url)['page_title'],
            'product_info': {}
        }
        
        # Extract product information
        product_info = extract_static_product_info(soup)
        product_info['interactive_content_found'] = False
        
        result['product_info'] = clean_product_info(product_info)
        
        return result
        
    except Exception as e:
        return {
            'success': False,
            'error': f'Static scraping error: {str(e)}',
            'method': 'static'
        }

@app.route('/debug_scrape', methods=['POST'])
@cross_origin()
def debug_scrape():
    """Debug endpoint to see what's actually happening during scraping"""
    try:
        data = request.get_json()
        url = data.get('url', '').strip()
        
        if not url:
            return jsonify({
                'success': False,
                'error': 'URL is required'
            }), 400
        
        driver = None
        try:
            print(f"Starting debug scrape for: {url}")
            
            # Setup driver in non-headless mode for debugging
            driver = setup_driver(headless=False)
            driver.set_page_load_timeout(30)
            
            # Step 1: Load page
            print("Step 1: Loading page...")
            driver.get(url)
            time.sleep(3)
            
            # Step 2: Handle popups
            print("Step 2: Handling popups...")
            popup_handled = handle_cookie_consent_and_popups(driver)
            print(f"Popup handled: {popup_handled}")
            
            # Step 3: Wait for page to be ready
            print("Step 3: Waiting for page to be ready...")
            wait = WebDriverWait(driver, 10)
            wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
            time.sleep(2)
            
            # Step 4: Get page source and basic info
            print("Step 4: Getting page info...")
            page_title = driver.title
            page_source_length = len(driver.page_source)
            
            # Step 5: Parse with BeautifulSoup and look for product content
            print("Step 5: Parsing with BeautifulSoup...")
            soup = BeautifulSoup(driver.page_source, 'html.parser')
            
            # Check for common product description selectors
            description_selectors = [
                '.product-description', '.pdp-description', '.product-details-description',
                '[data-testid="product-description"]', '.product-info-description',
                '.product-long-description', '.description-content', '.product-content',
                '.product-description-text', '.item-description', '.product-summary',
                '.product-details-content', '.product-information', '#product-description',
                # Add more generic selectors
                '.description', '.details', '.overview', '.about', '.info'
            ]
            
            found_descriptions = {}
            for selector in description_selectors:
                elements = soup.select(selector)
                if elements:
                    for i, element in enumerate(elements[:3]):  # Check first 3 matches
                        text = element.get_text(strip=True)
                        if text and len(text) > 20:  # Only meaningful text
                            found_descriptions[f"{selector}_{i}"] = {
                                'selector': selector,
                                'text_length': len(text),
                                'text_preview': text[:200] + '...' if len(text) > 200 else text
                            }
            
            # Check for any element with "description" in class or id
            description_elements = soup.find_all(attrs={'class': lambda x: x and 'description' in ' '.join(x).lower()})
            description_elements.extend(soup.find_all(attrs={'id': lambda x: x and 'description' in x.lower()}))
            
            generic_descriptions = {}
            for i, element in enumerate(description_elements[:5]):
                text = element.get_text(strip=True)
                if text and len(text) > 20:
                    generic_descriptions[f"generic_{i}"] = {
                        'classes': element.get('class', []),
                        'id': element.get('id', ''),
                        'tag': element.name,
                        'text_length': len(text),
                        'text_preview': text[:200] + '...' if len(text) > 200 else text
                    }
            
            # Check for paragraphs that might contain product info
            paragraphs = soup.find_all('p')
            long_paragraphs = []
            for p in paragraphs:
                text = p.get_text(strip=True)
                if len(text) > 100:  # Look for substantial paragraphs
                    long_paragraphs.append({
                        'text_length': len(text),
                        'text_preview': text[:150] + '...' if len(text) > 150 else text,
                        'classes': p.get('class', []),
                        'parent_classes': p.parent.get('class', []) if p.parent else []
                    })
            
            # Count different element types
            element_counts = {
                'total_paragraphs': len(soup.find_all('p')),
                'total_divs': len(soup.find_all('div')),
                'total_spans': len(soup.find_all('span')),
                'elements_with_description_class': len(description_elements),
                'long_paragraphs': len(long_paragraphs)
            }
            
            return jsonify({
                'success': True,
                'url': url,
                'page_title': page_title,
                'page_source_length': page_source_length,
                'popup_handled': popup_handled,
                'element_counts': element_counts,
                'found_descriptions': found_descriptions,
                'generic_descriptions': generic_descriptions,
                'long_paragraphs': long_paragraphs[:3],  # First 3 only
                'debug_info': {
                    'selectors_tested': len(description_selectors),
                    'soup_parsed': True,
                    'driver_title': page_title
                }
            })
            
        except Exception as e:
            return jsonify({
                'success': False,
                'error': f'Debug scraping error: {str(e)}'
            }), 500
        finally:
            if driver:
                try:
                    driver.quit()
                except:
                    pass
                    
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Request error: {str(e)}'
        }), 500

def clean_product_info(product_info):
    """Clean and format the extracted product information"""
    cleaned = {}
    
    for key, value in product_info.items():
        if value is None or value == '' or value == []:
            cleaned[key] = None
            continue
        
        if isinstance(value, str):
            # Clean whitespace
            value = ' '.join(value.split())
            
            # Remove HTML artifacts
            value = re.sub(r'<[^>]+>', '', value)
            
            # Clean up common text artifacts
            value = value.replace('\n', ' ').replace('\r', ' ')
            value = re.sub(r'\s+', ' ', value).strip()
            
            if len(value) > 10:  # Only keep meaningful text
                cleaned[key] = value
            else:
                cleaned[key] = None
        
        elif isinstance(value, dict):
            # For table data, keep as is but clean text
            if 'headers' in value and 'rows' in value:
                cleaned_table = {
                    'headers': [h.strip() for h in value['headers'] if h.strip()],
                    'rows': []
                }
                for row in value['rows']:
                    cleaned_row = [cell.strip() for cell in row if cell.strip()]
                    if cleaned_row:
                        cleaned_table['rows'].append(cleaned_row)
                
                if cleaned_table['headers'] or cleaned_table['rows']:
                    cleaned[key] = cleaned_table
                else:
                    cleaned[key] = None
            else:
                cleaned[key] = value
        
        else:
            cleaned[key] = value
    
    return cleaned

def generate(selected_chunks, query):
    client = openai.OpenAI()
    context = "\n\n".join(selected_chunks) 
    prompt = f"Answer the following query based on the provided text:\n\n{context}\n\nQuery: {query}\nAnswer:" 
    # response = client.chat.completions.create( 
    #     model="gpt-4", 
    #     messages=[ {"role": "system", "content": "You are a legal research and reasoning assistant trained in Indian income tax law, especially capital gains exemptions under the Income Tax Act. Your job is to analyze a user's scenario, determine applicability of specific sections (like Section 54F), and generate responses following a clear structure: Start with statutory interpretation â€” quote the relevant section (e.g., Section 54F) and clearly list the conditions in bullet points. Apply the law to the userâ€™s case â€” mention whether conditions are satisfied and explain eligibility for exemption. Cite relevant case law in support of the position taken. Choose cases that match the factual scenario and jurisdiction where possible. Include citation (e.g., ITA 4012/Mum/2023 - Abdul Nayab Shaikh). Quote only favourable rulings unless otherwise requested. Prefer recent, relevant, and jurisdictionally appropriate cases. Discuss any common exceptions or judicial deviations â€” e.g., benefit being allowed even when more than one residential unit is purchased, especially if adjacent or used as a single unit. Quote examples from case law or factual scenarios to support the interpretation or exception. Keep the examples precise and relevant. Format your response in a professional, advisory tone suitable for a tax consultantâ€™s opinion. Do not speculate â€” rely only on clear statutory provisions, circulars, and judicial precedents."}, {"role": "user", "content": prompt} ], 
    #     max_tokens=400, 
    #     temperature=0.1 ) 

    # response = client.chat.completions.create( 
    #     model="gpt-4", 
    #     messages=[ {"role": "system", "content": "You are a professional skills extractor"}, {"role": "user", "content": prompt} ], 
    #     max_tokens=400, 
    #     temperature=0.1 )
    # 

    # response = client.chat.completions.create( 
    #     model="gpt-4", 
    #      messages=[ {"role": "system", "content": "You are a PhD level research assistant that understands AI and its future and you also have a strong business acumen that will help you build a strong pitch for an AI startup"}, {"role": "user", "content": prompt} ], 
    #     max_tokens=400, 
    #     temperature=0.1 ) 

    response = client.chat.completions.create( 
    model="gpt-4", 
    messages=[ 
        {
            "role": "system", 
            "content": "You are a PhD-level research assistant with deep expertise in cryptocurrency markets, blockchain technology, and financial analysis. You provide insightful, data-driven analysis of crypto assets, including market trends, tokenomics, risk factors, and trading strategies."
        }, 
        {
            "role": "user", 
            "content": prompt
        } 
    ], 
    max_tokens=400, 
    temperature=0.1 
)
    
    
    print(response)
    answer = response.choices[0].message.content 
    # usage = response.usage
    return answer

@app.route('/upload', methods=['POST'])
def upload_file():
    # Accept folder_name from form data (for file uploads)
    folder_name = request.form.get('folder_name', '').strip()

    if 'file' not in request.files:
        return jsonify({"error": "No file part in the request"}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({"error": "No file selected for uploading"}), 400

    # Define the base upload directory
    base_upload_folder = "/Users/rhishikeshthakur/Enable/Software Development/enable_agents/data/uploaded_data"

    # If folder_name is provided, create/use subfolder
    if folder_name:
        upload_folder = os.path.join(base_upload_folder, folder_name)
    else:
        upload_folder = base_upload_folder

    os.makedirs(upload_folder, exist_ok=True)  # Create the directory if it doesn't exist

    # Check if the file already exists
    file_path = os.path.join(upload_folder, file.filename)
    if os.path.exists(file_path):
        return jsonify({"message": "File already exists", "file_path": file_path}), 200

    # Save the file
    file.save(file_path)

    return jsonify({"message": "File uploaded successfully", "file_path": file_path}), 200


@app.route('/scrape_product_info', methods=['POST'])
@cross_origin()
def scrape_product_info():
    """
    Enhanced product scraping API focused on descriptions and size information
    with interactive element handling
    """
    try:
        data = request.get_json()
        url = data.get('url', '').strip()
        interactive = data.get('interactive', True)  # Use interactive mode by default
        headless = data.get('headless', True)
        config = data.get('config', {})
        
        if not url:
            return jsonify({
                'success': False,
                'error': 'URL is required'
            }), 400
        
        # Validate URL
        try:
            parsed = urlparse(url)
            if not all([parsed.scheme, parsed.netloc]):
                return jsonify({
                    'success': False,
                    'error': 'Invalid URL format'
                }), 400
        except Exception as e:
            return jsonify({
                'success': False,
                'error': f'Invalid URL: {str(e)}'
            }), 400
        
        if interactive:
            # Use Selenium for interactive scraping
            result = scrape_with_interaction(url, headless, config)
        else:
            # Use regular HTTP scraping
            result = scrape_static_product(url, config)
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Scraping error: {str(e)}'
        }), 500


@app.route('/scrape', methods=['POST'])
@cross_origin()
def universal_scraper():
    """
    Universal scraping API - now includes enhanced product scraping
    """
    try:
        data = request.get_json()
        url = data.get('url', '').strip()
        scrape_type = data.get('type', 'basic')
        config = data.get('config', {})
        
        # Validation
        if not url:
            return jsonify({
                'success': False,
                'error': 'URL is required'
            }), 400
        
        # Validate URL format
        try:
            parsed = urlparse(url)
            if not all([parsed.scheme, parsed.netloc]):
                return jsonify({
                    'success': False,
                    'error': 'Invalid URL format. Include http:// or https://'
                }), 400
        except Exception as e:
            return jsonify({
                'success': False,
                'error': f'Invalid URL: {str(e)}'
            }), 400
        
        # Route to appropriate scraper based on type
        if scrape_type == 'product_info':
            # Use the new enhanced product scraper - but call it properly
            data['interactive'] = data.get('interactive', True)
            data['headless'] = data.get('headless', True)
            if data.get('interactive', True):
                result = scrape_with_interaction(url, data.get('headless', True), config)
            else:
                result = scrape_static_product(url, config)
            return jsonify(result)
        elif scrape_type == 'basic':
            result = scrape_basic_content(url, config)
            return jsonify(result)
        elif scrape_type == 'text':
            result = scrape_text_content(url, config)
            return jsonify(result)
        elif scrape_type == 'json_ld':
            result = scrape_json_ld_content(url, config)
            return jsonify(result)
        elif scrape_type == 'product':
            result = scrape_product_content(url, config)
            return jsonify(result)
        elif scrape_type == 'tables':
            result = scrape_table_content(url, config)
            return jsonify(result)
        elif scrape_type == 'custom':
            result = scrape_custom_selectors(url, config)
            return jsonify(result)
        else:
            return jsonify({
                'success': False,
                'error': f'Unknown scrape type: {scrape_type}'
            }), 400
            
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Scraping error: {str(e)}'
        }), 500

@app.route('/search_suggestions', methods=['POST'])
@cross_origin()
def openai_chat():
    """Simple OpenAI chat endpoint for matchmaking"""
    try:
        data = request.get_json()
        prompt = data.get('prompt', '')
        max_tokens = data.get('max_tokens', 500)
        temperature = data.get('temperature', 0.7)
        
        client = openai.OpenAI()
        client.api_key = os.environ['OPENAI_API_KEY']
        
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=max_tokens,
            temperature=temperature
        )
        
        return jsonify({
            'success': True,
            'response': response.choices[0].message.content.strip()
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

# @app.route('/rag_test', methods=['GET'])
# def rag_test():
#     query = request.args.get('query')
#     file_name = request.args.get('file_name')

#     upload_folder = "/Users/rhishikeshthakur/Enable/Software Development/enable_agents/data"
#     file_path = os.path.join(upload_folder, file_name)

#     # file_size = os.path.getsize(file_path)  # in bytes
    
#     if not os.path.exists(file_path):
#         return jsonify({"error": f"File '{file_name}' not found"}), 404

#     print(f"Absolute file path: {file_path}")
    
#     openai.api_key = get_credentials();

#     file_hash = get_file_hash(file_path)

#     if file_hash in cache:
#         print(f"Using cached embeddings for file hash: {file_hash}")
#         index, phrase_embeddings, page_chunks = load_embeddings(file_hash)
#     else:
#         print(f"Processing file and saving embeddings for file hash: {file_hash}")
#         pdf_doc = pdf_loader(file_path)
#         page_chunks = pdf_splitter(pdf_doc)

#         # print(page_chunks)

#         page_phrases = extract_keywords_from_pdf(pdf_doc)
#         chunk_phrases = extract_keywords_from_chunks(page_chunks)
        
#         index, phrase_embeddings = store_embeddings(page_phrases, chunk_phrases)
    
#         cache[file_hash] = (index, phrase_embeddings, page_chunks)
#         save_embeddings(file_hash, index, phrase_embeddings, page_chunks)
#         print(save_embeddings)
        
#     query_phrases = extract_phrases_from_query(query)
#     query_embeddings = get_embeddings_for_query(query_phrases)
#     selected_chunks = retrieve_similar_chunks(query_embeddings, index, phrase_embeddings, page_chunks)

#     max_chunks = 5
#     max_chunk_length = 1000  # characters
#     selected_chunks = [chunk[:max_chunk_length] for chunk in selected_chunks[:max_chunks]]


#     answer = generate(selected_chunks, query)


#     return jsonify({"answer": answer})  # Always return a JSON object

@app.route('/rag_test', methods=['POST'])
def rag_test():
    try:
        data = request.get_json()
        query = data.get('query')
        file_name = data.get('file_name')
        data_store = data.get('data_store')  # This is the folder name under uploaded_data

        if not query or not file_name:
            return jsonify({"error": "Missing query or file_name parameter"}), 400

        base_upload_folder = "/Users/rhishikeshthakur/Enable/Software Development/enable_agents/data/uploaded_data"
        if data_store:
            file_path = os.path.join(base_upload_folder, data_store, file_name)
        else:
            file_path = os.path.join(base_upload_folder, file_name)

        if not os.path.exists(file_path):
            return jsonify({"error": f"File '{file_path}' not found"}), 404

        print(f"Processing file: {file_path}")
        openai.api_key = get_credentials()
        file_hash = get_file_hash(file_path)

        try:
            if file_hash in cache:
                print(f"Using cached embeddings for file hash: {file_hash}")
                index, phrase_embeddings, page_chunks = load_embeddings(file_hash)
            else:
                print(f"Processing file and saving embeddings for file hash: {file_hash}")
                pdf_doc = pdf_loader(file_path)
                
                if not pdf_doc:
                    return jsonify({"error": "Could not extract text from PDF"}), 400
                
                page_chunks = pdf_splitter(pdf_doc)
                
                if not page_chunks:
                    return jsonify({"error": "Could not create chunks from PDF content"}), 400
                
                page_phrases = extract_keywords_from_pdf(pdf_doc)
                chunk_phrases = extract_keywords_from_chunks(page_chunks)
                
                index, phrase_embeddings = store_embeddings(page_phrases, chunk_phrases)
                cache[file_hash] = (index, phrase_embeddings, page_chunks)
                save_embeddings(file_hash, index, phrase_embeddings, page_chunks)

            query_phrases = extract_phrases_from_query(query)
            if not query_phrases:
                query_phrases = [query]  # Use the full query if no phrases extracted
            
            query_embeddings = get_embeddings_for_query(query_phrases)
            selected_chunks = retrieve_similar_chunks(query_embeddings, index, phrase_embeddings, page_chunks)

            max_chunks = 5
            max_chunk_length = 1000  # characters
            selected_chunks = [chunk[:max_chunk_length] for chunk in selected_chunks[:max_chunks]]

            answer = generate(selected_chunks, query)

            return jsonify({
                "answer": answer,
                "chunks_used": len(selected_chunks),
                "file_processed": file_name
            })
            
        except Exception as processing_error:
            print(f"Error processing PDF: {processing_error}")
            return jsonify({"error": f"Error processing PDF: {str(processing_error)}"}), 500

    except Exception as e:
        print(f"RAG test error: {e}")
        return jsonify({"error": f"Server error: {str(e)}"}), 500

@app.route('/enterprise_chat', methods=['POST', 'OPTIONS'])
@cross_origin()
def enterprise_chat():
    """
    Enterprise chat API that collects business context from the user.
    User can answer any of the 5 key questions first; the API will detect which question was answered,
    store it, and then ask the remaining unanswered questions.
    Once all questions are answered, sends chat_state to OpenAI to generate a summary and auto-fill missing fields.
    """
    data = request.get_json()
    chat_state = data.get('chat_state', {})
    last_answer = data.get('last_answer', '').strip()
    last_question_key = data.get('last_question_key', '').strip()

    # Define the sequence and mapping of questions
    questions = [
        {"key": "industry", "question": "To get a bit of context, which industry does your business operate in?"},
        {"key": "role_department", "question": "Whatâ€™s your role within the company, and what is your department mainly focused on right now?"},
        {"key": "tools", "question": "What tools or software do you and your team rely on most, and what do you use them for?"},
        {"key": "business_need", "question": "If you could change or improve one thing about how your team works today, what would it be?"}
    ]
    question_keys = [q['key'] for q in questions]

    # If the user answered a question, store it in chat_state
    if last_question_key in question_keys and last_answer:
        chat_state[last_question_key] = last_answer

    # If the user sent an answer but didn't specify which question, try to infer
    if not last_question_key and last_answer:
        for q in questions:
            if q['key'] not in chat_state or not chat_state.get(q['key']):
                chat_state[q['key']] = last_answer
                break

    # Find the next unanswered question
    for q in questions:
        if q['key'] not in chat_state or not chat_state[q['key']]:
            return jsonify({
                "success": True,
                "next_question": q['question'],
                "next_question_key": q['key'],
                "chat_state": chat_state,
                "completed": False
            })

    # If all questions answered, auto-fill and format chat_state using OpenAI
    try:
        import openai
        openai.api_key = os.environ.get("OPENAI_API_KEY")
        # Prompt to format and auto-fill chat_state
        autofill_prompt = (
            "Given the following user answers, format the business context as a JSON object with these keys: "
            "industry, role, department_context, business_need, and tools (as a list of objects with tool_name and description). "
            "If any field is missing or vague, infer and auto-fill it based on the other answers. "
            "Example format:\n"
            "{\n"
            '  "tools": [\n'
            '    {"tool_name": "Slack", "description": "Team communication and collaboration platform"},\n'
            '    {"tool_name": "Salesforce", "description": "CRM for managing customer relationships and sales pipeline"}\n'
            "  ],\n"
            '  "industry": "Technology",\n'
            '  "role": "Sales Manager",\n'
            '  "department_context": "Our sales department is focused on improving lead conversion and automating reporting.",\n'
            '  "business_need": "We want to integrate our communication and CRM tools, automate sales reporting, and identify missing modules for analytics."\n'
            "}\n"
            "User answers:\n"
            f"Industry: {chat_state.get('industry', '')}\n"
            f"Role: {chat_state.get('role', '')}\n"
            f"Department Context: {chat_state.get('department_context', '')}\n"
            f"Tools: {chat_state.get('tools', '')}\n"
            f"Business Need: {chat_state.get('business_need', '')}\n"
            "Return only valid JSON."
        )

        client = openai.OpenAI()
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a business analyst assistant."},
                {"role": "user", "content": autofill_prompt}
            ],
            max_tokens=300,
            temperature=0.2
        )
        # Extract JSON from response
        import re
        raw_content = response.choices[0].message.content.strip()
        match = re.search(r'\{[\s\S]*\}', raw_content)
        if match:
            formatted_state = json.loads(match.group(0))
        else:
            formatted_state = chat_state  # fallback

        # Summarize the context for search_summary
        summary_prompt = (
            "Summarize the following business context in 2-3 sentences for agent recommendation:\n\n"
            f"{json.dumps(formatted_state, indent=2)}"
        )
        summary_response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a business analyst assistant."},
                {"role": "user", "content": summary_prompt}
            ],
            max_tokens=150,
            temperature=0.3
        )
        search_summary = summary_response.choices[0].message.content.strip()
        print(formatted_state)

    except Exception as e:
        formatted_state = chat_state
        search_summary = f"Could not generate summary: {str(e)}"

    return jsonify({
        "success": True,
        "message": "Thank you! Here is the summary of your business context.",
        "chat_state": formatted_state,
        "completed": True,
        "search_summary": search_summary
    })


@app.route('/chat_api', methods=['POST', 'OPTIONS'])
@cross_origin()  # Allow CORS for this endpoint
def chat_api():
    import glob
    data = request.get_json()
    query = data.get('query')

    embeddings_folder = "/Users/rhishikeshthakur/Enable/Software Development/enable_agents/data/embeddings"
    embedding_files = glob.glob(os.path.join(embeddings_folder, "*_index.pkl"))
    file_hashes = [os.path.basename(f).split('_')[0] for f in embedding_files]

    if not file_hashes:
        return jsonify({"error": "No knowledge base available. Please upload and process at least one file first."}), 400

    # Aggregate all embeddings, phrase mappings, and page chunks
    all_indexes = []
    all_phrase_embeddings = {}
    all_page_chunks = {}

    for file_hash in file_hashes:
        try:
            index, phrase_embeddings, page_chunks = load_embeddings(file_hash)
        except Exception as e:
            print(f"Error loading embeddings for {file_hash}: {e}")
            continue  # Skip files that can't be loaded

        all_indexes.append(index)
        # Update keys to include file_hash for uniqueness
        for (page, chunk_number), phrases in phrase_embeddings.items():
            all_phrase_embeddings[(file_hash, page, chunk_number)] = phrases
        for page, chunks in page_chunks.items():
            all_page_chunks[(file_hash, page)] = chunks

    # For simplicity, use the first index (FAISS) for searching, or you can merge indexes if needed
    index = all_indexes[0] if all_indexes else None
    phrase_embeddings = all_phrase_embeddings
    page_chunks = all_page_chunks

    if not index or not phrase_embeddings or not page_chunks:
        return jsonify({"error": "No valid embeddings found."}), 400

    query_phrases = extract_phrases_from_query(query)
    query_embeddings = get_embeddings_for_query(query_phrases)
    selected_chunks = retrieve_similar_chunks(query_embeddings, index, phrase_embeddings, page_chunks)

    max_chunks = 5
    max_chunk_length = 1000  # characters
    selected_chunks = [chunk[:max_chunk_length] for chunk in selected_chunks[:max_chunks]]

    answer = generate(selected_chunks, query)

    return jsonify({"answer": answer})

@app.route('/save-prompt', methods=['POST'])
def save_prompt():
    data = request.get_json()
    prompt_id = str(uuid4())  # Generate a unique ID for the prompt
    data['id'] = prompt_id
    data['timestamp'] = datetime.now().isoformat()

    # Load existing prompts
    if os.path.exists(PROMPTS_FILE):
        with open(PROMPTS_FILE, 'r') as file:
            prompts = json.load(file)
    else:
        prompts = []

    # Add the new prompt
    prompts.append(data)

    # Save back to the file
    with open(PROMPTS_FILE, 'w') as file:
        json.dump(prompts, file, indent=4)

    return jsonify({"message": "Prompt saved successfully", "id": prompt_id})

@app.route('/previous-prompts', methods=['GET'])
def previous_prompts():
    if os.path.exists(PROMPTS_FILE):
        with open(PROMPTS_FILE, 'r') as file:
            prompts = json.load(file)
    else:
        prompts = []

    return jsonify({"prompts": prompts})

@app.route('/yfinance', methods=['POST'])
def yfinance_test():
    data = request.get_json()
    symbol = data.get('stock')
    region = data.get('region')

    if not symbol or not region:
        return "Missing required parameters: 'stock' and 'region'", 400

    conn = http.client.HTTPSConnection("yahoo-finance166.p.rapidapi.com")

    headers = {
        'x-rapidapi-key': "95cdd43379mshbd9483856442c47p1c2782jsn897449ebefb8",
        'x-rapidapi-host': "yahoo-finance166.p.rapidapi.com"
    }

    endpoint = f"/api/stock/get-financial-data?region={region}&symbol={symbol}"
    print(f"Requesting data from endpoint: {endpoint}")  # Debug statement
    conn.request("GET", endpoint, headers=headers)

    res = conn.getresponse()
    data = res.read()
    json_data = json.loads(data.decode("utf-8"))

    print(json_data)  # Debug statement to print the entire response

    if 'quoteSummary' not in json_data or 'result' not in json_data['quoteSummary'] or not json_data['quoteSummary']['result']:
        return jsonify({"error": "No data found for the given stock symbol and region"}), 404

    current_price = json_data['quoteSummary']['result'][0]['financialData']['currentPrice']['fmt']
    operating_margins = json_data['quoteSummary']['result'][0]['financialData']['operatingMargins']['fmt']
    netprofit_margins = json_data['quoteSummary']['result'][0]['financialData']['profitMargins']['fmt']
    gross_margins = json_data['quoteSummary']['result'][0]['financialData']['grossMargins']['fmt']
    revenue_growth = json_data['quoteSummary']['result'][0]['financialData']['revenueGrowth']['fmt']
    debt_to_equity = json_data['quoteSummary']['result'][0]['financialData']['debtToEquity']['fmt']
    quick_ratio = json_data['quoteSummary']['result'][0]['financialData']['quickRatio']['fmt']
    current_ratio = json_data['quoteSummary']['result'][0]['financialData']['currentRatio']['fmt']
    analyst_recommendation = json_data['quoteSummary']['result'][0]['financialData']['recommendationKey']
    number_of_analysts = json_data['quoteSummary']['result'][0]['financialData']['numberOfAnalystOpinions']['fmt']
    target_high_price = json_data['quoteSummary']['result'][0]['financialData']['targetHighPrice']['fmt']
    target_low_price = json_data['quoteSummary']['result'][0]['financialData']['targetLowPrice']['fmt']
    target_mean_price = json_data['quoteSummary']['result'][0]['financialData']['targetMeanPrice']['fmt']
    target_median_price = json_data['quoteSummary']['result'][0]['financialData']['targetMedianPrice']['fmt']

    financial_KPIs = {
        "current_price": current_price,
        "operating margin": operating_margins,
        "netprofit_margins": netprofit_margins,
        "gross_margins": gross_margins,
        "revenue_growth": revenue_growth,
        "debt_to_equity": debt_to_equity,
        "quick_ratio": quick_ratio,
        "current_ratio": current_ratio,
        "number_of_analysts": number_of_analysts,
        "analyst_recommendation": analyst_recommendation,
        "target_high_price": target_high_price,
        "target_low_price": target_low_price,
        "target_mean_price": target_mean_price,
        "target_median_price": target_median_price
    }

    return jsonify(financial_KPIs)

@app.route('/generate-requirements', methods=['POST'])
def generate_requirements():
    openai.api_key = get_credentials()

    data = request.get_json()
    overview = data.get('overview', '')
    context = data.get('context', '')  # Get the context from the payload
    country = data.get('countries', '')
    industries = data.get('industries', '')
    function = data.get('businessFunction', '')
    frameworks = data.get('frameworks', [])

    format = data.get('responseFormat', '')
    

    # prompt = f"""
    # Draft requirements based on the requirements {overview} that are specific, measurable, achievable, relevant, and time-bound (SMART).
    # Consider {context} as context for the requirements being asked for, focus on the market in {country} or region, 
    # consider {industries} for industry related insights, consider {function} as the role or business function of the requester,
    # and without mentioning the framework in the final response, conduct research taking into account these analysis frameworks: {frameworks} for one valuable and rare resource each using the VRIO, market forces for and against the startup using PESTLE, and product readiness using Mckinsey's 3 Horizon and use response format as reference: {format}.
    # """

    prompt = f"""
    You are a research assistant tasked with producing high-quality, insightful, and well-structured research on business opportunitiesand growth prospects. Your output should include a curated but accessible for free list of relevant academic papers, industry articles, expert quotations, market data, and other authoritative sources.

    Base your research on the following core requirement: {overview}.

    In addition, factor in the following contextual details where applicable:

    Geographic Market: Consider the business and technology landscape in {country}. Ignore if not specified.

    Industry Focus: Include insights, trends, and data from the following industries: {industries}. Ignore if not specified.

    Business Function: Tailor the analysis to the perspective or needs of a person working in {function}. Ignore if not specified.

    Strategic Frameworks: Incorporate or structure your research using the following analytical frameworks: {frameworks}.

    Use the following format as a guide for structuring your response: {format}.

    Your response should:

    Include direct citations or links where available.

    Be clear, logically organized, and easy to turn into a pitch or slide deck.

    Blend both technical insight (e.g., emerging technologies, R&D frontiers) and business relevance (e.g., market sizing, customer pain points, competitive dynamics).
    """

    print(prompt)

    client = openai.OpenAI()

    try:
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a research assistant."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=2000,
            temperature=0.6
        )

        answer = response.choices[0].message.content
        return jsonify({"requirements": answer})
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/simple_search', methods=['POST'])
@cross_origin()
def simple_search():
    """Enhanced search endpoint that handles both regular searches and special commands"""
    try:
        data = request.get_json()
        user_query = data.get('query', '').strip()
        json_data = data.get('data', [])
        user_id = data.get('user_id', 'default_user')  # For favorites
        
        if not user_query:
            return jsonify({'success': False, 'error': 'No search query provided'}), 400
        
        # Parse the query using enhanced function
        parse_result = parse_simple_query_enhanced(user_query)
        
        if not parse_result['success']:
            return jsonify({
                'success': False,
                'error': 'Could not parse search query',
                'query': user_query,
                'openai_error': parse_result.get('error', 'Unknown error')
            }), 400
        
        # Handle special commands
        if parse_result.get('special_command', False):
            command_type = parse_result.get('command_type')
            
            if command_type == 'show_companies':
                if not json_data:
                    return jsonify({'success': False, 'error': 'No data available to list companies'}), 400
                
                # Extract unique companies from data
                companies = set()
                for record in json_data:
                    company = record.get('company') or record.get('Company') or record.get('organization')
                    if company and str(company).strip():
                        companies.add(str(company).strip())
                
                company_list = sorted(list(companies))
                return jsonify({
                    "success": True,
                    "type": "companies_list",
                    "total_found": len(company_list),
                    "results": [{"company": comp} for comp in company_list],
                    "query": user_query,
                    "keywords": parse_result['keywords'],
                    "command_type": command_type
                })
            
            elif command_type == 'show_titles':
                if not json_data:
                    return jsonify({'success': False, 'error': 'No data available to list titles'}), 400
                
                # Extract unique titles from data
                titles = set()
                for record in json_data:
                    title = (record.get('title') or record.get('Title') or 
                            record.get('Job title') or record.get('position'))
                    if title and str(title).strip():
                        titles.add(str(title).strip())
                
                title_list = sorted(list(titles))
                return jsonify({
                    "success": True,
                    "type": "titles_list",
                    "total_found": len(title_list),
                    "results": [{"title": title} for title in title_list],
                    "query": user_query,
                    "keywords": parse_result['keywords'],
                    "command_type": command_type
                })
            
            elif command_type == 'show_locations':
                if not json_data:
                    return jsonify({'success': False, 'error': 'No data available to list locations'}), 400
                
                # Extract unique locations from data
                locations = set()
                for record in json_data:
                    location = (record.get('location') or record.get('Location') or 
                               record.get('city') or record.get('City'))
                    if location and str(location).strip():
                        locations.add(str(location).strip())
                
                location_list = sorted(list(locations))
                return jsonify({
                    "success": True,
                    "type": "locations_list",
                    "total_found": len(location_list),
                    "results": [{"location": loc} for loc in location_list],
                    "query": user_query,
                    "keywords": parse_result['keywords'],
                    "command_type": command_type
                })
            
            elif command_type == 'show_skills':
                if not json_data:
                    return jsonify({'success': False, 'error': 'No data available to list skills'}), 400
                
                # Extract unique skills from data
                skills = set()
                for record in json_data:
                    skill_fields = (record.get('skills') or record.get('Skills') or 
                                   record.get('technologies') or record.get('required_skills'))
                    if skill_fields:
                        if isinstance(skill_fields, str):
                            # Split by common delimiters
                            import re
                            skill_list = re.split(r'[,;|]+', skill_fields)
                            for skill in skill_list:
                                clean_skill = skill.strip()
                                if clean_skill:
                                    skills.add(clean_skill)
                        elif isinstance(skill_fields, list):
                            skills.update([str(s).strip() for s in skill_fields if str(s).strip()])
                
                skill_list = sorted(list(skills))
                return jsonify({
                    "success": True,
                    "type": "skills_list",
                    "total_found": len(skill_list),
                    "results": [{"skill": skill} for skill in skill_list],
                    "query": user_query,
                    "keywords": parse_result['keywords'],
                    "command_type": command_type
                })
            
            elif command_type == 'show_favorites':
                # Load user favorites (doesn't need json_data)
                favorites_file = os.path.join(
                    '/Users/rhishikeshthakur/Enable/Software Development/enable_agents/data', 
                    'user_data', user_id, 'favorites.json'
                )
                
                if os.path.exists(favorites_file):
                    with open(favorites_file, 'r', encoding='utf-8') as f:
                        favorites = json.load(f)
                    
                    return jsonify({
                        "success": True,
                        "type": "favorites_list",
                        "total_found": len(favorites),
                        "results": favorites,
                        "query": user_query,
                        "keywords": parse_result['keywords'],
                        "command_type": command_type
                    })
                else:
                    return jsonify({
                        "success": True,
                        "type": "favorites_list",
                        "total_found": 0,
                        "results": [],
                        "query": user_query,
                        "keywords": parse_result['keywords'],
                        "command_type": command_type,
                        "message": "No favorites saved yet"
                    })
            
            else:
                return jsonify({
                    'success': False,
                    'error': f'Unknown special command: {command_type}'
                }), 400
        
        # Regular search logic (existing)
        if not json_data:
            return jsonify({'success': False, 'error': 'No data provided for search'}), 400
        
        if not parse_result['keywords']:
            return jsonify({
                'success': False,
                'error': 'Could not extract search keywords from query',
                'query': user_query,
                'suggestion': 'Try being more specific with company names, job titles, or skills'
            }), 400
        
        # Search the data using existing function
        results = simple_search_json(json_data, parse_result['keywords'])
        
        return jsonify({
            'success': True,
            'type': 'search_results',
            'query': user_query,
            'keywords': parse_result['keywords'],
            'results': results,
            'total_found': len(results),
            'phrases': parse_result.get('phrases', [])
        })
        
    except Exception as e:
        print(f"Search error: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Search error: {str(e)}'
        }), 500


@app.route('/save_user_favorite', methods=['POST'])
@cross_origin()
def save_user_favorite():
    """Save a profile to user's favorites"""
    try:
        data = request.json
        user_id = data.get('user_id', 'default_user')
        profile_data = data.get('profile_data')
        
        if not profile_data:
            return jsonify({"success": False, "error": "No profile data provided"}), 400
        
        # Create user_data directory if it doesn't exist
        user_data_dir = os.path.join('/Users/rhishikeshthakur/Enable/Software Development/enable_agents/data', 'user_data')
        os.makedirs(user_data_dir, exist_ok=True)
        
        # Create user-specific subdirectory
        user_dir = os.path.join(user_data_dir, user_id)
        os.makedirs(user_dir, exist_ok=True)
        
        # File path for user's favorites
        favorites_file = os.path.join(user_dir, 'favorites.json')
        
        # Load existing favorites or create new list
        favorites = []
        if os.path.exists(favorites_file):
            with open(favorites_file, 'r', encoding='utf-8') as f:
                favorites = json.load(f)
        
        # Add metadata to profile
        profile_with_meta = {
            **profile_data,
            'saved_at': datetime.now().isoformat(),
            'favorite_id': len(favorites) + 1
        }
        
        # Check if already exists (by name and company)
        full_name = f"{profile_data.get('name', '')} {profile_data.get('lastname', '')}".strip()
        existing = next((fav for fav in favorites 
                        if fav.get('full_name') == full_name and 
                           fav.get('company') == profile_data.get('company')), None)
        
        if existing:
            return jsonify({"success": False, "error": "Profile already in favorites"})
        
        # Add to favorites
        favorites.append(profile_with_meta)
        
        # Save updated favorites
        with open(favorites_file, 'w', encoding='utf-8') as f:
            json.dump(favorites, f, indent=2, ensure_ascii=False)
        
        return jsonify({
            "success": True,
            "message": "Profile saved to favorites",
            "favorites_count": len(favorites)
        })
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/get_user_favorites', methods=['POST'])
@cross_origin()
def get_user_favorites():
    """Get user's saved favorites"""
    try:
        data = request.json
        user_id = data.get('user_id', 'default_user')
        
        # Path to user's favorites file
        favorites_file = os.path.join('/Users/rhishikeshthakur/Enable/Software Development/enable_agents/data', 'user_data', user_id, 'favorites.json')
        
        if not os.path.exists(favorites_file):
            return jsonify({"success": True, "favorites": [], "count": 0})
        
        # Load favorites
        with open(favorites_file, 'r', encoding='utf-8') as f:
            favorites = json.load(f)
        
        return jsonify({
            "success": True,
            "favorites": favorites,
            "count": len(favorites)
        })
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/remove_user_favorite', methods=['POST'])
@cross_origin()
def remove_user_favorite():
    """Remove a profile from user's favorites"""
    try:
        data = request.json
        user_id = data.get('user_id', 'default_user')
        favorite_id = data.get('favorite_id')
        
        if not favorite_id:
            return jsonify({"success": False, "error": "No favorite_id provided"}), 400
        
        # Path to user's favorites file
        favorites_file = os.path.join('/Users/rhishikeshthakur/Enable/Software Development/enable_agents/data', 'user_data', user_id, 'favorites.json')
        
        if not os.path.exists(favorites_file):
            return jsonify({"success": False, "error": "No favorites file found"}), 404
        
        # Load favorites
        with open(favorites_file, 'r', encoding='utf-8') as f:
            favorites = json.load(f)
        
        # Remove the favorite
        original_count = len(favorites)
        favorites = [fav for fav in favorites if fav.get('favorite_id') != favorite_id]
        
        if len(favorites) == original_count:
            return jsonify({"success": False, "error": "Favorite not found"}), 404
        
        # Save updated favorites
        with open(favorites_file, 'w', encoding='utf-8') as f:
            json.dump(favorites, f, indent=2, ensure_ascii=False)
        
        return jsonify({
            "success": True,
            "message": "Profile removed from favorites",
            "favorites_count": len(favorites)
        })
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/save-requirements', methods=['POST'])
def save_requirements():
    data = request.get_json()
    requirements = data.get('requirements', [])
    export_option = data.get('exportOption', 'Unknown')  # Get the export option

    if not requirements:
        return jsonify({"error": "No requirements to save"}), 400

    # Define the folder path
    folder_path = "/Users/rhishikeshthakur/Enable/Software Development/enable_agents/data/requirements_versions"
    os.makedirs(folder_path, exist_ok=True)  # Create the folder if it doesn't exist

    # Create a unique file name with the export option and timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    file_path = os.path.join(folder_path, f"requirements_{export_option}_{timestamp}.txt")

    # Save the requirements to the file
    with open(file_path, "w") as file:
        file.write("\n".join(requirements))

    return jsonify({"message": f"Requirements saved successfully via {export_option}", "file_path": file_path})

@app.route('/register', methods=['POST'])
def register():
    data = request.get_json()
    username = data.get('email') # Use email as username
    password = data.get('password')
    first_name = data.get('first_name')
    last_name = data.get('last_name')
    email = data.get('email')
    company = data.get('company')
    linkedin = data.get('linkedin')
    short_intro = data.get('short_intro')
    company_intro = data.get('company_intro')

    if not password or not email:
        return jsonify({'error': 'Email and password required'}), 400
    if User.query.filter_by(username=username).first():
        return jsonify({'error': 'Username already exists'}), 400
    if User.query.filter_by(email=email).first():
        return jsonify({'error': 'Email already registered'}), 400

    try:
        hashed_password = generate_password_hash(password)
        user = User(
            username=username,
            password=hashed_password,
            first_name=first_name,
            last_name=last_name,
            email=email,
            company=company,
            linkedin=linkedin,
            short_intro=short_intro,
            company_intro=company_intro
        )
        db.session.add(user)
        db.session.commit()
        return jsonify({'message': 'User registered successfully'}), 201
    except Exception as e:
        print("Registration error:", e)  # Add this line
        return jsonify({'error': str(e)}), 500

@app.route('/login', methods=['POST'])
def login():
    data = request.get_json()
    email = data.get('email')
    password = data.get('password')

    if not email or not password:
        return jsonify({'error': 'Email and password required'}), 400

    user = User.query.filter_by(email=email).first()
    if user and check_password_hash(user.password, password):
        # Login via email, but still return the username for frontend session storage if needed
        return jsonify({'message': 'Login successful', 'username': user.username, 'email': user.email}), 200
    else:
        return jsonify({'error': 'Invalid email or password'}), 401

GOOGLE_CLIENT_CONFIG = {
    "web": {
        "client_id": GOOGLE_CLIENT_ID,
        "project_id": "enable-agents",
        "auth_uri": "https://accounts.google.com/o/oauth2/auth",
        "token_uri": "https://oauth2.googleapis.com/token",
        "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
        "client_secret": GOOGLE_CLIENT_SECRET,
        "redirect_uris": [GOOGLE_REDIRECT_URI]
    }
}

SCOPES = [
    "openid",
    "https://www.googleapis.com/auth/userinfo.email",
    "https://www.googleapis.com/auth/userinfo.profile",
    "https://www.googleapis.com/auth/gmail.send"
]

@app.route('/auth/google/start', methods=['GET'])
def google_auth_start():
    # Use direct URL generation to avoid PKCE code_verifier state issues
    params = {
        'client_id': GOOGLE_CLIENT_ID,
        'redirect_uri': GOOGLE_REDIRECT_URI,
        'response_type': 'code',
        'scope': ' '.join(SCOPES),
        'access_type': 'offline',
        'prompt': 'consent',
        'state': 'user_login_flow'
    }
    auth_url = f"https://accounts.google.com/o/oauth2/v2/auth?{urlencode(params)}"
    return jsonify({"auth_url": auth_url, "state": "user_login_flow"})


@app.route('/emails/send_via_gmail', methods=['POST'])
def send_via_gmail():
    data = request.get_json()
    email = data.get('user_email')
    to_email = data.get('to')
    subject = data.get('subject')
    body = data.get('body')
    
    if not all([email, to_email, subject, body]):
        return jsonify({'error': 'Missing required fields'}), 400
    
    token_record = GoogleOAuthToken.query.filter_by(username=email).first()
    if not token_record:
        return jsonify({'error': 'Google account not connected'}), 401
        
    creds = Credentials(
        token=token_record.token,
        refresh_token=token_record.refresh_token,
        token_uri=token_record.token_uri,
        client_id=token_record.client_id,
        client_secret=token_record.client_secret,
        scopes=token_record.scopes.split(',') if token_record.scopes else SCOPES
    )
    
    try:
        service = googleapiclient.discovery.build('gmail', 'v1', credentials=creds)
        message = EmailMessage()
        message.set_content(body)
        message['To'] = to_email
        message['From'] = email
        message['Subject'] = subject
        
        encoded_message = base64.urlsafe_b64encode(message.as_bytes()).decode()
        create_message = {
            'raw': encoded_message
        }
        
        send_message = (service.users().messages().send(userId="me", body=create_message).execute())
        return jsonify({'message': 'Email sent successfully via Gmail API', 'id': send_message['id']})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/file_to_json_convert', methods=['POST'])
def convert_file():
    """Main endpoint to convert CSV/XLSX files to JSON"""
    
    # Check if file is present in request
    if 'file' not in request.files:
        return jsonify({
            'success': False,
            'error': 'No file provided',
            'data': []
        }), 400
    
    file = request.files['file']
    
    # Check if file is selected
    if file.filename == '':
        return jsonify({
            'success': False,
            'error': 'No file selected',
            'data': []
        }), 400
    
    # Check if file type is allowed
    if not allowed_file(file.filename):
        return jsonify({
            'success': False,
            'error': 'File type not allowed. Please upload CSV or XLSX files only.',
            'data': []
        }), 400
    
    try:
        # Create temporary upload folder for conversion
        temp_folder = "/Users/rhishikeshthakur/Enable/Software Development/enable_agents/data/temp_conversion"
        os.makedirs(temp_folder, exist_ok=True)
        
        # Save uploaded file
        filename = secure_filename(file.filename)
        file_path = os.path.join(temp_folder, filename)
        file.save(file_path)
        
        # Get conversion options
        multiple_sheets = request.form.get('multiple_sheets', 'false').lower() == 'true'
        
        # Convert based on file type
        file_extension = filename.rsplit('.', 1)[1].lower()
        
        if file_extension == 'csv':
            result = csv_to_json(file_path)
        elif file_extension in ['xlsx', 'xls']:
            if multiple_sheets:
                result = xlsx_to_json_multiple_sheets(file_path)
            else:
                result = xlsx_to_json(file_path)
        
        # Clean up - remove uploaded file
        try:
            os.remove(file_path)
        except:
            pass
        
        return jsonify(result)
        
    except Exception as e:
        # Clean up file if error occurs
        try:
            if 'file_path' in locals():
                os.remove(file_path)
        except:
            pass
        
        return jsonify({
            'success': False,
            'error': f'Server error: {str(e)}',
            'data': []
        }), 500
    

@app.route('/enrich_with_openai', methods=['POST'])
@cross_origin()
def enrich_with_openai():
    """API endpoint to enrich JSON data with required skills using OpenAI"""
    try:
        request_data = request.get_json()
        
        if not request_data or 'data' not in request_data:
            return jsonify({
                'success': False,
                'error': 'No data provided in request body'
            }), 400
        
        json_data = request_data['data']
        
        if not isinstance(json_data, list) or len(json_data) == 0:
            return jsonify({
                'success': False,
                'error': 'Data must be a non-empty list of objects'
            }), 400
        
        # Enrich data using the new workflow
        result = enrich_json_with_openai(json_data)
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Server error: {str(e)}'
        }), 500
    

@app.route('/chrome_history', methods=['GET'])
def get_chrome_history():
    """API endpoint to get Chrome browser history with better error handling"""
    try:
        user_id = request.args.get('user_id', 'default_user')
        result = read_chrome_history_safe()
        
        if result['success']:
            save_tools_landscape_for_user(user_id, result)
            return jsonify(result)
        else:
            return jsonify(result), 400
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Server error: {str(e)}'
        }), 500

@app.route('/chrome_status', methods=['GET'])
def check_chrome_status():
    """Check if Chrome is running"""
    try:
        
        chrome_processes = []
        for proc in psutil.process_iter(['pid', 'name']):
            try:
                if 'chrome' in proc.info['name'].lower():
                    chrome_processes.append({
                        'pid': proc.info['pid'],
                        'name': proc.info['name']
                    })
            except (psutil.NoSuchProcess, psutil.AccessDenied):
                pass
        
        return jsonify({
            'success': True,
            'chrome_running': len(chrome_processes) > 0,
            'processes': chrome_processes
        })
        
    except ImportError:
        return jsonify({
            'success': False,
            'error': 'psutil not installed. Cannot check Chrome status.'
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        })
    
@app.route('/check_existing_file', methods=['POST'])
def check_existing_file():
    try:
        data = request.json
        file_name = data.get('file_name')
        new_file_size = data.get('new_file_size')
        
        # Create file path - Updated to use the correct data folder structure
        json_file_name = file_name.replace('.csv', '.json').replace('.xlsx', '.json').replace('.xls', '.json')
        # Use the same structure as upload function
        file_path = os.path.join('/Users/rhishikeshthakur/Enable/Software Development/enable_agents/data/uploaded_data', 'alumni_data', json_file_name)
        
        if os.path.exists(file_path):
            existing_size = os.path.getsize(file_path)
            
            # If new file is not significantly larger (less than 10% increase), skip processing
            size_threshold = existing_size * 1.1  # 10% increase threshold
            should_skip = new_file_size <= size_threshold
            
            return jsonify({
                'exists': True,
                'existing_size': existing_size,
                'should_skip': should_skip,
                'message': f'File exists. Size: {existing_size} bytes vs new: {new_file_size} bytes'
            })
        
        return jsonify({
            'exists': False,
            'should_skip': False,
            'message': 'File does not exist'
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/save_json_file', methods=['POST'])
def save_json_file():
    try:
        data = request.json
        json_data = data.get('data')
        file_name = data.get('file_name')
        folder_name = data.get('folder_name', 'alumni_data')
        
        # Create directory using the correct data folder structure
        folder_path = os.path.join('/Users/rhishikeshthakur/Enable/Software Development/enable_agents/data/uploaded_data', folder_name)
        os.makedirs(folder_path, exist_ok=True)
        
        # Save JSON file
        file_path = os.path.join(folder_path, file_name)
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, indent=2, ensure_ascii=False)
        
        return jsonify({
            'success': True,
            'file_path': file_path,
            'file_name': file_name,
            'message': f'JSON file saved successfully: {file_name}'
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/load_json_file', methods=['POST'])
def load_json_file():
    try:
        data = request.json
        file_name = data.get('file_name')
        folder_name = data.get('folder_name', 'alumni_data')
        
        # Use the correct data folder structure
        file_path = os.path.join('/Users/rhishikeshthakur/Enable/Software Development/enable_agents/data/uploaded_data', folder_name, file_name)
        
        if not os.path.exists(file_path):
            return jsonify({'success': False, 'error': 'File not found'}), 404
        
        with open(file_path, 'r', encoding='utf-8') as f:
            json_data = json.load(f)
        
        return jsonify({
            'success': True,
            'data': json_data,
            'message': f'JSON file loaded successfully: {file_name}'
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

def read_chrome_history_safe():
    """
    Improved: Read Chrome browser history and return only unique domains
    classified as tools where the user has logged in or has a subscription.
    """
    try:
        history_path = get_chrome_history_path()
        if not os.path.exists(history_path):
            return {
                'success': False,
                'error': 'Chrome history file not found. Make sure Chrome is installed.'
            }

        now = datetime.now()
        seven_days_ago = now - timedelta(days=7)
        webkit_epoch = datetime(1601, 1, 1)
        seven_days_ago_webkit = int((seven_days_ago - webkit_epoch).total_seconds() * 1000000)

        # Copy Chrome history file to temp location
        temp_dir = tempfile.mkdtemp()
        temp_history = os.path.join(temp_dir, 'History')
        shutil.copy2(history_path, temp_history)

        conn = sqlite3.connect(temp_history)
        cursor = conn.cursor()
        query = """
        SELECT url, title, visit_count, last_visit_time,
            datetime(last_visit_time/1000000 + (strftime('%s', '1601-01-01')), 'unixepoch', 'localtime') as visit_date
        FROM urls
        WHERE last_visit_time >= ?
        ORDER BY last_visit_time DESC
        LIMIT 2000
        """
        cursor.execute(query, (seven_days_ago_webkit,))
        rows = cursor.fetchall()
        conn.close()
        os.remove(temp_history)
        os.rmdir(temp_dir)

        # Filter URLs for login/subscription/dashboard/account/profile/settings
        login_keywords = [
            '/login', '/signin', '/dashboard', '/account', '/settings', '/profile', '/subscription', '/user', '/me'
        ]
        domain_map = {}
        filtered_history = []
        for row in rows:
            url = row[0]
            parsed = urlparse(url)
            domain = parsed.netloc.lower()
            path = parsed.path.lower()
            # Only consider URLs with login/subscription/dashboard/account/profile/settings or base domain
            if any(kw in path for kw in login_keywords) or path in ['', '/']:
                if domain not in domain_map:
                    domain_map[domain] = {
                        'domain': domain,
                        'sample_url': url,
                        'title': row[1] if row[1] else 'No Title',
                        'visit_count': row[2],
                        'last_visit_time': row[3],
                        'visit_date': row[4]
                    }

        # Classify domains using OpenAI (reuse your identify_saas_tools_with_openai)
        unique_domains = list(domain_map.values())
        # Prepare for OpenAI classification
        history_data_for_ai = [{'url': item['sample_url']} for item in unique_domains]
        tools_result = identify_saas_tools_with_openai(history_data_for_ai)
        for i, item in enumerate(unique_domains):
            if tools_result['success'] and i in tools_result['mapping']:
                mapping = tools_result['mapping'][i]
                item.update({
                    'is_tool': mapping.get('is_tool', False),
                    'tool_name': mapping.get('tool_name'),
                    'category': mapping.get('category'),
                    'tool_type': mapping.get('type'),
                    'description': mapping.get('description'),
                })
            else:
                item.update({
                    'is_tool': None,
                    'tool_name': None,
                    'category': None,
                    'tool_type': None,
                    'description': 'Tool analysis unavailable'
                })

        # Only return classified tools (where is_tool is True)
        classified_tools = [item for item in unique_domains if item.get('is_tool')]

        return {
            'success': True,
            'unique_domains': len(unique_domains),
            'tools_found': len(classified_tools),
            'data': classified_tools,
            'date_range': {
                'from': seven_days_ago.strftime('%Y-%m-%d'),
                'to': now.strftime('%Y-%m-%d')
            }
        }
    except Exception as e:
        return {
            'success': False,
            'error': f'Unexpected error reading Chrome history: {str(e)}'
        }

def save_tools_landscape_for_user(user_id, tools_data):
    """
    Save the result of /chrome_history API into a JSON file called
    'SaaS & tools landscape.json' under user_data for the given user.
    Adds a timestamp for when the data was last updated.
    """
    try:
        user_folder = os.path.join(
            '/Users/rhishikeshthakur/Enable/Software Development/enable_agents/data/user_data/tools_landscape/'
        )
        os.makedirs(user_folder, exist_ok=True)
        file_path = os.path.join(user_folder, 'tools_landscape.json')

        # Add/update timestamp
        tools_data['last_updated'] = datetime.now().isoformat()
        tools_data['user'] = user_id

        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(tools_data, f, ensure_ascii=False, indent=2)

        return {
            'success': True,
            'message': f'Tools landscape saved for {user_id}',
            'file_path': file_path,
            'last_updated': tools_data['last_updated']
        }
    except Exception as e:
        return {'success': False, 'error': str(e)}

@app.route('/get_tools_landscape', methods=['GET'])
@cross_origin()
def get_tools_landscape():
    """
    GET API to read tools landscape from tools_landscape.json and return
    a list of tools with tool_name, description, and category.
    """
    try:
        file_path = '/Users/rhishikeshthakur/Enable/Software Development/enable_agents/data/user_data/tools_landscape/tools_landscape.json'
        if not os.path.exists(file_path):
            return jsonify({'success': False, 'error': 'tools_landscape.json not found'}), 404

        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Extract tools info
        tools = []
        for item in data.get('data', []):
            if item.get('tool_name'):
                tools.append({
                    'tool_name': item.get('tool_name'),
                    'description': item.get('description'),
                    'category': item.get('category')
                })

        return jsonify({'success': True, 'tools': tools})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500
    
import openai


@app.route('/recommend_agents', methods=['POST'])
@cross_origin()
def recommend_agents():
    """
    Recommend agents/modules based on user's tools, industry/domain, role, department/company context, and business need.
    Input JSON payload:
    {
        "tools": [ {"tool_name": "ToolA", "description": "..."}, ... ],
        "industry": "...",
        "role": "...",
        "department_context": "...",  # e.g. 'My company/department is doing X and is responsible for Y'
        "business_need": "..."         # e.g. 'I want to track this business task and generate insights'
    }
    Output JSON:
    {
        "success": true,
        "recommendations": {
            "recommended_tools": [ {"tool_name": "...", "description": "...", "why_recommended": "..."}, ... ],
            "integration_pairs": [ {"tools": ["ToolA", "ToolB"], "integration": "...", "data_shared": "..."}, ... ],
            "additional_tools": [ {"tool_name": "...", "description": "...", "why_needed": "..."}, ... ]
        }
    }
    """
    try:
        openai.api_key = get_credentials()
        data = request.json
        tools = data.get('tools', [])
        industry = data.get('industry', '')
        role = data.get('role', '')
        department_context = data.get('department_context', '')
        business_need = data.get('business_need', '')

        # Load available modules (from agents_modules.json)
        with open('/Users/rhishikeshthakur/Enable/Software Development/enable_agents/data/agents_modules.json', 'r', encoding='utf-8') as f:
            modules = json.load(f)

        # Prepare context for OpenAI
        context = {
            "tools": tools,
            "industry": industry,
            "role": role,
            "department_context": department_context,
            "business_need": business_need,
            "available_modules": modules
        }
        prompt = (
            "You are a technology consultant. Based on the following user context, "
            "recommend a set of software modules (tools/agents) that can cater to the business need, "
            "considering the existing tools, missing necessary tools, and possible integrations. "
            "For each recommendation, provide: "
            "1. recommended_tools: list of modules/tools with name, description, and why recommended. "
            "2. integration_pairs: pairs of tools/modules that should be integrated, with integration description and data shared. "
            "3. additional_tools: tools/modules that are needed but missing, with name, description, names of companies offering it and why needed. "
            "Return the output as a JSON object with a 'recommendations' key containing these three lists. "
            "Here is the user context and available modules:\n\n" + json.dumps(context, indent=2)
        )
        client = openai.OpenAI()

        response = response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a technology consultant for business software and workflow automation."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1000,
            temperature=0.3
        )

        # Try to parse the response as JSON
        import ast
        import re
        raw_content = response.choices[0].message.content
        # Extract JSON from response (in case LLM returns extra text)
        match = re.search(r'\{[\s\S]*\}', raw_content)
        if match:
            recommendations_json = match.group(0)
            try:
                recommendations = json.loads(recommendations_json)
            except Exception:
                recommendations = {"raw": raw_content}
        else:
            recommendations = {"raw": raw_content}

        print(recommendations)

        return jsonify({
            "success": True,
            "recommendations": recommendations.get('recommendations', recommendations)
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

# @app.route('/AI_ML', methods=['GET'])
# def yfinance_test():
    
#     return "Hello World!"


# @app.route('/Location', methods=['GET'])
# def yfinance_test():
    
#     return "Hello World!"

# @app.route('/Transportation', methods=['GET'])
# def yfinance_test():
    
#     return "Hello World!"

# @app.route('/Business- Enterprise', methods=['GET'])
# def yfinance_test():
    
#     return "Hello World!"

# @app.route('/Visual Recognition', methods=['GET'])
# def yfinance_test():
    
#     return "Hello World!"

# @app.route('/Small Tools', methods=['GET'])
# def yfinance_test():
    
#     return "Hello World!"

# @app.route('/Text Analysis', methods=['GET'])
# def yfinance_test():
    
#     return "Hello World!"

# @app.route('/Weather', methods=['GET'])
# def yfinance_test():
    
#     return "Hello World!"

# @app.route('/Messaging', methods=['GET'])
# def yfinance_test():
    
#     return "Hello World!"

# @app.route('/Logistics', methods=['GET'])
# def yfinance_test():
    
#     return "Hello World!"

# @app.route('/News', methods=['GET'])
# def yfinance_test():
    
#     return "Hello World!"

# @app.route('/Jobs', methods=['GET'])
# def yfinance_test():
    
#     return "Hello World!"

# @app.route('/yfinance', methods=['GET'])
# def yfinance_test():
#     symbol = request.args.get('stock')
#     region = request.args.get('region')

#     if not symbol or not region:
#         return "Missing required parameters: 'stock' and 'region'", 400

#     conn = http.client.HTTPSConnection("yahoo-finance166.p.rapidapi.com")

#     headers = {
#         'x-rapidapi-key': "95cdd43379mshbd9483856442c47p1c2782jsn897449ebefb8",
#         'x-rapidapi-host': "yahoo-finance166.p.rapidapi.com"
#     }

#     endpoint = f"/api/stock/get-financial-data?region={region}&symbol={symbol}"
#     print(f"Requesting data from endpoint: {endpoint}")  # Debug statement
#     conn.request("GET", endpoint, headers=headers)

#     res = conn.getresponse()
#     data = res.read()
#     json_data = json.loads(data.decode("utf-8"))

#     print(json_data)  # Debug statement to print the entire response

#     if 'quoteSummary' not in json_data or 'result' not in json_data['quoteSummary'] or not json_data['quoteSummary']['result']:
#         return jsonify({"error": "No data found for the given stock symbol and region"}), 404

#     current_price = json_data['quoteSummary']['result'][0]['financialData']['currentPrice']['fmt']
#     operating_margins = json_data['quoteSummary']['result'][0]['financialData']['operatingMargins']['fmt']
#     netprofit_margins = json_data['quoteSummary']['result'][0]['financialData']['profitMargins']['fmt']
#     gross_margins = json_data['quoteSummary']['result'][0]['financialData']['grossMargins']['fmt']
#     revenue_growth = json_data['quoteSummary']['result'][0]['financialData']['revenueGrowth']['fmt']
#     debt_to_equity = json_data['quoteSummary']['result'][0]['financialData']['debtToEquity']['fmt']
#     quick_ratio = json_data['quoteSummary']['result'][0]['financialData']['quickRatio']['fmt']
#     current_ratio = json_data['quoteSummary']['result'][0]['financialData']['currentRatio']['fmt']
#     analyst_recommendation = json_data['quoteSummary']['result'][0]['financialData']['recommendationKey']
#     number_of_analysts = json_data['quoteSummary']['result'][0]['financialData']['numberOfAnalystOpinions']['fmt']
#     target_high_price = json_data['quoteSummary']['result'][0]['financialData']['targetHighPrice']['fmt']
#     target_low_price = json_data['quoteSummary']['result'][0]['financialData']['targetLowPrice']['fmt']
#     target_mean_price = json_data['quoteSummary']['result'][0]['financialData']['targetMeanPrice']['fmt']
#     target_median_price = json_data['quoteSummary']['result'][0]['financialData']['targetMedianPrice']['fmt']

#     financial_KPIs = {
#         "current_price": current_price,
#         "operating margin": operating_margins,
#         "netprofit_margins": netprofit_margins,
#         "gross_margins": gross_margins,
#         "revenue_growth": revenue_growth,
#         "debt_to_equity": debt_to_equity,
#         "quick_ratio": quick_ratio,
#         "current_ratio": current_ratio,
#         "number_of_analysts": number_of_analysts,
#         "analyst_recommendation": analyst_recommendation,
#         "target_high_price": target_high_price,
#         "target_low_price": target_low_price,
#         "target_mean_price": target_mean_price,
#         "target_median_price": target_median_price
#     }

#     return jsonify(financial_KPIs)


# === KNOWLEDGE GRAPH + RAG API ===

import boto3
from docx import Document as DocxDocument
import networkx as nx

# Initialize S3 client for AWS operations
s3_client = boto3.client('s3')

def generate_cache_key(data):
    """Generate a hash key for caching based on input data"""
    data_string = json.dumps(data, sort_keys=True)
    return hashlib.md5(data_string.encode()).hexdigest()

def get_document_cache_key(documents):
    """Generate cache key for documents list"""
    doc_keys = []
    for doc in documents:
        key = f"{doc['source_type']}:{doc['path']}"
        if 'bucket' in doc:
            key += f":{doc['bucket']}"
        doc_keys.append(key)
    return generate_cache_key(sorted(doc_keys))

def get_kg_cache_key(nodes, edges):
    """Generate cache key for knowledge graph"""
    kg_data = {'nodes': nodes, 'edges': edges}
    return generate_cache_key(kg_data)

def load_document_from_source(source_type, source_path, bucket_name=None):
    """Load PDF or Word document from S3 or local machine"""
    if source_type == "s3":
        local_path = f"/tmp/{os.path.basename(source_path)}"
        s3_client.download_file(bucket_name, source_path, local_path)
        return local_path
    return source_path

def extract_text_from_pdf(file_path):
    """Extract text content from PDF file using PyMuPDF"""
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def extract_text_from_word(file_path):
    """Extract text content from Word document"""
    doc = DocxDocument(file_path)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return text

def extract_text_from_document(file_path):
    """Route extraction based on file extension"""
    if file_path.lower().endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_path.lower().endswith(('.docx', '.doc')):
        return extract_text_from_word(file_path)
    return ""

def build_knowledge_graph(nodes, edges):
    """Create a NetworkX graph from nodes and edges JSON input"""
    G = nx.DiGraph()
    for node in nodes:
        G.add_node(node['id'], **node.get('attributes', {}))
    for edge in edges:
        G.add_edge(edge['source'], edge['target'], **edge.get('attributes', {}))
    return G

def chunk_text(text, chunk_size=1000, overlap=200):
    """Split text into overlapping chunks for better context preservation"""
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start = end - overlap
    return chunks

def create_embeddings(chunks):
    """Generate OpenAI embeddings for text chunks"""
    embeddings_model = OpenAIEmbeddings()
    embeddings = embeddings_model.embed_documents(chunks)
    return np.array(embeddings)

def build_faiss_index(embeddings):
    """Create FAISS index for efficient similarity search"""
    dimension = embeddings.shape[1]
    index = faiss.IndexFlatL2(dimension)
    index.add(embeddings.astype('float32'))
    return index

def retrieve_relevant_chunks(query, index, chunks, embeddings_model, top_k=5):
    """Retrieve most relevant chunks using FAISS similarity search"""
    query_embedding = embeddings_model.embed_query(query)
    query_vector = np.array([query_embedding]).astype('float32')
    distances, indices = index.search(query_vector, top_k)
    return [chunks[i] for i in indices[0]]

def query_knowledge_graph(graph, query_type, node_id=None):
    """Query knowledge graph for specific information based on query type"""
    if query_type == "neighbors" and node_id:
        return list(graph.neighbors(node_id))
    elif query_type == "attributes" and node_id:
        return dict(graph.nodes[node_id])
    elif query_type == "all_nodes":
        return list(graph.nodes(data=True))
    elif query_type == "all_edges":
        return list(graph.edges(data=True))
    return None

def generate_answer_with_rag(query, relevant_chunks, kg_context):
    """Generate final answer using OpenAI with RAG context and KG information"""
    llm = ChatOpenAI(model="gpt-4", temperature=0)
    context_text = "\n\n".join(relevant_chunks)
    kg_text = json.dumps(kg_context, indent=2)
    
    prompt = f"""Based on the following document context and knowledge graph information, answer the query.

Document Context:
{context_text}

Knowledge Graph Context:
{kg_text}

Query: {query}

Provide a detailed answer:"""
    
    response = llm.predict(prompt)
    return response

def process_documents_with_kg_rag(documents, nodes, edges, query, include_context=False):
    """Main processing pipeline combining document loading, KG building, and RAG with caching"""
    # Generate cache keys
    doc_cache_key = get_document_cache_key(documents)
    kg_cache_key = get_kg_cache_key(nodes, edges)
    
    # Check if knowledge graph is cached
    if kg_cache_key in kg_rag_cache['knowledge_graphs']:
        kg = kg_rag_cache['knowledge_graphs'][kg_cache_key]
    else:
        # Build knowledge graph and cache it
        kg = build_knowledge_graph(nodes, edges)
        kg_rag_cache['knowledge_graphs'][kg_cache_key] = kg
    
    # Check if document embeddings are cached
    if doc_cache_key in kg_rag_cache['embeddings']:
        # Reuse cached data
        chunks = kg_rag_cache['chunks'][doc_cache_key]
        embeddings = kg_rag_cache['embeddings'][doc_cache_key]
        faiss_index = kg_rag_cache['faiss_indices'][doc_cache_key]
        embeddings_model = OpenAIEmbeddings()
    else:
        # Extract and combine text from all documents
        all_text = ""
        for doc_info in documents:
            local_path = load_document_from_source(
                doc_info['source_type'], 
                doc_info['path'], 
                doc_info.get('bucket')
            )
            text = extract_text_from_document(local_path)
            all_text += text + "\n\n"
        
        # Create chunks and embeddings
        chunks = chunk_text(all_text)
        embeddings_model = OpenAIEmbeddings()
        embeddings = create_embeddings(chunks)
        
        # Build FAISS index
        faiss_index = build_faiss_index(embeddings)
        
        # Cache all the expensive computations
        kg_rag_cache['chunks'][doc_cache_key] = chunks
        kg_rag_cache['embeddings'][doc_cache_key] = embeddings
        kg_rag_cache['faiss_indices'][doc_cache_key] = faiss_index
    
    # Retrieve relevant chunks (this is query-specific, not cached)
    relevant_chunks = retrieve_relevant_chunks(query, faiss_index, chunks, embeddings_model)
    
    # Query knowledge graph for additional context
    kg_context = {
        'nodes': query_knowledge_graph(kg, "all_nodes"),
        'edges': query_knowledge_graph(kg, "all_edges")
    }
    
    # Generate answer
    answer = generate_answer_with_rag(query, relevant_chunks, kg_context)
    
    # Return only answer by default (lightweight response)
    if include_context:
        return {
            'answer': answer,
            'relevant_chunks': relevant_chunks,
            'kg_context': kg_context
        }
    else:
        return {
            'answer': answer
        }

@app.route('/extract-with-kg-rag', methods=['POST'])
@cross_origin()
def extract_with_kg_rag():
    """API endpoint to extract information from documents using Knowledge Graph and RAG"""
    try:
        data = request.json
        
        # Validate input
        documents = data.get('documents', [])
        nodes = data.get('nodes', [])
        edges = data.get('edges', [])
        query = data.get('query', '')
        include_context = data.get('include_context', False)  # Optional: return chunks and KG context
        
        if not documents or not query:
            return jsonify({
                'success': False,
                'error': 'documents and query are required'
            }), 400
        
        # Process documents with KG and RAG
        result = process_documents_with_kg_rag(documents, nodes, edges, query, include_context)
        
        return jsonify({
            'success': True,
            'data': result
        }), 200
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/clear-kg-rag-cache', methods=['POST'])
@cross_origin()
def clear_kg_rag_cache():
    """Clear the KG+RAG cache to free up memory"""
    try:
        data = request.json or {}
        cache_type = data.get('cache_type', 'all')  # 'all', 'embeddings', 'graphs'
        
        if cache_type == 'all':
            kg_rag_cache['embeddings'].clear()
            kg_rag_cache['faiss_indices'].clear()
            kg_rag_cache['chunks'].clear()
            kg_rag_cache['knowledge_graphs'].clear()
            cleared = 'all caches'
        elif cache_type == 'embeddings':
            kg_rag_cache['embeddings'].clear()
            kg_rag_cache['faiss_indices'].clear()
            kg_rag_cache['chunks'].clear()
            cleared = 'embeddings cache'
        elif cache_type == 'graphs':
            kg_rag_cache['knowledge_graphs'].clear()
            cleared = 'knowledge graphs cache'
        else:
            return jsonify({
                'success': False,
                'error': 'Invalid cache_type. Use: all, embeddings, or graphs'
            }), 400
        
        return jsonify({
            'success': True,
            'message': f'Successfully cleared {cleared}'
        }), 200
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/kg-rag-cache-status', methods=['GET'])
@cross_origin()
def kg_rag_cache_status():
    """Get current cache status and statistics"""
    try:
        status = {
            'embeddings_cached': len(kg_rag_cache['embeddings']),
            'faiss_indices_cached': len(kg_rag_cache['faiss_indices']),
            'chunks_cached': len(kg_rag_cache['chunks']),
            'knowledge_graphs_cached': len(kg_rag_cache['knowledge_graphs']),
            'total_cached_items': (
                len(kg_rag_cache['embeddings']) + 
                len(kg_rag_cache['knowledge_graphs'])
            )
        }
        
        return jsonify({
            'success': True,
            'cache_status': status
        }), 200
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

# ====== ENTITY MANAGEMENT SYSTEM ======
# Domain-agnostic persistent knowledge graph + vector DB system
# Works for ANY domain: HR, Sales, Research, Healthcare, Legal, etc.
# Note: Commented out pending completion of entity_system module

# from entity_system import (
#     EntityKnowledgeGraphManager,
#     EntityVectorStoreManager,
#     process_entity_documents,
#     query_entity_profile,
#     chroma_client
# )
#
# # Initialize entity system managers
# kg_manager = EntityKnowledgeGraphManager(db)
# vector_manager = EntityVectorStoreManager(chroma_client)

# === ENTITY ENDPOINTS (Disabled - Pending entity_system module) ===
# 
# === ENTITY ENDPOINTS (Disabled - Pending entity_system module) ===
# These endpoints require the entity_system module which is not yet implemented
#
# All entity-related endpoints are disabled including:
# - POST /entity/upload
# - POST /entity/query
# - GET /entity/<entity_id>
# - DELETE /entity/<entity_id>
# - GET /entities/list
# - GET /system/health
#
# This module will be enabled once entity_system.py is properly implemented
# with EntityKnowledgeGraphManager, EntityVectorStoreManager, and related utilities
# ======


# === ENTITY ENDPOINTS (Disabled - Pending entity_system module) ===
# These endpoints require the entity_system module which is not yet implemented
#
# All entity-related endpoints disabled including:
# - entity_upload
# - entity_query
# - get_entity
# - delete_entity
# - list_entities
# - system_health
#
# ======


# ====== CONTENT MARKETING AGENT API ENDPOINTS ======

@app.route('/api/content-marketing/projects', methods=['POST'])
@cross_origin()
def create_content_marketing_project():
    """Create a new content marketing project"""
    try:
        data = request.json
        user_id = data.get('user_id', 'default_user')
        project_name = data.get('project_name', 'Untitled Project')
        industry = data.get('industry')
        sector = data.get('sector')
        
        project_id = f"project_{uuid4().hex[:12]}"
        
        conn = sqlite3.connect(CONTENT_MARKETING_DB_PATH)
        cursor = conn.cursor()
        
        cursor.execute('''
            INSERT INTO projects 
            (project_id, user_id, project_name, industry, sector, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        ''', (project_id, user_id, project_name, industry, sector, datetime.now(), datetime.now()))
        
        conn.commit()
        conn.close()
        
        return jsonify({
            'success': True,
            'project_id': project_id,
            'message': f'Project "{project_name}" created successfully'
        }), 201
    
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/content-marketing/projects/<project_id>', methods=['GET'])
@cross_origin()
def get_content_marketing_project(project_id):
    """Get project details"""
    try:
        conn = sqlite3.connect(CONTENT_MARKETING_DB_PATH)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        cursor.execute('SELECT * FROM projects WHERE project_id = ?', (project_id,))
        project = cursor.fetchone()
        
        if not project:
            conn.close()
            return jsonify({'success': False, 'error': 'Project not found'}), 404
        
        # Get documents count
        cursor.execute('SELECT COUNT(*) FROM documents WHERE project_id = ?', (project_id,))
        doc_count = cursor.fetchone()[0]
        
        # Get knowledge graph
        cursor.execute('SELECT kg_data FROM knowledge_graphs WHERE project_id = ?', (project_id,))
        kg_row = cursor.fetchone()
        has_kg = kg_row is not None
        
        conn.close()
        
        return jsonify({
            'success': True,
            'project': dict(project),
            'statistics': {
                'documents': doc_count,
                'has_knowledge_graph': has_kg
            }
        }), 200
    
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/content-marketing/documents/upload', methods=['POST'])
@cross_origin()
def upload_content_marketing_documents():
    """
    Upload documents to project
    Extracts text and creates initial knowledge graph using existing RAG
    """
    try:
        project_id = request.form.get('project_id')
        if not project_id:
            return jsonify({'success': False, 'error': 'project_id required'}), 400
        
        uploaded_files = request.files.getlist('files')
        if not uploaded_files:
            return jsonify({'success': False, 'error': 'No files provided'}), 400
        
        analyzer = DomainSpecializationAnalyzer()
        
        extracted_documents = []
        doc_ids = []
        
        conn = sqlite3.connect(CONTENT_MARKETING_DB_PATH)
        cursor = conn.cursor()
        
        for file in uploaded_files:
            if not file.filename:
                continue
            
            filename = secure_filename(file.filename)
            file_type = filename.split('.')[-1].lower()
            
            if file_type not in CONTENT_MARKETING_ALLOWED_EXTENSIONS:
                continue
            
            file_path = os.path.join(CONTENT_MARKETING_UPLOAD_FOLDER, project_id, filename)
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            file.save(file_path)
            
            # Extract text from document
            text_content = extract_text_from_file_content_marketing(file_path, file_type)
            
            # Store in database
            doc_id = f"doc_{uuid4().hex[:12]}"
            cursor.execute('''
                INSERT INTO documents
                (doc_id, project_id, file_name, file_type, file_path, 
                 file_size, upload_date, extracted_content)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            ''', (doc_id, project_id, filename, file_type, file_path,
                  os.path.getsize(file_path), datetime.now(), text_content))
            
            extracted_documents.append(text_content)
            doc_ids.append(doc_id)
        
        conn.commit()
        conn.close()
        
        # Analyze domain specialization
        domain_context = analyzer.analyze_documents(extracted_documents)
        
        # Build knowledge graph using existing RAG embeddings
        conn = sqlite3.connect(CONTENT_MARKETING_DB_PATH)
        cursor = conn.cursor()
        kg_id = f"kg_{uuid4().hex[:12]}"
        
        # Simple KG structure using the text
        kg_data = {
            'entities': [f'Entity_{i}' for i in range(min(10, len(extracted_documents)))],
            'relationships': [],
            'domain_context': domain_context,
            'documents_count': len(doc_ids)
        }
        
        cursor.execute('''
            INSERT INTO knowledge_graphs
            (kg_id, project_id, kg_data, entities, relationships, created_at)
            VALUES (?, ?, ?, ?, ?, ?)
        ''', (kg_id, project_id, json.dumps(kg_data),
              len(kg_data.get('entities', [])),
              len(kg_data.get('relationships', [])),
              datetime.now()))
        conn.commit()
        conn.close()
        
        return jsonify({
            'success': True,
            'uploaded_files': len(doc_ids),
            'document_ids': doc_ids,
            'knowledge_graph_id': kg_id,
            'domain_specialization': domain_context
        }), 201
    
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/content-marketing/documents/<project_id>', methods=['GET'])
@cross_origin()
def list_content_marketing_documents(project_id):
    """List all documents in a project"""
    try:
        conn = sqlite3.connect(CONTENT_MARKETING_DB_PATH)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT doc_id, file_name, file_type, upload_date, file_size
            FROM documents WHERE project_id = ?
            ORDER BY upload_date DESC
        ''', (project_id,))
        
        documents = [dict(row) for row in cursor.fetchall()]
        conn.close()
        
        return jsonify({
            'success': True,
            'documents': documents,
            'count': len(documents)
        }), 200
    
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/content-marketing/generate-content', methods=['POST'])
@cross_origin()
def generate_content_marketing():
    """
    Generate marketing content for specified channel
    Uses existing RAG endpoint with knowledge graph for content creation
    """
    try:
        data = request.json
        project_id = data.get('project_id')
        channel = data.get('channel', 'linkedin')
        content_type = data.get('content_type', 'post')
        user_context = data.get('context', '')
        
        if not project_id:
            return jsonify({'success': False, 'error': 'project_id required'}), 400
        
        # Get project and documents
        conn = sqlite3.connect(CONTENT_MARKETING_DB_PATH)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        cursor.execute('SELECT * FROM projects WHERE project_id = ?', (project_id,))
        project_row = cursor.fetchone()
        if not project_row:
            conn.close()
            return jsonify({'success': False, 'error': 'Project not found'}), 404
        
        project = dict(project_row)
        
        cursor.execute('''
            SELECT extracted_content FROM documents WHERE project_id = ?
        ''', (project_id,))
        
        documents = cursor.fetchall()
        doc_texts = [doc[0] for doc in documents if doc[0]]
        
        # Get knowledge graph
        cursor.execute('''
            SELECT kg_data FROM knowledge_graphs WHERE project_id = ?
            ORDER BY created_at DESC LIMIT 1
        ''', (project_id,))
        
        kg_row = cursor.fetchone()
        kg_data = json.loads(kg_row[0]) if kg_row else None
        
        conn.close()
        
        if not doc_texts:
            return jsonify({'success': False, 'error': 'No documents found in project'}), 400
        
        # Use existing RAG endpoint to generate content
        channel_config = {
            'linkedin': {
                'tone': 'professional',
                'max_length': 3000,
                'include_hashtags': True,
                'call_to_action': 'Connect with us'
            },
            'email': {
                'tone': 'persuasive',
                'max_length': 500,
                'include_subject': True,
                'call_to_action': 'Learn more'
            },
            'social': {
                'tone': 'casual',
                'max_length': 280,
                'include_hashtags': True,
                'call_to_action': 'Follow us'
            },
            'google_ads': {
                'tone': 'direct',
                'max_length': 150,
                'include_headline': True,
                'call_to_action': 'Click here'
            }
        }
        
        config = channel_config.get(channel, channel_config['linkedin'])
        
        # Simple content generation using context
        prompt = f"""Generate marketing content for {channel} channel.
        
Industry: {project.get('industry', 'General')}
Tone: {config['tone']}
Max Length: {config['max_length']} characters
Content Type: {content_type}
User Context: {user_context}

Documents Summary: {' '.join([doc[:200] for doc in doc_texts[:3]])}

Generate compelling marketing {content_type} content that is {config['tone']}, 
stays within {config['max_length']} characters, and resonates with the target audience."""
        
        llm = ChatOpenAI(model="gpt-4", temperature=0.7)
        response = llm.predict(prompt)
        
        # Store generated content
        content_id = f"content_{uuid4().hex[:12]}"
        conn = sqlite3.connect(CONTENT_MARKETING_DB_PATH)
        cursor = conn.cursor()
        
        cursor.execute('''
            INSERT INTO generated_content
            (content_id, project_id, channel, content_type, content, source_docs, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        ''', (content_id, project_id, channel, content_type,
              response, json.dumps([d[:100] for d in doc_texts]),
              datetime.now()))
        
        conn.commit()
        conn.close()
        
        return jsonify({
            'success': True,
            'content_id': content_id,
            'channel': channel,
            'content_type': content_type,
            'content': response,
            'variations': [response],  # Could generate multiple variations
            'metadata': config
        }), 201
    
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/content-marketing/chat', methods=['POST'])
@cross_origin()
def content_marketing_chat():
    """
    Conversational endpoint for iterative content refinement
    """
    try:
        data = request.json
        project_id = data.get('project_id')
        message = data.get('message')
        
        if not all([project_id, message]):
            return jsonify({'success': False, 'error': 'project_id and message required'}), 400
        
        # Retrieve context
        conn = sqlite3.connect(CONTENT_MARKETING_DB_PATH)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT extracted_content FROM documents WHERE project_id = ?
            LIMIT 10
        ''', (project_id,))
        documents = [row[0] for row in cursor.fetchall()]
        
        cursor.execute('''
            SELECT kg_data FROM knowledge_graphs WHERE project_id = ?
            ORDER BY created_at DESC LIMIT 1
        ''', (project_id,))
        kg_row = cursor.fetchone()
        kg_data = json.loads(kg_row[0]) if kg_row else None
        
        conn.close()
        
        # Generate response using LLM
        context_text = ' '.join([doc[:500] for doc in documents[:3]]) if documents else ''
        
        prompt = f"""Based on the following document context and knowledge graph, provide helpful marketing advice.

Document Context: {context_text}

Knowledge Graph: {json.dumps(kg_data)[:500] if kg_data else 'No KG available'}

User Question: {message}

Provide a helpful, concise response focused on marketing strategy and content improvement."""
        
        llm = ChatOpenAI(model="gpt-4", temperature=0.7)
        response = llm.predict(prompt)
        
        # Store conversation
        msg_id = f"msg_{uuid4().hex[:12]}"
        conn = sqlite3.connect(CONTENT_MARKETING_DB_PATH)
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO conversation_history
            (msg_id, project_id, user_message, agent_response, timestamp)
            VALUES (?, ?, ?, ?, ?)
        ''', (msg_id, project_id, message, response, datetime.now()))
        conn.commit()
        conn.close()
        
        return jsonify({
            'success': True,
            'response': response,
            'message_id': msg_id
        }), 200
    
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/content-marketing/knowledge-graph/<project_id>', methods=['GET'])
@cross_origin()
def get_content_marketing_knowledge_graph(project_id):
    """Retrieve knowledge graph for visualization"""
    try:
        conn = sqlite3.connect(CONTENT_MARKETING_DB_PATH)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT kg_data, entities, relationships, created_at
            FROM knowledge_graphs WHERE project_id = ?
            ORDER BY created_at DESC LIMIT 1
        ''', (project_id,))
        
        row = cursor.fetchone()
        conn.close()
        
        if not row:
            return jsonify({'success': False, 'error': 'Knowledge graph not found'}), 404
        
        kg_data = json.loads(row[0])
        
        return jsonify({
            'success': True,
            'graph': kg_data,
            'statistics': {
                'entities': row[1],
                'relationships': row[2],
                'created_at': str(row[3])
            }
        }), 200
    
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/get-google-credentials', methods=['GET'])
@cross_origin()
def get_google_credentials():
    """
    Get pre-configured Google OAuth credentials from environment
    Returns empty values if not configured
    """
    try:
        load_dotenv(ENV_FILE, override=True)
        has_oauth_credentials = all([
            os.getenv('GOOGLE_CLIENT_ID'),
            os.getenv('GOOGLE_CLIENT_SECRET'),
            os.getenv('GOOGLE_REDIRECT_URI')
        ])
        has_places_api_key = bool(os.getenv('GOOGLE_PLACES_API_KEY'))
        
        credentials = {
            'clientId': os.getenv('GOOGLE_CLIENT_ID', ''),
            'clientSecret': os.getenv('GOOGLE_CLIENT_SECRET', ''),
            'redirectUri': os.getenv('GOOGLE_REDIRECT_URI', ''),
            'hasCredentials': has_oauth_credentials or has_places_api_key,
            'hasPlacesApiKey': has_places_api_key
        }
        
        return jsonify({
            'success': True,
            'credentials': credentials
        }), 200
        
    except Exception as e:
        print(f"Error fetching Google credentials: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e),
            'credentials': {
                'clientId': '',
                'clientSecret': '',
                'redirectUri': '',
                'hasCredentials': False
            }
        }), 200


@app.route('/connect-google-business', methods=['POST'])
@cross_origin()
def connect_google_business():
    """
    Step 1: Save OAuth credentials and generate authorization URL
    Returns URL where user should go to authorize the app
    """
    try:
        data = request.get_json()
        
        required_fields = ['clientId', 'clientSecret', 'redirectUri']
        if not all(field in data for field in required_fields):
            return jsonify({
                'success': False,
                'error': 'Missing required fields'
            }), 400
        
        # Save credentials
        helper = GoogleBusinessHelper()
        if not helper.save_credentials(data):
            return jsonify({
                'success': False,
                'error': 'Failed to save credentials'
            }), 500
        
        # Generate Google OAuth authorization URL
        client_id = data.get('clientId')
        redirect_uri = data.get('redirectUri')
        
        auth_url = _generate_google_auth_url(client_id, redirect_uri)
        
        return jsonify({
            'success': True,
            'message': 'Credentials saved. Please authorize the app.',
            'authUrl': auth_url
        }), 200
        
    except Exception as e:
        print(f"Error in connect_google_business: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/auth/google/callback', methods=['GET'])
@cross_origin()
def google_auth_callback():
    """
    Step 2: OAuth callback endpoint
    Google redirects here with authorization code
    """
    try:
        auth_code = request.args.get('code')
        error = request.args.get('error')
        
        if error:
            return jsonify({
                'success': False,
                'error': f'Google authorization denied: {error}'
            }), 400
        
        if not auth_code:
            return jsonify({
                'success': False,
                'error': 'No authorization code received'
            }), 400
        
        state = request.args.get('state')
        
        if state == 'user_login_flow':
            # This is the login callback
            token_data = {
                'code': auth_code,
                'client_id': GOOGLE_CLIENT_ID,
                'client_secret': GOOGLE_CLIENT_SECRET,
                'redirect_uri': GOOGLE_REDIRECT_URI,
                'grant_type': 'authorization_code'
            }
            res = requests.post("https://oauth2.googleapis.com/token", data=token_data)
            
            if res.status_code != 200:
                print("Google Token Exchange Error:", res.text)
                return jsonify({'error': 'Failed to exchange token', 'details': res.json()}), 400
                
            token_response = res.json()
            access_token = token_response.get('access_token')
            refresh_token = token_response.get('refresh_token')
            scopes_received = token_response.get('scope', '')
            
            session_req = requests.Session()
            user_info = session_req.get('https://www.googleapis.com/oauth2/v1/userinfo', params={'access_token': access_token}).json()
            email = user_info.get('email')
            
            if not email:
                return jsonify({'error': 'Could not get email from Google'}), 400
                
            user = User.query.filter_by(email=email).first()
            if not user:
                user = User(
                    username=email,
                    email=email,
                    first_name=user_info.get('given_name', ''),
                    last_name=user_info.get('family_name', ''),
                    password=generate_password_hash(str(uuid4()))
                )
                db.session.add(user)
                db.session.commit()
                
            token_record = GoogleOAuthToken.query.filter_by(username=email).first()
            if not token_record:
                token_record = GoogleOAuthToken(username=email)
                db.session.add(token_record)
                
            token_record.token = access_token
            if refresh_token:
                token_record.refresh_token = refresh_token
            token_record.client_id = GOOGLE_CLIENT_ID
            token_record.client_secret = GOOGLE_CLIENT_SECRET
            token_record.token_uri = "https://oauth2.googleapis.com/token"
            token_record.scopes = scopes_received
            db.session.commit()
            
            # Redirecting to login page with params to automatically log in the user on the frontend
            return redirect(f"http://localhost:3000/login?google_auth=success&email={email}")

        # Exchange code for refresh token for Google Business

        client_id = os.getenv('GOOGLE_CLIENT_ID')

        client_secret = os.getenv('GOOGLE_CLIENT_SECRET')
        redirect_uri = os.getenv('GOOGLE_REDIRECT_URI')
        
        success = _exchange_auth_code_for_token(auth_code, client_id, client_secret, redirect_uri)
        
        if success:
            # Redirect back to app with success
            return redirect('http://localhost:3000?google_connected=true')
        else:
            return jsonify({
                'success': False,
                'error': 'Failed to exchange authorization code'
            }), 500
            
    except Exception as e:
        print(f"Error in google_auth_callback: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/get-google-business-data', methods=['GET'])
@cross_origin()
def get_google_business_data():
    """
    Fetch Google Business information using saved credentials
    Returns business profile, reviews, and metrics
    """
    try:
        helper = GoogleBusinessHelper()
        
        if not helper.is_connected():
            return jsonify({
                'success': False,
                'error': 'Google Business credentials not found. Please connect first.',
                'code': 'NOT_CONNECTED'
            }), 401
        
        # Get complete business data
        business_data = helper.get_complete_business_data()
        business_data['success'] = True
        
        return jsonify(business_data), 200
        
    except Exception as e:
        print(f"Error in get_google_business_data: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/get-requirements-with-google-data', methods=['POST'])
@cross_origin()
def get_requirements_with_google_data():
    """
    Combined endpoint to get user requirements along with Google Business data
    Returns both requirement inputs and business insights for context
    """
    try:
        data = request.get_json()
        
        # Extract user requirements from request
        user_requirements = {
            'overview': data.get('overview', ''),
            'context': data.get('context', ''),
            'region': data.get('region', ''),
            'countries': data.get('countries', []),
            'industries': data.get('industries', []),
            'businessFunctions': data.get('businessFunctions', []),
            'analysisFrameworks': data.get('analysisFrameworks', ''),
            'responseFormat': data.get('responseFormat', ''),
            'uploadedFile': data.get('uploadedFile', None)
        }
        
        # Check if Google Business is connected and fetch data
        helper = GoogleBusinessHelper()
        google_business_data = None
        
        if helper.is_connected():
            google_business_data = helper.get_complete_business_data()
        else:
            google_business_data = {
                'connected': False,
                'message': 'Google Business not connected'
            }
        
        response = {
            'success': True,
            'userRequirements': user_requirements,
            'googleBusinessData': google_business_data,
            'timestamp': datetime.now().isoformat()
        }
        
        return jsonify(response), 200
        
    except Exception as e:
        print(f"Error in get_requirements_with_google_data: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/search-google-businesses', methods=['POST'])
@cross_origin()
def search_google_businesses():
    """
    Search for Google businesses using Google Locations API
    Searches by business name and location with pagination support
    Supports up to 200 listings with configurable page size
    """
    try:
        data = request.get_json()
        
        query = data.get('query', '')
        location = data.get('location', '')
        page = data.get('page', 1)
        page_size = data.get('page_size', 20)
        
        if not query:
            return jsonify({
                'success': False,
                'error': 'Search query is required'
            }), 400
        
        if not location:
            return jsonify({
                'success': False,
                'error': 'Location is required'
            }), 400
        
        # Validate pagination parameters
        if page < 1:
            page = 1
        if page_size < 1:
            page_size = 20
        if page_size > 200:
            page_size = 200  # Cap at 200 per page
        
        print(f"[SEARCH] Query: {query}, Location: {location}, Page: {page}, Page Size: {page_size}")
        
        # OAuth is optional for this flow because Google Places API key is used for search.
        # If OAuth is configured, we still attempt token refresh and pass it through.
        client_id = os.getenv('GOOGLE_CLIENT_ID')
        client_secret = os.getenv('GOOGLE_CLIENT_SECRET')
        redirect_uri = os.getenv('GOOGLE_REDIRECT_URI')
        access_token = None
        if all([client_id, client_secret, redirect_uri]):
            print("[SEARCH] OAuth credentials detected, attempting token refresh...")
            access_token = _get_google_access_token(client_id, client_secret, redirect_uri)
            if access_token:
                print("[SEARCH] OAuth token refresh succeeded.")
            else:
                print("[SEARCH] OAuth token refresh failed; continuing with Places API key.")
        else:
            print("[SEARCH] OAuth credentials missing; continuing with Places API key.")
        
        # Use Google Locations API to search for businesses
        searcher = GoogleBusinessSearcher()
        if access_token:
            searcher.set_credentials(access_token)
        
        results = searcher.search_businesses(
            query=query,
            location=location,
            max_results=200,
            page=page,
            page_size=page_size
        )
        
        print(f"[SEARCH] Results: {results}")
        return jsonify(results), 200
        
    except Exception as e:
        print(f"[SEARCH] Error in search_google_businesses: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


def _generate_google_auth_url(client_id: str, redirect_uri: str) -> str:
    """
    Generate Google OAuth authorization URL
    User visits this URL to authorize the app
    """
    oauth_scopes = [
        'https://www.googleapis.com/auth/business.manage',
        'https://www.googleapis.com/auth/spreadsheets',
        'https://www.googleapis.com/auth/drive.file'
    ]
    params = {
        'client_id': client_id,
        'redirect_uri': redirect_uri,
        'response_type': 'code',
        'scope': ' '.join(oauth_scopes),
        'access_type': 'offline',
        'prompt': 'consent'
    }
    return f"https://accounts.google.com/o/oauth2/v2/auth?{urlencode(params)}"


@app.route('/export-google-sheet', methods=['POST'])
@cross_origin()
def export_google_sheet():
    """
    Create and populate a Google Sheet using OAuth refresh token.
    Expects JSON payload: { title: str, headers: [str], rows: [[...]] }
    """
    try:
        data = request.get_json() or {}
        title = data.get('title', f"Market Research {datetime.now().strftime('%Y-%m-%d')}")
        headers = data.get('headers', [])
        rows = data.get('rows', [])

        if not headers or not isinstance(headers, list):
            return jsonify({
                'success': False,
                'error': 'Invalid headers payload'
            }), 400

        if not isinstance(rows, list):
            return jsonify({
                'success': False,
                'error': 'Invalid rows payload'
            }), 400

        client_id = os.getenv('GOOGLE_CLIENT_ID')
        client_secret = os.getenv('GOOGLE_CLIENT_SECRET')
        redirect_uri = os.getenv('GOOGLE_REDIRECT_URI')

        if not all([client_id, client_secret, redirect_uri]):
            return jsonify({
                'success': False,
                'error': 'Google OAuth credentials not configured.',
                'code': 'CREDENTIALS_MISSING'
            }), 401

        access_token = _get_google_access_token(client_id, client_secret, redirect_uri)
        if not access_token:
            return jsonify({
                'success': False,
                'error': 'Unable to get Google access token. Please reconnect Google account.',
                'code': 'AUTH_FAILED',
                'authUrl': _generate_google_auth_url(client_id, redirect_uri)
            }), 401

        create_sheet_resp = requests.post(
            'https://sheets.googleapis.com/v4/spreadsheets',
            headers={
                'Authorization': f'Bearer {access_token}',
                'Content-Type': 'application/json'
            },
            json={
                'properties': {'title': title}
            },
            timeout=20
        )

        if create_sheet_resp.status_code != 200:
            error_payload = create_sheet_resp.json() if create_sheet_resp.text else {}
            error_message = error_payload.get('error', {}).get('message', create_sheet_resp.text)
            return jsonify({
                'success': False,
                'error': f'Failed to create Google Sheet: {error_message}',
                'code': 'SHEET_CREATE_FAILED',
                'authUrl': _generate_google_auth_url(client_id, redirect_uri)
            }), 400

        spreadsheet = create_sheet_resp.json()
        spreadsheet_id = spreadsheet.get('spreadsheetId')
        spreadsheet_url = spreadsheet.get('spreadsheetUrl')

        value_rows = [headers] + rows
        update_resp = requests.put(
            f'https://sheets.googleapis.com/v4/spreadsheets/{spreadsheet_id}/values/Sheet1!A1?valueInputOption=RAW',
            headers={
                'Authorization': f'Bearer {access_token}',
                'Content-Type': 'application/json'
            },
            json={
                'majorDimension': 'ROWS',
                'values': value_rows
            },
            timeout=20
        )

        if update_resp.status_code != 200:
            error_payload = update_resp.json() if update_resp.text else {}
            error_message = error_payload.get('error', {}).get('message', update_resp.text)
            return jsonify({
                'success': False,
                'error': f'Sheet created but data write failed: {error_message}',
                'code': 'SHEET_WRITE_FAILED',
                'spreadsheetUrl': spreadsheet_url,
                'authUrl': _generate_google_auth_url(client_id, redirect_uri)
            }), 400

        return jsonify({
            'success': True,
            'spreadsheetId': spreadsheet_id,
            'spreadsheetUrl': spreadsheet_url
        }), 200

    except Exception as e:
        print(f"Error in export_google_sheet: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


def _exchange_auth_code_for_token(auth_code: str, client_id: str, client_secret: str, redirect_uri: str) -> bool:
    """
    Exchange Google authorization code for refresh token
    Called when user returns from Google authorization page
    """
    try:
        token_url = "https://oauth2.googleapis.com/token"
        data = {
            'client_id': client_id,
            'client_secret': client_secret,
            'code': auth_code,
            'grant_type': 'authorization_code',
            'redirect_uri': redirect_uri
        }
        
        response = requests.post(token_url, data=data, timeout=10)
        
        if response.status_code == 200:
            token_data = response.json()
            refresh_token = token_data.get('refresh_token')
            
            if refresh_token:
                # Save refresh token to environment
                os.environ['GOOGLE_REFRESH_TOKEN'] = refresh_token
                
                # Save to .env file in the tools directory
                env_file = os.path.join(os.path.dirname(__file__), '.env')
                with open(env_file, 'a') as f:
                    f.write(f"\nGOOGLE_REFRESH_TOKEN={refresh_token}")
                
                print(f"Refresh token saved to {env_file}")
                load_dotenv()
                return True
            else:
                print("No refresh token in response - this might be the first time authorizing")
                print(f"Response data: {token_data}")
                return False
        else:
            print(f"Token exchange failed: {response.status_code} - {response.text}")
            return False
            
    except Exception as e:
        print(f"Error exchanging auth code: {str(e)}")
        import traceback
        traceback.print_exc()
        return False


def _get_google_access_token(client_id: str, client_secret: str, redirect_uri: str) -> str:
    """
    Get Google access token from stored refresh token
    """
    try:
        refresh_token = os.getenv('GOOGLE_REFRESH_TOKEN')
        
        if not refresh_token:
            print("No refresh token found. User must authorize the app first.")
            print(f"Check that GOOGLE_REFRESH_TOKEN is in .env file")
            return None
        
        print(f"Using refresh token: {refresh_token[:20]}...")
        
        # Use refresh token to get new access token
        token_url = "https://oauth2.googleapis.com/token"
        data = {
            'client_id': client_id,
            'client_secret': client_secret,
            'refresh_token': refresh_token,
            'grant_type': 'refresh_token'
        }
        
        response = requests.post(token_url, data=data, timeout=10)
        
        if response.status_code == 200:
            access_token = response.json().get('access_token')
            print(f"Successfully got access token: {access_token[:20]}...")
            return access_token
        else:
            print(f"Token refresh failed: {response.status_code} - {response.text}")
            return None
            
    except Exception as e:
        print(f"Error getting Google access token: {str(e)}")
        import traceback
        traceback.print_exc()
        return None


@app.route('/email-extraction-usage', methods=['GET'])
@cross_origin()
def get_email_extraction_usage():
    """Return extraction usage summary for a username."""
    try:
        _ensure_email_usage_tables()

        username = _normalize_username(
            request.args.get('username') or request.args.get('userId')
        )
        quota = _get_or_create_quota(username)

        return jsonify({
            'success': True,
            'usageSummary': _build_usage_summary(username, quota)
        }), 200
    except Exception as e:
        print(f"[EMAIL_USAGE] Failed to fetch usage summary: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

def generate_email_content(business, sender_name):
    import json
    import os
    from openai import OpenAI
    
    prompt = f"""We have a market research agent that scrapes company websites and 
extracts leads including company name, industry, decision maker name, 
email, and a short company summary from their website.

I need you to build an email personalization layer on top of this.

When the user clicks "Send Email" for a lead, before sending, the system 
should call the AI to fill in the following variables dynamically using 
the lead data we already have:

- first_name -> extract from the contact name we scraped
- company_name -> from lead data
- one_line_company_summary -> generate a one line summary of what the company does based on the website content we already scraped
- industry -> detected from the company description
- pain_point -> infer the most likely business pain point for this industry and company size
- value_proposition -> tailor this to the industry, e.g. for FinTech say something different than for SaaS
- sender_name -> from the logged in client's profile

The base template is:

Subject: Quick idea for {{{{company_name}}}}

Hi {{{{first_name}}}},

I came across {{{{company_name}}}} and noticed {{{{one_line_company_summary}}}}.

Companies in {{{{industry}}}} often struggle with {{{{pain_point}}}} - and that usually means lost time or missed opportunities.

We built a solution that helps {{{{industry}}}} teams {{{{value_proposition}}}}.

Would it make sense to connect for 15 minutes this week?

Best,
{sender_name}

Return the final filled email as JSON in this format:
{{
  "subject": "...",
  "body": "..."
}}

Lead Data:
Company Name: {business.get('name', 'Unknown')}
Content/Description: {business.get('description', '')} {business.get('summary', '')}
Website: {business.get('website', 'Unknown')}
Contact Name: {business.get('contact_name', 'There')}
Industry: {business.get('industry', 'Unknown')}

Do not add any explanation, just return the JSON.
"""

    client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
    response = client.chat.completions.create(
        model=os.getenv('OPENAI_MODEL', 'gpt-4o-mini'),
        messages=[
            {"role": "user", "content": prompt}
        ],
        response_format={ "type": "json_object" },
        temperature=0.7
    )
    
    return json.loads(response.choices[0].message.content)

@app.route('/api/generate-email', methods=['POST'])
@cross_origin()
def generate_email():
    """Generate a personalized email using an LLM."""
    try:
        from flask import request
        data = request.get_json()
        business = data.get('business', {})
        sender_name = data.get('sender_name', 'Alex')
        
        result = generate_email_content(business, sender_name)
        return jsonify(result), 200
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/send-bulk-emails', methods=['POST'])
@cross_origin()
def send_bulk_emails():
    """
    Send emails to a list of extracted businesses.
    Uses user's Google OAuth credentials if available, otherwise falls back to system SMTP.
    Expects JSON: { subject: str, body: str, businesses: list, userEmail: str }
    """
    import smtplib
    from email.message import EmailMessage
    import traceback
    
    try:
        data = request.get_json()
        subject = data.get('subject')
        body = data.get('body')
        businesses = data.get('businesses', [])
        user_email = data.get('userEmail')
        campaign_name = data.get('campaignName', 'Untitled Campaign')
        username = _normalize_username(data.get('username') or data.get('userId') or data.get('firstName'))

        use_ai_personalization = data.get('use_ai_personalization', False)
        if not use_ai_personalization and (not subject or not body):
            return jsonify({'success': False, 'error': 'Subject and body are required unless using AI personalization'}), 400

        valid_emails = [b.get('email') for b in businesses if b.get('email') and b.get('email') != 'N/A' and '@' in b.get('email')]

        if not valid_emails:
            return jsonify({'success': False, 'error': 'No valid emails found to send to'}), 400

        # Initialize DB Tables if needed
        _ensure_email_usage_tables()
        
        # Create Campaign Record
        import uuid
        campaign_id = str(uuid.uuid4())
        campaign = EmailCampaign(
            id=campaign_id,
            name=campaign_name,
            subject=subject,
            username=username
        )
        db.session.add(campaign)
        db.session.commit()
        
        # Check if user has Google credentials connected
        token_record = None
        if user_email:
            token_record = GoogleOAuthToken.query.filter_by(username=user_email).first()
        
        service = None
        server = None
        
        if token_record and token_record.token:
            # Use Gmail API
            from google.oauth2.credentials import Credentials
            import googleapiclient.discovery
            
            creds = Credentials(
                token=token_record.token,
                refresh_token=token_record.refresh_token,
                token_uri=token_record.token_uri,
                client_id=token_record.client_id,
                client_secret=token_record.client_secret,
                scopes=token_record.scopes.split(',') if token_record.scopes else SCOPES
            )
            service = googleapiclient.discovery.build('gmail', 'v1', credentials=creds)
        else:
            # Fallback to system SMTP
            email_host = os.getenv('EMAIL_HOST', 'smtp.gmail.com')
            email_port = int(os.getenv('EMAIL_PORT', 587))
            email_user = os.getenv('EMAIL_USER')
            email_pass = os.getenv('EMAIL_PASS')
            if not email_user or not email_pass:
                return jsonify({'success': False, 'error': 'Email credentials are not configured. Please sign in with Google or configure system SMTP.'}), 500
                
            server = smtplib.SMTP(email_host, email_port)
            server.starttls()
            server.login(email_user, email_pass)

        sent_count = 0
        for b in businesses:
            recipient = b.get('email')
            if not recipient or recipient == 'N/A' or '@' not in recipient:
                continue

            business_name = b.get('name', 'Business Owner')
            
            # Message personalization
            current_body = body
            if use_ai_personalization:
                try:
                    result = generate_email_content(b, username)
                    current_subject = result.get('subject', subject or 'Exclusive Offer')
                    current_body = result.get('body', current_body or '')
                except Exception as e:
                    print("Failed AI personalization for", business_name, e)
                    current_subject = subject or 'Exclusive Offer'
            else:
                current_subject = subject.replace('{{name}}', business_name) if subject else subject
                if current_body:
                    current_body = current_body.replace('{{name}}', business_name)

            msg = EmailMessage()
            msg.set_content(current_body)
            msg['Subject'] = current_subject
            msg['To'] = recipient
            
            if service:
                msg['From'] = user_email
                encoded_message = base64.urlsafe_b64encode(msg.as_bytes()).decode()
                create_message = {'raw': encoded_message}
                service.users().messages().send(userId="me", body=create_message).execute()
            else:
                msg['From'] = email_user
                server.send_message(msg)
                
            sent_count += 1
            
            # Record recipient for tracking
            recipient_record = EmailCampaignRecipient(
                campaign_id=campaign_id,
                receiver_email=recipient,
                receiver_name=business_name,
                status='SENT',
                reply_status='No Reply'
            )
            db.session.add(recipient_record)

        db.session.commit()
        
        if server:
            server.quit()
            
        # Commit any leftover recipient records
        db.session.commit()
        
        return jsonify({'success': True, 'count': sent_count, 'message': 'Emails successfully sent via user account!' if service else 'Emails sent via system account.'})
    except Exception as e:
        db.session.rollback()
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/webhook/zapier/email-reply', methods=['POST'])
@cross_origin()
def handle_email_reply():
    """
    Zapier Webhook Endpoint to log replies from emails sent in campaigns.
    Expects JSON: { "from_email": "example@domain.com", "subject": "Re: ...", "timestamp": "...", "snippet": "..." }
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'No JSON payload provided'}), 400
            
        from_email = data.get('from_email')
        
        if not from_email:
            return jsonify({'success': False, 'error': 'from_email is required'}), 400
            
        # Optional: Parse out name from format "Name <email@dom.com>"
        import re
        email_match = re.search(r'<(.+?)>', from_email)
        if email_match:
            clean_email = email_match.group(1).lower().strip()
        else:
            clean_email = from_email.lower().strip()
            
        # Find receiver in the database
        recipients = EmailCampaignRecipient.query.filter(
            EmailCampaignRecipient.receiver_email.ilike(f"%{clean_email}%")
        ).all()
        
        if not recipients:
            return jsonify({'success': False, 'message': f'Reply from {clean_email} logged, but not found in active campaigns.'}), 200
            
        updated_count = 0
        for rec in recipients:
            if rec.reply_status != 'Replied':
                rec.reply_status = 'Replied'
                rec.replied_at = datetime.utcnow()
                updated_count += 1
                
        db.session.commit()
        return jsonify({
            'success': True, 
            'message': f'Successfully updated {updated_count} recipient records to Replied status.'
        }), 200

    except Exception as e:
        print(f"[ZAPIER WEBHOOK ERROR] {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/campaigns/stats', methods=['GET'])
@cross_origin()
def get_campaign_stats():
    """Returns analytics for campaigns filtered by user."""
    try:
        username = request.args.get('username') or request.args.get('userId') or request.args.get('email')
        username = _normalize_username(username) if username else None
        
        if username:
            campaigns = EmailCampaign.query.filter_by(username=username).order_by(EmailCampaign.created_at.desc()).all()
        else:
            campaigns = EmailCampaign.query.order_by(EmailCampaign.created_at.desc()).all()
        results = []
        for c in campaigns:
            recipients = EmailCampaignRecipient.query.filter_by(campaign_id=c.id).all()
            total_sent = len(recipients)
            total_replied = sum(1 for r in recipients if r.reply_status == 'Replied')
            
            results.append({
                'id': c.id,
                'name': c.name,
                'subject': c.subject,
                'createdAt': c.created_at.isoformat(),
                'totalSent': total_sent,
                'totalReplied': total_replied,
                'replyRate': round((total_replied / total_sent * 100) if total_sent > 0 else 0, 1)
            })
            
        return jsonify({'success': True, 'campaigns': results}), 200
    except Exception as e:
        print(f"[CAMPAIGN STATS ERROR] {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/campaigns/<campaign_id>/recipients', methods=['GET'])
@cross_origin()
def get_campaign_recipients(campaign_id):
    """Returns recipients for a specific campaign."""
    try:
        recipients = EmailCampaignRecipient.query.filter_by(campaign_id=campaign_id).all()
        results = [{
            'email': r.receiver_email,
            'name': r.receiver_name,
            'status': r.status,
            'replyStatus': r.reply_status,
            'sentAt': r.sent_at.isoformat() if r.sent_at else None,
            'repliedAt': r.replied_at.isoformat() if r.replied_at else None
        } for r in recipients]
        return jsonify({'success': True, 'recipients': results}), 200
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/enrich-businesses-with-emails', methods=['POST'])
@cross_origin()
def enrich_businesses_with_emails():
    """
    Enrich business data with email addresses using scrap.io API
    Takes a list of businesses with website URLs and enriches them with email data
    """
    try:
        data = request.get_json()
        businesses = data.get('businesses', [])
        username = _normalize_username(
            data.get('username') or data.get('userId') or data.get('firstName')
        )
        scrap_io_api_key = os.getenv('SCRAP_IO_API_KEY')

        _ensure_email_usage_tables()
        quota = _get_or_create_quota(username)
        usage_before = _build_usage_summary(username, quota)

        if usage_before['remainingCount'] <= 0:
            return jsonify({
                'success': False,
                'error': 'Email extraction limit reached for this user.',
                'code': 'QUOTA_EXCEEDED',
                'usageSummary': usage_before
            }), 403
        
        if not businesses or len(businesses) == 0:
            return jsonify({
                'success': False,
                'error': 'No businesses provided'
            }), 400
        
        if not scrap_io_api_key:
            return jsonify({
                'success': False,
                'error': 'Scrap.io API key not configured',
                'code': 'API_KEY_MISSING'
            }), 401
        
        print(f"[EMAIL_ENRICHMENT] Enriching {len(businesses)} businesses with emails")
        
        def extract_emails_from_text(text):
            """Extract email addresses from text using regex"""
            if not text:
                return []
            email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
            emails = re.findall(email_pattern, str(text))
            return list(set(emails))  # Remove duplicates
        
        def get_email_from_google_places(business):
            """Try to extract email from Google Places data"""
            # Check if business data already contains formatted address, phone, website
            # that might have email info
            if business.get('formatted_address'):
                emails = extract_emails_from_text(business['formatted_address'])
                if emails:
                    return emails[0]
            
            return None
        
        def generate_common_email_patterns(domain, business_name):
            """Generate common email patterns for a domain"""
            # Remove www. if present for cleaner domain
            domain_clean = domain.replace('www.', '')
            
            # Common prefixes to try
            common_prefixes = [
                'info', 'contact', 'support', 'hello', 'business', 'sales', 
                'email', 'inquiry', 'admin', 'help', 'team'
            ]
            
            # Try with business name components
            name_parts = business_name.lower().split()[:2] if business_name else []
            
            patterns = []
            
            # Standard patterns
            for prefix in common_prefixes:
                patterns.append(f"{prefix}@{domain_clean}")
            
            # Business name patterns
            if name_parts:
                for part in name_parts:
                    part_clean = ''.join(c for c in part if c.isalnum())
                    if part_clean:
                        patterns.append(f"{part_clean}@{domain_clean}")
            
            return patterns
        
        def try_sync_website_scrape(website, business_name):
            """Try to scrape website for email addresses synchronously"""
            try:
                response = requests.get(
                    website if website.startswith('http') else f'https://{website}',
                    timeout=5,
                    headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
                )
                
                if response.status_code == 200:
                    # Look for email patterns in HTML
                    emails = extract_emails_from_text(response.text)
                    if emails:
                        # Filter out common auto-generated emails
                        filtered = [e for e in emails if not any(x in e.lower() for x in ['noreply', 'no-reply', 'postmaster'])]
                        return filtered[0] if filtered else None
            except:
                pass
            
            return None
        
        def extract_email_from_business_data(business_data, domain, website):
            """Extract email from scrap.io response with multiple fallback strategies"""
            email = 'N/A'
            
            website_data = business_data.get('website_data', {})
            
            if not website_data:
                return email
            
            # Strategy 1: Check emails array (primary)
            if website_data.get('emails') and len(website_data.get('emails', [])) > 0:
                email_item = website_data['emails'][0]
                if isinstance(email_item, dict) and 'email' in email_item:
                    extracted = email_item['email']
                    if extracted and extracted != 'N/A':
                        email = extracted
                        print(f"[EMAIL_ENRICHMENT] âœ… Strategy 1 (emails array) found: {email}")
                        return email
                elif isinstance(email_item, str) and email_item.strip():
                    email = email_item
                    print(f"[EMAIL_ENRICHMENT] âœ… Strategy 1 (emails string) found: {email}")
                    return email
            
            # Strategy 2: Try contact information
            if not email or email == 'N/A':
                contact_info = website_data.get('contact_info', {})
                if contact_info and contact_info.get('email'):
                    email = contact_info['email']
                    print(f"[EMAIL_ENRICHMENT] âœ… Strategy 2 (contact_info) found: {email}")
                    return email
            
            # Strategy 3: Extract from social profiles or other fields
            if not email or email == 'N/A':
                for key in ['socials', 'social_profiles', 'contact', 'business_info']:
                    if key in website_data:
                        data = website_data[key]
                        if isinstance(data, dict):
                            for field_key, field_value in data.items():
                                if field_value and isinstance(field_value, str):
                                    found_emails = extract_emails_from_text(field_value)
                                    if found_emails:
                                        email = found_emails[0]
                                        print(f"[EMAIL_ENRICHMENT] âœ… Strategy 3 ({key}.{field_key}) found: {email}")
                                        return email
            
            # Strategy 4: Check for contact URLs and other fields
            if not email or email == 'N/A':
                for key in ['contact_page', 'about_page', 'company_info', 'all_text']:
                    if key in website_data:
                        text_data = website_data[key]
                        if text_data:
                            found_emails = extract_emails_from_text(str(text_data))
                            if found_emails:
                                email = found_emails[0]
                                print(f"[EMAIL_ENRICHMENT] âœ… Strategy 4 ({key}) found: {email}")
                                return email
            
            # Strategy 5: Fallback to phone if no email found
            if (not email or email == 'N/A') and website_data.get('phones') and len(website_data.get('phones', [])) > 0:
                phone_obj = website_data['phones'][0]
                if isinstance(phone_obj, dict) and 'phone' in phone_obj:
                    phone = phone_obj['phone']
                    email = f"Phone: {phone}"
                    print(f"[EMAIL_ENRICHMENT] â„¹ï¸  Strategy 5 (phone fallback) found: {email}")
                    return email
            
            return email
        
        def call_scrap_io_with_retry(domain, max_retries=3, delay=2):
            """Call scrap.io with retries to handle async processing"""
            import time
            
            headers = {
                'Authorization': f'Bearer {scrap_io_api_key}',
                'Content-Type': 'application/json'
            }
            
            for attempt in range(max_retries):
                try:
                    # Use domain parameter
                    response = requests.get(
                        scrap_io_endpoint,
                        params={'domain': domain},
                        headers=headers,
                        timeout=15
                    )
                    
                    if response.status_code in [200, 202]:
                        result = response.json()
                        data = result.get('data', [])
                        meta = result.get('meta', {})
                        status = meta.get('status', 'unknown')
                        
                        # If we got data, return it
                        if data and len(data) > 0:
                            print(f"[EMAIL_ENRICHMENT] Got data on attempt {attempt + 1}")
                            return result, True
                        
                        # If status is incomplete, retry with delay
                        if status == 'incomplete' and attempt < max_retries - 1:
                            print(f"[EMAIL_ENRICHMENT] Status incomplete, retrying... (attempt {attempt + 1}/{max_retries})")
                            time.sleep(delay)
                            continue
                        
                        # Status is completed with no data, or this is the last attempt
                        return result, False
                    else:
                        print(f"[EMAIL_ENRICHMENT] API Error {response.status_code}: {response.text[:100]}")
                        return {}, False
                        
                except requests.exceptions.Timeout:
                    print(f"[EMAIL_ENRICHMENT] Timeout on attempt {attempt + 1}")
                    if attempt < max_retries - 1:
                        time.sleep(delay)
                        continue
                    return {}, False
                except Exception as e:
                    print(f"[EMAIL_ENRICHMENT] Error on attempt {attempt + 1}: {str(e)}")
                    return {}, False
            
            return {}, False
        
        enriched_businesses = []
        scrap_io_endpoint = "https://scrap.io/api/v1/gmap/enrich"
        
        for business in businesses:
            website = business.get('website')
            
            # If no website, add business as-is with empty email
            if not website:
                business_copy = business.copy()
                business_copy['email'] = 'N/A'
                enriched_businesses.append(business_copy)
                continue
            
            # Extract domain from website URL
            domain = website.replace('https://', '').replace('http://', '').split('/')[0]
            business_name = business.get('name', '')
            
            email = 'N/A'
            
            try:
                # First, try to extract from Google Places data itself
                google_email = get_email_from_google_places(business)
                if google_email:
                    email = google_email
                    print(f"[EMAIL_ENRICHMENT] âœ… Found email from Google Places: {email}")
                else:
                    # Try calling scrap.io with retry mechanism for async processing
                    result, has_data = call_scrap_io_with_retry(domain, max_retries=2, delay=1)
                    
                    if has_data:
                        data = result.get('data', [])
                        if data and len(data) > 0:
                            business_data = data[0]
                            email = extract_email_from_business_data(business_data, domain, website)
                    
                    # If still no email, try direct website scraping
                    if email == 'N/A':
                        print(f"[EMAIL_ENRICHMENT] Trying direct website scrape for {domain}...")
                        scraped_email = try_sync_website_scrape(website, business_name)
                        if scraped_email:
                            email = scraped_email
                            print(f"[EMAIL_ENRICHMENT] âœ… Found via website scrape: {email}")
                    
                    # If still no email, generate common patterns for user reference
                    if email == 'N/A':
                        patterns = generate_common_email_patterns(domain, business_name)
                        if patterns:
                            # Use the most common pattern as suggestion
                            email = patterns[0]
                            print(f"[EMAIL_ENRICHMENT] â„¹ï¸  Using common pattern suggestion: {email}")
                
                business_copy = business.copy()
                business_copy['email'] = email
                enriched_businesses.append(business_copy)
                
            except Exception as e:
                print(f"[EMAIL_ENRICHMENT] âŒ Error processing {business_name}: {str(e)}")
                business_copy = business.copy()
                business_copy['email'] = 'Error'
                enriched_businesses.append(business_copy)
        
        processed_count = len(enriched_businesses)
        billable_count = sum(
            1 for business in enriched_businesses if _is_billable_email(business.get('email'))
        )
        charged_count = min(billable_count, usage_before['remainingCount'])
        quota.used_count += charged_count
        db.session.add(quota)

        request_id = str(uuid4())
        cost_this_request = round(charged_count * EMAIL_EXTRACTION_UNIT_COST, 2)
        usage_after = _build_usage_summary(username, quota)
        usage_log = EmailExtractionUsageLog(
            request_id=request_id,
            username=username,
            processed_count=processed_count,
            billable_count=billable_count,
            charged_count=charged_count,
            cost_this_request=cost_this_request,
            total_cost_after=usage_after['totalCost']
        )
        db.session.add(usage_log)
        db.session.commit()

        print(f"[EMAIL_ENRICHMENT] Successfully enriched {processed_count} businesses")
        
        return jsonify({
            'success': True,
            'businesses': enriched_businesses,
            'enrichedCount': processed_count,
            'billableEmailCount': billable_count,
            'chargedEmailCount': charged_count,
            'costThisRequest': cost_this_request,
            'usageSummary': usage_after,
            'timestamp': datetime.now().isoformat()
        }), 200
        
    except Exception as e:
        db.session.rollback()
        print(f"[EMAIL_ENRICHMENT] Error in enrich_businesses_with_emails: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/enrich-businesses-with-linkedin', methods=['POST'])
@cross_origin()
def enrich_businesses_with_linkedin():
    try:
        data = request.get_json()
        businesses = data.get('businesses', [])
        
        # Use duckduckgo-search to intelligently grab the most relevant linkedin company URL based on business name
        
        try:
            from ddgs import DDGS
        except ImportError:
            DDGS = None
            
        enriched_businesses = []
        for b in businesses:
            raw_name = b.get('name', '')
            b['linkedin'] = "N/A"
            if raw_name and DDGS:
                try:
                    search_query = f"{raw_name} linkedin company"
                    results = DDGS().text(search_query, max_results=3)
                    for r in results:
                        if 'linkedin.com/company' in r.get('href', ''):
                            b['linkedin'] = r['href']
                            break
                except Exception as e:
                    print(f"Error searching linkedin for {raw_name}:", e)
            
            enriched_businesses.append(b)

        return jsonify({
            'success': True,
            'data': {'businesses': enriched_businesses}
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500



# === HEALTH CHECK ENDPOINT ===
@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint for Docker and load balancers"""
    return jsonify({
        'status': 'healthy',
        'service': 'enable-agents-api',
        'timestamp': datetime.now().isoformat()
    }), 200


@app.route('/test-connection', methods=['GET'])
def test_connection():
    """Frontend compatibility health endpoint."""
    return jsonify({
        'status': 'connected',
        'service': 'enable-agents-api',
        'timestamp': datetime.now().isoformat()
    }), 200


@app.route('/api/save-project', methods=['POST', 'OPTIONS'])
@cross_origin()
def save_project():
    if request.method == 'OPTIONS':
        return jsonify({}), 200
    try:
        data = request.json
        username = data.get('username')
        name = data.get('name')
        query_used = data.get('query', '')
        leads = data.get('businesses', [])
        
        if not username or not name:
            return jsonify({'success': False, 'error': 'Missing username or name'}), 400
            
        # Create Project
        project = SavedProject(username=username, name=name, query_used=query_used)
        db.session.add(project)
        db.session.flush() # get project id
        
        # Add leads
        import json
        for lead_data in leads:
            lead = SavedLead(
                project_id=project.id,
                name=lead_data.get('name', ''),
                website=lead_data.get('website', ''),
                phone=lead_data.get('phone', ''),
                address=lead_data.get('address', ''),
                emails=json.dumps(lead_data.get('emails', [])),
                linkedin_links=json.dumps(lead_data.get('linkedin_urls', [])),
                social_links=json.dumps(lead_data.get('social_links', {})),
                raw_data=json.dumps(lead_data)
            )
            db.session.add(lead)
            
        db.session.commit()
        return jsonify({'success': True, 'project_id': project.id, 'message': 'Leads saved successfully'})
    except Exception as e:
        db.session.rollback()
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/append-project', methods=['POST', 'OPTIONS'])
@cross_origin()
def append_project():
    if request.method == 'OPTIONS':
        return jsonify({}), 200
    try:
        data = request.json
        username = data.get('username')
        project_id = data.get('projectId')
        leads = data.get('businesses', [])
        
        if not username or not project_id:
            return jsonify({'success': False, 'error': 'Missing username or projectId'}), 400
            
        # verify
        project = db.session.get(SavedProject, project_id)
        if not project or project.username != username:
            return jsonify({'success': False, 'error': 'Project not found.'}), 404
        
        # Check current existing names in project to append uniquely
        existing_leads = db.session.query(SavedLead).filter_by(project_id=project_id).all()
        existing_names = {lead.name for lead in existing_leads if lead.name}
        
        import json
        added = 0
        for lead_data in leads:
            b_name = lead_data.get('name', '')
            if b_name and b_name in existing_names:
                continue
                
            lead = SavedLead(
                project_id=project.id,
                name=b_name,
                website=lead_data.get('website', ''),
                phone=lead_data.get('phone', ''),
                address=lead_data.get('address', ''),
                emails=json.dumps(lead_data.get('emails', [])),
                linkedin_links=json.dumps(lead_data.get('linkedin_urls', [])),
                social_links=json.dumps(lead_data.get('social_links', {})),
                raw_data=json.dumps(lead_data)
            )
            db.session.add(lead)
            existing_names.add(b_name)
            added += 1
            
        db.session.commit()
        return jsonify({'success': True, 'project_id': project.id, 'message': f'Appended {added} new leads.'})
    except Exception as e:
        db.session.rollback()
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/saved-projects', methods=['GET'])
@cross_origin()
def get_saved_projects():
    try:
        username = request.args.get('username')
        if not username:
             return jsonify({'success': False, 'error': 'Missing username'}), 400
             
        projects = db.session.query(SavedProject).filter_by(username=username).order_by(SavedProject.created_at.desc()).all()
        result = []
        for p in projects:
            lead_count = db.session.query(SavedLead).filter_by(project_id=p.id).count()
            result.append({
                'id': p.id,
                'name': p.name,
                'query_used': p.query_used,
                'created_at': p.created_at.isoformat(),
                'lead_count': lead_count
            })
            
        return jsonify({'success': True, 'projects': result})
    except Exception as e:
         import traceback
         traceback.print_exc()
         return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/saved-projects/<int:project_id>/leads', methods=['GET'])
@cross_origin()
def get_project_leads(project_id):
    try:
        username = request.args.get('username')
        project = db.session.get(SavedProject, project_id)
        
        if not project or (username and project.username != username):
            return jsonify({'success': False, 'error': 'Project not found'}), 404
            
        leads = db.session.query(SavedLead).filter_by(project_id=project_id).all()
        import json
        result = []
        for l in leads:
            result.append({
                'id': l.id,
                'name': l.name,
                'website': l.website,
                'phone': l.phone,
                'address': l.address,
                'emails': json.loads(l.emails) if l.emails else [],
                'linkedin_urls': json.loads(l.linkedin_links) if l.linkedin_links else [],
                'social_links': json.loads(l.social_links) if l.social_links else {},
            })
            
        return jsonify({'success': True, 'project': {'id': project.id, 'name': project.name}, 'leads': result})
    except Exception as e:
         import traceback
         traceback.print_exc()
         return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/saved-projects/<int:project_id>', methods=['DELETE', 'OPTIONS'])
@cross_origin()
def delete_saved_project(project_id):
    if request.method == 'OPTIONS':
        return jsonify({}), 200
    try:
        username = request.args.get('username')
        if not username:
             return jsonify({'success': False, 'error': 'Missing username'}), 400
             
        project = db.session.get(SavedProject, project_id)
        if not project or project.username != username:
            return jsonify({'success': False, 'error': 'Project not found or unauthorized'}), 404
            
        db.session.delete(project)
        db.session.commit()
        return jsonify({'success': True, 'message': 'Project deleted successfully'})
    except Exception as e:
         db.session.rollback()
         import traceback
         traceback.print_exc()
         return jsonify({'success': False, 'error': str(e)}), 500

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    app.run(debug=True, host='0.0.0.0', port=5000)
