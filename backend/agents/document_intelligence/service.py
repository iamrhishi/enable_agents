"""
Document Intelligence Agent — Service layer (orchestration).

Handles:
- Document upload and validation
- Text extraction (PDF, DOCX, TXT)
- Processing pipeline coordination
- Chat interface
- Context sharing with other agents via ContextStore
"""

import logging
import os
import re
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple
from uuid import uuid4

from werkzeug.utils import secure_filename

from core.context import ContextStore

logger = logging.getLogger(__name__)

AGENT_ID = "document_intelligence"

# Supported file types
ALLOWED_EXTENSIONS = {"pdf", "docx", "doc", "txt", "xlsx", "xls", "csv"}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB


def allowed_file(filename: str) -> bool:
    """Check if file extension is allowed."""
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def get_file_extension(filename: str) -> str:
    """Get lowercase file extension."""
    return filename.rsplit(".", 1)[1].lower() if "." in filename else ""


class DocumentService:
    """Main service for document intelligence operations."""

    def __init__(self):
        self.upload_dir = os.path.join(os.getcwd(), "data", "uploads")
        os.makedirs(self.upload_dir, exist_ok=True)

    def upload_document(
        self,
        file,
        user_id: str,
        project_id: str,
        document_type: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        Upload and validate a document to a project.

        Documents are project-scoped and shared among all project members.

        Args:
            file: Werkzeug FileStorage object
            user_id: User ID (who uploaded)
            project_id: Project ID (document owner)
            document_type: Optional document category

        Returns:
            Dict with document_id, project_id, and status
        """
        from core.database import db
        from agents.document_intelligence.models import ProcessedDocument

        if not file or not file.filename:
            raise ValueError("No file provided")

        if not project_id:
            raise ValueError("project_id is required")

        filename = secure_filename(file.filename)
        if not allowed_file(filename):
            raise ValueError(f"File type not allowed. Supported: {ALLOWED_EXTENSIONS}")

        # Generate document ID
        document_id = str(uuid4())
        file_ext = get_file_extension(filename)

        # Create project-specific upload directory
        project_dir = os.path.join(self.upload_dir, project_id)
        os.makedirs(project_dir, exist_ok=True)

        # Save file
        file_path = os.path.join(project_dir, f"{document_id}.{file_ext}")
        file.save(file_path)

        # Get file size
        file_size = os.path.getsize(file_path)
        if file_size > MAX_FILE_SIZE:
            os.remove(file_path)
            raise ValueError(f"File too large. Maximum size: {MAX_FILE_SIZE // (1024*1024)}MB")

        # Create database record
        doc = ProcessedDocument(
            document_id=document_id,
            project_id=project_id,
            uploaded_by=user_id,
            file_name=filename,
            file_type=file_ext,
            file_size=file_size,
            file_path=file_path,
            status="pending",
            document_type=document_type,
        )
        db.session.add(doc)
        db.session.commit()

        logger.info(f"Document uploaded: {document_id} to project {project_id} by {user_id}")

        return {
            "document_id": document_id,
            "project_id": project_id,
            "file_name": filename,
            "file_size": file_size,
            "status": "pending",
        }

    def extract_text(self, document_id: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from document.

        Returns:
            Tuple of (extracted_text, metadata)
        """
        from core.database import db
        from agents.document_intelligence.models import ProcessedDocument

        doc = ProcessedDocument.query.get(document_id)
        if not doc:
            raise ValueError(f"Document not found: {document_id}")

        file_path = doc.file_path
        file_type = doc.file_type

        text = ""
        metadata = {}

        try:
            if file_type == "pdf":
                text, metadata = self._extract_pdf(file_path)
            elif file_type in ("docx", "doc"):
                text, metadata = self._extract_docx(file_path)
            elif file_type == "txt":
                text, metadata = self._extract_txt(file_path)
            elif file_type in ("xlsx", "xls"):
                text, metadata = self._extract_xlsx(file_path)
            elif file_type == "csv":
                text, metadata = self._extract_csv(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")

            # Update document record
            doc.extracted_text = text
            doc.page_count = metadata.get("page_count")
            doc.word_count = len(text.split())
            db.session.commit()

            logger.info(f"Extracted {len(text)} characters from {document_id}")

        except Exception as e:
            doc.status = "failed"
            doc.error_message = str(e)
            db.session.commit()
            raise

        return text, metadata

    def _extract_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF using PyMuPDF."""
        import fitz  # PyMuPDF

        text_parts = []
        metadata = {"page_count": 0}

        with fitz.open(file_path) as pdf:
            metadata["page_count"] = len(pdf)

            for page_num, page in enumerate(pdf):
                page_text = page.get_text()
                if page_text.strip():
                    text_parts.append(f"[Page {page_num + 1}]\n{page_text}")

        return "\n\n".join(text_parts), metadata

    def _extract_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX using python-docx."""
        from docx import Document

        doc = Document(file_path)
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n\n".join(text_parts), {"page_count": None}

    def _extract_txt(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from plain text file."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        return text, {"page_count": None}

    def _extract_xlsx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from Excel file using openpyxl."""
        import openpyxl

        workbook = openpyxl.load_workbook(file_path, data_only=True)
        text_parts = []
        sheet_count = len(workbook.sheetnames)

        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            rows = []

            for row in sheet.iter_rows(values_only=True):
                # Filter out completely empty rows
                row_values = [str(cell) if cell is not None else "" for cell in row]
                if any(v.strip() for v in row_values):
                    rows.append(" | ".join(row_values))

            if rows:
                text_parts.append(f"[Sheet: {sheet_name}]")
                text_parts.extend(rows)
                text_parts.append("")  # Blank line between sheets

        workbook.close()
        return "\n".join(text_parts), {"page_count": sheet_count, "sheet_count": sheet_count}

    def _extract_csv(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from CSV file."""
        import csv

        text_parts = []
        row_count = 0

        # Try different encodings
        for encoding in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
            try:
                with open(file_path, "r", encoding=encoding, newline="") as f:
                    reader = csv.reader(f)
                    for row in reader:
                        if any(cell.strip() for cell in row):
                            text_parts.append(" | ".join(row))
                            row_count += 1
                break
            except UnicodeDecodeError:
                continue

        return "\n".join(text_parts), {"page_count": None, "row_count": row_count}

    def process_document(self, document_id: str) -> Dict[str, Any]:
        """
        Full document processing pipeline.

        Stages:
        1. Extract text
        2. Chunk content
        3. Generate embeddings
        4. Extract entities

        Returns:
            Processing result dict
        """
        from core.database import db
        from core.vector_store import VectorStore
        from agents.document_intelligence.models import ProcessedDocument
        from agents.document_intelligence.chunking import create_semantic_chunks, get_adaptive_chunk_config

        doc = ProcessedDocument.query.get(document_id)
        if not doc:
            raise ValueError(f"Document not found: {document_id}")

        doc.status = "processing"
        doc.processing_stage = "extracting"
        doc.processing_progress = 0.1
        db.session.commit()

        try:
            # Stage 1: Extract text
            text, metadata = self.extract_text(document_id)

            doc.processing_stage = "chunking"
            doc.processing_progress = 0.3
            db.session.commit()

            # Stage 2: Chunk content
            chunk_size, overlap = get_adaptive_chunk_config(len(text))
            chunks = create_semantic_chunks(text, chunk_size=chunk_size, overlap=overlap)

            doc.chunk_count = len(chunks)
            doc.processing_stage = "embedding"
            doc.processing_progress = 0.5
            db.session.commit()

            # Stage 3: Generate and store embeddings
            vector_store = VectorStore()
            chunk_dicts = [{"content": c["content"], "metadata": c} for c in chunks]
            chunk_ids = vector_store.store_chunks(document_id, chunk_dicts, user_id=doc.project_id)

            doc.processing_stage = "entities"
            doc.processing_progress = 0.8
            db.session.commit()

            # Stage 4: Extract entities (basic implementation)
            entities = self._extract_entities(document_id, text, chunk_ids)
            doc.entity_count = len(entities)

            # Complete
            doc.status = "completed"
            doc.processing_stage = "complete"
            doc.processing_progress = 1.0
            doc.processed_at = datetime.utcnow()
            db.session.commit()

            logger.info(f"Document {document_id} processed: {len(chunks)} chunks, {len(entities)} entities")

            # Share document context with other agents via ContextStore
            self._update_context(doc, chunks, entities)

            return {
                "document_id": document_id,
                "status": "completed",
                "chunk_count": len(chunks),
                "entity_count": len(entities),
                "word_count": doc.word_count,
            }

        except Exception as e:
            doc.status = "failed"
            doc.error_message = str(e)
            db.session.commit()
            logger.error(f"Document processing failed: {document_id} - {e}")
            raise

    def _extract_entities(
        self,
        document_id: str,
        text: str,
        chunk_ids: List[str],
    ) -> List[Dict[str, Any]]:
        """
        Basic entity extraction using regex patterns.

        For production, this should use an LLM or NER model.
        """
        from core.database import db
        from agents.document_intelligence.models import DocumentEntity

        entities = []

        # Simple pattern-based extraction
        patterns = {
            "email": r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b",
            "phone": r"\b\d{3}[-.]?\d{3}[-.]?\d{4}\b",
            "url": r"https?://[^\s<>\"{}|\\^`\[\]]+",
            "money": r"\$[\d,]+(?:\.\d{2})?",
            "percentage": r"\d+(?:\.\d+)?%",
            "date": r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b",
        }

        for entity_type, pattern in patterns.items():
            matches = re.findall(pattern, text, re.IGNORECASE)
            unique_matches = list(set(matches))

            for match in unique_matches[:50]:  # Limit per type
                entity = DocumentEntity(
                    entity_id=str(uuid4()),
                    document_id=document_id,
                    entity_type=entity_type,
                    entity_name=match,
                    normalized_name=match.lower(),
                    chunk_ids=chunk_ids[:3],  # Associate with first chunks
                    confidence=0.8,
                )
                db.session.add(entity)
                entities.append(entity.to_dict())

        db.session.commit()
        return entities

    def _update_context(
        self,
        doc,
        chunks: List[Dict[str, Any]],
        entities: List[Dict[str, Any]],
    ) -> None:
        """
        Update ContextStore with document metadata for cross-agent access.

        Documents are project-scoped, so context is stored per project.
        Other agents can read this context to:
        - Know what documents are available in a project
        - Get document summaries without re-querying
        - Access extracted entities for their workflows
        """
        try:
            ctx = ContextStore()

            # Use project_id for context scoping (documents are project-owned)
            scope_id = doc.project_id or doc.uploaded_by

            # Store individual document metadata
            doc_context = {
                "document_id": doc.document_id,
                "project_id": doc.project_id,
                "file_name": doc.file_name,
                "file_type": doc.file_type,
                "document_type": doc.document_type,
                "word_count": doc.word_count,
                "chunk_count": len(chunks),
                "entity_count": len(entities),
                "status": "completed",
                "processed_at": datetime.utcnow().isoformat(),
            }

            # Add preview (first 500 chars of content)
            if chunks:
                preview = chunks[0].get("content", "")[:500]
                doc_context["preview"] = preview

            # Add entity summary (grouped by type)
            entity_summary = {}
            for e in entities[:100]:  # Limit to avoid huge context
                etype = e.get("entity_type", "unknown")
                if etype not in entity_summary:
                    entity_summary[etype] = []
                if len(entity_summary[etype]) < 10:
                    entity_summary[etype].append(e.get("entity_name"))
            doc_context["entities"] = entity_summary

            ctx.set(
                user_id=scope_id,
                agent_id=AGENT_ID,
                key=f"document:{doc.document_id}",
                value=doc_context,
            )

            # Update project's document index (list of available docs)
            existing_index = ctx.get(scope_id, AGENT_ID, "documents:index", default=[])
            # Add or update this document in the index
            doc_entry = {
                "document_id": doc.document_id,
                "file_name": doc.file_name,
                "document_type": doc.document_type,
                "updated_at": datetime.utcnow().isoformat(),
            }
            # Remove old entry if exists
            existing_index = [d for d in existing_index if d.get("document_id") != doc.document_id]
            # Add new entry at front
            existing_index.insert(0, doc_entry)
            # Keep last 50 documents in index
            existing_index = existing_index[:50]

            ctx.set(
                user_id=scope_id,
                agent_id=AGENT_ID,
                key="documents:index",
                value=existing_index,
            )

            logger.info(f"Updated ContextStore for document {doc.document_id} (project: {doc.project_id})")

        except Exception as e:
            # Don't fail document processing if context update fails
            logger.warning(f"Failed to update context for document {doc.document_id}: {e}")

    def get_document_status(
        self, document_id: str, project_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """Get document processing status.

        Args:
            document_id: Document ID
            project_id: Optional project ID to verify ownership
        """
        from agents.document_intelligence.models import ProcessedDocument

        doc = ProcessedDocument.query.get(document_id)

        if not doc:
            raise ValueError(f"Document not found: {document_id}")

        # Verify project ownership if project_id provided
        if project_id and doc.project_id != project_id:
            raise ValueError(f"Document not found in project: {document_id}")

        return doc.to_dict()

    def list_documents(
        self, project_id: str, limit: int = 50
    ) -> List[Dict[str, Any]]:
        """List project's documents.

        Documents are project-scoped and shared among all project members.
        """
        from agents.document_intelligence.models import ProcessedDocument

        docs = (
            ProcessedDocument.query.filter_by(project_id=project_id)
            .order_by(ProcessedDocument.created_at.desc())
            .limit(limit)
            .all()
        )

        return [doc.to_dict() for doc in docs]

    def delete_document(self, document_id: str, project_id: str) -> bool:
        """Delete a document and all associated data.

        Args:
            document_id: Document ID
            project_id: Project ID (for ownership verification)
        """
        from core.database import db
        from core.vector_store import VectorStore
        from agents.document_intelligence.models import ProcessedDocument

        doc = ProcessedDocument.query.filter_by(
            document_id=document_id, project_id=project_id
        ).first()

        if not doc:
            raise ValueError(f"Document not found: {document_id}")

        # Delete file
        if doc.file_path and os.path.exists(doc.file_path):
            os.remove(doc.file_path)

        # Delete embeddings
        vector_store = VectorStore()
        vector_store.delete_document_chunks(document_id)

        # Delete document (cascades to chunks and entities)
        db.session.delete(doc)
        db.session.commit()

        logger.info(f"Document deleted: {document_id} from project {project_id}")
        return True

    def chat(
        self,
        query: str,
        user_id: str,
        document_ids: Optional[List[str]] = None,
        use_entity_boost: bool = False,
        project_id: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        Chat with documents using RAG.

        Args:
            query: User question
            user_id: User ID
            document_ids: Optional filter to specific documents
            use_entity_boost: Whether to use entity-enhanced retrieval

        Returns:
            Dict with 'answer', 'sources', 'chunk_count'
        """
        import openai
        from agents.document_intelligence.retrieval import DocumentRetriever

        retriever = DocumentRetriever()

        # Scope to the active project's documents when the caller didn't
        # already pin specific document_ids - otherwise chat searches across
        # every document the user has ever uploaded, in any project.
        if not document_ids and project_id:
            project_docs = self.list_documents(project_id=project_id)
            document_ids = [
                d["document_id"] for d in project_docs if d.get("status") == "completed"
            ]
            if not document_ids:
                return {
                    "answer": "This project doesn't have any processed documents yet.",
                    "sources": [],
                    "chunk_count": 0,
                }

        # Get relevant context
        if use_entity_boost:
            context_result = retriever.get_entity_enhanced_context(
                query=query,
                user_id=user_id,
                document_ids=document_ids,
            )
        else:
            context_result = retriever.get_context_for_query(
                query=query,
                user_id=user_id,
                document_ids=document_ids,
            )

        if not context_result["context"]:
            return {
                "answer": "I couldn't find relevant information in your documents to answer this question.",
                "sources": [],
                "chunk_count": 0,
            }

        enriched_sources = self._enrich_sources(context_result["sources"])

        # Generate response using OpenAI
        client = openai.OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

        system_prompt = """You are a helpful assistant that answers questions based on the provided document context.
Use only the information from the context to answer. If the context doesn't contain enough information, say so.
Be concise but thorough. Cite relevant parts of the context when helpful."""

        user_prompt = f"""Context from documents:
{context_result['context']}

Question: {query}

Answer based on the context above:"""

        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                max_tokens=1000,
                temperature=0.3,
            )

            answer = response.choices[0].message.content

        except Exception as e:
            logger.error(f"Chat completion failed: {e}")
            return {
                "answer": f"Error generating response: {str(e)}",
                "sources": enriched_sources,
                "chunk_count": context_result["chunk_count"],
            }

        return {
            "answer": answer,
            "sources": enriched_sources,
            "chunk_count": context_result["chunk_count"],
        }

    def _enrich_sources(self, sources: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Attach real page number and content snippet to retrieval sources.

        The raw retriever only returns chunk_id/chunk_index/score - callers
        (chat UI, Ask AI) need an actual page number and a text excerpt to
        show a real citation instead of a hardcoded fallback.
        """
        from agents.document_intelligence.models import DocumentChunk

        chunk_ids = [s.get("chunk_id") for s in sources if s.get("chunk_id")]
        if not chunk_ids:
            return sources

        chunks_by_id = {
            c.chunk_id: c
            for c in DocumentChunk.query.filter(DocumentChunk.chunk_id.in_(chunk_ids)).all()
        }

        enriched = []
        for s in sources:
            chunk = chunks_by_id.get(s.get("chunk_id"))
            page_number = (chunk.metadata_dict.get("page") if chunk else None) or (
                (s.get("chunk_index") or 0) + 1
            )
            enriched.append(
                {
                    **s,
                    "page_number": page_number,
                    "content": chunk.content[:200] if chunk else "",
                }
            )
        return enriched

    def get_document_insight(
        self,
        document_id: str,
        project_id: Optional[str] = None,
        user_id: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        Generate a structured analysis (summary, key facts, recommendations,
        source citations) for a single processed document.

        Computed on demand from the document's real extracted text via an LLM
        call, plus real vector-similarity source citations - nothing here is
        canned/demo data.
        """
        import json as json_lib

        import openai

        from agents.document_intelligence.models import DocumentChunk, ProcessedDocument

        doc = ProcessedDocument.query.get(document_id)
        if not doc:
            raise ValueError(f"Document not found: {document_id}")
        if project_id and doc.project_id != project_id:
            raise ValueError(f"Document not found in project: {document_id}")

        if doc.status != "completed":
            return {
                "status": doc.status,
                "summary": None,
                "keyFacts": [],
                "recommendations": [],
                "sources": [],
            }

        chunks = (
            DocumentChunk.query.filter_by(document_id=document_id)
            .order_by(DocumentChunk.chunk_index)
            .all()
        )
        full_text = "\n\n".join(c.content for c in chunks)[:12000]

        client = openai.OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
        system_prompt = (
            "You analyze a business document and respond with strict JSON only, "
            'matching this shape: {"summary": "one paragraph", '
            '"key_facts": [{"fact": "...", "confidence": 0.0-1.0}], '
            '"recommendations": ["...", "..."]}. '
            "Base everything only on the provided document text - do not invent "
            "facts. Keep key_facts to the 5-8 most important, specific, and "
            "verifiable facts. Keep recommendations to 2-4 concrete, actionable "
            "items."
        )

        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {
                        "role": "user",
                        "content": f"Document: {doc.file_name}\n\n{full_text}",
                    },
                ],
                max_tokens=900,
                temperature=0.3,
                response_format={"type": "json_object"},
            )
            parsed = json_lib.loads(response.choices[0].message.content)
        except Exception as e:
            logger.error(f"Insight generation failed: {e}")
            parsed = {}

        key_facts = [
            {
                "fact": kf.get("fact", ""),
                "confidence": kf.get("confidence", 0.8),
                "source": doc.file_name,
            }
            for kf in parsed.get("key_facts", [])
            if kf.get("fact")
        ]
        summary = parsed.get("summary") or f"Analysis of {doc.file_name}."

        # Real retrieval-backed source citations - reuse vector search against
        # the generated summary so relevance scores are genuine similarity,
        # not fabricated numbers.
        sources = []
        if user_id and chunks:
            from agents.document_intelligence.retrieval import DocumentRetriever

            retriever = DocumentRetriever()
            context_result = retriever.get_context_for_query(
                query=summary,
                user_id=user_id,
                document_ids=[document_id],
                max_chunks=5,
            )
            chunk_by_id = {c.chunk_id: c for c in chunks}
            for src in context_result["sources"]:
                chunk = chunk_by_id.get(src.get("chunk_id"))
                if not chunk:
                    continue
                sources.append(
                    {
                        "page": chunk.metadata_dict.get("page", (chunk.chunk_index or 0) + 1),
                        "text": chunk.content[:220],
                        "relevance": src.get("score") or 0.5,
                    }
                )

        return {
            "status": doc.status,
            "summary": summary,
            "keyFacts": key_facts,
            "recommendations": parsed.get("recommendations") or [],
            "sources": sources,
        }


# ─────────────────────────────────────────────────────────────────────────────
# Cross-agent helpers: Allow other agents to access document intelligence
# ─────────────────────────────────────────────────────────────────────────────


def get_project_documents(project_id: str) -> List[Dict[str, Any]]:
    """
    Get list of project's processed documents from ContextStore.

    Documents are project-scoped and shared among all project members.

    Usage in other agents:
        from agents.document_intelligence.service import get_project_documents
        docs = get_project_documents(project_id)
    """
    ctx = ContextStore()
    return ctx.get(project_id, AGENT_ID, "documents:index", default=[])


# Alias for backwards compatibility
def get_user_documents(user_id: str) -> List[Dict[str, Any]]:
    """Deprecated: Use get_project_documents instead."""
    return get_project_documents(user_id)


def get_document_context(project_id: str, document_id: str) -> Optional[Dict[str, Any]]:
    """
    Get document metadata and entity summary from ContextStore.

    Args:
        project_id: Project ID that owns the document
        document_id: Document ID

    Returns None if document not found or not accessible.
    """
    ctx = ContextStore()
    return ctx.get(project_id, AGENT_ID, f"document:{document_id}", default=None)


def search_documents(
    project_id: str,
    query: str,
    document_ids: Optional[List[str]] = None,
    top_k: int = 5,
) -> List[Dict[str, Any]]:
    """
    Search project's documents for relevant content.

    Documents are project-scoped.

    Usage in other agents:
        from agents.document_intelligence.service import search_documents
        results = search_documents(project_id, "What is our pricing?")
    """
    from agents.document_intelligence.retrieval import DocumentRetriever

    retriever = DocumentRetriever()
    return retriever.search(
        query=query,
        user_id=project_id,  # project_id used as scope
        document_ids=document_ids,
        top_k=top_k,
    )


def get_document_context_for_prompt(
    project_id: str,
    query: str,
    document_ids: Optional[List[str]] = None,
    max_tokens: int = 2000,
) -> str:
    """
    Get assembled context string ready for LLM prompt injection.

    Usage in other agents:
        from agents.document_intelligence.service import get_document_context_for_prompt
        context = get_document_context_for_prompt(project_id, query)
        prompt = f"Based on these documents:\n{context}\n\nAnswer: {query}"
    """
    from agents.document_intelligence.retrieval import DocumentRetriever

    retriever = DocumentRetriever()
    result = retriever.get_context_for_query(
        query=query,
        user_id=project_id,  # project_id used as scope
        document_ids=document_ids,
        max_tokens=max_tokens,
    )
    return result.get("context", "")
