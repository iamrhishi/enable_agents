from flask import Flask, request, jsonify, redirect, g
import requests
import sqlite3
import shutil
import psutil
import tempfile
from datetime import datetime, timedelta
from uuid import uuid4
from time import time
import json
import os
import bs4
import fitz
import faiss
from typing import Dict, List
from typing_extensions import TypedDict
from dotenv import load_dotenv
from rake_nltk import Rake
import numpy as np
from scipy.spatial.distance import cosine
import openai
import nltk
import pandas as pd
import pickle
import hashlib
import zipfile
import http.client
import glob
from google_auth_oauthlib.flow import Flow
from google.oauth2.credentials import Credentials
import googleapiclient.discovery
from email.message import EmailMessage
import base64
from flask_cors import CORS, cross_origin
from flask_migrate import Migrate
from sqlalchemy import inspect, text
from werkzeug.security import generate_password_hash as _gen_pw_hash, check_password_hash

# Use pbkdf2 instead of scrypt for Python 3.9 compatibility
def generate_password_hash(password):
    return _gen_pw_hash(password, method='pbkdf2:sha256')
from werkzeug.utils import secure_filename
import openpyxl
from urllib.parse import urlencode, urlparse
import logging

from core.database import db
from core.context import ContextStore
from core.auth import require_auth

# LangChain imports with fallbacks for version compatibility
try:
    from langchain_openai import OpenAIEmbeddings, ChatOpenAI
except (ImportError, Exception) as e:
    try:
        from langchain.embeddings.openai import OpenAIEmbeddings
        from langchain.chat_models import ChatOpenAI
    except (ImportError, Exception):
        # Set placeholders for testing without full dependencies
        OpenAIEmbeddings = None
        ChatOpenAI = None
        print(f"Warning: LangChain OpenAI modules not available: {e}")

try:
    from langchain_community.vectorstores import FAISS
except (ImportError, Exception):
    try:
        from langchain.vectorstores import FAISS
    except (ImportError, Exception):
        FAISS = None
        print("Warning: FAISS not available")

try:
    from langchain_community.document_loaders import WebBaseLoader, PyPDFLoader
except (ImportError, Exception):
    try:
        from langchain.document_loaders import WebBaseLoader, PyPDFLoader
    except (ImportError, Exception):
        WebBaseLoader = None
        PyPDFLoader = None
        print("Warning: Document loaders not available")

try:
    from langchain_core.documents import Document
except (ImportError, Exception):
    try:
        from langchain.schema import Document
    except (ImportError, Exception):
        Document = None
        print("Warning: Document class not available")

try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except (ImportError, Exception):
    try:
        from langchain.text_splitter import RecursiveCharacterTextSplitter
    except (ImportError, Exception):
        RecursiveCharacterTextSplitter = None
        print("Warning: RecursiveCharacterTextSplitter not available")

try:
    from langchain_core.prompts import ChatPromptTemplate
except (ImportError, Exception):
    try:
        from langchain.prompts import ChatPromptTemplate
    except (ImportError, Exception):
        ChatPromptTemplate = None
        print("Warning: ChatPromptTemplate not available")
from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.common.exceptions import TimeoutException, NoSuchElementException
from webdriver_manager.chrome import ChromeDriverManager
import time
import re
from docx import Document as DocxDocument
import networkx as nx
# NOTE: GoogleBusinessHelper imports moved to functions to avoid circular import

PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))


def _compact_text(value):
    text = re.sub(r'\s+', ' ', str(value or '')).strip()
    return text
ENV_FILES = [
    os.path.join(PROJECT_ROOT, '.env'),
    os.path.join(PROJECT_ROOT, 'tools', '.env'),
    os.path.join(os.path.dirname(__file__), '.env'),
    os.path.join(PROJECT_ROOT, '.env.docker'),
]

# Docker Compose injects `.env.docker` before Python starts; `load_dotenv(..., override=True)`
# below can replace those URIs from repo/backend `.env` and break Google/LinkedIn redirect matching.
_oauth_google_redirect_from_process_env = os.environ.get('GOOGLE_REDIRECT_URI')
_oauth_linkedin_redirect_from_process_env = os.environ.get('LINKEDIN_REDIRECT_URI')

for env_file in ENV_FILES:
    if os.path.exists(env_file):
        load_dotenv(env_file, override=True)

if _oauth_google_redirect_from_process_env:
    os.environ['GOOGLE_REDIRECT_URI'] = _oauth_google_redirect_from_process_env
if _oauth_linkedin_redirect_from_process_env:
    os.environ['LINKEDIN_REDIRECT_URI'] = _oauth_linkedin_redirect_from_process_env

LINKEDIN_CLIENT_ID = os.getenv('LINKEDIN_CLIENT_ID')
LINKEDIN_CLIENT_SECRET = os.getenv('LINKEDIN_CLIENT_SECRET')
LINKEDIN_REDIRECT_URI = os.getenv('LINKEDIN_REDIRECT_URI', 'http://localhost:5000/linkedin/callback')

def _ensure_nltk_resource(resource_path, download_name):
    """Fetch NLTK data if missing; repair corrupt caches (e.g. partial zip → BadZipFile)."""
    force = False
    try:
        nltk.data.find(resource_path)
        return
    except LookupError:
        pass
    except (zipfile.BadZipFile, OSError) as exc:
        force = True
        print(
            f"Warning: NLTK {resource_path!r} unreadable ({type(exc).__name__}: {exc}); "
            f"re-downloading {download_name!r}."
        )
    try:
        nltk.download(download_name, quiet=True, force=force)
    except TypeError:
        nltk.download(download_name, quiet=True)
    except Exception as download_error:
        print(f"Warning: could not download NLTK resource {download_name}: {download_error}")
        return
    if force:
        try:
            nltk.data.find(resource_path)
        except Exception as err:
            print(f"Warning: NLTK {resource_path!r} still unavailable after re-download: {err}")


_ensure_nltk_resource('corpora/stopwords', 'stopwords')
_ensure_nltk_resource('tokenizers/punkt_tab', 'punkt_tab')

DATA_DIR = os.path.join(os.path.dirname(__file__), 'data')
PROMPTS_FILE = os.path.join(DATA_DIR, 'prompts.json')


def _default_frontend_url():
    """Browser origin for the SPA (CORS, defaults). Prefer FRONTEND_URL, then PUBLIC_URL for real hosts."""
    fe = (os.getenv('FRONTEND_URL') or '').strip().rstrip('/')
    if fe:
        return fe
    pub = (os.getenv('PUBLIC_URL') or '').strip().rstrip('/')
    if pub:
        pl = pub.lower()
        if pl.startswith('https://'):
            return pub
        if pl.startswith('http://') and 'localhost' not in pl and '127.0.0.1' not in pl:
            return pub
    if (os.getenv('DEPLOY_MODE') or '').strip().lower() == 'remote' and pub:
        return pub
    return 'http://localhost:3000'


def _spa_redirect_base():
    """OAuth/callback redirects: use FRONTEND_URL if set, otherwise derive from proxy headers."""
    # If FRONTEND_URL is explicitly set, use it
    fe = (os.getenv('FRONTEND_URL') or '').strip().rstrip('/')
    if fe:
        return fe
    # In dev mode (localhost:3000), always redirect to frontend, not to request.host (which could be backend port)
    base = _default_frontend_url()
    if base == 'http://localhost:3000':
        return base
    # Otherwise try to derive from reverse proxy headers (for remote deployments)
    scheme = (request.headers.get('X-Forwarded-Proto') or request.scheme or 'https').split(',')[0].strip()
    host = (request.headers.get('X-Forwarded-Host') or '').split(',')[0].strip()
    if host:
        return f'{scheme}://{host}'.rstrip('/')
    # Fallback to default frontend URL
    return _default_frontend_url()


app = Flask(__name__)
_sk = (os.getenv('SECRET_KEY') or '').strip()
if not _sk:
    _sk = 'local-dev-insecure-enable-agents-set-SECRET_KEY-in-production'
app.config['SECRET_KEY'] = _sk

FRONTEND_URL = _default_frontend_url()
_cors_raw = os.getenv('CORS_ORIGINS', '').strip()
if _cors_raw:
    _cors_origins = [o.strip() for o in _cors_raw.split(',') if o.strip()]
else:
    _cors_origins = [
        FRONTEND_URL,
    ]
_seen_cors = set()
CORS_ORIGINS = []
for _o in _cors_origins:
    if _o not in _seen_cors:
        _seen_cors.add(_o)
        CORS_ORIGINS.append(_o)
CORS(app, origins=CORS_ORIGINS, supports_credentials=False)

GOOGLE_CLIENT_ID = (os.getenv('GOOGLE_CLIENT_ID') or '').strip()
GOOGLE_CLIENT_SECRET = (os.getenv('GOOGLE_CLIENT_SECRET') or '').strip()
GOOGLE_REDIRECT_URI = (os.getenv('GOOGLE_REDIRECT_URI') or 'http://localhost:5000/auth/google/callback').strip()
if os.getenv('ENVIRONMENT') != 'production':
    os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1' # allow HTTP for local dev only

# Database config (env override + local fallback)
app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv('DATABASE_URI')
if not app.config['SQLALCHEMY_DATABASE_URI']:
    raise ValueError("DATABASE_URI environment variable is required. PostgreSQL connection string expected.")
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Initialize shared db instance from core.database
db.init_app(app)

# Import models for Flask-Migrate to detect them
from core import models as core_models  # noqa: F401

migrate = Migrate(app, db)

# Initialize Celery
from core.celery_app import make_celery
celery = make_celery(app)

# Enable pgvector extension for PostgreSQL
db_uri = app.config.get('SQLALCHEMY_DATABASE_URI', '')
if 'postgresql' in db_uri:
    with app.app_context():
        try:
            db.session.execute(text("CREATE EXTENSION IF NOT EXISTS vector"))
            db.session.commit()
        except Exception:
            db.session.rollback()

# Register SSE notification routes
from core.notifications import register_sse_routes
register_sse_routes(app)

# NOTE: Agent registration moved to end of file to avoid circular imports

ALLOWED_EXTENSIONS = {'csv', 'xlsx', 'xls'}
MAX_FILE_SIZE = 16 * 1024 * 1024  # 16MB
app.config['MAX_CONTENT_LENGTH'] = MAX_FILE_SIZE

# Content Marketing Agent Configuration
CONTENT_MARKETING_UPLOAD_FOLDER = os.environ.get('CONTENT_MARKETING_UPLOAD_FOLDER', os.path.join(os.path.dirname(__file__), 'data', 'content_marketing_uploads'))
CONTENT_MARKETING_ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt', 'xlsx', 'html', 'md'}
os.makedirs(CONTENT_MARKETING_UPLOAD_FOLDER, exist_ok=True)

# Content Marketing tables are now in PostgreSQL via SQLAlchemy models
# See: agents/content_marketing/models.py

class User(db.Model):
    __tablename__ = 'users'
    __table_args__ = {'extend_existing': True}
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password = db.Column(db.String(512), nullable=False)
    first_name = db.Column(db.String(80))
    last_name = db.Column(db.String(80))
    email = db.Column(db.String(120))
    company = db.Column(db.String(120))
    linkedin = db.Column(db.String(256))
    short_intro = db.Column(db.String(256))
    company_intro = db.Column(db.String(256))

class GoogleOAuthToken(db.Model):
    __tablename__ = 'google_oauth_tokens'
    __table_args__ = {'extend_existing': True}
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), db.ForeignKey('users.username'), nullable=False)
    token = db.Column(db.Text, nullable=False)
    refresh_token = db.Column(db.Text)
    token_uri = db.Column(db.String(512))
    client_id = db.Column(db.String(512))
    client_secret = db.Column(db.String(512))
    scopes = db.Column(db.Text)


EMAIL_EXTRACTION_UNIT_COST = float(os.getenv('EMAIL_EXTRACTION_UNIT_COST', '0.20'))
DEFAULT_EMAIL_EXTRACTION_LIMIT = int(os.getenv('EMAIL_EXTRACTION_DEFAULT_LIMIT', '500'))


# Email models imported from centralized location
from agents.email_outreach.models import (
    EmailCampaign,
    EmailCampaignRecipient,
    EmailExtractionQuota,
)


class EmailExtractionUsageLog(db.Model):
    __tablename__ = 'email_extraction_usage_logs'
    __table_args__ = {'extend_existing': True}
    id = db.Column(db.Integer, primary_key=True)
    request_id = db.Column(db.String(64), unique=True, nullable=False, index=True)
    username = db.Column(db.String(120), nullable=False, index=True)
    processed_count = db.Column(db.Integer, nullable=False, default=0)
    billable_count = db.Column(db.Integer, nullable=False, default=0)
    charged_count = db.Column(db.Integer, nullable=False, default=0)
    cost_this_request = db.Column(db.Float, nullable=False, default=0.0)
    total_cost_after = db.Column(db.Float, nullable=False, default=0.0)
    created_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)

class SavedProject(db.Model):
    __tablename__ = 'saved_projects'
    __table_args__ = {'extend_existing': True}
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(120), nullable=False, index=True)
    name = db.Column(db.String(255), nullable=False)
    query_used = db.Column(db.String(512), nullable=True)
    created_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)

class SavedLead(db.Model):
    __tablename__ = 'saved_leads'
    __table_args__ = {'extend_existing': True}
    id = db.Column(db.Integer, primary_key=True)
    project_id = db.Column(db.Integer, db.ForeignKey('saved_projects.id', ondelete='CASCADE'), nullable=False)
    
    name = db.Column(db.String(255))
    website = db.Column(db.String(512))
    phone = db.Column(db.String(100))
    address = db.Column(db.String(512))
    
    emails = db.Column(db.Text) 
    linkedin_links = db.Column(db.Text)
    social_links = db.Column(db.Text)
    
    has_extracted = db.Column(db.Boolean, default=False)
    
    raw_data = db.Column(db.Text) 
    
    created_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)


def _ensure_email_usage_tables():
    """Create usage tracking tables if they don't exist yet."""
    db.create_all()


def _ensure_campaign_reply_tracking_columns():
    """Ensure campaign/reply tracking columns exist."""
    try:
        inspector = inspect(db.engine)
        recipient_columns = {col['name'] for col in inspector.get_columns('email_campaign_recipients')}
        campaign_columns = {col['name'] for col in inspector.get_columns('email_campaigns')}
        statements = []

        if 'message_id' not in recipient_columns:
            statements.append("ALTER TABLE email_campaign_recipients ADD COLUMN message_id VARCHAR(255)")
        if 'thread_id' not in recipient_columns:
            statements.append("ALTER TABLE email_campaign_recipients ADD COLUMN thread_id VARCHAR(255)")
        if 'reply_subject' not in recipient_columns:
            statements.append("ALTER TABLE email_campaign_recipients ADD COLUMN reply_subject VARCHAR(512)")
        if 'reply_snippet' not in recipient_columns:
            statements.append("ALTER TABLE email_campaign_recipients ADD COLUMN reply_snippet TEXT")
        if 'reply_body' not in recipient_columns:
            statements.append("ALTER TABLE email_campaign_recipients ADD COLUMN reply_body TEXT")
        if 'sender_email' not in campaign_columns:
            statements.append("ALTER TABLE email_campaigns ADD COLUMN sender_email VARCHAR(255)")

        for stmt in statements:
            db.session.execute(text(stmt))

        if statements:
            db.session.commit()
            print(f"[DB MIGRATION] Applied tracking schema updates: {', '.join(statements)}")
    except Exception as migration_error:
        db.session.rollback()
        print(f"[DB MIGRATION] Could not ensure reply-tracking columns: {migration_error}")


def _extract_gmail_message_text(message):
    """Best-effort extraction of plain text from a Gmail message payload."""
    try:
        payload = message.get('payload') or {}
        if message.get('snippet'):
            text_value = str(message.get('snippet')).strip()
        else:
            text_value = ''

        def _walk_parts(part):
            nonlocal text_value
            if not isinstance(part, dict):
                return
            mime_type = (part.get('mimeType') or '').lower()
            body = part.get('body') or {}
            data = body.get('data')
            if data and mime_type in {'text/plain', 'text/html'}:
                try:
                    decoded = base64.urlsafe_b64decode(data.encode('utf-8')).decode('utf-8', errors='ignore')
                    decoded = re.sub(r'<[^>]+>', ' ', decoded)
                    decoded = re.sub(r'\s+', ' ', decoded).strip()
                    if decoded:
                        text_value = f"{text_value} {decoded}".strip() if text_value else decoded
                except Exception:
                    pass
            for child in part.get('parts', []) or []:
                _walk_parts(child)

        _walk_parts(payload)
        return re.sub(r'\s+', ' ', text_value).strip()
    except Exception:
        return (message.get('snippet') or '').strip()


def _extract_gmail_message_subject(message):
    headers = (message.get('payload') or {}).get('headers', [])
    return next((h['value'] for h in headers if h.get('name', '').lower() == 'subject'), '')


def _resolve_google_token_for_campaign(campaign, fallback_email=None, fallback_username=None):
    """Resolve the best Google token for a campaign using available identifiers."""
    candidates = []

    if fallback_email:
        candidates.append(fallback_email)
    if campaign.sender_email:
        candidates.append(campaign.sender_email)
    if campaign.username and '@' in campaign.username:
        candidates.append(campaign.username)
    if fallback_username and '@' in fallback_username:
        candidates.append(fallback_username)

    # Try mapping app username/first name to user email as a fallback.
    if campaign.username:
        user_by_username = User.query.filter_by(username=campaign.username).first()
        if user_by_username and user_by_username.email:
            candidates.append(user_by_username.email)

        user_by_first_name = User.query.filter_by(first_name=campaign.username).first()
        if user_by_first_name and user_by_first_name.email:
            candidates.append(user_by_first_name.email)

    # Deduplicate while preserving order.
    seen = set()
    ordered_candidates = []
    for c in candidates:
        key = (c or '').strip().lower()
        if key and key not in seen:
            seen.add(key)
            ordered_candidates.append(key)

    for candidate in ordered_candidates:
        token_record = GoogleOAuthToken.query.filter_by(username=candidate).first()
        if token_record and token_record.token:
            return token_record

    return None


def _sync_replies_for_campaign(campaign, fallback_email=None, fallback_username=None):
    """Best-effort Gmail reply sync for one campaign. Returns number of updates."""
    token_record = _resolve_google_token_for_campaign(campaign, fallback_email, fallback_username)
    recipients = EmailCampaignRecipient.query.filter_by(campaign_id=campaign.id).all()
    updated_count = 0

    if token_record and token_record.token:
        from google.oauth2.credentials import Credentials
        import googleapiclient.discovery

        try:
            creds = Credentials(
                token=token_record.token,
                refresh_token=token_record.refresh_token,
                token_uri=token_record.token_uri,
                client_id=token_record.client_id,
                client_secret=token_record.client_secret,
                scopes=token_record.scopes.split(',') if token_record.scopes else SCOPES
            )
            service = googleapiclient.discovery.build('gmail', 'v1', credentials=creds)

            for recipient in recipients:
                if not recipient.thread_id:
                    continue

                needs_text_repair = not _compact_text(recipient.reply_body) and not _compact_text(recipient.reply_snippet) and not _compact_text(recipient.reply_subject)
                if recipient.reply_status == 'Replied' and not needs_text_repair:
                    continue

                try:
                    thread = service.users().threads().get(userId='me', id=recipient.thread_id).execute()
                    messages = thread.get('messages', [])

                    for msg in messages:
                        headers = msg.get('payload', {}).get('headers', [])
                        from_header = next((h['value'] for h in headers if h['name'].lower() == 'from'), '')

                        if recipient.receiver_email and recipient.receiver_email.lower() in from_header.lower():
                            reply_subject = _extract_gmail_message_subject(msg)
                            reply_snippet = (msg.get('snippet') or '').strip()
                            reply_body = _extract_gmail_message_text(msg)
                            recipient.reply_status = 'Replied'
                            internal_date = int(msg.get('internalDate', 0)) / 1000.0
                            recipient.replied_at = datetime.fromtimestamp(internal_date) if internal_date > 0 else datetime.utcnow()
                            recipient.reply_subject = reply_subject or recipient.reply_subject
                            recipient.reply_snippet = reply_snippet or recipient.reply_snippet
                            recipient.reply_body = reply_body or recipient.reply_body
                            updated_count += 1
                            break
                except Exception as thread_err:
                    print(f"[REPLY SYNC] Error checking thread {recipient.thread_id}: {thread_err}")
        except Exception as gmail_sync_err:
            print(f"[REPLY SYNC] Gmail sync unavailable for campaign {campaign.id}: {gmail_sync_err}")

    try:
        imap_updates = _sync_replies_via_imap(campaign)
        updated_count += imap_updates
    except Exception as imap_err:
        print(f"[REPLY SYNC] IMAP sync unavailable for campaign {campaign.id}: {imap_err}")

    if updated_count > 0:
        db.session.commit()

    return updated_count


def _sync_replies_via_imap(campaign):
    """Fallback reply sync using the system mailbox via IMAP."""
    import imaplib
    import email as py_email
    from email import policy

    imap_host = (os.getenv('IMAP_HOST') or '').strip() or 'imap.gmail.com'
    imap_port = int(os.getenv('IMAP_PORT', '993'))
    imap_user = (os.getenv('EMAIL_USER') or '').strip()
    imap_pass = (os.getenv('EMAIL_PASS') or '').strip()

    if not imap_user or not imap_pass:
        return 0

    try:
        mailbox = imaplib.IMAP4_SSL(imap_host, imap_port)
        mailbox.login(imap_user, imap_pass)
        mailbox.select('INBOX')

        recipients = EmailCampaignRecipient.query.filter_by(campaign_id=campaign.id).all()
        pending = [
            r for r in recipients
            if r.reply_status != 'Replied'
            or not _compact_text(r.reply_body)
            or not _compact_text(r.reply_snippet)
            or not _compact_text(r.reply_subject)
        ]
        if not pending:
            mailbox.logout()
            return 0

        status, data = mailbox.search(None, 'ALL')
        if status != 'OK':
            mailbox.logout()
            return 0

        message_ids = (data[0] or b'').split()
        updated_count = 0

        # Check recent messages first.
        for msg_id in reversed(message_ids[-100:]):
            try:
                _, msg_data = mailbox.fetch(msg_id, '(RFC822)')
                if not msg_data or not msg_data[0]:
                    continue
                raw_email = msg_data[0][1]
                message = py_email.message_from_bytes(raw_email, policy=policy.default)
                from_header = (message.get('From') or '').lower()
                subject = (message.get('Subject') or '').lower()
                in_reply_to = (message.get('In-Reply-To') or '').strip()
                references = (message.get('References') or '').strip()

                def _extract_message_text(email_message):
                    if email_message.is_multipart():
                        parts = []
                        for part in email_message.walk():
                            content_type = (part.get_content_type() or '').lower()
                            disposition = (part.get_content_disposition() or '').lower()
                            if part.is_multipart() or disposition == 'attachment':
                                continue
                            if content_type in {'text/plain', 'text/html'}:
                                try:
                                    payload = part.get_content()
                                except Exception:
                                    raw_part = part.get_payload(decode=True) or b''
                                    charset = part.get_content_charset() or 'utf-8'
                                    payload = raw_part.decode(charset, errors='ignore')
                                if payload:
                                    parts.append(str(payload))
                        return _compact_text('\n'.join(parts))

                    try:
                        payload = email_message.get_content()
                    except Exception:
                        raw_payload = email_message.get_payload(decode=True) or b''
                        charset = email_message.get_content_charset() or 'utf-8'
                        payload = raw_payload.decode(charset, errors='ignore')
                    return _compact_text(payload or '')

                for recipient in pending:
                    recipient_email = (recipient.receiver_email or '').lower().strip()
                    message_id = (recipient.message_id or '').strip()
                    if not recipient_email:
                        continue

                    matches_sender = recipient_email in from_header
                    matches_thread = bool(message_id) and (message_id in in_reply_to or message_id in references)
                    is_reply_subject = subject.startswith('re:') or subject.startswith('fw:')

                    if matches_sender and (matches_thread or is_reply_subject):
                        reply_subject = (message.get('Subject') or '').strip()
                        reply_snippet = _compact_text((message.get_body(preferencelist=('plain', 'html')).get_content() if message.get_body(preferencelist=('plain', 'html')) else '') or '')
                        reply_body = _extract_message_text(message)
                        recipient.reply_status = 'Replied'
                        if not recipient.replied_at:
                            recipient.replied_at = datetime.utcnow()
                        recipient.reply_subject = reply_subject or recipient.reply_subject
                        recipient.reply_snippet = reply_snippet or recipient.reply_snippet or reply_body[:220]
                        recipient.reply_body = reply_body or recipient.reply_body or reply_snippet
                        updated_count += 1
                        break
            except Exception as imap_msg_error:
                print(f"[IMAP SYNC] Error parsing message {msg_id}: {imap_msg_error}")

        if updated_count > 0:
            db.session.commit()

        mailbox.logout()
        return updated_count
    except Exception as imap_error:
        print(f"[IMAP SYNC] Error checking inbox for campaign {campaign.id}: {imap_error}")
        return 0


def _normalize_username(value):
    candidate = (value or '').strip()
    return candidate if candidate else 'anonymous'


def _require_context_service_key_or_401():
    """
    When CONTEXT_API_SECRET is set, require matching X-Context-Api-Key.
    Session identity must use Authorization: Bearer <browser_session_token> (signed with SECRET_KEY).
    """
    secret = (os.getenv('CONTEXT_API_SECRET') or '').strip()
    if not secret:
        return None
    header_key = request.headers.get('X-Context-Api-Key', '').strip()
    if header_key != secret:
        return jsonify({
            'success': False,
            'error': 'Invalid or missing X-Context-Api-Key (CONTEXT_API_SECRET is set)',
        }), 401
    return None


def _production_auth_strict():
    return os.getenv('ENVIRONMENT', '').strip().lower() == 'production'


def _bearer_session_user_id():
    auth = (request.headers.get('Authorization') or '').strip()
    if not auth.lower().startswith('bearer '):
        return None
    raw = auth[7:].strip()
    if not raw:
        return None
    from core.session_token import verify_browser_session_token

    return verify_browser_session_token(app.config.get('SECRET_KEY') or '', raw)


def _resolve_session_user_id(fallback_claimed=None):
    """
    Prefer Authorization: Bearer <signed session_token> from login/Google/register.
    In production without a valid bearer → 401.
    In development/test, falls back to client-supplied user id (legacy).
    """
    uid = _bearer_session_user_id()
    if uid:
        return uid, None
    if _production_auth_strict():
        return None, (
            jsonify({
                'success': False,
                'error': 'Missing or invalid session. Sign in and send Authorization: Bearer <session_token>.',
            }),
            401,
        )
    return _normalize_username(fallback_claimed), None


def _effective_context_api_user_id(payload_dict):
    """Bearer always wins; with CONTEXT_API_SECRET, bearer is mandatory (ignore body principal)."""
    uid = _bearer_session_user_id()
    if uid:
        return uid, None
    if (os.getenv('CONTEXT_API_SECRET') or '').strip():
        return None, (
            jsonify({
                'success': False,
                'error': 'Bearer session token required when CONTEXT_API_SECRET is set',
            }),
            401,
        )
    claimed = (
        (payload_dict or {}).get('user_id') or (payload_dict or {}).get('username')
    )
    return _resolve_session_user_id(claimed)


def _shared_context_entry_key(source_agent, source_action, entry_type, payload):
    digest_source = json.dumps(
        {
            'source_agent': source_agent,
            'source_action': source_action,
            'entry_type': entry_type,
            'payload': payload,
        },
        sort_keys=True,
        default=str,
    )
    return hashlib.sha1(digest_source.encode('utf-8')).hexdigest()


def _shared_context_payload_dict(
    source_agent, source_action, entry_type, project_id, data_obj, text_value, metadata_obj
):
    return {
        'source_agent': source_agent,
        'source_action': source_action,
        'entry_type': entry_type,
        'project_id': project_id,
        'data': data_obj,
        'text': text_value,
        'entry_metadata': metadata_obj or {},
    }


def _is_billable_email(value):
    if not value:
        return False

    email_str = str(value).strip()
    lowered = email_str.lower()
    if lowered in {'n/a', 'na', 'none', 'error'}:
        return False
    if lowered.startswith('phone:'):
        return False

    return re.fullmatch(r'[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}', email_str) is not None


def _get_or_create_quota(username):
    quota = EmailExtractionQuota.query.filter_by(username=username).first()
    if quota:
        return quota

    quota = EmailExtractionQuota(
        username=username,
        monthly_limit=DEFAULT_EMAIL_EXTRACTION_LIMIT,
        emails_used_this_month=0
    )
    db.session.add(quota)
    db.session.commit()
    return quota


def _build_usage_summary(username, quota=None):
    quota = quota or _get_or_create_quota(username)
    used = max(quota.emails_used_this_month, 0)
    remaining = max(quota.monthly_limit - used, 0)

    return {
        'username': username,
        'totalAllowed': quota.monthly_limit,
        'usedCount': used,
        'remainingCount': remaining,
        'unitCost': EMAIL_EXTRACTION_UNIT_COST,
        'totalCost': round(used * EMAIL_EXTRACTION_UNIT_COST, 2)
    }

# 1. Load: First we need to load our data. This is done with Document Loaders.
# 2. Split: Text splitters break large Documents into smaller chunks. This is useful both for indexing data and passing it into a model, as large chunks are harder to search over and won't fit in a model's finite context window.
# 3. Store: We need somewhere to store and index our splits, so that they can be searched over later. This is often done using a VectorStore and Embeddings model.
# 4. Retrieve: Given a user input, relevant splits are retrieved from storage using a Retriever.
# 5. Generate: A ChatModel / LLM produces an answer using a prompt that includes both the question with the retrieved data

cache = {}

# Cache for KG+RAG to avoid recreating embeddings and graphs for same documents
kg_rag_cache = {
    'embeddings': {},      # Store embeddings by document hash
    'faiss_indices': {},   # Store FAISS indices by document hash
    'chunks': {},          # Store text chunks by document hash
    'knowledge_graphs': {} # Store knowledge graphs by nodes/edges hash
}

# Define state for application
class State(TypedDict):
    question: str
    context: 'List'  # List of Document objects
    answer: str


# ====== CONTENT MARKETING AGENT CLASSES & HELPERS ======

class DomainSpecializationAnalyzer:
    """Analyzes documents to extract domain specialization information"""
    
    def __init__(self):
        self.llm = ChatOpenAI(model="gpt-4", temperature=0)
        self.industry_keywords = self._load_industry_keywords()
    
    def _load_industry_keywords(self) -> Dict[str, List[str]]:
        """Load industry-specific keywords"""
        return {
            'Technology': ['software', 'cloud', 'api', 'infrastructure', 'devops', 'saas'],
            'Healthcare': ['medical', 'patient', 'pharmaceutical', 'clinical', 'health', 'disease'],
            'Finance': ['banking', 'investment', 'portfolio', 'trading', 'compliance', 'regulatory'],
            'Retail': ['ecommerce', 'inventory', 'customer', 'sales', 'purchase', 'product'],
            'Manufacturing': ['production', 'supply chain', 'logistics', 'quality', 'automation'],
            'Real Estate': ['property', 'tenant', 'lease', 'valuation', 'construction'],
            'Education': ['student', 'curriculum', 'learning', 'course', 'assessment'],
        }
    
    def analyze_documents(self, documents: List[str]) -> Dict:
        """
        Analyze documents to extract domain specialization
        
        Args:
            documents: List of document texts
            
        Returns:
            Dictionary with industry, sector, function, role analysis
        """
        combined_text = ' '.join(documents[:3]) if documents else ''
        
        prompt = ChatPromptTemplate.from_template("""
        Analyze the following business documents and extract domain specialization information.
        
        Documents:
        {documents}
        
        Provide a JSON response with:
        {{
            "industry": "identified industry",
            "sector": "business sector",
            "function": "primary business function",
            "role": "primary role/persona",
            "target_audience": "target customer/audience",
            "value_proposition": "key value proposition",
            "tone": "recommended tone (professional/casual/formal)",
            "key_themes": ["theme1", "theme2", ...]
        }}
        """)
        
        try:
            chain = prompt | self.llm
            response = chain.invoke({"documents": combined_text[:2000]})
            
            import re
            json_match = re.search(r'\{.*\}', response.content, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
        except:
            pass
        
        return {
            "industry": "General",
            "sector": "Unknown",
            "function": "Marketing",
            "role": "Marketing Manager",
            "target_audience": "Business Professionals",
            "value_proposition": "Enhanced marketing through AI",
            "tone": "professional",
            "key_themes": ["innovation", "value", "efficiency"]
        }


def extract_text_from_file_content_marketing(file_path: str, file_type: str) -> str:
    """Extract text content from various file formats"""
    try:
        if file_type == 'pdf':
            text = []
            pdf_document = fitz.open(file_path)
            for page in pdf_document:
                text.append(page.get_text())
            pdf_document.close()
            return '\n'.join(text)
        
        elif file_type == 'docx':
            doc = DocxDocument(file_path)
            return '\n'.join([para.text for para in doc.paragraphs])
        
        elif file_type == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        elif file_type == 'html':
            with open(file_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
                return soup.get_text()
        
        elif file_type == 'md':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        else:
            return ''
    
    except Exception as e:
        print(f"Error extracting text from {file_path}: {str(e)}")
        return ''


def setup_driver(headless=True):
    """Setup Chrome WebDriver for interactive scraping with popup handling"""
    chrome_options = Options()
    
    if headless:
        chrome_options.add_argument('--headless')
    
    chrome_options.add_argument('--no-sandbox')
    chrome_options.add_argument('--disable-dev-shm-usage')
    chrome_options.add_argument('--disable-gpu')
    chrome_options.add_argument('--window-size=1920,1080')
    chrome_options.add_argument('--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36')
    
    # Disable notifications and popups
    prefs = {
        "profile.default_content_setting_values.notifications": 2,
        "profile.default_content_settings.popups": 0,
        "profile.managed_default_content_settings.images": 2  # Block images for faster loading
    }
    chrome_options.add_experimental_option("prefs", prefs)
    
    # Additional options to handle consent and cookies
    chrome_options.add_argument('--disable-features=VizDisplayCompositor')
    chrome_options.add_argument('--disable-extensions')
    chrome_options.add_argument('--disable-plugins')
    
    service = Service(ChromeDriverManager().install())
    return webdriver.Chrome(service=service, options=chrome_options)

def get_chrome_history_path():
    """Get Chrome history file path based on OS"""
    if os.name == 'nt':  # Windows
        return os.path.expanduser('~\\AppData\\Local\\Google\\Chrome\\User Data\\Default\\History')
    elif 'darwin' in os.sys.platform.lower():  # macOS
        return os.path.expanduser('~/Library/Application Support/Google/Chrome/Default/History')
    else:  # Linux
        return os.path.expanduser('~/.config/google-chrome/Default/History')

def identify_saas_tools_with_openai(history_data):
    """Use OpenAI to identify which URLs are web applications, tools, SaaS, PaaS, or productivity platforms"""
    try:
        # Extract URLs for analysis (limit to avoid token limits)
        urls_to_analyze = [item['url'] for item in history_data[:50]]  # Analyze top 50 URLs
        urls_text = "\n".join([f"{i+1}. {url}" for i, url in enumerate(urls_to_analyze)])
        
        prompt = f"""Analyze the following URLs and identify which ones are web applications, tools, websites, or platforms. This includes:

- Web Applications
- Websites
- Tools
- Platforms
- PaaS (Platform as a Service) platforms of any kind
- Web-based productivity tools 
- Cloud platforms and services
- Development tools and platforms
- Business applications and tools
- Design and creative tools
- Communication and collaboration platforms
- Analytics and monitoring tools
- Project management tools
- CRM and business software
- Educational and learning platforms
- Entertainment and media platforms (if they're tools/apps)
- Social media platforms (if used as business tools)
- Storage and file-sharing platforms
- Any other web-based tools or platforms not mentioned above

For each URL, return a JSON array with this structure:
[
  {{"url_index": 1, "is_tool": true, "tool_name": "Google Docs", "category": "Productivity", "type": "SaaS", "description": "Document creation and collaboration"}},
  {{"url_index": 2, "is_tool": true, "tool_name": "AWS Console", "category": "Cloud Platform", "type": "PaaS", "description": "Cloud computing services"}},
  {{"url_index": 3, "is_tool": false, "tool_name": null, "category": null, "type": null, "description": "Regular website"}}
]

URLs to analyze:
{urls_text}

Rules:
- Identify ANY website, web-based tool, application, or platform
- Include Google Workspace (Docs, Sheets, Drive, Gmail), Microsoft 365, Slack, Zoom, etc.
- Include development platforms (GitHub, GitLab, Heroku, Vercel)
- Include cloud platforms (AWS, Azure, GCP)  
- Include design tools (Figma, Canva, Adobe Creative Cloud)
- Include business tools (Salesforce, HubSpot, Trello, Asana)
- Include social media platforms if used as business tools
- Categories: Development, Communication, Productivity, Design, Analytics, Cloud Platform, CRM, Project Management, Storage, Entertainment, Education, Social Media, Other
- Types: PaaS, Web App, Platform, Tool, Website
- Differentiate between a tool, WebApp, Website, and Platform
- Keep descriptions short (under 60 characters)
- Only mark as "false" if it's clearly a regular informational website
- Return valid JSON only"""

        client = openai.OpenAI()
        client.api_key = os.environ['OPENAI_API_KEY']
        
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a web application and tool identifier. You identify ALL types of web-based tools, applications, and platforms. Return only valid JSON arrays."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=3000,  # Increased for more detailed analysis
            temperature=0.1
        )
        
        response_text = response.choices[0].message.content.strip()
        response_text = response_text.replace('```json', '').replace('```', '').strip()
        
        print(response_text)

        try:
            tools_analysis = json.loads(response_text)

            
    
            # Create a mapping from URL index to tool info
            tools_mapping = {}
            for item in tools_analysis:
                if 'url_index' in item:
                    tools_mapping[item['url_index'] - 1] = {  # Convert to 0-based index
                        'is_tool': item.get('is_tool', False),
                        'tool_name': item.get('tool_name'),
                        'category': item.get('category'),
                        'type': item.get('type'),  # SaaS, PaaS, Web App, etc.
                        'description': item.get('description')
                    }
            
            return {
                'success': True,
                'mapping': tools_mapping
            }
            
        except json.JSONDecodeError as e:
            print(f"JSON parsing failed for tools analysis: {str(e)}")
            return {
                'success': False,
                'error': f'JSON parsing failed: {str(e)}'
            }
            
    except Exception as e:
        print(f"OpenAI tools analysis error: {str(e)}")
        return {
            'success': False,
            'error': f'OpenAI error: {str(e)}'
        }



def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def clean_dataframe_strict(df, required_field_groups=None):
    """Clean dataframe with strict field group requirements"""
    # ... existing cleaning code ...
    
    if required_field_groups:
        condition = pd.Series([True] * len(df))
        
        for group_name, field_list in required_field_groups.items():
            # Find which fields from this group exist in the dataframe
            existing_fields = [field for field in field_list if field in df.columns]
            
            if existing_fields:
                # For this group, at least one field must have data
                group_condition = pd.Series([False] * len(df))
                
                for field in existing_fields:
                    field_condition = (df[field] != '') & (df[field].str.lower() != 'n/a') & (df[field].str.lower() != 'na') & (df[field].str.lower() != 'none') & (df[field].str.lower() != 'null')
                    group_condition = group_condition | field_condition
                
                # All groups must have at least one field with data
                condition = condition & group_condition
        
        df = df[condition]
        print(f"After strict filtering: {len(df)} rows remaining")
    
    return df


def clean_dataframe(df, required_columns=None):
    """Clean dataframe by removing rows with missing required columns"""
    if required_columns is None:
        return df
    
    # Filter rows where at least one of the required columns has a value
    if required_columns:
        condition = pd.Series([False] * len(df))
        for col in required_columns:
            if col in df.columns:
                condition = condition | (df[col].notna() & (df[col] != ''))
        df = df[condition]
    
    return df


def clean_csv(df, required_columns=None):
    """Remove rows where any specified columns have blank/missing values"""
    if required_columns is None:
        required_columns = df.columns
    
    # Replace empty strings and common null representations with NaN
    df[required_columns] = df[required_columns].replace(['', ' ', 'N/A', 'n/a', 'NA', 'na', 'null', 'NULL'], pd.NA)
    
    # Drop rows with any NaN in required columns
    return df.dropna(subset=required_columns, how='any')


def csv_to_json(file_path):
    """Convert CSV file to JSON object with row filtering"""
    try:
        # Read CSV with error handling for different encodings
        try:
            df = pd.read_csv(file_path, encoding='utf-8')
        except UnicodeDecodeError:
            try:
                df = pd.read_csv(file_path, encoding='latin-1')
            except UnicodeDecodeError:
                df = pd.read_csv(file_path, encoding='cp1252')
        
        original_count = len(df)
        
        # Clean the dataframe - specify required columns or clean all
        required_columns = ['company', 'Company', 'name', 'Name', 'title', 'Title']
        existing_required = [col for col in required_columns if col in df.columns]
        
        if existing_required:
            df = clean_csv(df, existing_required)
        else:
            df = clean_csv(df)  # Clean all columns if no specific ones found
        
        filtered_count = len(df)
        
        # Convert to JSON
        json_data = df.to_dict('records')
        
        return {
            'success': True,
            'data': json_data,
            'total_records': len(json_data),
            'original_records': original_count,
            'filtered_records': filtered_count,
            'rows_removed': original_count - filtered_count,
            'columns': list(df.columns),
            'message': f'Successfully converted {len(json_data)} records (removed {original_count - filtered_count} empty rows)'
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': f'Error processing CSV file: {str(e)}',
            'data': []
        }

def xlsx_to_json(file_path):
    """Convert XLSX file to JSON object with row filtering"""
    try:
        # Read Excel file (first sheet by default)
        df = pd.read_excel(file_path, engine='openpyxl')
        
        original_count = len(df)
        
        # Define important columns that should not be empty
        required_columns = [
            'company', 'Company', 'organization', 'Organization', 'employer', 'Employer',
            'title', 'Title', 'position', 'Position', 'job_title', 'Job Title', 'role', 'Role',
            'name', 'Name', 'full_name', 'Full Name', 'first_name', 'First Name', 'last_name', 'Last Name'
        ]
        
        # Clean the dataframe with filtering
        df = clean_dataframe(df, required_columns)
        
        filtered_count = len(df)
        
        # Convert to JSON
        json_data = df.to_dict('records')
        
        return {
            'success': True,
            'data': json_data,
            'total_records': len(json_data),
            'original_records': original_count,
            'filtered_records': filtered_count,
            'rows_removed': original_count - filtered_count,
            'columns': list(df.columns),
            'message': f'Successfully converted {len(json_data)} records (removed {original_count - filtered_count} empty rows)'
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': f'Error processing XLSX file: {str(e)}',
            'data': []
        }

def xlsx_to_json_multiple_sheets(file_path):
    """Convert XLSX file with multiple sheets to JSON object with row filtering"""
    try:
        # Read all sheets
        excel_file = pd.ExcelFile(file_path)
        sheets_data = {}
        total_original = 0
        total_filtered = 0
        
        # Define important columns that should not be empty
        required_columns = [
            'company', 'Company', 'organization', 'Organization', 'employer', 'Employer',
            'title', 'Title', 'position', 'Position', 'job_title', 'Job Title', 'role', 'Role',
            'name', 'Name', 'full_name', 'Full Name', 'first_name', 'First Name', 'last_name', 'Last Name'
        ]
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name, engine='openpyxl')
            original_count = len(df)
            total_original += original_count
            
            df = clean_dataframe(df, required_columns)
            filtered_count = len(df)
            total_filtered += filtered_count
            
            sheets_data[sheet_name] = df.to_dict('records')
        
        return {
            'success': True,
            'data': sheets_data,
            'sheets': list(excel_file.sheet_names),
            'total_records': total_filtered,
            'original_records': total_original,
            'filtered_records': total_filtered,
            'rows_removed': total_original - total_filtered,
            'message': f'Successfully converted {len(excel_file.sheet_names)} sheets with {total_filtered} total records (removed {total_original - total_filtered} empty rows)'
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': f'Error processing XLSX file: {str(e)}',
            'data': {}
        }
    
def extract_unique_companies(json_data):
    """Extract unique company names from JSON data"""
    companies = set()
    
    for record in json_data:
        # Check various possible company field names
        company_fields = ['company', 'Company', 'organization', 'Organization', 'employer', 'Employer']
        
        for field in company_fields:
            if field in record and record[field]:
                company_name = str(record[field]).strip()
                if company_name and company_name.lower() not in ['', 'n/a', 'na', 'none', 'null']:
                    companies.add(company_name)
                break
    
    return list(companies)

def get_company_skills_from_openai(company_list):
    """Send company list to OpenAI and get required skills for each company"""
    try:
        companies_text = ", ".join(company_list)
        
        # Simplified prompt with clearer instructions
        prompt = f"""For the following companies, provide required skills in JSON format:

Companies: {companies_text}

Return a JSON array with this exact structure:
[
  {{"company": "Company1", "required_skills": ["skill1", "skill2", "skill3", "skill4", "skill5", "skill6"]}},
  {{"company": "Company2", "required_skills": ["skill1", "skill2", "skill3", "skill4", "skill5", "skill6"]}}
]

Rules:
- Exactly 6 skills per company
- Mix of technical and soft skills
- Valid JSON only, no explanations
- Keep skills concise (1-3 words each)"""

        client = openai.OpenAI()
        client.api_key=os.environ['OPENAI_API_KEY']
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a JSON generator. Return only valid JSON arrays. Keep responses concise."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1500,  # Reduced token limit to prevent truncation
            temperature=0.1   # Lower temperature for more consistent output
        )
        
        response_text = response.choices[0].message.content.strip()
        
        # More robust cleaning
        response_text = response_text.replace('```json', '').replace('```', '').strip()
        
        # Remove any trailing incomplete content
        if response_text.endswith(','):
            response_text = response_text[:-1]
        if response_text.endswith('}'):
            response_text += ']'
        elif not response_text.endswith(']'):
            # Find the last complete object and truncate there
            last_complete = response_text.rfind('}')
            if last_complete > 0:
                response_text = response_text[:last_complete + 1] + ']'
        
        # Try to fix common JSON issues
        response_text = fix_json_issues(response_text)
        
        try:
            company_skills_data = json.loads(response_text)
            
            # Validate the structure
            if not isinstance(company_skills_data, list):
                raise ValueError("Response is not a list")
            
            for item in company_skills_data:
                if not isinstance(item, dict) or 'company' not in item or 'required_skills' not in item:
                    raise ValueError("Invalid item structure")
                if not isinstance(item['required_skills'], list):
                    raise ValueError("required_skills is not a list")
            
            return {
                'success': True,
                'data': company_skills_data
            }
            
        except (json.JSONDecodeError, ValueError) as e:
            print(f"JSON parsing failed: {str(e)}")
            print(f"Raw response: {response_text[:500]}...")
            
            # Fallback: try to extract data manually
            fallback_result = extract_skills_manually(companies_text, response_text)
            if fallback_result['success']:
                return fallback_result
            
            return {
                'success': False,
                'error': f'Invalid JSON response: {str(e)}',
                'raw_response': response_text[:300]
            }
        
    except Exception as e:
        print(f"OpenAI API error: {str(e)}")
        return {
            'success': False,
            'error': f'OpenAI API error: {str(e)}'
        }
    

def fix_json_issues(json_text):
    """Fix common JSON formatting issues"""
    # Remove any text before the first [
    start_idx = json_text.find('[')
    if start_idx > 0:
        json_text = json_text[start_idx:]
    
    # Remove any text after the last ]
    end_idx = json_text.rfind(']')
    if end_idx > 0:
        json_text = json_text[:end_idx + 1]
    
    # Fix common escape issues
    json_text = json_text.replace('\n', ' ').replace('\r', ' ')
    json_text = json_text.replace('"', '"').replace('"', '"')  # Fix smart quotes
    
    return json_text

def extract_skills_manually(companies_text, response_text):
    """Fallback method to extract skills manually if JSON parsing fails"""
    try:
        import re
        
        company_list = [c.strip() for c in companies_text.split(',')]
        result_data = []
        
        # Default skills for different company types
        default_skills = {
            'tech': ['Python', 'JavaScript', 'Problem Solving', 'Communication', 'Teamwork', 'Leadership'],
            'finance': ['Excel', 'Financial Analysis', 'Risk Management', 'Communication', 'Attention to Detail', 'Leadership'],
            'healthcare': ['Patient Care', 'Medical Knowledge', 'Communication', 'Empathy', 'Attention to Detail', 'Teamwork'],
            'default': ['Communication', 'Problem Solving', 'Leadership', 'Teamwork', 'Analytical Thinking', 'Adaptability']
        }
        
        for company in company_list:
            # Try to determine company type and assign appropriate skills
            company_lower = company.lower()
            if any(tech_word in company_lower for tech_word in ['google', 'microsoft', 'apple', 'facebook', 'amazon', 'tech', 'software']):
                skills = default_skills['tech']
            elif any(fin_word in company_lower for fin_word in ['bank', 'finance', 'capital', 'investment', 'goldman', 'morgan']):
                skills = default_skills['finance']
            elif any(health_word in company_lower for health_word in ['hospital', 'medical', 'health', 'pharma', 'clinic']):
                skills = default_skills['healthcare']
            else:
                skills = default_skills['default']
            
            result_data.append({
                'company': company,
                'required_skills': skills
            })
        
        return {
            'success': True,
            'data': result_data
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': f'Manual extraction failed: {str(e)}'
        }

def enrich_json_with_openai(json_data):
    """Enhanced function: Extract companies -> Get skills from OpenAI -> Enrich original data"""
    try:
        # Step 1: Extract unique company names from JSON data
        unique_companies = extract_unique_companies(json_data)
        
        if not unique_companies:
            return {
                'success': False,
                'error': 'No company names found in the data'
            }
        
        print(f"Found {len(unique_companies)} unique companies: {unique_companies}")
        
        # Step 2: Send company list to OpenAI to get required skills
        openai_result = get_company_skills_from_openai(unique_companies)
        
        if not openai_result['success']:
            return {
                'success': False,
                'error': f'Failed to get skills from OpenAI: {openai_result.get("error", "Unknown error")}'
            }
        
        # Step 3: Create company-skills mapping from OpenAI response
        company_skills_map = {}
        for item in openai_result['data']:
            if 'company' in item and 'required_skills' in item:
                company_skills_map[item['company']] = item['required_skills']
        
        print(f"Created skills mapping for {len(company_skills_map)} companies")
        
        # Step 4: Enrich original JSON data by matching company names
        enriched_data = []
        skills_added_count = 0
        
        for record in json_data:
            # Create a copy of the original record
            enriched_record = record.copy()
            
            # Find company name in this record
            company_name = None
            company_fields = ['company', 'Company', 'organization', 'Organization', 'employer', 'Employer']
            
            for field in company_fields:
                if field in record and record[field]:
                    company_name = str(record[field]).strip()
                    break
            
            # Add required_skills based on company match
            if company_name and company_name in company_skills_map:
                enriched_record['required_skills'] = company_skills_map[company_name]
                skills_added_count += 1
            else:
                enriched_record['required_skills'] = []
            
            enriched_data.append(enriched_record)
        
        return {
            'success': True,
            'data': enriched_data,
            'message': f'Successfully enriched {len(enriched_data)} profiles. Added skills to {skills_added_count} records based on {len(unique_companies)} companies.',
            'companies_processed': len(unique_companies),
            'records_enriched': skills_added_count
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': f'Enrichment process error: {str(e)}'
        }

def get_credentials():
    load_dotenv()
    return os.getenv("OPENAI_API_KEY")

def get_file_hash(file_path):
    hasher = hashlib.md5()
    with open(file_path, 'rb') as f:
        buf = f.read()
        hasher.update(buf)
    return hasher.hexdigest()

def save_embeddings(file_hash, index, phrase_embeddings, page_chunks):
    embeddings_folder = os.path.join(DATA_DIR, 'embeddings')
    os.makedirs(embeddings_folder, exist_ok=True)
    with open(os.path.join(embeddings_folder, f"{file_hash}_index.pkl"), "wb") as f:
        pickle.dump(index, f)
    with open(os.path.join(embeddings_folder, f"{file_hash}_phrase_embeddings.pkl"), "wb") as f:
        pickle.dump(phrase_embeddings, f)
    with open(os.path.join(embeddings_folder, f"{file_hash}_page_chunks.pkl"), "wb") as f:
        pickle.dump(page_chunks, f)

def load_embeddings(file_hash):
    embeddings_folder = os.path.join(DATA_DIR, 'embeddings')
    with open(os.path.join(embeddings_folder, f"{file_hash}_index.pkl"), "rb") as f:
        index = pickle.load(f)
    with open(os.path.join(embeddings_folder, f"{file_hash}_phrase_embeddings.pkl"), "rb") as f:
        phrase_embeddings = pickle.load(f)
    with open(os.path.join(embeddings_folder, f"{file_hash}_page_chunks.pkl"), "rb") as f:
        page_chunks = pickle.load(f)
    print("Embeddings loaded successfully.")
    return index, phrase_embeddings, page_chunks

def init_llm():
    llm = ChatOpenAI(model="gpt-4o-mini", temperature=0)
    return llm

def init_embeddings():
    embeddings = OpenAIEmbeddings(model="text-embedding-3-large")
    return embeddings

def init_vector_store(embeddings):
    return None

def pdf_loader(file_path):
    pdf_document = fitz.open(file_path)
    pdf_text ={}
    for page_number in range(pdf_document.page_count):
        page = pdf_document.load_page(page_number)
        pdf_text[page_number + 1] = page.get_text()
    pdf_document.close()
    return pdf_text

def web_loader():
    loader = WebBaseLoader(
        web_paths=("https://lilianweng.github.io/posts/2023-06-23-agent/",),
        bs_kwargs=dict(parse_only=bs4.SoupStrainer(class_=("post-content", "post-title", "post-header"))),
    )
    docs = loader.load()
    return docs

def pdf_splitter(pdf_text):
    number_of_characters = sum(len(text) for text in pdf_text.values())
    print(f"Total number of characters in PDF: {number_of_characters}")  # Debug print for total characters
    # Set chunk size and overlap based on the total number of characters
    # For example, if the total number of characters is less than 100,000, use smaller chunks
    # Adjust chunk size and overlap based on the total number of characters
    if number_of_characters < 100000:
        chunk_size = 1000
        chunk_overlap = 200
    elif number_of_characters < 500000:
        chunk_size = 800
        chunk_overlap = 200
    elif number_of_characters < 1000000:
        chunk_size = 600
        chunk_overlap = 200
    elif number_of_characters < 2000000:
        chunk_size = 500
        chunk_overlap = 200
    elif number_of_characters < 5000000:
        chunk_size = 400
        chunk_overlap = 200
    elif number_of_characters < 10000000:
        chunk_size = 300
        chunk_overlap = 200
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    page_chunks = {}
    for page, text in pdf_text.items():
        # print(f"Page {page} length: {len(text)}")  # Debug print for text length
        chunks = text_splitter.split_text(text)
        # print(f"Page {page} chunks: {len(chunks)}")  # Debug print for number of chunks
        page_chunks[page] = chunks
    return page_chunks

def extract_keywords_from_pdf(pdf_text):
    rake = Rake()
    page_phrases = {}
    for page, text in pdf_text.items():
        rake.extract_keywords_from_text(text)
        phrases = rake.get_ranked_phrases()[:5]
        page_phrases[page] = phrases
    return page_phrases

def extract_keywords_from_chunks(page_chunks):
    rake = Rake()
    chunk_phrases = {}
    for page, chunks in page_chunks.items():
        for chunk_number, chunk in enumerate(chunks, start=1):
            rake.extract_keywords_from_text(chunk)
            phrases = rake.get_ranked_phrases()[:5]
            chunk_phrases[(page, chunk_number)] = phrases
    return chunk_phrases

def get_embeddings(phrase):
    client = openai.OpenAI()
    client.api_key = get_credentials()
    response = client.embeddings.create(model="text-embedding-ada-002", input=phrase)
    return response.data[0].embedding

def get_embeddings_batch(phrases):
    client = openai.OpenAI()
    client.api_key = get_credentials()
    response = client.embeddings.create(model="text-embedding-ada-002", input=phrases)
    return [data.embedding for data in response.data]

def store_embeddings(page_phrases, chunk_phrases):
    """Store embeddings with better error handling"""
    print("Processing embeddings...")
    print(f"Page phrases: {page_phrases}")
    print(f"Chunk phrases keys: {list(chunk_phrases.keys())}")
    
    phrase_embeddings = {}
    all_embeddings = []  # Store all embeddings for FAISS index
    
    # Process chunk phrases and get embeddings
    for (page, chunk_number), phrases in chunk_phrases.items():
        if phrases:  # Only process if there are phrases
            try:
                embeddings = get_embeddings_batch(phrases)
                phrase_embeddings[(page, chunk_number)] = list(zip(phrases, embeddings))
                # Collect all embeddings for FAISS index
                all_embeddings.extend(embeddings)
            except Exception as e:
                print(f"Error getting embeddings for page {page}, chunk {chunk_number}: {e}")
                phrase_embeddings[(page, chunk_number)] = []
        else:
            phrase_embeddings[(page, chunk_number)] = []
    
    # Check if we have any embeddings
    if not all_embeddings:
        print("Warning: No embeddings were created. Using dummy embeddings.")
        # Create a dummy embedding for testing
        dummy_embedding = [0.0] * 1536  # OpenAI ada-002 embedding dimension
        all_embeddings = [dummy_embedding]
        # Add dummy phrase_embeddings entry
        phrase_embeddings[(1, 1)] = [("no content found", dummy_embedding)]
    
    # Initialize FAISS index
    dimension = len(all_embeddings[0])
    print(f"Creating FAISS index with dimension: {dimension}")
    index = faiss.IndexFlatIP(dimension)
    
    # Add all embeddings to the index
    if all_embeddings:
        embeddings_array = np.array(all_embeddings, dtype=np.float32)
        index.add(embeddings_array)
        print(f"Added {len(all_embeddings)} embeddings to FAISS index")
    
    return index, phrase_embeddings

def extract_phrases_from_query(query):
    rake = Rake()
    rake.extract_keywords_from_text(query)
    return rake.get_ranked_phrases()

def get_embeddings_for_query(phrases):
    client = openai.OpenAI()
    client.api_key = get_credentials()
    return [client.embeddings.create(model="text-embedding-ada-002", input=phrase).data[0].embedding for phrase in phrases]

def get_cosine_similarity(embedding1, embedding2):
    return 1 - cosine(embedding1, embedding2)

def store_cosine_similarities(query_embeddings, phrase_embeddings, page_chunks):
    chunk_similarities = {}
    for (page, chunk_number), phrases in phrase_embeddings.items():
        similarities = []
        for phrase, embedding in phrases:
            phrase_similarities = [get_cosine_similarity(embedding, query_embedding) for query_embedding in query_embeddings] 
        similarities.append(max(phrase_similarities)) 
        # Choose the highest similarity for each phrase 
        average_similarity = np.mean(similarities) 
        # Average similarity for the chunk 
        chunk_similarities[(page, chunk_number)] = average_similarity 
    # Get top 5 chunks by similarity 
    top_chunks = sorted(chunk_similarities.items(), key=lambda x: x[1], reverse=True)[:5] 
    # Output top 5 chunks 
    print("Top 5 most relatable chunks:") 
    selected_chunks = []
    for (page, chunk_number), similarity in top_chunks: 
        print(f"Page: {page}, Chunk: {chunk_number}, Similarity: {similarity}") 
        print(f"Chunk text:\n{page_chunks[page][chunk_number-1]}\n")
        selected_chunks.append(page_chunks[page][chunk_number-1])
    return selected_chunks

def retrieve_similar_chunks(query_embeddings, index, phrase_embeddings, page_chunks):
    """Retrieve similar chunks with better error handling"""
    try:
        if not query_embeddings or not phrase_embeddings:
            print("Warning: No query embeddings or phrase embeddings available")
            return ["No relevant content found."]
        
        query_embeddings_np = np.array(query_embeddings, dtype=np.float32)
        
        # Ensure we don't search for more chunks than available
        available_chunks = index.ntotal
        k = min(5, available_chunks) if available_chunks > 0 else 1
        
        print(f"Searching for {k} similar chunks from {available_chunks} total")
        
        if available_chunks == 0:
            return ["No indexed content available for search."]
        
        D, I = index.search(query_embeddings_np, k=k)
        
        selected_chunks = []
        processed_indices = set()
        
        # Map FAISS indices back to chunks
        embedding_index = 0
        index_to_chunk_map = {}
        
        # Create mapping from FAISS index to chunk content
        for (page, chunk_number), phrases in phrase_embeddings.items():
            for phrase, embedding in phrases:
                index_to_chunk_map[embedding_index] = (page, chunk_number)
                embedding_index += 1
        
        # Retrieve chunks based on similarity
        for i in range(len(I)):
            for j in range(len(I[i])):
                chunk_idx = int(I[i][j])
                
                if chunk_idx in index_to_chunk_map and chunk_idx not in processed_indices:
                    page, chunk_number = index_to_chunk_map[chunk_idx]
                    
                    # Get the actual chunk content
                    if page in page_chunks and chunk_number - 1 < len(page_chunks[page]):
                        chunk_content = page_chunks[page][chunk_number - 1]
                        selected_chunks.append(chunk_content)
                        processed_indices.add(chunk_idx)
                        
                        if len(selected_chunks) >= 5:  # Limit to 5 chunks
                            break
            
            if len(selected_chunks) >= 5:
                break
        
        # Fallback: if no chunks found, return some content from page_chunks
        if not selected_chunks and page_chunks:
            print("Using fallback: returning first available chunks")
            for page, chunks in page_chunks.items():
                for chunk in chunks[:2]:  # Take first 2 chunks from each page
                    selected_chunks.append(chunk)
                    if len(selected_chunks) >= 3:
                        break
                if len(selected_chunks) >= 3:
                    break
        
        if not selected_chunks:
            selected_chunks = ["Unable to find relevant content in the document."]
        
        print(f"Retrieved {len(selected_chunks)} chunks for context")
        return selected_chunks
        
    except Exception as e:
        print(f"Error in retrieve_similar_chunks: {e}")
        return [f"Error retrieving content: {str(e)}"]

def extract_keywords_from_chunks(page_chunks):
    """Extract keywords from chunks with better error handling"""
    rake = Rake()
    chunk_phrases = {}
    
    for page, chunks in page_chunks.items():
        for chunk_number, chunk in enumerate(chunks, start=1):
            try:
                if chunk and chunk.strip():  # Only process non-empty chunks
                    rake.extract_keywords_from_text(chunk)
                    phrases = rake.get_ranked_phrases()[:5]
                    chunk_phrases[(page, chunk_number)] = phrases
                else:
                    chunk_phrases[(page, chunk_number)] = []
            except Exception as e:
                print(f"Error extracting keywords from page {page}, chunk {chunk_number}: {e}")
                chunk_phrases[(page, chunk_number)] = []
    
    print(f"Extracted keywords from {len(chunk_phrases)} chunks")
    return chunk_phrases


def parse_simple_query_enhanced(user_query):
    """Enhanced keyword extraction using OpenAI function calling with special commands support"""
    try:
        # Define the function schema for OpenAI with special commands
        extract_function = {
            "name": "extract_search_criteria",
            "description": "Extract search criteria from natural language query, including special commands",
            "parameters": {
                "type": "object",
                "properties": {
                    "company": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "Company names, organizations, or employers mentioned"
                    },
                    "title": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "Job titles, positions, or roles mentioned"
                    },
                    "name": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "Person names mentioned"
                    },
                    "skills": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "Technical skills, technologies, or expertise mentioned"
                    },
                    "location": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "Cities, countries, or geographic locations mentioned"
                    },
                    "phrases": {
                        "type": "array",
                        "items": {
                            "type": "string",
                            "enum": [
                                "show_all_companies",
                                "show_all_titles", 
                                "show_all_locations",
                                "show_all_skills",
                                "show_favorites",
                                "list_companies",
                                "list_titles",
                                "list_locations", 
                                "list_skills",
                                "my_favorites",
                                "saved_profiles",
                                "all_companies",
                                "all_titles",
                                "all_locations",
                                "all_skills"
                            ]
                        },
                        "description": "Special command phrases detected in the query like 'show all companies', 'show my favorites', 'list all titles', etc."
                    }
                },
                "required": []
            }
        }

        client = openai.OpenAI()
        client.api_key = os.environ['OPENAI_API_KEY']
        
        # Enhanced system prompt to detect special commands while maintaining your existing logic
        system_prompt = """You are a search query parser that extracts search criteria accurately from user queries. Extract search criteria accurately from user queries. Include information that is explicitly mentioned but also matching terms that are close to what is mentioned. They could be singular or plural forms, sub-strings or variations, etc.

SPECIAL COMMANDS TO DETECT:
- "show all companies" / "list companies" / "all companies" â†’ show_all_companies
- "show all titles" / "list titles" / "all titles" / "job titles" â†’ show_all_titles  
- "show all locations" / "list locations" / "all locations" â†’ show_all_locations
- "show all skills" / "list skills" / "all skills" â†’ show_all_skills
- "show favorites" / "my favorites" / "show saved" / "saved profiles" â†’ show_favorites

Extract both specific search terms AND any special commands detected. If the query contains both search terms and special commands, include both in the response."""

        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {
                    "role": "system", 
                    "content": system_prompt
                },
                {
                    "role": "user", 
                    "content": f"Extract search criteria and special commands from this query: '{user_query}'"
                }
            ],
            functions=[extract_function],
            function_call={"name": "extract_search_criteria"},
            temperature=0.1
        )
        
        # Extract function call result
        function_call = response.choices[0].message.function_call
        keywords = json.loads(function_call.arguments)
        
        # Clean up empty arrays
        keywords = {k: v for k, v in keywords.items() if v and len(v) > 0}
        
        # Check if this is a special command
        special_phrases = keywords.get('phrases', [])
        is_special_command = len(special_phrases) > 0
        
        # Determine command type from phrases
        command_type = None
        if is_special_command:
            phrase = special_phrases[0]  # Use first detected phrase
            if phrase in ['show_all_companies', 'list_companies', 'all_companies']:
                command_type = 'show_companies'
            elif phrase in ['show_all_titles', 'list_titles', 'all_titles']:
                command_type = 'show_titles'
            elif phrase in ['show_all_locations', 'list_locations', 'all_locations']:
                command_type = 'show_locations'
            elif phrase in ['show_all_skills', 'list_skills', 'all_skills']:
                command_type = 'show_skills'
            elif phrase in ['show_favorites', 'my_favorites', 'saved_profiles']:
                command_type = 'show_favorites'
        
        print(f"OpenAI extracted keywords: {keywords}")
        print(f"Special command detected: {is_special_command}, Type: {command_type}")
        
        return {
            'success': True, 
            'keywords': keywords,
            'special_command': is_special_command,
            'command_type': command_type,
            'phrases': special_phrases
        }
        
    except Exception as e:
        print(f"OpenAI parsing error: {e}")
        return {
            'success': False, 
            'keywords': {}, 
            'special_command': False,
            'command_type': None,
            'phrases': [],
            'error': str(e)
        }


def simple_search_json(json_data, keywords):
    """Simple search function with substring matching for both keys and values"""
    results = []
    print(f"Searching with keywords: {keywords}")
    
    for record in json_data:
        match = True
        
        # Check each keyword type
        for field, values in keywords.items():
            if not values:  # Skip empty fields
                continue
                
            field_match = False
            
            # Map search fields to possible record fields with substring matching
            if field == 'company':
                record_fields = ['company', 'Company', 'organization', 'Organization', 'employer', 'Employer']
            elif field == 'title':
                record_fields = ['title', 'Title', 'position', 'Position', 'job_title', 'Job Title', 'role', 'Role']
            elif field == 'name':
                record_fields = ['name', 'Name', 'full_name', 'Full Name', 'first_name', 'First Name', 'last_name', 'Last Name']
            elif field == 'location':
                record_fields = ['location', 'Location', 'city', 'City', 'country', 'Country']
            elif field == 'skills':
                record_fields = ['skills', 'Skills', 'required_skills', 'technologies', 'Technologies']
            else:
                record_fields = [field]
            
            # First try exact field matching
            for record_field in record_fields:
                if record_field in record and record[record_field]:
                    record_value = str(record[record_field]).lower()
                    
                    for search_value in values:
                        search_value_lower = search_value.lower()
                        
                        # Check if search value is substring of record value
                        if search_value_lower in record_value:
                            field_match = True
                            print(f"Exact match: '{search_value}' found in '{record_value}' (field: {record_field})")
                            break
                        
                        # Check if record value is substring of search value (reverse match)
                        if record_value in search_value_lower:
                            field_match = True
                            print(f"Reverse match: '{record_value}' found in '{search_value}' (field: {record_field})")
                            break
                    
                    if field_match:
                        break
            
            # If no exact field match, try substring matching on field names
            if not field_match:
                for record_field in record.keys():
                    # Check if search field is substring of record field or vice versa
                    field_lower = field.lower()
                    record_field_lower = record_field.lower()
                    
                    # Match if field names have substring relationship
                    if (field_lower in record_field_lower or record_field_lower in field_lower) and record[record_field]:
                        record_value = str(record[record_field]).lower()
                        
                        for search_value in values:
                            search_value_lower = search_value.lower()
                            
                            # Check substring matches in both directions
                            if search_value_lower in record_value:
                                field_match = True
                                print(f"Field substring match: '{search_value}' found in '{record_value}' (field: {record_field})")
                                break
                            
                            if record_value in search_value_lower:
                                field_match = True
                                print(f"Field reverse substring match: '{record_value}' found in '{search_value}' (field: {record_field})")
                                break
                        
                        if field_match:
                            break
            
            # If still no match, try word-level substring matching
            if not field_match:
                for record_field in record_fields:
                    if record_field in record and record[record_field]:
                        record_value = str(record[record_field]).lower()
                        
                        for search_value in values:
                            search_value_lower = search_value.lower()
                            
                            # Split search value into words and check each word
                            search_words = [word.strip() for word in search_value_lower.replace('-', ' ').split() if len(word.strip()) > 2]
                            
                            for word in search_words:
                                if word in record_value:
                                    field_match = True
                                    print(f"Word substring match: '{word}' from '{search_value}' found in '{record_value}' (field: {record_field})")
                                    break
                            
                            # Also split record value and check against search value
                            if not field_match:
                                record_words = [word.strip() for word in record_value.replace('-', ' ').split() if len(word.strip()) > 2]
                                
                                for word in record_words:
                                    if word in search_value_lower:
                                        field_match = True
                                        print(f"Record word substring match: '{word}' from '{record_value}' found in '{search_value}' (field: {record_field})")
                                        break
                            
                            if field_match:
                                break
                    
                    if field_match:
                        break
            
            # If this field didn't match, exclude the record
            if not field_match:
                match = False
                break
        
        
        if match:
            results.append(record)
    print(match)
    return results

def make_request(url, config):
    """Centralized HTTP request handler"""
    headers = config.get('headers', {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
    })
    
    timeout = config.get('timeout', 30)
    
    response = requests.get(url, headers=headers, timeout=timeout)
    response.raise_for_status()
    
    return response

def parse_html(response):
    """Centralized HTML parsing"""
    return BeautifulSoup(response.text, 'html.parser')

def get_page_metadata(soup, url):
    """Extract basic page metadata"""
    title = None
    title_tag = soup.find('title')
    if title_tag:
        title = title_tag.get_text(strip=True)
    
    meta_description = None
    meta_desc = soup.find('meta', attrs={'name': 'description'})
    if meta_desc:
        meta_description = meta_desc.get('content', '')
    
    return {
        'page_title': title,
        'meta_description': meta_description,
        'url': url
    }

def scrape_basic_content(url, config):
    """Basic content scraping - common elements"""
    try:
        response = make_request(url, config)
        soup = parse_html(response)
        
        # Get page metadata
        result = get_page_metadata(soup, url)
        result['status_code'] = response.status_code
        
        # Get selectors from config or use defaults
        selectors = config.get('selectors', {})
        
        if selectors:
            # Use custom selectors
            for name, selector_config in selectors.items():
                result[name] = extract_with_selector(soup, selector_config)
        else:
            # Default extraction
            result.update(extract_common_elements(soup))
        
        return {
            'success': True,
            'type': 'basic',
            'data': result,
            'scraped_at': datetime.now().isoformat()
        }
        
    except requests.exceptions.RequestException as e:
        return {
            'success': False,
            'error': f'Request failed: {str(e)}',
            'url': url
        }
    except Exception as e:
        return {
            'success': False,
            'error': f'Scraping error: {str(e)}'
        }

def scrape_text_content(url, config):
    """Text-only content scraping"""
    try:
        response = make_request(url, config)
        soup = parse_html(response)
        
        selector = config.get('selector')
        max_length = config.get('max_length')
        clean_text = config.get('clean_text', True)
        
        if selector:
            # Extract text from specific selector
            element = soup.select_one(selector)
            text = element.get_text(strip=clean_text) if element else None
        else:
            # Extract all text
            text = soup.get_text()
            if clean_text:
                text = ' '.join(text.split())
        
        # Apply length limit
        if text and max_length and len(text) > max_length:
            text = text[:max_length] + '...'
        
        return {
            'success': True,
            'type': 'text',
            'url': url,
            'selector': selector if selector else 'entire_page',
            'text': text,
            'text_length': len(text) if text else 0,
            'scraped_at': datetime.now().isoformat()
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }

def scrape_json_ld_content(url, config):
    """JSON-LD structured data scraping"""
    try:
        response = make_request(url, config)
        soup = parse_html(response)
        
        # Extract JSON-LD scripts
        json_ld_data = []
        scripts = soup.find_all("script", type="application/ld+json")
        
        filters = config.get('filters', {})
        schema_type = filters.get('schema_type')
        keywords = filters.get('keywords', [])
        
        for script in scripts:
            try:
                if script.string:
                    data_obj = json.loads(script.string)
                    
                    # Apply filters
                    if schema_type and data_obj.get('@type', '').lower() != schema_type.lower():
                        continue
                    
                    if keywords:
                        data_str = str(data_obj).lower()
                        if not any(keyword.lower() in data_str for keyword in keywords):
                            continue
                    
                    json_ld_data.append(data_obj)
            except json.JSONDecodeError:
                continue
        
        return {
            'success': True,
            'type': 'json_ld',
            'url': url,
            'json_ld': json_ld_data,
            'count': len(json_ld_data),
            'filters_applied': filters,
            'scraped_at': datetime.now().isoformat()
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }

def scrape_product_content(url, config):
    """Enhanced product data scraping"""
    try:
        response = make_request(url, config)
        soup = parse_html(response)
        
        result = get_page_metadata(soup, url)
        
        # Product-specific selectors
        product_selectors = {
            'name': [
                'h1.product-title', 'h1.product-name', '.product-title', '.pdp-product-name',
                'h1[data-testid="product-title"]', '.product-info h1'
            ],
            'price': [
                '.price', '.current-price', '.product-price', '[data-testid="price"]',
                '.price-current', '.final-price'
            ],
            'description': [
                '.product-description', '.pdp-description', '.product-details-description',
                '[data-testid="product-description"]', '.product-info-description'
            ],
            'brand': [
                '.brand', '.product-brand', '[data-testid="brand"]', '.manufacturer'
            ],
            'availability': [
                '.availability', '.stock-status', '[data-testid="availability"]', '.in-stock'
            ],
            'images': [
                '.product-image img', '.product-gallery img', '.pdp-images img'
            ]
        }
        
        # Override with custom selectors if provided
        custom_selectors = config.get('selectors', {})
        if custom_selectors:
            product_selectors.update(custom_selectors)
        
        # Extract product data
        product_data = {}
        for field, selector_list in product_selectors.items():
            if field == 'images':
                # Handle images specially
                images = []
                for selector in selector_list:
                    elements = soup.select(selector)
                    for img in elements[:5]:  # Limit to 5 images
                        src = img.get('src') or img.get('data-src')
                        if src:
                            images.append({
                                'src': src,
                                'alt': img.get('alt', '')
                            })
                    if images:
                        break
                product_data[field] = images
            else:
                # Handle text fields
                for selector in selector_list:
                    element = soup.select_one(selector)
                    if element:
                        product_data[field] = element.get_text(strip=True)
                        break
        
        result['product_data'] = product_data
        
        return {
            'success': True,
            'type': 'product',
            'data': result,
            'scraped_at': datetime.now().isoformat()
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }

def scrape_table_content(url, config):
    """Table data scraping"""
    try:
        response = make_request(url, config)
        soup = parse_html(response)
        
        table_selector = config.get('table_selector', 'table')
        include_headers = config.get('include_headers', True)
        max_rows = config.get('max_rows')
        
        tables = soup.select(table_selector)
        tables_data = []
        
        for i, table in enumerate(tables):
            table_data = {
                'table_index': i + 1,
                'headers': [],
                'rows': []
            }
            
            # Extract headers
            if include_headers:
                header_row = table.find('thead') or table.find('tr')
                if header_row:
                    headers = header_row.find_all(['th', 'td'])
                    table_data['headers'] = [th.get_text(strip=True) for th in headers]
            
            # Extract rows
            body = table.find('tbody') or table
            rows = body.find_all('tr')
            
            if include_headers and table_data['headers']:
                rows = rows[1:]  # Skip header row
            
            for row_idx, row in enumerate(rows):
                if max_rows and row_idx >= max_rows:
                    break
                    
                cells = row.find_all(['td', 'th'])
                if cells:  # Skip empty rows
                    row_data = [cell.get_text(strip=True) for cell in cells]
                    table_data['rows'].append(row_data)
            
            if table_data['headers'] or table_data['rows']:
                tables_data.append(table_data)
        
        return {
            'success': True,
            'type': 'tables',
            'url': url,
            'tables': tables_data,
            'total_tables': len(tables_data),
            'scraped_at': datetime.now().isoformat()
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }

def scrape_custom_selectors(url, config):
    """Custom selector-based scraping"""
    try:
        response = make_request(url, config)
        soup = parse_html(response)
        
        selectors = config.get('selectors', {})
        if not selectors:
            return {
                'success': False,
                'error': 'No selectors provided for custom scraping'
            }
        
        result = get_page_metadata(soup, url)
        
        for name, selector_config in selectors.items():
            result[name] = extract_with_selector(soup, selector_config)
        
        return {
            'success': True,
            'type': 'custom',
            'data': result,
            'scraped_at': datetime.now().isoformat()
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }

def extract_with_selector(soup, selector_config):
    """Universal selector extraction handler"""
    try:
        if isinstance(selector_config, str):
            # Simple string selector
            element = soup.select_one(selector_config)
            return element.get_text(strip=True) if element else None
        
        elif isinstance(selector_config, dict):
            # Advanced selector with options
            selector = selector_config.get('selector')
            attribute = selector_config.get('attribute', 'text')
            multiple = selector_config.get('multiple', False)
            
            if multiple:
                elements = soup.select(selector)
                if attribute == 'text':
                    return [el.get_text(strip=True) for el in elements]
                else:
                    return [el.get(attribute) for el in elements if el.get(attribute)]
            else:
                element = soup.select_one(selector)
                if element:
                    if attribute == 'text':
                        return element.get_text(strip=True)
                    else:
                        return element.get(attribute)
                return None
        
        return None
        
    except Exception as e:
        return f"Error: {str(e)}"

def extract_common_elements(soup):
    """Extract common page elements when no specific selectors provided"""
    common_data = {}
    
    try:
        # Headings (limit to first 10)
        headings = []
        for i in range(1, 7):  # h1 to h6
            heading_elements = soup.find_all(f'h{i}')
            for h in heading_elements[:3]:  # Max 3 per level
                text = h.get_text(strip=True)
                if text:
                    headings.append({
                        'level': i,
                        'text': text
                    })
        common_data['headings'] = headings[:10]
        
        # Paragraphs (first 5 meaningful ones)
        paragraphs = []
        p_elements = soup.find_all('p')
        for p in p_elements:
            text = p.get_text(strip=True)
            if text and len(text) > 30:  # Only meaningful paragraphs
                paragraphs.append(text)
                if len(paragraphs) >= 5:
                    break
        common_data['paragraphs'] = paragraphs
        
        # Links (first 10)
        links = []
        a_elements = soup.find_all('a', href=True)
        for a in a_elements[:10]:
            href = a.get('href')
            text = a.get_text(strip=True)
            if href and text:
                links.append({
                    'url': href,
                    'text': text
                })
        common_data['links'] = links
        
        # Images (first 5)
        images = []
        img_elements = soup.find_all('img')
        for img in img_elements[:5]:
            src = img.get('src') or img.get('data-src')
            if src:
                images.append({
                    'src': src,
                    'alt': img.get('alt', '')
                })
        common_data['images'] = images
        
    except Exception as e:
        common_data['error'] = f"Error extracting common elements: {str(e)}"
    
    return common_data

def handle_cookie_consent_and_popups(driver):
    """Enhanced cookie consent handler specifically for ZARA and similar sites"""
    try:
        # Wait a moment for popups to appear
        time.sleep(3)
        
        # Specific selectors for ZARA and other fashion sites
        zara_selectors = [
            # ZARA specific
            'button[data-qa-action="accept-all"]',
            'button[data-qa-action="accept"]',
            '#onetrust-accept-btn-handler',
            '.ot-pc-refuse-all-handler',
            
            # Common cookie consent buttons
            'button:contains("Accept All")',
            'button:contains("Alle akzeptieren")',  # German
            'button:contains("Aceptar todo")',      # Spanish
            'button:contains("Accepter tout")',     # French
            'button:contains("Accept")',
            'button:contains("Akzeptieren")',
            'button:contains("I Accept")',
            'button:contains("Agree")',
            'button:contains("Einverstanden")',
            'button:contains("OK")',
            'button:contains("Continue")',
            'button:contains("Weiter")',
            'button:contains("Got it")',
            
            # Common class names and IDs
            '.accept-all',
            '.accept-cookies',
            '.accept-btn',
            '.cookie-accept',
            '.consent-accept',
            '.terms-accept',
            '#accept-all',
            '#accept-cookies',
            '#cookie-accept',
            '.onetrust-close-btn-handler',
            '.optanon-allow-all',
            
            # Data attributes
            '[data-accept="all"]',
            '[data-cookie-accept]',
            '[data-consent="accept"]',
            '[data-testid="accept"]',
            '[data-testid="accept-all"]',
            '[data-qa-action="accept-all"]',
            
            # OneTrust specific (used by ZARA)
            '#onetrust-accept-btn-handler',
            '.onetrust-close-btn-handler',
            '.ot-pc-refuse-all-handler',
            
            # ARIA labels
            '[aria-label*="Accept"]',
            '[aria-label*="Agree"]',
            '[aria-label*="akzeptieren"]',
            '[role="button"][aria-label*="Accept"]'
        ]
        
        print("Looking for cookie consent buttons...")
        
        # Try each selector
        for selector in zara_selectors:
            try:
                if ':contains(' in selector:
                    # Handle text-based selectors with XPath
                    text = selector.split(':contains("')[1].split('")')[0]
                    tag = selector.split(':contains(')[0]
                    xpath = f"//{tag}[contains(translate(text(), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), '{text.lower()}')]"
                    elements = driver.find_elements(By.XPATH, xpath)
                else:
                    elements = driver.find_elements(By.CSS_SELECTOR, selector)
                
                for element in elements:
                    if element.is_displayed() and element.is_enabled():
                        print(f"Found and clicking accept button: {selector}")
                        # Scroll to element first
                        driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", element)
                        time.sleep(1)
                        
                        # Try multiple click methods
                        try:
                            element.click()
                        except Exception:
                            try:
                                driver.execute_script("arguments[0].click();", element)
                            except Exception:
                                # Force click using coordinates
                                from selenium.webdriver.common.action_chains import ActionChains
                                ActionChains(driver).move_to_element(element).click().perform()
                        
                        time.sleep(3)  # Wait for dialog to close
                        print(f"Successfully clicked consent button")
                        return True
                        
            except Exception as e:
                print(f"Error trying selector {selector}: {e}")
                continue
        
        print("No cookie consent button found or clicked")
        return False
        
    except Exception as e:
        print(f"Error handling popups: {e}")
        return False

def setup_driver(headless=True):
    """Setup Chrome WebDriver for interactive scraping with popup handling"""
    chrome_options = Options()
    
    if headless:
        chrome_options.add_argument('--headless')
    
    chrome_options.add_argument('--no-sandbox')
    chrome_options.add_argument('--disable-dev-shm-usage')
    chrome_options.add_argument('--disable-gpu')
    chrome_options.add_argument('--window-size=1920,1080')
    chrome_options.add_argument('--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36')
    
    # Disable notifications and popups
    prefs = {
        "profile.default_content_setting_values.notifications": 2,
        "profile.default_content_settings.popups": 0,
        "profile.managed_default_content_settings.images": 2  # Block images for faster loading
    }
    chrome_options.add_experimental_option("prefs", prefs)
    
    # Additional options to handle consent and cookies
    chrome_options.add_argument('--disable-features=VizDisplayCompositor')
    chrome_options.add_argument('--disable-extensions')
    chrome_options.add_argument('--disable-plugins')
    
    service = Service(ChromeDriverManager().install())
    return webdriver.Chrome(service=service, options=chrome_options)

def scrape_with_interaction(url, headless, config):
    """Scrape product information with interactive capabilities and popup handling"""
    driver = None
    try:
        driver = setup_driver(headless)
        
        # Set page load timeout
        driver.set_page_load_timeout(30)
        
        print(f"Loading page: {url}")
        driver.get(url)
        
        # Handle cookie consent and terms acceptance first
        print("Handling cookie consent and popups...")
        popup_handled = handle_cookie_consent_and_popups(driver)
        if popup_handled:
            print("Successfully handled popup/consent dialog")
        else:
            print("No popup found or couldn't handle it")
        
        # Wait for page to load after handling popups
        wait = WebDriverWait(driver, 15)
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        
        # Additional wait for dynamic content
        time.sleep(3)
        
        result = {
            'success': True,
            'url': url,
            'scraped_at': datetime.now().isoformat(),
            'method': 'interactive',
            'popup_handled': popup_handled,
            'product_info': {}
        }
        
        # Get basic page info
        page_title = driver.title
        result['page_title'] = page_title
        print(f"Page title: {page_title}")
        
        # Extract product information
        product_info = extract_product_information(driver, wait, config)
        result['product_info'] = product_info
        
        return result
        
    except Exception as e:
        error_msg = f'Interactive scraping error: {str(e)}'
        print(error_msg)
        return {
            'success': False,
            'error': error_msg,
            'method': 'interactive'
        }
    finally:
        if driver:
            try:
                driver.quit()
            except:
                pass

def extract_product_information(driver, wait, config):
    """Extract product description and size information with enhanced interaction"""
    product_info = {
        'description': None,
        'size_info': None,
        'size_chart': None,
        'size_measurements': None,  # Add this new field
        'materials': None,
        'care_instructions': None,
        'specifications': None,
        'interactive_content_found': False,
        'extraction_methods_used': []
    }
    
    try:
        # Step 1: Try to extract visible content first
        print("Extracting visible content...")
        soup = BeautifulSoup(driver.page_source, 'html.parser')
        visible_content = extract_static_product_info(soup)
        product_info.update(visible_content)
        product_info['extraction_methods_used'].append('static_html')
        
        # Step 2: Look for and interact with clickable elements
        print("Looking for interactive elements...")
        interactive_elements = find_interactive_elements(driver)
        
        if interactive_elements:
            product_info['interactive_content_found'] = True
            product_info['extraction_methods_used'].append('interactive_clicks')
            
            for element_info in interactive_elements:
                try:
                    print(f"Interacting with {element_info['type']}: {element_info['text']}")
                    
                    # Click the element
                    element = element_info['element']
                    element_type = element_info['type']
                    
                    # Scroll to element and click
                    driver.execute_script("arguments[0].scrollIntoView(true);", element)
                    time.sleep(1)
                    
                    # Get page source before click for comparison
                    before_click_content = driver.page_source
                    
                    # Try different click methods
                    try:
                        element.click()
                    except:
                        # If regular click fails, try JavaScript click
                        driver.execute_script("arguments[0].click();", element)
                    
                    # Wait longer for content to load
                    time.sleep(4)  # Increased wait time
                    
                    # Get page source after click
                    after_click_content = driver.page_source
                    
                    # Check if content changed
                    if len(after_click_content) != len(before_click_content):
                        print(f"Page content changed after clicking {element_type}")
                    
                    # Extract content after interaction
                    new_soup = BeautifulSoup(driver.page_source, 'html.parser')
                    new_content = extract_dynamic_content(new_soup, element_type)
                    
                    # Merge new content
                    for key, value in new_content.items():
                        if value and (not product_info.get(key) or product_info.get(key) == element_info['text']):
                            product_info[key] = value
                            print(f"Updated {key} with new content from {element_type}")
                    
                    # Close popup/modal if it opened
                    close_popup(driver)
                    time.sleep(1)  # Wait after closing
                    
                except Exception as e:
                    print(f"Error interacting with element {element_info['type']}: {e}")
                    continue
        
        # Step 3: Try scrolling to load more content
        print("Scrolling to load additional content...")
        driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
        time.sleep(2)
        driver.execute_script("window.scrollTo(0, 0);")
        time.sleep(1)
        
        # Extract content after scrolling
        final_soup = BeautifulSoup(driver.page_source, 'html.parser')
        scroll_content = extract_static_product_info(final_soup)
        
        # Merge scroll content
        for key, value in scroll_content.items():
            if value and not product_info.get(key):
                product_info[key] = value
        
        if scroll_content:
            product_info['extraction_methods_used'].append('scroll_content')
        
        # Step 4: Clean and format the extracted data
        product_info = clean_product_info(product_info)
        
        print(f"Extraction complete. Methods used: {product_info['extraction_methods_used']}")
        print(f"Final product info keys: {list(product_info.keys())}")
        
        return product_info
        
    except Exception as e:
        print(f"Error in extract_product_information: {e}")
        product_info['error'] = str(e)
        return product_info



def find_interactive_elements(driver):
    """Enhanced function to find clickable elements for multiple brands"""
    interactive_elements = []
    
    # Enhanced selectors for different brands and languages
    element_selectors = {
        'size_guide': [
            # English selectors
            'a:contains("Size Guide")', 'button:contains("Size Guide")',
            'a:contains("Size Chart")', 'button:contains("Size Chart")',
            'a:contains("Sizing")', 'button:contains("Sizing")',
            'a:contains("Measurements")', 'button:contains("Measurements")',
            'a:contains("Product Measurements")', 'button:contains("Product Measurements")',
            
            # German selectors
            'a:contains("GrÃ¶ÃŸentabelle")', 'button:contains("GrÃ¶ÃŸentabelle")',
            'a:contains("GrÃ¶ÃŸenfÃ¼hrung")', 'button:contains("GrÃ¶ÃŸenfÃ¼hrung")',
            'a:contains("MaÃŸe")', 'button:contains("MaÃŸe")',
            
            # French selectors
            'a:contains("Guide des tailles")', 'button:contains("Guide des tailles")',
            'a:contains("Tableau des tailles")', 'button:contains("Tableau des tailles")',
            
            # Spanish selectors
            'a:contains("GuÃ­a de tallas")', 'button:contains("GuÃ­a de tallas")',
            'a:contains("Tabla de tallas")', 'button:contains("Tabla de tallas")',
            
            # Generic attribute selectors
            'a[href*="size"]', 'a[href*="sizing"]', 'a[href*="measurement"]',
            'button[class*="size"]', 'a[class*="size"]',
            '[data-size-guide]', '[data-testid*="size"]',
            '[data-qa-action*="size"]', '[data-qa*="size"]',
            '[data-modal*="size"]', '[data-popup*="size"]',
            
            # Brand specific selectors
            'button[data-qa-action="size-guide"]',
            'a[data-qa-action="size-guide"]',
            '.product-size-guide-link',
            '.size-guide-trigger',
            'button.size-chart-btn',
            'a.size-chart-link',
            '.size-guide', '.size-chart', '.sizing-info',
            
            # Common class patterns
            '.size-guide-link', '.sizing-link', '.measurement-link',
            '[class*="size-guide"]', '[class*="sizing"]', '[class*="measurement"]'
        ],
        'description': [
            # English
            'a:contains("Description")', 'button:contains("Description")',
            'a:contains("More Details")', 'button:contains("More Details")',
            'a:contains("Product Details")', 'button:contains("Product Details")',
            'a:contains("Details")', 'button:contains("Details")',
            
            # German
            'a:contains("Beschreibung")', 'button:contains("Beschreibung")',
            'a:contains("Mehr Details")', 'button:contains("Mehr Details")',
            'a:contains("Produktdetails")', 'button:contains("Produktdetails")',
            
            # Generic
            '.description-toggle', '.product-description-toggle',
            '.more-info', '.view-more', '.expand-description',
            '[data-qa-action*="description"]', '[data-testid*="description"]'
        ],
        'specifications': [
            'a:contains("Specifications")', 'button:contains("Specifications")',
            'a:contains("Specs")', 'button:contains("Specs")',
            'a:contains("Technical Details")', 'button:contains("Technical Details")',
            'a:contains("Spezifikationen")', 'button:contains("Spezifikationen")',
            '.specs-toggle', '.specifications-toggle', '.product-specs',
            '[data-qa-action*="specs"]', '[data-qa-action*="details"]'
        ],
        'care': [
            'a:contains("Care")', 'button:contains("Care Instructions")',
            'a:contains("Care Instructions")', 'a:contains("Washing Instructions")',
            'button:contains("Pflege")', 'button:contains("Pflegeanleitung")',
            '.care-instructions', '.washing-instructions', '.garment-care',
            '[data-qa-action*="care"]'
        ]
    }
    
    print("Searching for interactive elements...")
    
    for element_type, selectors in element_selectors.items():
        print(f"Searching for {element_type} elements...")
        
        for selector in selectors:
            try:
                if ':contains(' in selector:
                    # Handle text-based selectors with XPath
                    text = selector.split(':contains("')[1].split('")')[0]
                    tag = selector.split(':contains(')[0]
                    xpath = f"//{tag}[contains(translate(text(), 'ABCDEFGHIJKLMNOPQRSTUVWXYZ', 'abcdefghijklmnopqrstuvwxyz'), '{text.lower()}')]"
                    elements = driver.find_elements(By.XPATH, xpath)
                else:
                    elements = driver.find_elements(By.CSS_SELECTOR, selector)
                
                for element in elements:
                    if element.is_displayed() and element.is_enabled():
                        element_text = element.text.strip()
                        element_tag = element.tag_name
                        
                        print(f"Found interactive element: {element_type} - '{element_text}' ({element_tag}) (selector: {selector})")
                        
                        interactive_elements.append({
                            'element': element,
                            'type': element_type,
                            'selector': selector,
                            'text': element_text,
                            'tag': element_tag
                        })
                        break  # Only take first working element of each type
                
                if any(e['type'] == element_type for e in interactive_elements):
                    break  # Found element for this type, move to next
                    
            except Exception as e:
                print(f"Error trying selector {selector}: {e}")
                continue
    
    print(f"Found {len(interactive_elements)} interactive elements total")
    return interactive_elements

def extract_static_product_info(soup):
    """Enhanced static product extraction for multiple brands"""
    info = {}
    
    # Enhanced description selectors for various brands
    description_selectors = [
        # ZARA specific
        '.expandable-text__inner-content p',
        '.expandable-text__inner-content',
        '.product-detail-info__description',
        
        # H&M specific
        '.product-description-text',
        '.pdp-product-description',
        '.product-description-content',
        
        # Uniqlo specific
        '.product-description',
        '.pdp-description',
        
        # Generic brand selectors
        '.product-detail-view__description',
        '.product-info__description',
        '[data-testid="product-description"]',
        '.product-info-description',
        '.product-long-description',
        '.description-content',
        '.product-content',
        '.item-description',
        '.product-summary',
        '.product-details-content',
        '.product-information',
        '#product-description',
        '.description',
        '.details',
        '.overview',
        '.about-product',
        '.product-info',
        '.product-detail',
        '.product-overview',
        '.item-details',
        '.product-specs',
        
        # Common patterns
        '[class*="description"]',
        '[class*="product-info"]',
        '[id*="description"]'
    ]
    
    # Extract description with enhanced search
    for selector in description_selectors:
        try:
            element = soup.select_one(selector)
            if element:
                text = element.get_text(strip=True)
                if text and len(text) > 20:
                    info['description'] = text
                    print(f"Found description using selector: {selector}")
                    break
        except Exception as e:
            continue
    
    # Enhanced fallback for description
    if not info.get('description'):
        # Look for expandable content patterns
        expandable_patterns = [
            '[class*="expandable"]',
            '[class*="collapsible"]',
            '[class*="toggle"]'
        ]
        
        for pattern in expandable_patterns:
            elements = soup.select(pattern)
            for element in elements:
                text = element.get_text(strip=True)
                if (text and len(text) > 50 and 
                    not any(exclude in text.lower() for exclude in ['cookie', 'privacy', 'terms', 'size', 'shipping'])):
                    info['description'] = text
                    print(f"Found description in expandable content: {pattern}")
                    break
            if info.get('description'):
                break
    
    # Enhanced size information extraction
    size_selectors = [
        # Direct size guide links/buttons
        '.size-guide', '.size-chart', '.sizing-info', '.size-information',
        '.product-sizing', '.fit-guide', '.measurements', '.dimensions',
        '[data-testid="size-info"]', '.size-details', '.sizing-chart',
        
        # Brand specific
        '.product-detail-size-info', '.size-fit-info', '.size-guide-trigger',
        '.size-chart-link', '.product-size-info', '.sizing-information',
        
        # Interactive elements that might contain size info
        'button[class*="size"]', 'a[class*="size"]',
        '[data-qa-action*="size"]', '[data-testid*="size"]',
        '[data-modal*="size"]', '[data-popup*="size"]'
    ]
    
    for selector in size_selectors:
        try:
            element = soup.select_one(selector)
            if element:
                text = element.get_text(strip=True)
                if text and len(text) > 5:
                    info['size_info'] = text
                    print(f"Found size info using selector: {selector}")
                    break
        except Exception as e:
            continue
    
    # Enhanced materials extraction
    material_selectors = [
        '.materials', '.fabric', '.composition', '.material-composition',
        '.product-materials', '.fabric-composition', '[data-testid="materials"]',
        '.product-detail-composition', '.composition-info',
        '[class*="composition"]', '[class*="material"]', '[class*="fabric"]'
    ]
    
    for selector in material_selectors:
        element = soup.select_one(selector)
        if element:
            text = element.get_text(strip=True)
            if text and any(keyword in text.lower() for keyword in ['cotton', 'polyester', 'material', 'fabric', '%', 'composition']):
                info['materials'] = text
                print(f"Found materials using selector: {selector}")
                break
    
    print(f"Static extraction completed. Found: {list(info.keys())}")
    return info

def extract_dynamic_content(soup, content_type):
    """Extract content that appeared after interaction - Enhanced for multiple brands"""
    content = {}
    
    # Enhanced modal/popup selectors for different brands and platforms
    modal_selectors = [
        # Generic modal selectors
        '.modal-body', '.popup-content', '.modal-content', '.modal-dialog',
        '.dialog-content', '.lightbox-content', '.overlay-content',
        '.modal', '.popup', '.dialog', '.lightbox', '.overlay',
        
        # Size guide specific
        '.size-guide-modal', '.size-modal', '.sizing-modal',
        '.measurements-modal', '.product-measurements', '.size-chart-modal',
        
        # Brand specific selectors
        '.zara-modal', '.h-and-m-modal', '.uniqlo-modal',
        
        # Role-based selectors
        '[role="dialog"]', '[role="modal"]', '[role="alertdialog"]',
        
        # Data attribute selectors
        '[data-modal]', '[data-popup]', '[data-size-guide]',
        
        # Class patterns
        '[class*="modal"]', '[class*="popup"]', '[class*="dialog"]',
        '[class*="size-guide"]', '[class*="measurement"]',
        
        # ID patterns
        '#modal', '#popup', '#size-guide', '#measurements',
        
        # Recently appeared content (might not be in modal)
        '.size-guide-content', '.measurements-content', '.sizing-content',
        '.product-measurements-content', '.size-chart-content'
    ]
    
    modal_content = None
    modal_selector_used = None
    
    # Try each selector to find modal content
    for selector in modal_selectors:
        try:
            element = soup.select_one(selector)
            if element:
                text = element.get_text(strip=True)
                if text and len(text) > 20:  # Must have meaningful content
                    modal_content = element
                    modal_selector_used = selector
                    print(f"Found modal content using selector: {selector}")
                    break
        except Exception as e:
            continue
    
    # If no modal found, look for any recently appeared content with size keywords
    if not modal_content and content_type == 'size_guide':
        print("No modal found, searching for size-related content...")
        
        # Look for any visible element containing size information
        size_keywords = ['size', 'measurement', 'dimension', 'length', 'width', 'chest', 'waist', 'hip']
        
        # Search through all elements for size-related content
        all_elements = soup.find_all(['div', 'section', 'article', 'table', 'ul', 'ol'])
        
        for element in all_elements:
            text = element.get_text(strip=True).lower()
            
            # Check if element contains size-related keywords
            if (len(text) > 50 and 
                any(keyword in text for keyword in size_keywords) and
                # Exclude navigation and header content
                not any(exclude in text for exclude in ['cookie', 'privacy', 'terms', 'navigation', 'menu'])):
                
                modal_content = element
                modal_selector_used = f"size_keyword_search_{element.name}"
                print(f"Found size content in {element.name} element: {text[:100]}...")
                break
    
    if modal_content:
        modal_text = modal_content.get_text(strip=True)
        print(f"Modal content found ({len(modal_text)} chars): {modal_text[:200]}...")
        
        if content_type == 'size_guide':
            # Extract size chart table from modal
            tables = modal_content.select('table')
            if tables:
                print(f"Found {len(tables)} tables in modal")
                for i, table in enumerate(tables):
                    table_data = extract_table_data(table)
                    if table_data and (table_data.get('headers') or table_data.get('rows')):
                        content['size_chart'] = table_data
                        print(f"Extracted size chart table {i}")
                        break
            
            # Look for structured measurement data
            measurements = extract_measurements_from_text(modal_text)
            if measurements:
                content['size_measurements'] = measurements
                print(f"Extracted measurements: {list(measurements.keys())}")
            
            # Always include the full text if it's substantial
            if modal_text and len(modal_text) > 20:
                content['size_info'] = modal_text
                print("Added full modal text as size_info")
            
            # Look for specific measurement patterns
            size_patterns = extract_size_patterns(modal_text)
            if size_patterns:
                content['size_patterns'] = size_patterns
                print(f"Extracted size patterns: {size_patterns}")
        
        elif content_type == 'description':
            if modal_text and len(modal_text) > 20:
                content['description'] = modal_text
                print("Added modal text as description")
        
        elif content_type == 'specifications':
            if modal_text and len(modal_text) > 20:
                content['specifications'] = modal_text
        
        elif content_type == 'care':
            if modal_text and len(modal_text) > 10:
                content['care_instructions'] = modal_text
        
        # Add metadata about extraction
        content['extraction_method'] = 'modal_extraction'
        content['modal_selector'] = modal_selector_used
        content['modal_text_length'] = len(modal_text)
    
    # If still no content found for size guide, try alternative approaches
    if not content and content_type == 'size_guide':
        print("Trying alternative size extraction methods...")
        
        # Method 1: Look for any new tables that might have appeared
        all_tables = soup.select('table')
        for table in all_tables:
            table_text = table.get_text(strip=True).lower()
            if any(keyword in table_text for keyword in ['size', 'measurement', 'cm', 'inch', 'xs', 'sm', 'md', 'lg', 'xl']):
                table_data = extract_table_data(table)
                if table_data:
                    content['size_chart'] = table_data
                    content['extraction_method'] = 'table_scan'
                    print("Found size table through table scan")
                    break
        
        # Method 2: Look for lists with size information
        all_lists = soup.select('ul, ol')
        for list_elem in all_lists:
            list_text = list_elem.get_text(strip=True)
            if (len(list_text) > 50 and 
                any(keyword in list_text.lower() for keyword in ['size', 'measurement', 'dimension'])):
                content['size_info'] = list_text
                content['extraction_method'] = 'list_extraction'
                print("Found size info in list element")
                break
        
        # Method 3: Check page for any content changes
        page_text = soup.get_text()
        if any(keyword in page_text.lower() for keyword in ['size guide', 'measurements', 'sizing chart']):
            measurements = extract_measurements_from_text(page_text)
            if measurements:
                content['size_measurements'] = measurements
                content['size_info'] = "Measurements found in page content"
                content['extraction_method'] = 'page_scan'
                print("Extracted measurements from full page")
    
    print(f"Dynamic content extraction result: {list(content.keys())}")
    return content

def extract_size_patterns(text):
    """Extract common size patterns from text"""
    import re
    
    patterns = {}
    
    # Pattern for size ranges (e.g., "XS: 32-34", "Size S: 36-38")
    size_range_pattern = r'(?:Size\s+)?([XS]{1,2}|[SML]{1,2}|\d{2})\s*:?\s*(\d{1,3}(?:\.\d)?)\s*[-â€“]\s*(\d{1,3}(?:\.\d)?)'
    matches = re.finditer(size_range_pattern, text, re.IGNORECASE)
    
    for match in matches:
        size, min_val, max_val = match.groups()
        patterns[f"size_{size.upper()}_range"] = f"{min_val}-{max_val}"
    
    # Pattern for measurements with units
    measurement_pattern = r'(\w+(?:\s+\w+)?)\s*:?\s*(\d{1,3}(?:\.\d)?)\s*(cm|inches?|in)'
    matches = re.finditer(measurement_pattern, text, re.IGNORECASE)
    
    for match in matches:
        measurement_type, value, unit = match.groups()
        clean_type = measurement_type.strip().lower().replace(' ', '_')
        patterns[f"{clean_type}_{unit.lower()}"] = f"{value} {unit}"
    
    return patterns if patterns else None

def extract_measurements_from_text(text):
    """Extract structured measurements from text content"""
    import re
    
    measurements = {}
    
    # Common measurement patterns
    measurement_patterns = [
        # Pattern: "Size XS: Length 95cm, Waist 68cm"
        r'Size\s+(\w+):\s*([^,]+(?:,\s*[^,]+)*)',
        # Pattern: "XS - Length: 95cm, Waist: 68cm"
        r'(\w+)\s*-\s*([^-]+(?=\w+\s*-|$))',
        # Pattern: "Length: 95cm" 
        r'(\w+):\s*(\d+(?:\.\d+)?\s*cm|\d+(?:\.\d+)?\s*inches?)',
        # Pattern: "95cm length"
        r'(\d+(?:\.\d+)?\s*cm|\d+(?:\.\d+)?\s*inches?)\s+(\w+)',
    ]
    
    # Size categories to look for
    size_categories = ['XS', 'S', 'M', 'L', 'XL', 'XXL', '34', '36', '38', '40', '42', '44', '46']
    measurement_types = ['length', 'width', 'chest', 'waist', 'hips', 'shoulder', 'sleeve', 'inseam', 'rise']
    
    for pattern in measurement_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            groups = match.groups()
            if len(groups) >= 2:
                key, value = groups[0], groups[1]
                
                # Clean and structure the data
                key = key.strip()
                value = value.strip()
                
                if key.upper() in size_categories:
                    # This is a size with measurements
                    measurements[f"size_{key.upper()}"] = value
                elif any(mtype in key.lower() for mtype in measurement_types):
                    # This is a measurement type
                    measurements[key.lower()] = value
    
    return measurements if measurements else None

def extract_table_data(table):
    """Extract data from size chart tables"""
    try:
        table_data = {
            'headers': [],
            'rows': []
        }
        
        # Extract headers
        headers = table.select('thead th, tr:first-child th, tr:first-child td')
        if headers:
            table_data['headers'] = [th.get_text(strip=True) for th in headers]
        
        # Extract rows
        rows = table.select('tbody tr, tr')
        for i, row in enumerate(rows):
            if i == 0 and table_data['headers']:
                continue  # Skip header row
            
            cells = row.select('td, th')
            row_data = [cell.get_text(strip=True) for cell in cells]
            if row_data and any(row_data):  # Skip empty rows
                table_data['rows'].append(row_data)
        
        return table_data
        
    except Exception as e:
        return {'error': f'Error extracting table: {str(e)}'}

def close_popup(driver):
    """Enhanced popup closing with more selectors"""
    close_selectors = [
        '.modal-close', '.close', '.popup-close', '[aria-label="Close"]',
        '.close-btn', '.modal-dismiss', 'button[data-dismiss="modal"]',
        '.overlay-close', '.lightbox-close',
        # ZARA specific close buttons
        '.zara-modal-close', '.modal-backdrop',
        # Generic close patterns
        'button:contains("Close")', 'button:contains("Ã—")', 
        'a:contains("Close")', '[title="Close"]',
        # ESC key simulation
        '.modal.show', '.modal.in'  # For backdrop click
    ]
    
    for selector in close_selectors:
        try:
            if ':contains(' in selector:
                # Handle text-based selectors with XPath
                text = selector.split(':contains("')[1].split('")')[0]
                tag = selector.split(':contains(')[0]
                xpath = f"//{tag}[contains(text(), '{text}')]"
                close_elements = driver.find_elements(By.XPATH, xpath)
            else:
                close_elements = driver.find_elements(By.CSS_SELECTOR, selector)
            
            for close_btn in close_elements:
                if close_btn.is_displayed():
                    try:
                        close_btn.click()
                        time.sleep(1)
                        print(f"Closed popup using selector: {selector}")
                        return True
                    except:
                        # Try JavaScript click
                        try:
                            driver.execute_script("arguments[0].click();", close_btn)
                            time.sleep(1)
                            print(f"Closed popup with JS using selector: {selector}")
                            return True
                        except:
                            continue
                            
        except Exception as e:
            continue
    
    # Try pressing ESC key as fallback
    try:
        from selenium.webdriver.common.keys import Keys
        from selenium.webdriver.common.action_chains import ActionChains
        ActionChains(driver).send_keys(Keys.ESCAPE).perform()
        time.sleep(1)
        print("Closed popup using ESC key")
        return True
    except:
        pass
    
    return False

def scrape_static_product(url, config):
    """Fallback: scrape product info using regular HTTP requests"""
    try:
        response = make_request(url, config)
        soup = parse_html(response)
        
        result = {
            'success': True,
            'url': url,
            'scraped_at': datetime.now().isoformat(),
            'method': 'static',
            'page_title': get_page_metadata(soup, url)['page_title'],
            'product_info': {}
        }
        
        # Extract product information
        product_info = extract_static_product_info(soup)
        product_info['interactive_content_found'] = False
        
        result['product_info'] = clean_product_info(product_info)
        
        return result
        
    except Exception as e:
        return {
            'success': False,
            'error': f'Static scraping error: {str(e)}',
            'method': 'static'
        }

@app.route('/debug_scrape', methods=['POST'])
@cross_origin()
@require_auth
def debug_scrape():
    """Debug endpoint to see what's actually happening during scraping"""
    try:
        data = request.get_json()
        url = data.get('url', '').strip()
        
        if not url:
            return jsonify({
                'success': False,
                'error': 'URL is required'
            }), 400
        
        driver = None
        try:
            print(f"Starting debug scrape for: {url}")
            
            # Setup driver in non-headless mode for debugging
            driver = setup_driver(headless=False)
            driver.set_page_load_timeout(30)
            
            # Step 1: Load page
            print("Step 1: Loading page...")
            driver.get(url)
            time.sleep(3)
            
            # Step 2: Handle popups
            print("Step 2: Handling popups...")
            popup_handled = handle_cookie_consent_and_popups(driver)
            print(f"Popup handled: {popup_handled}")
            
            # Step 3: Wait for page to be ready
            print("Step 3: Waiting for page to be ready...")
            wait = WebDriverWait(driver, 10)
            wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
            time.sleep(2)
            
            # Step 4: Get page source and basic info
            print("Step 4: Getting page info...")
            page_title = driver.title
            page_source_length = len(driver.page_source)
            
            # Step 5: Parse with BeautifulSoup and look for product content
            print("Step 5: Parsing with BeautifulSoup...")
            soup = BeautifulSoup(driver.page_source, 'html.parser')
            
            # Check for common product description selectors
            description_selectors = [
                '.product-description', '.pdp-description', '.product-details-description',
                '[data-testid="product-description"]', '.product-info-description',
                '.product-long-description', '.description-content', '.product-content',
                '.product-description-text', '.item-description', '.product-summary',
                '.product-details-content', '.product-information', '#product-description',
                # Add more generic selectors
                '.description', '.details', '.overview', '.about', '.info'
            ]
            
            found_descriptions = {}
            for selector in description_selectors:
                elements = soup.select(selector)
                if elements:
                    for i, element in enumerate(elements[:3]):  # Check first 3 matches
                        text = element.get_text(strip=True)
                        if text and len(text) > 20:  # Only meaningful text
                            found_descriptions[f"{selector}_{i}"] = {
                                'selector': selector,
                                'text_length': len(text),
                                'text_preview': text[:200] + '...' if len(text) > 200 else text
                            }
            
            # Check for any element with "description" in class or id
            description_elements = soup.find_all(attrs={'class': lambda x: x and 'description' in ' '.join(x).lower()})
            description_elements.extend(soup.find_all(attrs={'id': lambda x: x and 'description' in x.lower()}))
            
            generic_descriptions = {}
            for i, element in enumerate(description_elements[:5]):
                text = element.get_text(strip=True)
                if text and len(text) > 20:
                    generic_descriptions[f"generic_{i}"] = {
                        'classes': element.get('class', []),
                        'id': element.get('id', ''),
                        'tag': element.name,
                        'text_length': len(text),
                        'text_preview': text[:200] + '...' if len(text) > 200 else text
                    }
            
            # Check for paragraphs that might contain product info
            paragraphs = soup.find_all('p')
            long_paragraphs = []
            for p in paragraphs:
                text = p.get_text(strip=True)
                if len(text) > 100:  # Look for substantial paragraphs
                    long_paragraphs.append({
                        'text_length': len(text),
                        'text_preview': text[:150] + '...' if len(text) > 150 else text,
                        'classes': p.get('class', []),
                        'parent_classes': p.parent.get('class', []) if p.parent else []
                    })
            
            # Count different element types
            element_counts = {
                'total_paragraphs': len(soup.find_all('p')),
                'total_divs': len(soup.find_all('div')),
                'total_spans': len(soup.find_all('span')),
                'elements_with_description_class': len(description_elements),
                'long_paragraphs': len(long_paragraphs)
            }
            
            return jsonify({
                'success': True,
                'url': url,
                'page_title': page_title,
                'page_source_length': page_source_length,
                'popup_handled': popup_handled,
                'element_counts': element_counts,
                'found_descriptions': found_descriptions,
                'generic_descriptions': generic_descriptions,
                'long_paragraphs': long_paragraphs[:3],  # First 3 only
                'debug_info': {
                    'selectors_tested': len(description_selectors),
                    'soup_parsed': True,
                    'driver_title': page_title
                }
            })
            
        except Exception as e:
            return jsonify({
                'success': False,
                'error': f'Debug scraping error: {str(e)}'
            }), 500
        finally:
            if driver:
                try:
                    driver.quit()
                except:
                    pass
                    
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Request error: {str(e)}'
        }), 500

def clean_product_info(product_info):
    """Clean and format the extracted product information"""
    cleaned = {}
    
    for key, value in product_info.items():
        if value is None or value == '' or value == []:
            cleaned[key] = None
            continue
        
        if isinstance(value, str):
            # Clean whitespace
            value = ' '.join(value.split())
            
            # Remove HTML artifacts
            value = re.sub(r'<[^>]+>', '', value)
            
            # Clean up common text artifacts
            value = value.replace('\n', ' ').replace('\r', ' ')
            value = re.sub(r'\s+', ' ', value).strip()
            
            if len(value) > 10:  # Only keep meaningful text
                cleaned[key] = value
            else:
                cleaned[key] = None
        
        elif isinstance(value, dict):
            # For table data, keep as is but clean text
            if 'headers' in value and 'rows' in value:
                cleaned_table = {
                    'headers': [h.strip() for h in value['headers'] if h.strip()],
                    'rows': []
                }
                for row in value['rows']:
                    cleaned_row = [cell.strip() for cell in row if cell.strip()]
                    if cleaned_row:
                        cleaned_table['rows'].append(cleaned_row)
                
                if cleaned_table['headers'] or cleaned_table['rows']:
                    cleaned[key] = cleaned_table
                else:
                    cleaned[key] = None
            else:
                cleaned[key] = value
        
        else:
            cleaned[key] = value
    
    return cleaned

def generate(selected_chunks, query):
    client = openai.OpenAI()
    context = "\n\n".join(selected_chunks) 
    prompt = f"Answer the following query based on the provided text:\n\n{context}\n\nQuery: {query}\nAnswer:" 
    # response = client.chat.completions.create( 
    #     model="gpt-4", 
    #     messages=[ {"role": "system", "content": "You are a legal research and reasoning assistant trained in Indian income tax law, especially capital gains exemptions under the Income Tax Act. Your job is to analyze a user's scenario, determine applicability of specific sections (like Section 54F), and generate responses following a clear structure: Start with statutory interpretation â€” quote the relevant section (e.g., Section 54F) and clearly list the conditions in bullet points. Apply the law to the userâ€™s case â€” mention whether conditions are satisfied and explain eligibility for exemption. Cite relevant case law in support of the position taken. Choose cases that match the factual scenario and jurisdiction where possible. Include citation (e.g., ITA 4012/Mum/2023 - Abdul Nayab Shaikh). Quote only favourable rulings unless otherwise requested. Prefer recent, relevant, and jurisdictionally appropriate cases. Discuss any common exceptions or judicial deviations â€” e.g., benefit being allowed even when more than one residential unit is purchased, especially if adjacent or used as a single unit. Quote examples from case law or factual scenarios to support the interpretation or exception. Keep the examples precise and relevant. Format your response in a professional, advisory tone suitable for a tax consultantâ€™s opinion. Do not speculate â€” rely only on clear statutory provisions, circulars, and judicial precedents."}, {"role": "user", "content": prompt} ], 
    #     max_tokens=400, 
    #     temperature=0.1 ) 

    # response = client.chat.completions.create( 
    #     model="gpt-4", 
    #     messages=[ {"role": "system", "content": "You are a professional skills extractor"}, {"role": "user", "content": prompt} ], 
    #     max_tokens=400, 
    #     temperature=0.1 )
    # 

    # response = client.chat.completions.create( 
    #     model="gpt-4", 
    #      messages=[ {"role": "system", "content": "You are a PhD level research assistant that understands AI and its future and you also have a strong business acumen that will help you build a strong pitch for an AI startup"}, {"role": "user", "content": prompt} ], 
    #     max_tokens=400, 
    #     temperature=0.1 ) 

    response = client.chat.completions.create( 
    model="gpt-4", 
    messages=[ 
        {
            "role": "system", 
            "content": "You are a PhD-level research assistant with deep expertise in cryptocurrency markets, blockchain technology, and financial analysis. You provide insightful, data-driven analysis of crypto assets, including market trends, tokenomics, risk factors, and trading strategies."
        }, 
        {
            "role": "user", 
            "content": prompt
        } 
    ], 
    max_tokens=400, 
    temperature=0.1 
)
    
    
    print(response)
    answer = response.choices[0].message.content 
    # usage = response.usage
    return answer

@app.route('/upload', methods=['POST'])
@require_auth
def upload_file():
    # Accept folder_name from form data (for file uploads)
    folder_name = request.form.get('folder_name', '').strip()

    if 'file' not in request.files:
        return jsonify({"error": "No file part in the request"}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({"error": "No file selected for uploading"}), 400

    # Define the base upload directory
    base_upload_folder = os.path.join(DATA_DIR, 'uploaded_data')

    # If folder_name is provided, create/use subfolder
    if folder_name:
        upload_folder = os.path.join(base_upload_folder, folder_name)
    else:
        upload_folder = base_upload_folder

    os.makedirs(upload_folder, exist_ok=True)  # Create the directory if it doesn't exist

    # Check if the file already exists
    file_path = os.path.join(upload_folder, file.filename)
    if os.path.exists(file_path):
        return jsonify({"message": "File already exists", "file_path": file_path}), 200

    # Save the file
    file.save(file_path)

    return jsonify({"message": "File uploaded successfully", "file_path": file_path}), 200


@app.route('/scrape_product_info', methods=['POST'])
@cross_origin()
@require_auth
def scrape_product_info():
    """
    Enhanced product scraping API focused on descriptions and size information
    with interactive element handling
    """
    try:
        data = request.get_json()
        url = data.get('url', '').strip()
        interactive = data.get('interactive', True)  # Use interactive mode by default
        headless = data.get('headless', True)
        config = data.get('config', {})
        
        if not url:
            return jsonify({
                'success': False,
                'error': 'URL is required'
            }), 400
        
        # Validate URL
        try:
            parsed = urlparse(url)
            if not all([parsed.scheme, parsed.netloc]):
                return jsonify({
                    'success': False,
                    'error': 'Invalid URL format'
                }), 400
        except Exception as e:
            return jsonify({
                'success': False,
                'error': f'Invalid URL: {str(e)}'
            }), 400
        
        if interactive:
            # Use Selenium for interactive scraping
            result = scrape_with_interaction(url, headless, config)
        else:
            # Use regular HTTP scraping
            result = scrape_static_product(url, config)
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Scraping error: {str(e)}'
        }), 500


@app.route('/scrape', methods=['POST'])
@cross_origin()
@require_auth
def universal_scraper():
    """
    Universal scraping API - now includes enhanced product scraping.

    DEPRECATION NOTICE: This endpoint is maintained for backwards compatibility.
    New code should use the Connector API instead:

        POST /api/connectors/web_scraper/fetch
        {
            "resource": "page|text|tables|product|json_ld|links|custom",
            "params": {"url": "https://example.com", ...}
        }

    The Connector API provides:
    - Automatic context storage (retrieved data is available to all agents)
    - User-specific proxy/rate-limit settings from Settings UI
    - Unified error handling
    """
    try:
        data = request.get_json()
        url = data.get('url', '').strip()
        scrape_type = data.get('type', 'basic')
        config = data.get('config', {})
        
        # Validation
        if not url:
            return jsonify({
                'success': False,
                'error': 'URL is required'
            }), 400
        
        # Validate URL format
        try:
            parsed = urlparse(url)
            if not all([parsed.scheme, parsed.netloc]):
                return jsonify({
                    'success': False,
                    'error': 'Invalid URL format. Include http:// or https://'
                }), 400
        except Exception as e:
            return jsonify({
                'success': False,
                'error': f'Invalid URL: {str(e)}'
            }), 400
        
        # Route to appropriate scraper based on type
        if scrape_type == 'product_info':
            # Use the new enhanced product scraper - but call it properly
            data['interactive'] = data.get('interactive', True)
            data['headless'] = data.get('headless', True)
            if data.get('interactive', True):
                result = scrape_with_interaction(url, data.get('headless', True), config)
            else:
                result = scrape_static_product(url, config)
            return jsonify(result)
        elif scrape_type == 'basic':
            result = scrape_basic_content(url, config)
            return jsonify(result)
        elif scrape_type == 'text':
            result = scrape_text_content(url, config)
            return jsonify(result)
        elif scrape_type == 'json_ld':
            result = scrape_json_ld_content(url, config)
            return jsonify(result)
        elif scrape_type == 'product':
            result = scrape_product_content(url, config)
            return jsonify(result)
        elif scrape_type == 'tables':
            result = scrape_table_content(url, config)
            return jsonify(result)
        elif scrape_type == 'custom':
            result = scrape_custom_selectors(url, config)
            return jsonify(result)
        else:
            return jsonify({
                'success': False,
                'error': f'Unknown scrape type: {scrape_type}'
            }), 400
            
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Scraping error: {str(e)}'
        }), 500

@app.route('/search_suggestions', methods=['POST'])
@cross_origin()
@require_auth
def openai_chat():
    """Simple OpenAI chat endpoint for matchmaking"""
    try:
        data = request.get_json()
        prompt = data.get('prompt', '')
        max_tokens = data.get('max_tokens', 500)
        temperature = data.get('temperature', 0.7)
        
        client = openai.OpenAI()
        client.api_key = os.environ['OPENAI_API_KEY']
        
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=max_tokens,
            temperature=temperature
        )
        
        return jsonify({
            'success': True,
            'response': response.choices[0].message.content.strip()
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

# @app.route('/rag_test', methods=['GET'])
# def rag_test():
#     query = request.args.get('query')
#     file_name = request.args.get('file_name')

#     upload_folder = "/Users/rhishikeshthakur/Enable/Software Development/enable_agents/data"
#     file_path = os.path.join(upload_folder, file_name)

#     # file_size = os.path.getsize(file_path)  # in bytes
    
#     if not os.path.exists(file_path):
#         return jsonify({"error": f"File '{file_name}' not found"}), 404

#     print(f"Absolute file path: {file_path}")
    
#     openai.api_key = get_credentials();

#     file_hash = get_file_hash(file_path)

#     if file_hash in cache:
#         print(f"Using cached embeddings for file hash: {file_hash}")
#         index, phrase_embeddings, page_chunks = load_embeddings(file_hash)
#     else:
#         print(f"Processing file and saving embeddings for file hash: {file_hash}")
#         pdf_doc = pdf_loader(file_path)
#         page_chunks = pdf_splitter(pdf_doc)

#         # print(page_chunks)

#         page_phrases = extract_keywords_from_pdf(pdf_doc)
#         chunk_phrases = extract_keywords_from_chunks(page_chunks)
        
#         index, phrase_embeddings = store_embeddings(page_phrases, chunk_phrases)
    
#         cache[file_hash] = (index, phrase_embeddings, page_chunks)
#         save_embeddings(file_hash, index, phrase_embeddings, page_chunks)
#         print(save_embeddings)
        
#     query_phrases = extract_phrases_from_query(query)
#     query_embeddings = get_embeddings_for_query(query_phrases)
#     selected_chunks = retrieve_similar_chunks(query_embeddings, index, phrase_embeddings, page_chunks)

#     max_chunks = 5
#     max_chunk_length = 1000  # characters
#     selected_chunks = [chunk[:max_chunk_length] for chunk in selected_chunks[:max_chunks]]


#     answer = generate(selected_chunks, query)


#     return jsonify({"answer": answer})  # Always return a JSON object

@app.route('/rag_test', methods=['POST'])
@require_auth
def rag_test():
    try:
        data = request.get_json()
        query = data.get('query')
        file_name = data.get('file_name')
        data_store = data.get('data_store')  # This is the folder name under uploaded_data

        if not query or not file_name:
            return jsonify({"error": "Missing query or file_name parameter"}), 400

        base_upload_folder = os.path.join(DATA_DIR, 'uploaded_data')
        if data_store:
            file_path = os.path.join(base_upload_folder, data_store, file_name)
        else:
            file_path = os.path.join(base_upload_folder, file_name)

        if not os.path.exists(file_path):
            return jsonify({"error": f"File '{file_path}' not found"}), 404

        print(f"Processing file: {file_path}")
        openai.api_key = get_credentials()
        file_hash = get_file_hash(file_path)

        try:
            if file_hash in cache:
                print(f"Using cached embeddings for file hash: {file_hash}")
                index, phrase_embeddings, page_chunks = load_embeddings(file_hash)
            else:
                print(f"Processing file and saving embeddings for file hash: {file_hash}")
                pdf_doc = pdf_loader(file_path)
                
                if not pdf_doc:
                    return jsonify({"error": "Could not extract text from PDF"}), 400
                
                page_chunks = pdf_splitter(pdf_doc)
                
                if not page_chunks:
                    return jsonify({"error": "Could not create chunks from PDF content"}), 400
                
                page_phrases = extract_keywords_from_pdf(pdf_doc)
                chunk_phrases = extract_keywords_from_chunks(page_chunks)
                
                index, phrase_embeddings = store_embeddings(page_phrases, chunk_phrases)
                cache[file_hash] = (index, phrase_embeddings, page_chunks)
                save_embeddings(file_hash, index, phrase_embeddings, page_chunks)

            query_phrases = extract_phrases_from_query(query)
            if not query_phrases:
                query_phrases = [query]  # Use the full query if no phrases extracted
            
            query_embeddings = get_embeddings_for_query(query_phrases)
            selected_chunks = retrieve_similar_chunks(query_embeddings, index, phrase_embeddings, page_chunks)

            max_chunks = 5
            max_chunk_length = 1000  # characters
            selected_chunks = [chunk[:max_chunk_length] for chunk in selected_chunks[:max_chunks]]

            answer = generate(selected_chunks, query)

            return jsonify({
                "answer": answer,
                "chunks_used": len(selected_chunks),
                "file_processed": file_name
            })
            
        except Exception as processing_error:
            print(f"Error processing PDF: {processing_error}")
            return jsonify({"error": f"Error processing PDF: {str(processing_error)}"}), 500

    except Exception as e:
        print(f"RAG test error: {e}")
        return jsonify({"error": f"Server error: {str(e)}"}), 500

@app.route('/enterprise_chat', methods=['POST', 'OPTIONS'])
@cross_origin()
@require_auth
def enterprise_chat():
    """
    Enterprise chat API that collects business context from the user.
    User can answer any of the key questions first; the API will detect which question was answered,
    store it, and then ask the remaining unanswered questions.
    Once all questions are answered, sends chat_state to OpenAI to generate a summary and auto-fill missing fields.
    """
    data = request.get_json()
    chat_state = data.get('chat_state', {})
    last_answer = data.get('last_answer', '').strip()
    last_question_key = data.get('last_question_key', '').strip()

    # Define the sequence and mapping of questions
    questions = [
        {"key": "industry", "question": "To get a bit of context, which industry does your business operate in?"},
        {"key": "product_service", "question": "What primary product or service does your business offer to customers?"},
        {"key": "role_department", "question": "What is your role within the company, and what is your department mainly focused on right now?"},
        {"key": "tools", "question": "What tools or software do you and your team rely on most, and what do you use them for?"},
        {"key": "business_need", "question": "If you could change or improve one thing about how your team works today, what would it be?"}
    ]
    question_keys = [q['key'] for q in questions]

    # If the user answered a question, store it in chat_state
    if last_question_key in question_keys and last_answer:
        chat_state[last_question_key] = last_answer

    # If the user sent an answer but didn't specify which question, try to infer
    if not last_question_key and last_answer:
        for q in questions:
            if q['key'] not in chat_state or not chat_state.get(q['key']):
                chat_state[q['key']] = last_answer
                break

    # Find the next unanswered question
    for q in questions:
        if q['key'] not in chat_state or not chat_state[q['key']]:
            return jsonify({
                "success": True,
                "next_question": q['question'],
                "next_question_key": q['key'],
                "chat_state": chat_state,
                "completed": False
            })

    # If all questions answered, auto-fill and format chat_state using OpenAI
    try:
        import openai
        openai.api_key = os.environ.get("OPENAI_API_KEY")
        # Prompt to format and auto-fill chat_state
        autofill_prompt = (
            "Given the following user answers, format the business context as a JSON object with these keys: "
            "industry, product_service, role, department_context, business_need, and tools (as a list of objects with tool_name and description). "
            "If any field is missing or vague, infer and auto-fill it based on the other answers. "
            "Example format:\n"
            "{\n"
            '  "tools": [\n'
            '    {"tool_name": "Slack", "description": "Team communication and collaboration platform"},\n'
            '    {"tool_name": "Salesforce", "description": "CRM for managing customer relationships and sales pipeline"}\n'
            "  ],\n"
            '  "industry": "Technology",\n'
            '  "product_service": "B2B workflow automation software for sales and operations teams",\n'
            '  "role": "Sales Manager",\n'
            '  "department_context": "Our sales department is focused on improving lead conversion and automating reporting.",\n'
            '  "business_need": "We want to integrate our communication and CRM tools, automate sales reporting, and identify missing modules for analytics."\n'
            "}\n"
            "User answers:\n"
            f"Industry: {chat_state.get('industry', '')}\n"
            f"Product/Service: {chat_state.get('product_service', '')}\n"
            f"Role and Department Context: {chat_state.get('role_department', '')}\n"
            f"Tools: {chat_state.get('tools', '')}\n"
            f"Business Need: {chat_state.get('business_need', '')}\n"
            "Return only valid JSON."
        )

        client = openai.OpenAI()
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a business analyst assistant."},
                {"role": "user", "content": autofill_prompt}
            ],
            max_tokens=300,
            temperature=0.2
        )
        # Extract JSON from response
        import re
        raw_content = response.choices[0].message.content.strip()
        match = re.search(r'\{[\s\S]*\}', raw_content)
        if match:
            formatted_state = json.loads(match.group(0))
        else:
            formatted_state = chat_state  # fallback

        # Summarize the context for search_summary
        summary_prompt = (
            "Summarize the following business context in 2-3 sentences for agent recommendation:\n\n"
            f"{json.dumps(formatted_state, indent=2)}"
        )
        summary_response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a business analyst assistant."},
                {"role": "user", "content": summary_prompt}
            ],
            max_tokens=150,
            temperature=0.3
        )
        search_summary = summary_response.choices[0].message.content.strip()
        print(formatted_state)

    except Exception as e:
        formatted_state = chat_state
        search_summary = f"Could not generate summary: {str(e)}"

    return jsonify({
        "success": True,
        "message": "Thank you! Here is the summary of your business context.",
        "chat_state": formatted_state,
        "completed": True,
        "search_summary": search_summary
    })


@app.route('/chat_api', methods=['POST', 'OPTIONS'])
@cross_origin()  # Allow CORS for this endpoint
@require_auth
def chat_api():
    import glob
    data = request.get_json()
    query = data.get('query')

    embeddings_folder = os.path.join(DATA_DIR, 'embeddings')
    embedding_files = glob.glob(os.path.join(embeddings_folder, "*_index.pkl"))
    file_hashes = [os.path.basename(f).split('_')[0] for f in embedding_files]

    if not file_hashes:
        return jsonify({"error": "No knowledge base available. Please upload and process at least one file first."}), 400

    # Aggregate all embeddings, phrase mappings, and page chunks
    all_indexes = []
    all_phrase_embeddings = {}
    all_page_chunks = {}

    for file_hash in file_hashes:
        try:
            index, phrase_embeddings, page_chunks = load_embeddings(file_hash)
        except Exception as e:
            print(f"Error loading embeddings for {file_hash}: {e}")
            continue  # Skip files that can't be loaded

        all_indexes.append(index)
        # Update keys to include file_hash for uniqueness
        for (page, chunk_number), phrases in phrase_embeddings.items():
            all_phrase_embeddings[(file_hash, page, chunk_number)] = phrases
        for page, chunks in page_chunks.items():
            all_page_chunks[(file_hash, page)] = chunks

    # For simplicity, use the first index (FAISS) for searching, or you can merge indexes if needed
    index = all_indexes[0] if all_indexes else None
    phrase_embeddings = all_phrase_embeddings
    page_chunks = all_page_chunks

    if not index or not phrase_embeddings or not page_chunks:
        return jsonify({"error": "No valid embeddings found."}), 400

    query_phrases = extract_phrases_from_query(query)
    query_embeddings = get_embeddings_for_query(query_phrases)
    selected_chunks = retrieve_similar_chunks(query_embeddings, index, phrase_embeddings, page_chunks)

    max_chunks = 5
    max_chunk_length = 1000  # characters
    selected_chunks = [chunk[:max_chunk_length] for chunk in selected_chunks[:max_chunks]]

    answer = generate(selected_chunks, query)

    return jsonify({"answer": answer})

@app.route('/save-prompt', methods=['POST'])
@require_auth
def save_prompt():
    data = request.get_json()
    prompt_id = str(uuid4())  # Generate a unique ID for the prompt
    data['id'] = prompt_id
    data['timestamp'] = datetime.now().isoformat()

    # Load existing prompts
    if os.path.exists(PROMPTS_FILE):
        with open(PROMPTS_FILE, 'r') as file:
            prompts = json.load(file)
    else:
        prompts = []

    # Add the new prompt
    prompts.append(data)

    # Save back to the file
    with open(PROMPTS_FILE, 'w') as file:
        json.dump(prompts, file, indent=4)

    return jsonify({"message": "Prompt saved successfully", "id": prompt_id})

@app.route('/previous-prompts', methods=['GET'])
@require_auth
def previous_prompts():
    if os.path.exists(PROMPTS_FILE):
        with open(PROMPTS_FILE, 'r') as file:
            prompts = json.load(file)
    else:
        prompts = []

    return jsonify({"prompts": prompts})

@app.route('/yfinance', methods=['POST'])
@require_auth
def yfinance_test():
    data = request.get_json()
    symbol = data.get('stock')
    region = data.get('region')

    if not symbol or not region:
        return "Missing required parameters: 'stock' and 'region'", 400

    conn = http.client.HTTPSConnection("yahoo-finance166.p.rapidapi.com")

    headers = {
        'x-rapidapi-key': "95cdd43379mshbd9483856442c47p1c2782jsn897449ebefb8",
        'x-rapidapi-host': "yahoo-finance166.p.rapidapi.com"
    }

    endpoint = f"/api/stock/get-financial-data?region={region}&symbol={symbol}"
    print(f"Requesting data from endpoint: {endpoint}")  # Debug statement
    conn.request("GET", endpoint, headers=headers)

    res = conn.getresponse()
    data = res.read()
    json_data = json.loads(data.decode("utf-8"))

    print(json_data)  # Debug statement to print the entire response

    if 'quoteSummary' not in json_data or 'result' not in json_data['quoteSummary'] or not json_data['quoteSummary']['result']:
        return jsonify({"error": "No data found for the given stock symbol and region"}), 404

    current_price = json_data['quoteSummary']['result'][0]['financialData']['currentPrice']['fmt']
    operating_margins = json_data['quoteSummary']['result'][0]['financialData']['operatingMargins']['fmt']
    netprofit_margins = json_data['quoteSummary']['result'][0]['financialData']['profitMargins']['fmt']
    gross_margins = json_data['quoteSummary']['result'][0]['financialData']['grossMargins']['fmt']
    revenue_growth = json_data['quoteSummary']['result'][0]['financialData']['revenueGrowth']['fmt']
    debt_to_equity = json_data['quoteSummary']['result'][0]['financialData']['debtToEquity']['fmt']
    quick_ratio = json_data['quoteSummary']['result'][0]['financialData']['quickRatio']['fmt']
    current_ratio = json_data['quoteSummary']['result'][0]['financialData']['currentRatio']['fmt']
    analyst_recommendation = json_data['quoteSummary']['result'][0]['financialData']['recommendationKey']
    number_of_analysts = json_data['quoteSummary']['result'][0]['financialData']['numberOfAnalystOpinions']['fmt']
    target_high_price = json_data['quoteSummary']['result'][0]['financialData']['targetHighPrice']['fmt']
    target_low_price = json_data['quoteSummary']['result'][0]['financialData']['targetLowPrice']['fmt']
    target_mean_price = json_data['quoteSummary']['result'][0]['financialData']['targetMeanPrice']['fmt']
    target_median_price = json_data['quoteSummary']['result'][0]['financialData']['targetMedianPrice']['fmt']

    financial_KPIs = {
        "current_price": current_price,
        "operating margin": operating_margins,
        "netprofit_margins": netprofit_margins,
        "gross_margins": gross_margins,
        "revenue_growth": revenue_growth,
        "debt_to_equity": debt_to_equity,
        "quick_ratio": quick_ratio,
        "current_ratio": current_ratio,
        "number_of_analysts": number_of_analysts,
        "analyst_recommendation": analyst_recommendation,
        "target_high_price": target_high_price,
        "target_low_price": target_low_price,
        "target_mean_price": target_mean_price,
        "target_median_price": target_median_price
    }

    return jsonify(financial_KPIs)

@app.route('/generate-requirements', methods=['POST'])
@require_auth
def generate_requirements():
    openai.api_key = get_credentials()

    data = request.get_json()
    overview = data.get('overview', '')
    context = data.get('context', '')  # Get the context from the payload
    country = data.get('countries', '')
    industries = data.get('industries', '')
    function = data.get('businessFunction', '')
    frameworks = data.get('frameworks', [])

    format = data.get('responseFormat', '')
    

    # prompt = f"""
    # Draft requirements based on the requirements {overview} that are specific, measurable, achievable, relevant, and time-bound (SMART).
    # Consider {context} as context for the requirements being asked for, focus on the market in {country} or region, 
    # consider {industries} for industry related insights, consider {function} as the role or business function of the requester,
    # and without mentioning the framework in the final response, conduct research taking into account these analysis frameworks: {frameworks} for one valuable and rare resource each using the VRIO, market forces for and against the startup using PESTLE, and product readiness using Mckinsey's 3 Horizon and use response format as reference: {format}.
    # """

    prompt = f"""
    You are a research assistant tasked with producing high-quality, insightful, and well-structured research on business opportunitiesand growth prospects. Your output should include a curated but accessible for free list of relevant academic papers, industry articles, expert quotations, market data, and other authoritative sources.

    Base your research on the following core requirement: {overview}.

    In addition, factor in the following contextual details where applicable:

    Geographic Market: Consider the business and technology landscape in {country}. Ignore if not specified.

    Industry Focus: Include insights, trends, and data from the following industries: {industries}. Ignore if not specified.

    Business Function: Tailor the analysis to the perspective or needs of a person working in {function}. Ignore if not specified.

    Strategic Frameworks: Incorporate or structure your research using the following analytical frameworks: {frameworks}.

    Use the following format as a guide for structuring your response: {format}.

    Your response should:

    Include direct citations or links where available.

    Be clear, logically organized, and easy to turn into a pitch or slide deck.

    Blend both technical insight (e.g., emerging technologies, R&D frontiers) and business relevance (e.g., market sizing, customer pain points, competitive dynamics).
    """

    print(prompt)

    client = openai.OpenAI()

    try:
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a research assistant."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=2000,
            temperature=0.6
        )

        answer = response.choices[0].message.content
        return jsonify({"requirements": answer})
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/simple_search', methods=['POST'])
@cross_origin()
@require_auth
def simple_search():
    """Enhanced search endpoint that handles both regular searches and special commands"""
    try:
        data = request.get_json()
        user_query = data.get('query', '').strip()
        json_data = data.get('data', [])
        user_id = data.get('user_id', 'default_user')  # For favorites
        
        if not user_query:
            return jsonify({'success': False, 'error': 'No search query provided'}), 400
        
        # Parse the query using enhanced function
        parse_result = parse_simple_query_enhanced(user_query)
        
        if not parse_result['success']:
            return jsonify({
                'success': False,
                'error': 'Could not parse search query',
                'query': user_query,
                'openai_error': parse_result.get('error', 'Unknown error')
            }), 400
        
        # Handle special commands
        if parse_result.get('special_command', False):
            command_type = parse_result.get('command_type')
            
            if command_type == 'show_companies':
                if not json_data:
                    return jsonify({'success': False, 'error': 'No data available to list companies'}), 400
                
                # Extract unique companies from data
                companies = set()
                for record in json_data:
                    company = record.get('company') or record.get('Company') or record.get('organization')
                    if company and str(company).strip():
                        companies.add(str(company).strip())
                
                company_list = sorted(list(companies))
                return jsonify({
                    "success": True,
                    "type": "companies_list",
                    "total_found": len(company_list),
                    "results": [{"company": comp} for comp in company_list],
                    "query": user_query,
                    "keywords": parse_result['keywords'],
                    "command_type": command_type
                })
            
            elif command_type == 'show_titles':
                if not json_data:
                    return jsonify({'success': False, 'error': 'No data available to list titles'}), 400
                
                # Extract unique titles from data
                titles = set()
                for record in json_data:
                    title = (record.get('title') or record.get('Title') or 
                            record.get('Job title') or record.get('position'))
                    if title and str(title).strip():
                        titles.add(str(title).strip())
                
                title_list = sorted(list(titles))
                return jsonify({
                    "success": True,
                    "type": "titles_list",
                    "total_found": len(title_list),
                    "results": [{"title": title} for title in title_list],
                    "query": user_query,
                    "keywords": parse_result['keywords'],
                    "command_type": command_type
                })
            
            elif command_type == 'show_locations':
                if not json_data:
                    return jsonify({'success': False, 'error': 'No data available to list locations'}), 400
                
                # Extract unique locations from data
                locations = set()
                for record in json_data:
                    location = (record.get('location') or record.get('Location') or 
                               record.get('city') or record.get('City'))
                    if location and str(location).strip():
                        locations.add(str(location).strip())
                
                location_list = sorted(list(locations))
                return jsonify({
                    "success": True,
                    "type": "locations_list",
                    "total_found": len(location_list),
                    "results": [{"location": loc} for loc in location_list],
                    "query": user_query,
                    "keywords": parse_result['keywords'],
                    "command_type": command_type
                })
            
            elif command_type == 'show_skills':
                if not json_data:
                    return jsonify({'success': False, 'error': 'No data available to list skills'}), 400
                
                # Extract unique skills from data
                skills = set()
                for record in json_data:
                    skill_fields = (record.get('skills') or record.get('Skills') or 
                                   record.get('technologies') or record.get('required_skills'))
                    if skill_fields:
                        if isinstance(skill_fields, str):
                            # Split by common delimiters
                            import re
                            skill_list = re.split(r'[,;|]+', skill_fields)
                            for skill in skill_list:
                                clean_skill = skill.strip()
                                if clean_skill:
                                    skills.add(clean_skill)
                        elif isinstance(skill_fields, list):
                            skills.update([str(s).strip() for s in skill_fields if str(s).strip()])
                
                skill_list = sorted(list(skills))
                return jsonify({
                    "success": True,
                    "type": "skills_list",
                    "total_found": len(skill_list),
                    "results": [{"skill": skill} for skill in skill_list],
                    "query": user_query,
                    "keywords": parse_result['keywords'],
                    "command_type": command_type
                })
            
            elif command_type == 'show_favorites':
                # Load user favorites (doesn't need json_data)
                favorites_file = os.path.join(
                    DATA_DIR, 'user_data', user_id, 'favorites.json'
                )
                
                if os.path.exists(favorites_file):
                    with open(favorites_file, 'r', encoding='utf-8') as f:
                        favorites = json.load(f)
                    
                    return jsonify({
                        "success": True,
                        "type": "favorites_list",
                        "total_found": len(favorites),
                        "results": favorites,
                        "query": user_query,
                        "keywords": parse_result['keywords'],
                        "command_type": command_type
                    })
                else:
                    return jsonify({
                        "success": True,
                        "type": "favorites_list",
                        "total_found": 0,
                        "results": [],
                        "query": user_query,
                        "keywords": parse_result['keywords'],
                        "command_type": command_type,
                        "message": "No favorites saved yet"
                    })
            
            else:
                return jsonify({
                    'success': False,
                    'error': f'Unknown special command: {command_type}'
                }), 400
        
        # Regular search logic (existing)
        if not json_data:
            return jsonify({'success': False, 'error': 'No data provided for search'}), 400
        
        if not parse_result['keywords']:
            return jsonify({
                'success': False,
                'error': 'Could not extract search keywords from query',
                'query': user_query,
                'suggestion': 'Try being more specific with company names, job titles, or skills'
            }), 400
        
        # Search the data using existing function
        results = simple_search_json(json_data, parse_result['keywords'])
        
        return jsonify({
            'success': True,
            'type': 'search_results',
            'query': user_query,
            'keywords': parse_result['keywords'],
            'results': results,
            'total_found': len(results),
            'phrases': parse_result.get('phrases', [])
        })
        
    except Exception as e:
        print(f"Search error: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Search error: {str(e)}'
        }), 500


@app.route('/save_user_favorite', methods=['POST'])
@cross_origin()
@require_auth
def save_user_favorite():
    """Save a profile to user's favorites"""
    try:
        data = request.json
        user_id = data.get('user_id', 'default_user')
        profile_data = data.get('profile_data')
        
        if not profile_data:
            return jsonify({"success": False, "error": "No profile data provided"}), 400
        
        # Create user_data directory if it doesn't exist
        user_data_dir = os.path.join(DATA_DIR, 'user_data')
        os.makedirs(user_data_dir, exist_ok=True)
        
        # Create user-specific subdirectory
        user_dir = os.path.join(user_data_dir, user_id)
        os.makedirs(user_dir, exist_ok=True)
        
        # File path for user's favorites
        favorites_file = os.path.join(user_dir, 'favorites.json')
        
        # Load existing favorites or create new list
        favorites = []
        if os.path.exists(favorites_file):
            with open(favorites_file, 'r', encoding='utf-8') as f:
                favorites = json.load(f)
        
        # Add metadata to profile
        profile_with_meta = {
            **profile_data,
            'saved_at': datetime.now().isoformat(),
            'favorite_id': len(favorites) + 1
        }
        
        # Check if already exists (by name and company)
        full_name = f"{profile_data.get('name', '')} {profile_data.get('lastname', '')}".strip()
        existing = next((fav for fav in favorites 
                        if fav.get('full_name') == full_name and 
                           fav.get('company') == profile_data.get('company')), None)
        
        if existing:
            return jsonify({"success": False, "error": "Profile already in favorites"})
        
        # Add to favorites
        favorites.append(profile_with_meta)
        
        # Save updated favorites
        with open(favorites_file, 'w', encoding='utf-8') as f:
            json.dump(favorites, f, indent=2, ensure_ascii=False)
        
        return jsonify({
            "success": True,
            "message": "Profile saved to favorites",
            "favorites_count": len(favorites)
        })
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/get_user_favorites', methods=['POST'])
@cross_origin()
@require_auth
def get_user_favorites():
    """Get user's saved favorites"""
    try:
        data = request.json
        user_id = data.get('user_id', 'default_user')
        
        # Path to user's favorites file
        favorites_file = os.path.join(DATA_DIR, 'user_data', user_id, 'favorites.json')
        
        if not os.path.exists(favorites_file):
            return jsonify({"success": True, "favorites": [], "count": 0})
        
        # Load favorites
        with open(favorites_file, 'r', encoding='utf-8') as f:
            favorites = json.load(f)
        
        return jsonify({
            "success": True,
            "favorites": favorites,
            "count": len(favorites)
        })
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/remove_user_favorite', methods=['POST'])
@cross_origin()
@require_auth
def remove_user_favorite():
    """Remove a profile from user's favorites"""
    try:
        data = request.json
        user_id = data.get('user_id', 'default_user')
        favorite_id = data.get('favorite_id')
        
        if not favorite_id:
            return jsonify({"success": False, "error": "No favorite_id provided"}), 400
        
        # Path to user's favorites file
        favorites_file = os.path.join(DATA_DIR, 'user_data', user_id, 'favorites.json')
        
        if not os.path.exists(favorites_file):
            return jsonify({"success": False, "error": "No favorites file found"}), 404
        
        # Load favorites
        with open(favorites_file, 'r', encoding='utf-8') as f:
            favorites = json.load(f)
        
        # Remove the favorite
        original_count = len(favorites)
        favorites = [fav for fav in favorites if fav.get('favorite_id') != favorite_id]
        
        if len(favorites) == original_count:
            return jsonify({"success": False, "error": "Favorite not found"}), 404
        
        # Save updated favorites
        with open(favorites_file, 'w', encoding='utf-8') as f:
            json.dump(favorites, f, indent=2, ensure_ascii=False)
        
        return jsonify({
            "success": True,
            "message": "Profile removed from favorites",
            "favorites_count": len(favorites)
        })
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/save-requirements', methods=['POST'])
@require_auth
def save_requirements():
    data = request.get_json()
    requirements = data.get('requirements', [])
    export_option = data.get('exportOption', 'Unknown')  # Get the export option

    if not requirements:
        return jsonify({"error": "No requirements to save"}), 400

    # Define the folder path
    folder_path = os.path.join(DATA_DIR, 'requirements_versions')
    os.makedirs(folder_path, exist_ok=True)  # Create the folder if it doesn't exist

    # Create a unique file name with the export option and timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    file_path = os.path.join(folder_path, f"requirements_{export_option}_{timestamp}.txt")

    # Save the requirements to the file
    with open(file_path, "w") as file:
        file.write("\n".join(requirements))

    return jsonify({"message": f"Requirements saved successfully via {export_option}", "file_path": file_path})

@app.route('/register', methods=['POST'])
def register():
    data = request.get_json()
    username = data.get('email') # Use email as username
    password = data.get('password')
    first_name = data.get('first_name')
    last_name = data.get('last_name')
    email = data.get('email')
    company = data.get('company')
    linkedin = data.get('linkedin')
    short_intro = data.get('short_intro')
    company_intro = data.get('company_intro')

    if not password or not email:
        return jsonify({'error': 'Email and password required'}), 400
    if User.query.filter_by(username=username).first():
        return jsonify({'error': 'Username already exists'}), 400
    if User.query.filter_by(email=email).first():
        return jsonify({'error': 'Email already registered'}), 400

    try:
        hashed_password = generate_password_hash(password)
        user = User(
            username=username,
            password=hashed_password,
            first_name=first_name,
            last_name=last_name,
            email=email,
            company=company,
            linkedin=linkedin,
            short_intro=short_intro,
            company_intro=company_intro
        )
        db.session.add(user)
        db.session.commit()
        from core.session_token import issue_browser_session_token

        token = issue_browser_session_token(app.config['SECRET_KEY'], username)
        return jsonify({
            'message': 'User registered successfully',
            'session_token': token,
            'email': email,
            'username': username,
        }), 201
    except Exception as e:
        print("Registration error:", e)  # Add this line
        return jsonify({'error': str(e)}), 500

@app.route('/login', methods=['POST'])
def login():
    data = request.get_json()
    email = data.get('email')
    password = data.get('password')

    if not email or not password:
        return jsonify({'error': 'Email and password required'}), 400

    # DEV MODE: Allow login with password "dev123" for local testing
    if password == 'dev123' and os.environ.get('FLASK_ENV') != 'production':
        from core.session_token import issue_browser_session_token
        username = email.split('@')[0]
        token = issue_browser_session_token(app.config['SECRET_KEY'], username)
        return jsonify({
            'message': 'Login successful (dev mode)',
            'username': username,
            'email': email,
            'session_token': token,
        }), 200

    user = User.query.filter_by(email=email).first()
    if user and check_password_hash(user.password, password):
        # Login via email, but still return the username for frontend session storage if needed
        from core.session_token import issue_browser_session_token

        token = issue_browser_session_token(app.config['SECRET_KEY'], user.username)
        return jsonify({
            'message': 'Login successful',
            'username': user.username,
            'email': user.email,
            'session_token': token,
        }), 200
    else:
        return jsonify({'error': 'Invalid email or password'}), 401

GOOGLE_CLIENT_CONFIG = {
    "web": {
        "client_id": GOOGLE_CLIENT_ID,
        "project_id": "enable-agents",
        "auth_uri": "https://accounts.google.com/o/oauth2/auth",
        "token_uri": "https://oauth2.googleapis.com/token",
        "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
        "client_secret": GOOGLE_CLIENT_SECRET,
        "redirect_uris": [GOOGLE_REDIRECT_URI]
    }
}

SCOPES = [
    "openid",
    "https://www.googleapis.com/auth/userinfo.email",
    "https://www.googleapis.com/auth/userinfo.profile",
    "https://www.googleapis.com/auth/gmail.send"
]

@app.route('/auth/google/start', methods=['GET'])
def google_auth_start():
    # Use direct URL generation to avoid PKCE code_verifier state issues
    params = {
        'client_id': GOOGLE_CLIENT_ID,
        'redirect_uri': GOOGLE_REDIRECT_URI,
        'response_type': 'code',
        'scope': ' '.join(SCOPES),
        'access_type': 'offline',
        'prompt': 'consent',
        'state': 'user_login_flow'
    }
    auth_url = f"https://accounts.google.com/o/oauth2/v2/auth?{urlencode(params)}"
    # redirect_uri + client_id must exactly match one row in Google Cloud Console for this OAuth client
    out = {
        "auth_url": auth_url,
        "state": "user_login_flow",
        "redirect_uri_used": GOOGLE_REDIRECT_URI,
        "oauth_client_id": GOOGLE_CLIENT_ID,
    }
    return jsonify(out)


def send_platform_email(sender_user_id, to_email, subject, body):
    """
    Send a real email on behalf of sender_user_id: tries their connected
    Gmail account first (refreshing the token if needed), then falls back
    to system SMTP (EMAIL_USER/EMAIL_PASS) if Gmail isn't connected or
    fails. Returns (success: bool, error: str | None) - never raises, so
    callers can always show the user an accurate status instead of
    silently pretending an email went out.
    """
    token_record = GoogleOAuthToken.query.filter_by(username=sender_user_id).first()

    if token_record and token_record.token:
        try:
            creds = Credentials(
                token=token_record.token,
                refresh_token=token_record.refresh_token,
                token_uri=token_record.token_uri,
                client_id=token_record.client_id,
                client_secret=token_record.client_secret,
                scopes=token_record.scopes.split(',') if token_record.scopes else SCOPES,
            )
            if creds.refresh_token and (not creds.valid or creds.expired):
                creds.refresh(Request())
                token_record.token = creds.token
                db.session.commit()

            service = googleapiclient.discovery.build('gmail', 'v1', credentials=creds)
            message = EmailMessage()
            message.set_content(body)
            message['To'] = to_email
            message['From'] = sender_user_id
            message['Subject'] = subject
            encoded_message = base64.urlsafe_b64encode(message.as_bytes()).decode()
            service.users().messages().send(userId="me", body={'raw': encoded_message}).execute()
            return True, None
        except Exception as gmail_error:
            gmail_error_summary = str(gmail_error).split('.', 1)[0][:200]
            print(f"[send_platform_email] Gmail send failed for {sender_user_id}, trying SMTP: {gmail_error}")
    else:
        gmail_error_summary = None

    # SMTP fallback (or primary path if Gmail was never connected)
    email_host = os.getenv('EMAIL_HOST', 'smtp.gmail.com')
    email_port = int(os.getenv('EMAIL_PORT', 587))
    email_user = os.getenv('EMAIL_USER')
    email_pass = os.getenv('EMAIL_PASS')
    if not email_user or not email_pass:
        if gmail_error_summary:
            return False, f'Gmail send failed ({gmail_error_summary}) and SMTP fallback is unavailable.'
        return False, 'No email account connected. Connect Google in Settings, or ask an admin to configure SMTP.'

    try:
        smtp_server = smtplib.SMTP(email_host, email_port, timeout=15)
        smtp_server.starttls()
        smtp_server.login(email_user, email_pass)
        message = EmailMessage()
        message.set_content(body)
        message['To'] = to_email
        message['From'] = email_user
        message['Reply-To'] = sender_user_id
        message['Subject'] = subject
        smtp_server.send_message(message)
        smtp_server.quit()
        return True, None
    except Exception as smtp_error:
        return False, f'Email send failed: {str(smtp_error).split(".", 1)[0][:200]}'


@app.route('/emails/send_via_gmail', methods=['POST'])
@require_auth
def send_via_gmail():
    data = request.get_json()
    to_email = data.get('to')
    subject = data.get('subject')
    body = data.get('body')

    if not all([to_email, subject, body]):
        return jsonify({'error': 'Missing required fields'}), 400

    success, error = send_platform_email(g.user_id, to_email, subject, body)
    if not success:
        return jsonify({'error': error}), 500
    return jsonify({'message': 'Email sent successfully'})


@app.route('/file_to_json_convert', methods=['POST'])
@require_auth
def convert_file():
    """Main endpoint to convert CSV/XLSX files to JSON"""
    
    # Check if file is present in request
    if 'file' not in request.files:
        return jsonify({
            'success': False,
            'error': 'No file provided',
            'data': []
        }), 400
    
    file = request.files['file']
    
    # Check if file is selected
    if file.filename == '':
        return jsonify({
            'success': False,
            'error': 'No file selected',
            'data': []
        }), 400
    
    # Check if file type is allowed
    if not allowed_file(file.filename):
        return jsonify({
            'success': False,
            'error': 'File type not allowed. Please upload CSV or XLSX files only.',
            'data': []
        }), 400
    
    try:
        # Create temporary upload folder for conversion
        temp_folder = os.path.join(DATA_DIR, 'temp_conversion')
        os.makedirs(temp_folder, exist_ok=True)
        
        # Save uploaded file
        filename = secure_filename(file.filename)
        file_path = os.path.join(temp_folder, filename)
        file.save(file_path)
        
        # Get conversion options
        multiple_sheets = request.form.get('multiple_sheets', 'false').lower() == 'true'
        
        # Convert based on file type
        file_extension = filename.rsplit('.', 1)[1].lower()
        
        if file_extension == 'csv':
            result = csv_to_json(file_path)
        elif file_extension in ['xlsx', 'xls']:
            if multiple_sheets:
                result = xlsx_to_json_multiple_sheets(file_path)
            else:
                result = xlsx_to_json(file_path)
        
        # Clean up - remove uploaded file
        try:
            os.remove(file_path)
        except:
            pass
        
        return jsonify(result)
        
    except Exception as e:
        # Clean up file if error occurs
        try:
            if 'file_path' in locals():
                os.remove(file_path)
        except:
            pass
        
        return jsonify({
            'success': False,
            'error': f'Server error: {str(e)}',
            'data': []
        }), 500
    

@app.route('/enrich_with_openai', methods=['POST'])
@cross_origin()
@require_auth
def enrich_with_openai():
    """API endpoint to enrich JSON data with required skills using OpenAI"""
    try:
        request_data = request.get_json()
        
        if not request_data or 'data' not in request_data:
            return jsonify({
                'success': False,
                'error': 'No data provided in request body'
            }), 400
        
        json_data = request_data['data']
        
        if not isinstance(json_data, list) or len(json_data) == 0:
            return jsonify({
                'success': False,
                'error': 'Data must be a non-empty list of objects'
            }), 400
        
        # Enrich data using the new workflow
        result = enrich_json_with_openai(json_data)
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Server error: {str(e)}'
        }), 500
    

@app.route('/chrome_history', methods=['GET'])
@require_auth
def get_chrome_history():
    """API endpoint to get Chrome browser history with better error handling"""
    try:
        user_id = request.args.get('user_id', 'default_user')
        result = read_chrome_history_safe()
        
        if result['success']:
            save_tools_landscape_for_user(user_id, result)
            return jsonify(result)
        else:
            return jsonify(result), 400
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Server error: {str(e)}'
        }), 500

@app.route('/chrome_status', methods=['GET'])
@require_auth
def check_chrome_status():
    """Check if Chrome is running"""
    try:
        
        chrome_processes = []
        for proc in psutil.process_iter(['pid', 'name']):
            try:
                if 'chrome' in proc.info['name'].lower():
                    chrome_processes.append({
                        'pid': proc.info['pid'],
                        'name': proc.info['name']
                    })
            except (psutil.NoSuchProcess, psutil.AccessDenied):
                pass
        
        return jsonify({
            'success': True,
            'chrome_running': len(chrome_processes) > 0,
            'processes': chrome_processes
        })
        
    except ImportError:
        return jsonify({
            'success': False,
            'error': 'psutil not installed. Cannot check Chrome status.'
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        })
    
@app.route('/check_existing_file', methods=['POST'])
@require_auth
def check_existing_file():
    try:
        data = request.json
        file_name = data.get('file_name')
        new_file_size = data.get('new_file_size')
        
        # Create file path - Updated to use the correct data folder structure
        json_file_name = file_name.replace('.csv', '.json').replace('.xlsx', '.json').replace('.xls', '.json')
        # Use the same structure as upload function
        file_path = os.path.join(DATA_DIR, 'uploaded_data', 'alumni_data', json_file_name)
        
        if os.path.exists(file_path):
            existing_size = os.path.getsize(file_path)
            
            # If new file is not significantly larger (less than 10% increase), skip processing
            size_threshold = existing_size * 1.1  # 10% increase threshold
            should_skip = new_file_size <= size_threshold
            
            return jsonify({
                'exists': True,
                'existing_size': existing_size,
                'should_skip': should_skip,
                'message': f'File exists. Size: {existing_size} bytes vs new: {new_file_size} bytes'
            })
        
        return jsonify({
            'exists': False,
            'should_skip': False,
            'message': 'File does not exist'
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/save_json_file', methods=['POST'])
@require_auth
def save_json_file():
    try:
        data = request.json
        json_data = data.get('data')
        file_name = data.get('file_name')
        folder_name = data.get('folder_name', 'alumni_data')
        
        # Create directory using the correct data folder structure
        folder_path = os.path.join(DATA_DIR, 'uploaded_data', folder_name)
        os.makedirs(folder_path, exist_ok=True)
        
        # Save JSON file
        file_path = os.path.join(folder_path, file_name)
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, indent=2, ensure_ascii=False)
        
        return jsonify({
            'success': True,
            'file_path': file_path,
            'file_name': file_name,
            'message': f'JSON file saved successfully: {file_name}'
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/load_json_file', methods=['POST'])
@require_auth
def load_json_file():
    try:
        data = request.json
        file_name = data.get('file_name')
        folder_name = data.get('folder_name', 'alumni_data')
        
        # Use the correct data folder structure
        file_path = os.path.join(DATA_DIR, 'uploaded_data', folder_name, file_name)
        
        if not os.path.exists(file_path):
            return jsonify({'success': False, 'error': 'File not found'}), 404
        
        with open(file_path, 'r', encoding='utf-8') as f:
            json_data = json.load(f)
        
        return jsonify({
            'success': True,
            'data': json_data,
            'message': f'JSON file loaded successfully: {file_name}'
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

def read_chrome_history_safe():
    """
    Improved: Read Chrome browser history and return only unique domains
    classified as tools where the user has logged in or has a subscription.
    """
    try:
        history_path = get_chrome_history_path()
        if not os.path.exists(history_path):
            return {
                'success': False,
                'error': 'Chrome history file not found. Make sure Chrome is installed.'
            }

        now = datetime.now()
        seven_days_ago = now - timedelta(days=7)
        webkit_epoch = datetime(1601, 1, 1)
        seven_days_ago_webkit = int((seven_days_ago - webkit_epoch).total_seconds() * 1000000)

        # Copy Chrome history file to temp location
        temp_dir = tempfile.mkdtemp()
        temp_history = os.path.join(temp_dir, 'History')
        shutil.copy2(history_path, temp_history)

        conn = sqlite3.connect(temp_history)
        cursor = conn.cursor()
        query = """
        SELECT url, title, visit_count, last_visit_time,
            datetime(last_visit_time/1000000 + (strftime('%s', '1601-01-01')), 'unixepoch', 'localtime') as visit_date
        FROM urls
        WHERE last_visit_time >= ?
        ORDER BY last_visit_time DESC
        LIMIT 2000
        """
        cursor.execute(query, (seven_days_ago_webkit,))
        rows = cursor.fetchall()
        conn.close()
        os.remove(temp_history)
        os.rmdir(temp_dir)

        # Filter URLs for login/subscription/dashboard/account/profile/settings
        login_keywords = [
            '/login', '/signin', '/dashboard', '/account', '/settings', '/profile', '/subscription', '/user', '/me'
        ]
        domain_map = {}
        filtered_history = []
        for row in rows:
            url = row[0]
            parsed = urlparse(url)
            domain = parsed.netloc.lower()
            path = parsed.path.lower()
            # Only consider URLs with login/subscription/dashboard/account/profile/settings or base domain
            if any(kw in path for kw in login_keywords) or path in ['', '/']:
                if domain not in domain_map:
                    domain_map[domain] = {
                        'domain': domain,
                        'sample_url': url,
                        'title': row[1] if row[1] else 'No Title',
                        'visit_count': row[2],
                        'last_visit_time': row[3],
                        'visit_date': row[4]
                    }

        # Classify domains using OpenAI (reuse your identify_saas_tools_with_openai)
        unique_domains = list(domain_map.values())
        # Prepare for OpenAI classification
        history_data_for_ai = [{'url': item['sample_url']} for item in unique_domains]
        tools_result = identify_saas_tools_with_openai(history_data_for_ai)
        for i, item in enumerate(unique_domains):
            if tools_result['success'] and i in tools_result['mapping']:
                mapping = tools_result['mapping'][i]
                item.update({
                    'is_tool': mapping.get('is_tool', False),
                    'tool_name': mapping.get('tool_name'),
                    'category': mapping.get('category'),
                    'tool_type': mapping.get('type'),
                    'description': mapping.get('description'),
                })
            else:
                item.update({
                    'is_tool': None,
                    'tool_name': None,
                    'category': None,
                    'tool_type': None,
                    'description': 'Tool analysis unavailable'
                })

        # Only return classified tools (where is_tool is True)
        classified_tools = [item for item in unique_domains if item.get('is_tool')]

        return {
            'success': True,
            'unique_domains': len(unique_domains),
            'tools_found': len(classified_tools),
            'data': classified_tools,
            'date_range': {
                'from': seven_days_ago.strftime('%Y-%m-%d'),
                'to': now.strftime('%Y-%m-%d')
            }
        }
    except Exception as e:
        return {
            'success': False,
            'error': f'Unexpected error reading Chrome history: {str(e)}'
        }

def save_tools_landscape_for_user(user_id, tools_data):
    """
    Save the result of /chrome_history API into a JSON file called
    'SaaS & tools landscape.json' under user_data for the given user.
    Adds a timestamp for when the data was last updated.
    """
    try:
        user_folder = os.path.join(DATA_DIR, 'user_data', 'tools_landscape')
        os.makedirs(user_folder, exist_ok=True)
        file_path = os.path.join(user_folder, 'tools_landscape.json')

        # Add/update timestamp
        tools_data['last_updated'] = datetime.now().isoformat()
        tools_data['user'] = user_id

        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(tools_data, f, ensure_ascii=False, indent=2)

        return {
            'success': True,
            'message': f'Tools landscape saved for {user_id}',
            'file_path': file_path,
            'last_updated': tools_data['last_updated']
        }
    except Exception as e:
        return {'success': False, 'error': str(e)}

@app.route('/get_tools_landscape', methods=['GET'])
@cross_origin()
@require_auth
def get_tools_landscape():
    """
    GET API to read tools landscape from tools_landscape.json and return
    a list of tools with tool_name, description, and category.
    """
    try:
        file_path = os.path.join(DATA_DIR, 'user_data', 'tools_landscape', 'tools_landscape.json')
        if not os.path.exists(file_path):
            return jsonify({'success': False, 'error': 'tools_landscape.json not found'}), 404

        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Extract tools info
        tools = []
        for item in data.get('data', []):
            if item.get('tool_name'):
                tools.append({
                    'tool_name': item.get('tool_name'),
                    'description': item.get('description'),
                    'category': item.get('category')
                })

        return jsonify({'success': True, 'tools': tools})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500
    
import openai


@app.route('/recommend_agents', methods=['POST'])
@cross_origin()
@require_auth
def recommend_agents():
    """
    Recommend agents/modules based on user's tools, industry/domain, role, department/company context, and business need.
    Input JSON payload:
    {
        "tools": [ {"tool_name": "ToolA", "description": "..."}, ... ],
        "industry": "...",
        "product_service": "...",
        "role": "...",
        "department_context": "...",  # e.g. 'My company/department is doing X and is responsible for Y'
        "business_need": "..."         # e.g. 'I want to track this business task and generate insights'
    }
    Output JSON:
    {
        "success": true,
        "recommendations": {
            "recommended_tools": [ {"tool_name": "...", "description": "...", "why_recommended": "..."}, ... ],
            "integration_pairs": [ {"tools": ["ToolA", "ToolB"], "integration": "...", "data_shared": "..."}, ... ],
            "additional_tools": [ {"tool_name": "...", "description": "...", "why_needed": "..."}, ... ]
        }
    }
    """
    try:
        openai.api_key = get_credentials()
        data = request.json
        tools = data.get('tools', [])
        industry = data.get('industry', '')
        product_service = data.get('product_service', '')
        role = data.get('role', '')
        department_context = data.get('department_context', '')
        business_need = data.get('business_need', '')

        # Load available modules (from agents_modules.json). If missing, use a safe fallback catalog.
        modules_file = os.path.join(DATA_DIR, 'agents_modules.json')
        if os.path.exists(modules_file):
            with open(modules_file, 'r', encoding='utf-8') as f:
                modules = json.load(f)
        else:
            modules = [
                {"name": "Market Research", "description": "Market analysis, competitor research, and customer insights."},
                {"name": "Sales Helper Agent", "description": "Lead management, sales enablement, and CRM support."},
                {"name": "Content Marketing Agent", "description": "Content strategy and campaign execution support."},
                {"name": "Executive Assistant Agent", "description": "Task coordination, reminders, and stakeholder updates."},
                {"name": "Supply Chain Agent", "description": "Supply chain monitoring and impact analysis."},
                {"name": "Data Discovery", "description": "Data exploration and business insight generation."},
                {"name": "AI Chatbot", "description": "Conversational workflow automation and user support."},
                {"name": "Dashboards", "description": "KPI dashboards, reporting, and decision support."},
                {"name": "Integration", "description": "Connect tools and automate cross-system data flow."},
                {"name": "Automation", "description": "Automate repetitive workflows and approvals."}
            ]

        # Prepare context for OpenAI
        context = {
            "tools": tools,
            "industry": industry,
            "product_service": product_service,
            "role": role,
            "department_context": department_context,
            "business_need": business_need,
            "available_modules": modules
        }
        prompt = (
            "You are a technology consultant. Based on the following user context, "
            "recommend a set of software modules (tools/agents) that can cater to the business need, "
            "considering the existing tools, missing necessary tools, and possible integrations. "
            "For each recommendation, provide: "
            "1. recommended_tools: list of modules/tools with name, description, and why recommended. "
            "2. integration_pairs: pairs of tools/modules that should be integrated, with integration description and data shared. "
            "3. additional_tools: tools/modules that are needed but missing, with name, description, names of companies offering it and why needed. "
            "Return the output as a JSON object with a 'recommendations' key containing these three lists. "
            "Here is the user context and available modules:\n\n" + json.dumps(context, indent=2)
        )
        client = openai.OpenAI()

        response = response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a technology consultant for business software and workflow automation."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1000,
            temperature=0.3
        )

        # Try to parse the response as JSON
        import ast
        import re
        raw_content = response.choices[0].message.content
        # Extract JSON from response (in case LLM returns extra text)
        match = re.search(r'\{[\s\S]*\}', raw_content)
        if match:
            recommendations_json = match.group(0)
            try:
                recommendations = json.loads(recommendations_json)
            except Exception:
                recommendations = {"raw": raw_content}
        else:
            recommendations = {"raw": raw_content}

        print(recommendations)

        return jsonify({
            "success": True,
            "recommendations": recommendations.get('recommendations', recommendations)
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

# @app.route('/AI_ML', methods=['GET'])
# def yfinance_test():
    
#     return "Hello World!"


# @app.route('/Location', methods=['GET'])
# def yfinance_test():
    
#     return "Hello World!"

# @app.route('/Transportation', methods=['GET'])
# def yfinance_test():
    
#     return "Hello World!"

# @app.route('/Business- Enterprise', methods=['GET'])
# def yfinance_test():
    
#     return "Hello World!"

# @app.route('/Visual Recognition', methods=['GET'])
# def yfinance_test():
    
#     return "Hello World!"

# @app.route('/Small Tools', methods=['GET'])
# def yfinance_test():
    
#     return "Hello World!"

# @app.route('/Text Analysis', methods=['GET'])
# def yfinance_test():
    
#     return "Hello World!"

# @app.route('/Weather', methods=['GET'])
# def yfinance_test():
    
#     return "Hello World!"

# @app.route('/Messaging', methods=['GET'])
# def yfinance_test():
    
#     return "Hello World!"

# @app.route('/Logistics', methods=['GET'])
# def yfinance_test():
    
#     return "Hello World!"

# @app.route('/News', methods=['GET'])
# def yfinance_test():
    
#     return "Hello World!"

# @app.route('/Jobs', methods=['GET'])
# def yfinance_test():
    
#     return "Hello World!"

# @app.route('/yfinance', methods=['GET'])
# def yfinance_test():
#     symbol = request.args.get('stock')
#     region = request.args.get('region')

#     if not symbol or not region:
#         return "Missing required parameters: 'stock' and 'region'", 400

#     conn = http.client.HTTPSConnection("yahoo-finance166.p.rapidapi.com")

#     headers = {
#         'x-rapidapi-key': "95cdd43379mshbd9483856442c47p1c2782jsn897449ebefb8",
#         'x-rapidapi-host': "yahoo-finance166.p.rapidapi.com"
#     }

#     endpoint = f"/api/stock/get-financial-data?region={region}&symbol={symbol}"
#     print(f"Requesting data from endpoint: {endpoint}")  # Debug statement
#     conn.request("GET", endpoint, headers=headers)

#     res = conn.getresponse()
#     data = res.read()
#     json_data = json.loads(data.decode("utf-8"))

#     print(json_data)  # Debug statement to print the entire response

#     if 'quoteSummary' not in json_data or 'result' not in json_data['quoteSummary'] or not json_data['quoteSummary']['result']:
#         return jsonify({"error": "No data found for the given stock symbol and region"}), 404

#     current_price = json_data['quoteSummary']['result'][0]['financialData']['currentPrice']['fmt']
#     operating_margins = json_data['quoteSummary']['result'][0]['financialData']['operatingMargins']['fmt']
#     netprofit_margins = json_data['quoteSummary']['result'][0]['financialData']['profitMargins']['fmt']
#     gross_margins = json_data['quoteSummary']['result'][0]['financialData']['grossMargins']['fmt']
#     revenue_growth = json_data['quoteSummary']['result'][0]['financialData']['revenueGrowth']['fmt']
#     debt_to_equity = json_data['quoteSummary']['result'][0]['financialData']['debtToEquity']['fmt']
#     quick_ratio = json_data['quoteSummary']['result'][0]['financialData']['quickRatio']['fmt']
#     current_ratio = json_data['quoteSummary']['result'][0]['financialData']['currentRatio']['fmt']
#     analyst_recommendation = json_data['quoteSummary']['result'][0]['financialData']['recommendationKey']
#     number_of_analysts = json_data['quoteSummary']['result'][0]['financialData']['numberOfAnalystOpinions']['fmt']
#     target_high_price = json_data['quoteSummary']['result'][0]['financialData']['targetHighPrice']['fmt']
#     target_low_price = json_data['quoteSummary']['result'][0]['financialData']['targetLowPrice']['fmt']
#     target_mean_price = json_data['quoteSummary']['result'][0]['financialData']['targetMeanPrice']['fmt']
#     target_median_price = json_data['quoteSummary']['result'][0]['financialData']['targetMedianPrice']['fmt']

#     financial_KPIs = {
#         "current_price": current_price,
#         "operating margin": operating_margins,
#         "netprofit_margins": netprofit_margins,
#         "gross_margins": gross_margins,
#         "revenue_growth": revenue_growth,
#         "debt_to_equity": debt_to_equity,
#         "quick_ratio": quick_ratio,
#         "current_ratio": current_ratio,
#         "number_of_analysts": number_of_analysts,
#         "analyst_recommendation": analyst_recommendation,
#         "target_high_price": target_high_price,
#         "target_low_price": target_low_price,
#         "target_mean_price": target_mean_price,
#         "target_median_price": target_median_price
#     }

#     return jsonify(financial_KPIs)


# === KNOWLEDGE GRAPH + RAG API ===

import boto3
from docx import Document as DocxDocument
import networkx as nx

# Initialize S3 client for AWS operations
s3_client = boto3.client('s3')

def generate_cache_key(data):
    """Generate a hash key for caching based on input data"""
    data_string = json.dumps(data, sort_keys=True)
    return hashlib.md5(data_string.encode()).hexdigest()

def get_document_cache_key(documents):
    """Generate cache key for documents list"""
    doc_keys = []
    for doc in documents:
        key = f"{doc['source_type']}:{doc['path']}"
        if 'bucket' in doc:
            key += f":{doc['bucket']}"
        doc_keys.append(key)
    return generate_cache_key(sorted(doc_keys))

def get_kg_cache_key(nodes, edges):
    """Generate cache key for knowledge graph"""
    kg_data = {'nodes': nodes, 'edges': edges}
    return generate_cache_key(kg_data)

def load_document_from_source(source_type, source_path, bucket_name=None):
    """Load PDF or Word document from S3 or local machine"""
    if source_type == "s3":
        local_path = f"/tmp/{os.path.basename(source_path)}"
        s3_client.download_file(bucket_name, source_path, local_path)
        return local_path
    return source_path

def extract_text_from_pdf(file_path):
    """Extract text content from PDF file using PyMuPDF"""
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def extract_text_from_word(file_path):
    """Extract text content from Word document"""
    doc = DocxDocument(file_path)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return text

def extract_text_from_document(file_path):
    """Route extraction based on file extension"""
    if file_path.lower().endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_path.lower().endswith(('.docx', '.doc')):
        return extract_text_from_word(file_path)
    return ""

def build_knowledge_graph(nodes, edges):
    """Create a NetworkX graph from nodes and edges JSON input"""
    G = nx.DiGraph()
    for node in nodes:
        G.add_node(node['id'], **node.get('attributes', {}))
    for edge in edges:
        G.add_edge(edge['source'], edge['target'], **edge.get('attributes', {}))
    return G

def chunk_text(text, chunk_size=1000, overlap=200):
    """Split text into overlapping chunks for better context preservation"""
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start = end - overlap
    return chunks

def create_embeddings(chunks):
    """Generate OpenAI embeddings for text chunks"""
    embeddings_model = OpenAIEmbeddings()
    embeddings = embeddings_model.embed_documents(chunks)
    return np.array(embeddings)

def build_faiss_index(embeddings):
    """Create FAISS index for efficient similarity search"""
    dimension = embeddings.shape[1]
    index = faiss.IndexFlatL2(dimension)
    index.add(embeddings.astype('float32'))
    return index

def retrieve_relevant_chunks(query, index, chunks, embeddings_model, top_k=5):
    """Retrieve most relevant chunks using FAISS similarity search"""
    query_embedding = embeddings_model.embed_query(query)
    query_vector = np.array([query_embedding]).astype('float32')
    distances, indices = index.search(query_vector, top_k)
    return [chunks[i] for i in indices[0]]

def query_knowledge_graph(graph, query_type, node_id=None):
    """Query knowledge graph for specific information based on query type"""
    if query_type == "neighbors" and node_id:
        return list(graph.neighbors(node_id))
    elif query_type == "attributes" and node_id:
        return dict(graph.nodes[node_id])
    elif query_type == "all_nodes":
        return list(graph.nodes(data=True))
    elif query_type == "all_edges":
        return list(graph.edges(data=True))
    return None

def generate_answer_with_rag(query, relevant_chunks, kg_context):
    """Generate final answer using OpenAI with RAG context and KG information"""
    llm = ChatOpenAI(model="gpt-4", temperature=0)
    context_text = "\n\n".join(relevant_chunks)
    kg_text = json.dumps(kg_context, indent=2)
    
    prompt = f"""Based on the following document context and knowledge graph information, answer the query.

Document Context:
{context_text}

Knowledge Graph Context:
{kg_text}

Query: {query}

Provide a detailed answer:"""
    
    response = llm.invoke(prompt).content
    return response

def process_documents_with_kg_rag(documents, nodes, edges, query, include_context=False):
    """Main processing pipeline combining document loading, KG building, and RAG with caching"""
    # Generate cache keys
    doc_cache_key = get_document_cache_key(documents)
    kg_cache_key = get_kg_cache_key(nodes, edges)
    
    # Check if knowledge graph is cached
    if kg_cache_key in kg_rag_cache['knowledge_graphs']:
        kg = kg_rag_cache['knowledge_graphs'][kg_cache_key]
    else:
        # Build knowledge graph and cache it
        kg = build_knowledge_graph(nodes, edges)
        kg_rag_cache['knowledge_graphs'][kg_cache_key] = kg
    
    # Check if document embeddings are cached
    if doc_cache_key in kg_rag_cache['embeddings']:
        # Reuse cached data
        chunks = kg_rag_cache['chunks'][doc_cache_key]
        embeddings = kg_rag_cache['embeddings'][doc_cache_key]
        faiss_index = kg_rag_cache['faiss_indices'][doc_cache_key]
        embeddings_model = OpenAIEmbeddings()
    else:
        # Extract and combine text from all documents
        all_text = ""
        for doc_info in documents:
            local_path = load_document_from_source(
                doc_info['source_type'], 
                doc_info['path'], 
                doc_info.get('bucket')
            )
            text = extract_text_from_document(local_path)
            all_text += text + "\n\n"
        
        # Create chunks and embeddings
        chunks = chunk_text(all_text)
        embeddings_model = OpenAIEmbeddings()
        embeddings = create_embeddings(chunks)
        
        # Build FAISS index
        faiss_index = build_faiss_index(embeddings)
        
        # Cache all the expensive computations
        kg_rag_cache['chunks'][doc_cache_key] = chunks
        kg_rag_cache['embeddings'][doc_cache_key] = embeddings
        kg_rag_cache['faiss_indices'][doc_cache_key] = faiss_index
    
    # Retrieve relevant chunks (this is query-specific, not cached)
    relevant_chunks = retrieve_relevant_chunks(query, faiss_index, chunks, embeddings_model)
    
    # Query knowledge graph for additional context
    kg_context = {
        'nodes': query_knowledge_graph(kg, "all_nodes"),
        'edges': query_knowledge_graph(kg, "all_edges")
    }
    
    # Generate answer
    answer = generate_answer_with_rag(query, relevant_chunks, kg_context)
    
    # Return only answer by default (lightweight response)
    if include_context:
        return {
            'answer': answer,
            'relevant_chunks': relevant_chunks,
            'kg_context': kg_context
        }
    else:
        return {
            'answer': answer
        }

@app.route('/extract-with-kg-rag', methods=['POST'])
@cross_origin()
@require_auth
def extract_with_kg_rag():
    """API endpoint to extract information from documents using Knowledge Graph and RAG"""
    try:
        data = request.json
        
        # Validate input
        documents = data.get('documents', [])
        nodes = data.get('nodes', [])
        edges = data.get('edges', [])
        query = data.get('query', '')
        include_context = data.get('include_context', False)  # Optional: return chunks and KG context
        
        if not documents or not query:
            return jsonify({
                'success': False,
                'error': 'documents and query are required'
            }), 400
        
        # Process documents with KG and RAG
        result = process_documents_with_kg_rag(documents, nodes, edges, query, include_context)
        
        return jsonify({
            'success': True,
            'data': result
        }), 200
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/clear-kg-rag-cache', methods=['POST'])
@cross_origin()
@require_auth
def clear_kg_rag_cache():
    """Clear the KG+RAG cache to free up memory"""
    try:
        data = request.json or {}
        cache_type = data.get('cache_type', 'all')  # 'all', 'embeddings', 'graphs'
        
        if cache_type == 'all':
            kg_rag_cache['embeddings'].clear()
            kg_rag_cache['faiss_indices'].clear()
            kg_rag_cache['chunks'].clear()
            kg_rag_cache['knowledge_graphs'].clear()
            cleared = 'all caches'
        elif cache_type == 'embeddings':
            kg_rag_cache['embeddings'].clear()
            kg_rag_cache['faiss_indices'].clear()
            kg_rag_cache['chunks'].clear()
            cleared = 'embeddings cache'
        elif cache_type == 'graphs':
            kg_rag_cache['knowledge_graphs'].clear()
            cleared = 'knowledge graphs cache'
        else:
            return jsonify({
                'success': False,
                'error': 'Invalid cache_type. Use: all, embeddings, or graphs'
            }), 400
        
        return jsonify({
            'success': True,
            'message': f'Successfully cleared {cleared}'
        }), 200
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/kg-rag-cache-status', methods=['GET'])
@cross_origin()
@require_auth
def kg_rag_cache_status():
    """Get current cache status and statistics"""
    try:
        status = {
            'embeddings_cached': len(kg_rag_cache['embeddings']),
            'faiss_indices_cached': len(kg_rag_cache['faiss_indices']),
            'chunks_cached': len(kg_rag_cache['chunks']),
            'knowledge_graphs_cached': len(kg_rag_cache['knowledge_graphs']),
            'total_cached_items': (
                len(kg_rag_cache['embeddings']) + 
                len(kg_rag_cache['knowledge_graphs'])
            )
        }
        
        return jsonify({
            'success': True,
            'cache_status': status
        }), 200
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

# ====== ENTITY MANAGEMENT SYSTEM ======
# Domain-agnostic persistent knowledge graph + vector DB system
# Works for ANY domain: HR, Sales, Research, Healthcare, Legal, etc.
# Note: Commented out pending completion of entity_system module

# from entity_system import (
#     EntityKnowledgeGraphManager,
#     EntityVectorStoreManager,
#     process_entity_documents,
#     query_entity_profile,
#     chroma_client
# )
#
# # Initialize entity system managers
# kg_manager = EntityKnowledgeGraphManager(db)
# vector_manager = EntityVectorStoreManager(chroma_client)

# === ENTITY ENDPOINTS (Disabled - Pending entity_system module) ===
# 
# === ENTITY ENDPOINTS (Disabled - Pending entity_system module) ===
# These endpoints require the entity_system module which is not yet implemented
#
# All entity-related endpoints are disabled including:
# - POST /entity/upload
# - POST /entity/query
# - GET /entity/<entity_id>
# - DELETE /entity/<entity_id>
# - GET /entities/list
# - GET /system/health
#
# This module will be enabled once entity_system.py is properly implemented
# with EntityKnowledgeGraphManager, EntityVectorStoreManager, and related utilities
# ======


# === ENTITY ENDPOINTS (Disabled - Pending entity_system module) ===
# These endpoints require the entity_system module which is not yet implemented
#
# All entity-related endpoints disabled including:
# - entity_upload
# - entity_query
# - get_entity
# - delete_entity
# - list_entities
# - system_health
#
# ======


# ====== CONTENT MARKETING AGENT API ENDPOINTS ======
# All SQLite code migrated to PostgreSQL via SQLAlchemy
# See: agents/content_marketing/service.py and agents/content_marketing/models.py

from agents.content_marketing import service as cm_service
from agents.content_marketing.models import CMProject, CMDocument, CMKnowledgeGraph, CMGeneratedContent, CMConversation

@app.route('/api/content-marketing/projects', methods=['POST'])
@cross_origin()
@require_auth
def create_content_marketing_project():
    """Create a new content marketing project"""
    return cm_service.create_project()


@app.route('/api/content-marketing/projects/<project_id>', methods=['GET'])
@cross_origin()
@require_auth
def get_content_marketing_project(project_id):
    """Get project details"""
    return cm_service.get_project(project_id)


@app.route('/api/content-marketing/documents/upload', methods=['POST'])
@cross_origin()
@require_auth
def upload_content_marketing_documents():
    """Upload documents to project and build knowledge graph"""
    try:
        project_id = request.form.get('project_id')
        if not project_id:
            return jsonify({'success': False, 'error': 'project_id required'}), 400

        project = CMProject.query.filter_by(project_id=project_id).first()
        if not project or project.user_id != g.user_id:
            return jsonify({'success': False, 'error': 'Project not found'}), 404

        uploaded_files = request.files.getlist('files')
        if not uploaded_files:
            return jsonify({'success': False, 'error': 'No files provided'}), 400

        analyzer = DomainSpecializationAnalyzer()
        extracted_documents = []
        doc_ids = []

        for file in uploaded_files:
            if not file.filename:
                continue

            filename = secure_filename(file.filename)
            file_type = filename.split('.')[-1].lower()

            if file_type not in CONTENT_MARKETING_ALLOWED_EXTENSIONS:
                continue

            file_path = os.path.join(CONTENT_MARKETING_UPLOAD_FOLDER, project_id, filename)
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            file.save(file_path)

            # Extract text
            text_content = extract_text_from_file_content_marketing(file_path, file_type)

            # Store in PostgreSQL
            doc_id = f"doc_{uuid4().hex[:12]}"
            doc = CMDocument(
                doc_id=doc_id,
                project_id=project_id,
                file_name=filename,
                file_type=file_type,
                file_path=file_path,
                file_size=os.path.getsize(file_path),
                extracted_content=text_content
            )
            db.session.add(doc)
            extracted_documents.append(text_content)
            doc_ids.append(doc_id)

        db.session.commit()

        # Analyze domain specialization
        domain_context = analyzer.analyze_documents(extracted_documents)

        # Build knowledge graph
        kg_id = f"kg_{uuid4().hex[:12]}"
        kg_data = {
            'entities': [f'Entity_{i}' for i in range(min(10, len(extracted_documents)))],
            'relationships': [],
            'domain_context': domain_context,
            'documents_count': len(doc_ids)
        }

        kg = CMKnowledgeGraph(kg_id=kg_id, project_id=project_id)
        kg.kg_data = kg_data
        kg.entities = len(kg_data.get('entities', []))
        kg.relationships = len(kg_data.get('relationships', []))
        db.session.add(kg)
        db.session.commit()

        return jsonify({
            'success': True,
            'uploaded_files': len(doc_ids),
            'document_ids': doc_ids,
            'knowledge_graph_id': kg_id,
            'domain_specialization': domain_context
        }), 201

    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/content-marketing/documents/<project_id>', methods=['GET'])
@cross_origin()
@require_auth
def list_content_marketing_documents(project_id):
    """List all documents in a project"""
    return cm_service.list_documents(project_id)


@app.route('/api/content-marketing/generate-content', methods=['POST'])
@cross_origin()
@require_auth
def generate_content_marketing():
    """Generate marketing content for specified channel"""
    try:
        data = request.json
        project_id = data.get('project_id')
        channel = data.get('channel', 'linkedin')
        content_type = data.get('content_type', 'post')
        user_context = data.get('context', '')

        if not project_id:
            return jsonify({'success': False, 'error': 'project_id required'}), 400

        project = CMProject.query.filter_by(project_id=project_id).first()
        if not project or project.user_id != g.user_id:
            return jsonify({'success': False, 'error': 'Project not found'}), 404

        docs = CMDocument.query.filter_by(project_id=project_id).all()
        doc_texts = [d.extracted_content for d in docs if d.extracted_content]

        if not doc_texts:
            return jsonify({'success': False, 'error': 'No documents found in project'}), 400

        kg = CMKnowledgeGraph.query.filter_by(project_id=project_id).order_by(CMKnowledgeGraph.created_at.desc()).first()

        channel_config = {
            'linkedin': {'tone': 'professional', 'max_length': 3000},
            'email': {'tone': 'persuasive', 'max_length': 500},
            'social': {'tone': 'casual', 'max_length': 280},
            'google_ads': {'tone': 'direct', 'max_length': 150}
        }
        config = channel_config.get(channel, channel_config['linkedin'])

        prompt = f"""Generate marketing content for {channel} channel.
Industry: {project.industry or 'General'}
Tone: {config['tone']}
Max Length: {config['max_length']} characters
Content Type: {content_type}
User Context: {user_context}
Documents Summary: {' '.join([doc[:200] for doc in doc_texts[:3]])}

Generate compelling marketing {content_type} content."""

        llm = ChatOpenAI(model="gpt-4", temperature=0.7)
        response = llm.invoke(prompt).content

        # Store in PostgreSQL
        content_id = f"content_{uuid4().hex[:12]}"
        content = CMGeneratedContent(
            content_id=content_id,
            project_id=project_id,
            channel=channel,
            content_type=content_type,
            content=response
        )
        content.source_docs = [d.doc_id for d in docs]
        content.domain_context = {"industry": project.industry, "prompt": user_context}
        db.session.add(content)
        db.session.commit()

        return jsonify({
            'success': True,
            'content_id': content_id,
            'channel': channel,
            'content_type': content_type,
            'content': response,
            'variations': [response],
            'metadata': config
        }), 201

    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/content-marketing/chat', methods=['POST'])
@cross_origin()
@require_auth
def content_marketing_chat():
    """Conversational endpoint for iterative content refinement"""
    try:
        data = request.json
        project_id = data.get('project_id')
        message = data.get('message')

        if not all([project_id, message]):
            return jsonify({'success': False, 'error': 'project_id and message required'}), 400

        project = CMProject.query.filter_by(project_id=project_id).first()
        if not project or project.user_id != g.user_id:
            return jsonify({'success': False, 'error': 'Project not found'}), 404

        docs = CMDocument.query.filter_by(project_id=project_id).limit(10).all()
        kg = CMKnowledgeGraph.query.filter_by(project_id=project_id).order_by(CMKnowledgeGraph.created_at.desc()).first()

        context_text = ' '.join([d.extracted_content[:500] for d in docs if d.extracted_content])

        prompt = f"""Based on the following document context and knowledge graph, provide helpful marketing advice.
Document Context: {context_text}
Knowledge Graph: {json.dumps(kg.kg_data)[:500] if kg else 'No KG available'}
User Question: {message}
Provide a helpful, concise response focused on marketing strategy and content improvement."""

        llm = ChatOpenAI(model="gpt-4", temperature=0.7)
        response = llm.invoke(prompt).content

        # Store in PostgreSQL
        msg_id = f"msg_{uuid4().hex[:12]}"
        conv = CMConversation(
            msg_id=msg_id,
            project_id=project_id,
            user_message=message,
            agent_response=response
        )
        conv.context = {"project_name": project.project_name, "doc_count": len(docs)}
        db.session.add(conv)
        db.session.commit()

        return jsonify({
            'success': True,
            'response': response,
            'message_id': msg_id
        }), 200

    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/content-marketing/knowledge-graph/<project_id>', methods=['GET'])
@cross_origin()
@require_auth
def get_content_marketing_knowledge_graph(project_id):
    """Retrieve knowledge graph for visualization"""
    return cm_service.get_knowledge_graph(project_id)


@app.route('/api/get-google-credentials', methods=['GET'])
@cross_origin()
def get_google_credentials():
    """
    Get pre-configured Google OAuth credentials status from environment.
    Requires authentication. Never exposes client secret.
    """
    # Require authenticated session
    user_id, err = _resolve_session_user_id(None)
    if err:
        return err
    if not user_id:
        return jsonify({
            'success': False,
            'error': 'Authentication required'
        }), 401

    try:
        env_file = os.path.join(os.path.dirname(__file__), '.env')
        load_dotenv(env_file, override=True)
        has_oauth_credentials = all([
            os.getenv('GOOGLE_CLIENT_ID'),
            os.getenv('GOOGLE_CLIENT_SECRET'),
            os.getenv('GOOGLE_REDIRECT_URI')
        ])
        has_places_api_key = bool(os.getenv('GOOGLE_PLACES_API_KEY'))

        # Only expose public info - NEVER expose clientSecret
        credentials = {
            'clientId': os.getenv('GOOGLE_CLIENT_ID', ''),
            'redirectUri': os.getenv('GOOGLE_REDIRECT_URI', ''),
            'hasCredentials': has_oauth_credentials or has_places_api_key,
            'hasPlacesApiKey': has_places_api_key
        }

        return jsonify({
            'success': True,
            'credentials': credentials
        }), 200

    except Exception as e:
        print(f"Error fetching Google credentials: {str(e)}")
        return jsonify({
            'success': False,
            'error': 'Failed to fetch credentials status',
            'credentials': {
                'clientId': '',
                'redirectUri': '',
                'hasCredentials': False
            }
        }), 200


@app.route('/connect-google-business', methods=['POST'])
@cross_origin()
@require_auth
def connect_google_business():
    """
    Step 1: Save OAuth credentials and generate authorization URL
    Returns URL where user should go to authorize the app
    """
    from agents.market_research.google_business_helper import GoogleBusinessHelper
    try:
        data = request.get_json()

        required_fields = ['clientId', 'clientSecret', 'redirectUri']
        if not all(field in data for field in required_fields):
            return jsonify({
                'success': False,
                'error': 'Missing required fields'
            }), 400

        # Save credentials
        helper = GoogleBusinessHelper()
        if not helper.save_credentials(data):
            return jsonify({
                'success': False,
                'error': 'Failed to save credentials'
            }), 500
        
        # Generate Google OAuth authorization URL
        client_id = data.get('clientId')
        redirect_uri = data.get('redirectUri')
        
        auth_url = _generate_google_auth_url(client_id, redirect_uri)
        
        return jsonify({
            'success': True,
            'message': 'Credentials saved. Please authorize the app.',
            'authUrl': auth_url
        }), 200
        
    except Exception as e:
        print(f"Error in connect_google_business: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/auth/google/callback', methods=['GET'])
@cross_origin()
def google_auth_callback():
    """
    Step 2: OAuth callback endpoint
    Google redirects here with authorization code
    """
    try:
        auth_code = request.args.get('code')
        error = request.args.get('error')
        
        if error:
            return jsonify({
                'success': False,
                'error': f'Google authorization denied: {error}'
            }), 400
        
        if not auth_code:
            return jsonify({
                'success': False,
                'error': 'No authorization code received'
            }), 400
        
        state = request.args.get('state')
        
        if state == 'user_login_flow':
            # This is the login callback
            token_data = {
                'code': auth_code,
                'client_id': GOOGLE_CLIENT_ID,
                'client_secret': GOOGLE_CLIENT_SECRET,
                'redirect_uri': GOOGLE_REDIRECT_URI,
                'grant_type': 'authorization_code'
            }
            res = requests.post("https://oauth2.googleapis.com/token", data=token_data)
            
            if res.status_code != 200:
                print("Google Token Exchange Error:", res.text)
                return jsonify({'error': 'Failed to exchange token', 'details': res.json()}), 400
                
            token_response = res.json()
            access_token = token_response.get('access_token')
            refresh_token = token_response.get('refresh_token')
            scopes_received = token_response.get('scope', '')
            
            session_req = requests.Session()
            user_info = session_req.get('https://www.googleapis.com/oauth2/v1/userinfo', params={'access_token': access_token}).json()
            email = user_info.get('email')
            
            if not email:
                return jsonify({'error': 'Could not get email from Google'}), 400
                
            user = User.query.filter_by(email=email).first()
            if not user:
                user = User(
                    username=email,
                    email=email,
                    first_name=user_info.get('given_name', ''),
                    last_name=user_info.get('family_name', ''),
                    password=generate_password_hash(str(uuid4()))
                )
                db.session.add(user)
                db.session.commit()
                
            token_record = GoogleOAuthToken.query.filter_by(username=email).first()
            if not token_record:
                token_record = GoogleOAuthToken(username=email)
                db.session.add(token_record)
                
            token_record.token = access_token
            if refresh_token:
                token_record.refresh_token = refresh_token
            token_record.client_id = GOOGLE_CLIENT_ID
            token_record.client_secret = GOOGLE_CLIENT_SECRET
            token_record.token_uri = "https://oauth2.googleapis.com/token"
            token_record.scopes = scopes_received
            db.session.commit()
            
            from core.session_token import issue_browser_session_token

            session_tok = issue_browser_session_token(app.config['SECRET_KEY'], email)
            qs = urlencode(
                {
                    'google_auth': 'success',
                    'email': email,
                    'session_token': session_tok,
                }
            )
            return redirect(f"{_spa_redirect_base()}/login?{qs}")

        # Exchange code for refresh token for Google Business

        client_id = os.getenv('GOOGLE_CLIENT_ID')

        client_secret = os.getenv('GOOGLE_CLIENT_SECRET')
        redirect_uri = os.getenv('GOOGLE_REDIRECT_URI')
        
        success = _exchange_auth_code_for_token(auth_code, client_id, client_secret, redirect_uri)
        
        if success:
            # Redirect back to app with success
            return redirect(f'{_spa_redirect_base()}?google_connected=true')
        else:
            return jsonify({
                'success': False,
                'error': 'Failed to exchange authorization code'
            }), 500
            
    except Exception as e:
        print(f"Error in google_auth_callback: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/get-google-business-data', methods=['GET'])
@cross_origin()
@require_auth
def get_google_business_data():
    """
    Fetch Google Business information using saved credentials
    Returns business profile, reviews, and metrics
    """
    from agents.market_research.google_business_helper import GoogleBusinessHelper
    try:
        helper = GoogleBusinessHelper()
        
        if not helper.is_connected():
            return jsonify({
                'success': False,
                'error': 'Google Business credentials not found. Please connect first.',
                'code': 'NOT_CONNECTED'
            }), 401
        
        # Get complete business data
        business_data = helper.get_complete_business_data()
        business_data['success'] = True
        
        return jsonify(business_data), 200
        
    except Exception as e:
        print(f"Error in get_google_business_data: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/get-requirements-with-google-data', methods=['POST'])
@cross_origin()
@require_auth
def get_requirements_with_google_data():
    """
    Combined endpoint to get user requirements along with Google Business data
    Returns both requirement inputs and business insights for context
    """
    from agents.market_research.google_business_helper import GoogleBusinessHelper
    try:
        data = request.get_json()

        # Extract user requirements from request
        user_requirements = {
            'overview': data.get('overview', ''),
            'context': data.get('context', ''),
            'region': data.get('region', ''),
            'countries': data.get('countries', []),
            'industries': data.get('industries', []),
            'businessFunctions': data.get('businessFunctions', []),
            'analysisFrameworks': data.get('analysisFrameworks', ''),
            'responseFormat': data.get('responseFormat', ''),
            'uploadedFile': data.get('uploadedFile', None)
        }
        
        # Check if Google Business is connected and fetch data
        helper = GoogleBusinessHelper()
        google_business_data = None
        
        if helper.is_connected():
            google_business_data = helper.get_complete_business_data()
        else:
            google_business_data = {
                'connected': False,
                'message': 'Google Business not connected'
            }
        
        response = {
            'success': True,
            'userRequirements': user_requirements,
            'googleBusinessData': google_business_data,
            'timestamp': datetime.now().isoformat()
        }
        
        return jsonify(response), 200
        
    except Exception as e:
        print(f"Error in get_requirements_with_google_data: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/search-google-businesses', methods=['POST'])
@app.route('/api/search-google-businesses', methods=['POST'])
@cross_origin()
@require_auth
def search_google_businesses():
    """
    Search for Google businesses using Google Locations API
    Searches by business name and location with pagination support
    Supports up to 200 listings with configurable page size
    """
    from agents.market_research.google_business_helper import GoogleBusinessSearcher
    try:
        data = request.get_json()

        query = data.get('query', '')
        location = data.get('location', '')
        page = data.get('page', 1)
        page_size = data.get('page_size', 20)
        
        if not query:
            return jsonify({
                'success': False,
                'error': 'Search query is required'
            }), 400
        
        if not location:
            return jsonify({
                'success': False,
                'error': 'Location is required'
            }), 400
        
        # Validate pagination parameters
        if page < 1:
            page = 1
        if page_size < 1:
            page_size = 20
        if page_size > 200:
            page_size = 200  # Cap at 200 per page
        
        print(f"[SEARCH] Query: {query}, Location: {location}, Page: {page}, Page Size: {page_size}")
        
        # OAuth is optional for this flow because Google Places API key is used for search.
        # If OAuth is configured, we still attempt token refresh and pass it through.
        client_id = os.getenv('GOOGLE_CLIENT_ID')
        client_secret = os.getenv('GOOGLE_CLIENT_SECRET')
        redirect_uri = os.getenv('GOOGLE_REDIRECT_URI')
        access_token = None
        if all([client_id, client_secret, redirect_uri]):
            print("[SEARCH] OAuth credentials detected, attempting token refresh...")
            access_token = _get_google_access_token(client_id, client_secret, redirect_uri)
            if access_token:
                print("[SEARCH] OAuth token refresh succeeded.")
            else:
                print("[SEARCH] OAuth token refresh failed; continuing with Places API key.")
        else:
            print("[SEARCH] OAuth credentials missing; continuing with Places API key.")
        
        # Use Google Locations API to search for businesses
        searcher = GoogleBusinessSearcher()
        if access_token:
            searcher.set_credentials(access_token)
        
        results = searcher.search_businesses(
            query=query,
            location=location,
            max_results=200,
            page=page,
            page_size=page_size
        )
        
        print(f"[SEARCH] Results: {results}")
        return jsonify(results), 200
        
    except Exception as e:
        print(f"[SEARCH] Error in search_google_businesses: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


def _generate_google_auth_url(client_id: str, redirect_uri: str) -> str:
    """
    Generate Google OAuth authorization URL
    User visits this URL to authorize the app
    """
    oauth_scopes = [
        'https://www.googleapis.com/auth/business.manage',
        'https://www.googleapis.com/auth/spreadsheets',
        'https://www.googleapis.com/auth/drive.file'
    ]
    params = {
        'client_id': client_id,
        'redirect_uri': redirect_uri,
        'response_type': 'code',
        'scope': ' '.join(oauth_scopes),
        'access_type': 'offline',
        'prompt': 'consent'
    }
    return f"https://accounts.google.com/o/oauth2/v2/auth?{urlencode(params)}"


@app.route('/export-google-sheet', methods=['POST'])
@cross_origin()
@require_auth
def export_google_sheet():
    """
    Create and populate a Google Sheet using OAuth refresh token.
    Expects JSON payload: { title: str, headers: [str], rows: [[...]] }
    """
    try:
        data = request.get_json() or {}
        title = data.get('title', f"Market Research {datetime.now().strftime('%Y-%m-%d')}")
        headers = data.get('headers', [])
        rows = data.get('rows', [])

        if not headers or not isinstance(headers, list):
            return jsonify({
                'success': False,
                'error': 'Invalid headers payload'
            }), 400

        if not isinstance(rows, list):
            return jsonify({
                'success': False,
                'error': 'Invalid rows payload'
            }), 400

        client_id = os.getenv('GOOGLE_CLIENT_ID')
        client_secret = os.getenv('GOOGLE_CLIENT_SECRET')
        redirect_uri = os.getenv('GOOGLE_REDIRECT_URI')

        if not all([client_id, client_secret, redirect_uri]):
            return jsonify({
                'success': False,
                'error': 'Google OAuth credentials not configured.',
                'code': 'CREDENTIALS_MISSING'
            }), 401

        access_token = _get_google_access_token(client_id, client_secret, redirect_uri)
        if not access_token:
            return jsonify({
                'success': False,
                'error': 'Unable to get Google access token. Please reconnect Google account.',
                'code': 'AUTH_FAILED',
                'authUrl': _generate_google_auth_url(client_id, redirect_uri)
            }), 401

        create_sheet_resp = requests.post(
            'https://sheets.googleapis.com/v4/spreadsheets',
            headers={
                'Authorization': f'Bearer {access_token}',
                'Content-Type': 'application/json'
            },
            json={
                'properties': {'title': title}
            },
            timeout=20
        )

        if create_sheet_resp.status_code != 200:
            error_payload = create_sheet_resp.json() if create_sheet_resp.text else {}
            error_message = error_payload.get('error', {}).get('message', create_sheet_resp.text)
            return jsonify({
                'success': False,
                'error': f'Failed to create Google Sheet: {error_message}',
                'code': 'SHEET_CREATE_FAILED',
                'authUrl': _generate_google_auth_url(client_id, redirect_uri)
            }), 400

        spreadsheet = create_sheet_resp.json()
        spreadsheet_id = spreadsheet.get('spreadsheetId')
        spreadsheet_url = spreadsheet.get('spreadsheetUrl')

        value_rows = [headers] + rows
        update_resp = requests.put(
            f'https://sheets.googleapis.com/v4/spreadsheets/{spreadsheet_id}/values/Sheet1!A1?valueInputOption=RAW',
            headers={
                'Authorization': f'Bearer {access_token}',
                'Content-Type': 'application/json'
            },
            json={
                'majorDimension': 'ROWS',
                'values': value_rows
            },
            timeout=20
        )

        if update_resp.status_code != 200:
            error_payload = update_resp.json() if update_resp.text else {}
            error_message = error_payload.get('error', {}).get('message', update_resp.text)
            return jsonify({
                'success': False,
                'error': f'Sheet created but data write failed: {error_message}',
                'code': 'SHEET_WRITE_FAILED',
                'spreadsheetUrl': spreadsheet_url,
                'authUrl': _generate_google_auth_url(client_id, redirect_uri)
            }), 400

        return jsonify({
            'success': True,
            'spreadsheetId': spreadsheet_id,
            'spreadsheetUrl': spreadsheet_url
        }), 200

    except Exception as e:
        print(f"Error in export_google_sheet: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


def _exchange_auth_code_for_token(auth_code: str, client_id: str, client_secret: str, redirect_uri: str) -> bool:
    """
    Exchange Google authorization code for refresh token
    Called when user returns from Google authorization page
    """
    try:
        token_url = "https://oauth2.googleapis.com/token"
        data = {
            'client_id': client_id,
            'client_secret': client_secret,
            'code': auth_code,
            'grant_type': 'authorization_code',
            'redirect_uri': redirect_uri
        }
        
        response = requests.post(token_url, data=data, timeout=10)
        
        if response.status_code == 200:
            token_data = response.json()
            refresh_token = token_data.get('refresh_token')
            
            if refresh_token:
                # Save refresh token to environment
                os.environ['GOOGLE_REFRESH_TOKEN'] = refresh_token
                
                # Save to .env file in the tools directory
                env_file = os.path.join(os.path.dirname(__file__), '.env')
                with open(env_file, 'a') as f:
                    f.write(f"\nGOOGLE_REFRESH_TOKEN={refresh_token}")
                
                print(f"Refresh token saved to {env_file}")
                load_dotenv()
                return True
            else:
                print("No refresh token in response - this might be the first time authorizing")
                print(f"Response data: {token_data}")
                return False
        else:
            print(f"Token exchange failed: {response.status_code} - {response.text}")
            return False
            
    except Exception as e:
        print(f"Error exchanging auth code: {str(e)}")
        import traceback
        traceback.print_exc()
        return False


def _get_google_access_token(client_id: str, client_secret: str, redirect_uri: str) -> str:
    """
    Get Google access token from stored refresh token
    """
    try:
        refresh_token = os.getenv('GOOGLE_REFRESH_TOKEN')
        
        if not refresh_token:
            print("No refresh token found. User must authorize the app first.")
            print(f"Check that GOOGLE_REFRESH_TOKEN is in .env file")
            return None
        
        print(f"Using refresh token: {refresh_token[:20]}...")
        
        # Use refresh token to get new access token
        token_url = "https://oauth2.googleapis.com/token"
        data = {
            'client_id': client_id,
            'client_secret': client_secret,
            'refresh_token': refresh_token,
            'grant_type': 'refresh_token'
        }
        
        response = requests.post(token_url, data=data, timeout=10)
        
        if response.status_code == 200:
            access_token = response.json().get('access_token')
            print(f"Successfully got access token: {access_token[:20]}...")
            return access_token
        else:
            print(f"Token refresh failed: {response.status_code} - {response.text}")
            return None
            
    except Exception as e:
        print(f"Error getting Google access token: {str(e)}")
        import traceback
        traceback.print_exc()
        return None


@app.route('/api/email-extraction-usage', methods=['GET'])
@cross_origin()
@require_auth
def get_email_extraction_usage():
    """Return extraction usage summary for a username."""
    try:
        _ensure_email_usage_tables()

        username = _normalize_username(
            request.args.get('username') or request.args.get('userId')
        )
        quota = _get_or_create_quota(username)

        return jsonify({
            'success': True,
            'usageSummary': _build_usage_summary(username, quota)
        }), 200
    except Exception as e:
        print(f"[EMAIL_USAGE] Failed to fetch usage summary: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

def generate_email_content(business, sender_name):
    import json
    import os
    from openai import OpenAI
    
    prompt = f"""We have a market research agent that scrapes company websites and 
extracts leads including company name, industry, decision maker name, 
email, and a short company summary from their website.

I need you to build an email personalization layer on top of this.

When the user clicks "Send Email" for a lead, before sending, the system 
should call the AI to fill in the following variables dynamically using 
the lead data we already have:

- first_name -> extract from the contact name we scraped
- company_name -> from lead data
- one_line_company_summary -> generate a one line summary of what the company does based on the website content we already scraped
- industry -> detected from the company description
- pain_point -> infer the most likely business pain point for this industry and company size
- value_proposition -> tailor this to the industry, e.g. for FinTech say something different than for SaaS
- sender_name -> from the logged in client's profile

The base template is:

Subject: Quick idea for {{{{company_name}}}}

Hi {{{{first_name}}}},

I came across {{{{company_name}}}} and noticed {{{{one_line_company_summary}}}}.

Companies in {{{{industry}}}} often struggle with {{{{pain_point}}}} - and that usually means lost time or missed opportunities.

We built a solution that helps {{{{industry}}}} teams {{{{value_proposition}}}}.

Would it make sense to connect for 15 minutes this week?

Best,
{sender_name}

Return the final filled email as JSON in this format:
{{
  "subject": "...",
  "body": "..."
}}

Lead Data:
Company Name: {business.get('name', 'Unknown')}
Content/Description: {business.get('description', '')} {business.get('summary', '')}
Website: {business.get('website', 'Unknown')}
Contact Name: {business.get('contact_name', 'There')}
Industry: {business.get('industry', 'Unknown')}

Do not add any explanation, just return the JSON.
"""

    client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
    response = client.chat.completions.create(
        model=os.getenv('OPENAI_MODEL', 'gpt-4o-mini'),
        messages=[
            {"role": "user", "content": prompt}
        ],
        response_format={ "type": "json_object" },
        temperature=0.7
    )
    
    return json.loads(response.choices[0].message.content)

@app.route('/api/generate-email', methods=['POST'])
@cross_origin()
@require_auth
def generate_email():
    """Generate a personalized email using an LLM."""
    try:
        from flask import request
        data = request.get_json()
        business = data.get('business', {})
        sender_name = data.get('sender_name', 'Alex')
        
        result = generate_email_content(business, sender_name)
        return jsonify(result), 200
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/send-bulk-emails', methods=['POST'])
@app.route('/api/send-bulk-emails', methods=['POST', 'OPTIONS'])
@cross_origin()
@require_auth
def send_bulk_emails():
    """
    Send emails to a list of extracted businesses.
    Uses user's Google OAuth credentials if available, otherwise falls back to system SMTP.
    Expects JSON: { subject: str, body: str, businesses: list, userEmail: str }
    """
    import smtplib
    from email.message import EmailMessage
    import traceback
    
    try:
        data = request.get_json()
        subject = data.get('subject')
        body = data.get('body')
        businesses = data.get('businesses', [])
        # The sender identity (which Google OAuth token gets used, whose
        # name the campaign is filed under) must come from the verified
        # session, never from the request body - otherwise any caller could
        # supply someone else's email and send through their connected
        # Gmail account without their consent.
        user_email = g.user_id
        campaign_name = data.get('campaignName', 'Untitled Campaign')
        username = _normalize_username(g.user_id)

        if not user_email or '@' not in str(user_email):
            return jsonify({'success': False, 'error': 'Registered user email is required to send campaign mail.'}), 400

        use_ai_personalization = data.get('use_ai_personalization', False)
        if not use_ai_personalization and (not subject or not body):
            return jsonify({'success': False, 'error': 'Subject and body are required unless using AI personalization'}), 400

        valid_emails = [b.get('email') for b in businesses if b.get('email') and b.get('email') != 'N/A' and '@' in b.get('email')]

        if not valid_emails:
            return jsonify({'success': False, 'error': 'No valid emails found to send to'}), 400

        # Initialize DB tables and apply lightweight runtime migrations if needed.
        _ensure_email_usage_tables()
        _ensure_campaign_reply_tracking_columns()
        
        # Create Campaign Record
        import uuid
        campaign_id = str(uuid.uuid4())
        campaign = EmailCampaign(
            id=campaign_id,
            name=campaign_name,
            subject=subject,
            username=username,
            sender_email=user_email
        )
        db.session.add(campaign)
        db.session.commit()
        
        # Check if user has Google credentials connected
        token_record = None
        if user_email:
            token_record = GoogleOAuthToken.query.filter_by(username=user_email).first()
        
        service = None
        server = None
        smtp_sender_email = os.getenv('EMAIL_USER') or user_email or ''

        def _connect_smtp_server():
            email_host = os.getenv('EMAIL_HOST', 'smtp.gmail.com')
            email_port = int(os.getenv('EMAIL_PORT', 587))
            email_user = os.getenv('EMAIL_USER')
            email_pass = os.getenv('EMAIL_PASS')
            if not email_user or not email_pass:
                return None, None, 'Email credentials are not configured. Please sign in with Google or configure system SMTP.'

            smtp_server = smtplib.SMTP(email_host, email_port)
            smtp_server.starttls()
            smtp_server.login(email_user, email_pass)
            return smtp_server, email_user, None
        
        if token_record and token_record.token:
            # Use Gmail API
            from google.oauth2.credentials import Credentials
            from google.auth.transport.requests import Request
            import googleapiclient.discovery
            
            creds = Credentials(
                token=token_record.token,
                refresh_token=token_record.refresh_token,
                token_uri=token_record.token_uri,
                client_id=token_record.client_id,
                client_secret=token_record.client_secret,
                scopes=token_record.scopes.split(',') if token_record.scopes else SCOPES
            )
            if creds.refresh_token and (not creds.valid or creds.expired):
                creds.refresh(Request())
                token_record.token = creds.token
                db.session.commit()
            service = googleapiclient.discovery.build('gmail', 'v1', credentials=creds)
        else:
            # Fallback to system SMTP
            server, email_user, smtp_err = _connect_smtp_server()
            if smtp_err:
                return jsonify({'success': False, 'error': smtp_err}), 500

        sent_count = 0
        for b in businesses:
            recipient = b.get('email')
            if not recipient or recipient == 'N/A' or '@' not in recipient:
                continue

            business_name = b.get('name', 'Business Owner')
            
            # Message personalization
            current_body = body
            if use_ai_personalization:
                try:
                    result = generate_email_content(b, username)
                    current_subject = result.get('subject', subject or 'Exclusive Offer')
                    current_body = result.get('body', current_body or '')
                except Exception as e:
                    print("Failed AI personalization for", business_name, e)
                    current_subject = subject or 'Exclusive Offer'
            else:
                current_subject = subject.replace('{{name}}', business_name) if subject else subject
                if current_body:
                    current_body = current_body.replace('{{name}}', business_name)

            msg = EmailMessage()
            msg.set_content(current_body)
            msg['Subject'] = current_subject
            msg['To'] = recipient

            def _set_from_header(message, sender_value):
                if 'From' in message:
                    del message['From']
                message['From'] = sender_value

            def _set_reply_to_header(message, reply_to_value):
                if 'Reply-To' in message:
                    del message['Reply-To']
                message['Reply-To'] = reply_to_value
            
            thread_id = None
            msg_id = None
            generated_message_id = f"<{uuid.uuid4().hex}@enable-agents.local>"
            _set_from_header(msg, user_email or smtp_sender_email or recipient)
            _set_reply_to_header(msg, smtp_sender_email or user_email or recipient)
            msg['Message-ID'] = generated_message_id
            if service:
                try:
                    encoded_message = base64.urlsafe_b64encode(msg.as_bytes()).decode()
                    create_message = {'raw': encoded_message}
                    sent_msg = service.users().messages().send(userId="me", body=create_message).execute()
                    thread_id = sent_msg.get('threadId')
                    msg_id = sent_msg.get('id')
                except Exception as send_error:
                    # Any Gmail API failure (expired/invalid creds, API not
                    # enabled on the project, quota, etc.) should fall back to
                    # SMTP rather than only specific credential error strings -
                    # narrowly matching text meant only one failure mode ever
                    # got a second chance.
                    print(f"[SEND_EMAILS] Gmail API failed, falling back to SMTP: {send_error}")
                    gmail_error_summary = str(send_error).split('.', 1)[0][:200]
                    service = None
                    server, smtp_sender_email, smtp_err = _connect_smtp_server()
                    if smtp_err:
                        return jsonify({'success': False, 'error': f'Gmail send failed ({gmail_error_summary}) and SMTP fallback is unavailable: {smtp_err}'}), 500
                    _set_from_header(msg, user_email or smtp_sender_email or recipient)
                    _set_reply_to_header(msg, smtp_sender_email or user_email or recipient)
                    server.send_message(msg)
            else:
                _set_from_header(msg, user_email or smtp_sender_email or recipient)
                _set_reply_to_header(msg, smtp_sender_email or user_email or recipient)
                server.send_message(msg)
                
            sent_count += 1
            
            # Record recipient for tracking
            recipient_record = EmailCampaignRecipient(
                campaign_id=campaign_id,
                receiver_email=recipient,
                receiver_name=business_name,
                status='SENT',
                reply_status='No Reply',
                message_id=msg_id or generated_message_id,
                thread_id=thread_id
            )
            db.session.add(recipient_record)

        db.session.commit()
        
        if server:
            server.quit()
            
        # Commit any leftover recipient records
        db.session.commit()
        
        return jsonify({'success': True, 'count': sent_count, 'message': 'Emails successfully sent via user account!' if service else 'Emails sent via system account.'})
    except Exception as e:
        db.session.rollback()
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/webhook/zapier/email-reply', methods=['POST'])
@cross_origin()
def handle_email_reply():
    """
    Zapier Webhook Endpoint to log replies from emails sent in campaigns.
    Expects JSON: { "from_email": "example@domain.com", "subject": "Re: ...", "timestamp": "...", "snippet": "..." }
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'No JSON payload provided'}), 400
            
        from_email = data.get('from_email')
        
        if not from_email:
            return jsonify({'success': False, 'error': 'from_email is required'}), 400
            
        # Optional: Parse out name from format "Name <email@dom.com>"
        import re
        email_match = re.search(r'<(.+?)>', from_email)
        if email_match:
            clean_email = email_match.group(1).lower().strip()
        else:
            clean_email = from_email.lower().strip()
            
        # Find receiver in the database
        recipients = EmailCampaignRecipient.query.filter(
            EmailCampaignRecipient.receiver_email.ilike(f"%{clean_email}%")
        ).all()
        
        if not recipients:
            return jsonify({'success': False, 'message': f'Reply from {clean_email} logged, but not found in active campaigns.'}), 200
            
        updated_count = 0
        for rec in recipients:
            if rec.reply_status != 'Replied':
                rec.reply_status = 'Replied'
                rec.replied_at = datetime.utcnow()
                updated_count += 1
                
        db.session.commit()
        return jsonify({
            'success': True, 
            'message': f'Successfully updated {updated_count} recipient records to Replied status.'
        }), 200

    except Exception as e:
        print(f"[ZAPIER WEBHOOK ERROR] {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/campaigns', methods=['GET'])
@cross_origin()
@require_auth
def list_campaigns():
    """Lightweight campaign list for cross-agent pickers (e.g. Content
    Marketing's 'Send to Email Campaign'). Distinct from /api/campaigns/stats,
    which includes reply-rate analytics this picker doesn't need."""
    try:
        query = EmailCampaign.query.filter_by(username=g.user_id)
        campaigns = query.order_by(EmailCampaign.created_at.desc()).all()

        results = [{
            'id': c.id,
            'name': c.name,
            'subject': c.subject,
            'status': c.status,
            'lead_count': len(c.recipients),
            'createdAt': c.created_at.isoformat() if c.created_at else None,
        } for c in campaigns]

        return jsonify({'success': True, 'campaigns': results}), 200
    except Exception as e:
        print(f"[LIST CAMPAIGNS ERROR] {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/campaigns/<campaign_id>/content', methods=['PUT'])
@cross_origin()
@require_auth
def update_campaign_content(campaign_id):
    """Update a campaign's email body - used by Content Marketing's
    'Send to Email Campaign' to hand off generated content."""
    try:
        campaign = EmailCampaign.query.get(campaign_id)
        if not campaign or campaign.username != g.user_id:
            return jsonify({'success': False, 'error': 'Campaign not found'}), 404

        data = request.get_json() or {}
        email_body = data.get('email_body')
        if not email_body:
            return jsonify({'success': False, 'error': 'email_body is required'}), 400

        campaign.body_template = email_body
        db.session.commit()

        return jsonify({'success': True, 'campaign_id': campaign_id}), 200
    except Exception as e:
        print(f"[UPDATE CAMPAIGN CONTENT ERROR] {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/campaigns/stats', methods=['GET'])
@cross_origin()
@require_auth
def get_campaign_stats():
    """Returns analytics for campaigns filtered by user."""
    try:
        _ensure_email_usage_tables()
        _ensure_campaign_reply_tracking_columns()

        username = g.user_id
        user_email = request.args.get('email')

        campaigns = EmailCampaign.query.filter_by(username=username).order_by(EmailCampaign.created_at.desc()).all()

        # Keep the stats endpoint fast. Reply sync is available through the
        # dedicated check-replies route so the dashboard does not block on inbox scans.
        if (request.args.get('syncReplies') or '').strip().lower() in {'1', 'true', 'yes'}:
            for campaign in campaigns:
                try:
                    _sync_replies_for_campaign(campaign, fallback_email=user_email, fallback_username=username)
                except Exception as sync_err:
                    print(f"[AUTO SYNC] Skipping campaign {campaign.id}: {sync_err}")

        results = []
        for c in campaigns:
            recipients = EmailCampaignRecipient.query.filter_by(campaign_id=c.id).all()
            total_sent = len(recipients)
            total_replied = sum(1 for r in recipients if r.reply_status == 'Replied')
            
            results.append({
                'id': c.id,
                'name': c.name,
                'subject': c.subject,
                'createdAt': c.created_at.isoformat(),
                'totalSent': total_sent,
                'totalReplied': total_replied,
                'replyRate': round((total_replied / total_sent * 100) if total_sent > 0 else 0, 1)
            })
            
        return jsonify({'success': True, 'campaigns': results}), 200
    except Exception as e:
        print(f"[CAMPAIGN STATS ERROR] {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/campaigns/<campaign_id>/recipients', methods=['GET'])
@cross_origin()
@require_auth
def get_campaign_recipients(campaign_id):
    """Returns recipients for a specific campaign."""
    try:
        _ensure_email_usage_tables()
        _ensure_campaign_reply_tracking_columns()

        campaign = EmailCampaign.query.get(campaign_id)
        if not campaign or campaign.username != g.user_id:
            return jsonify({'success': False, 'error': 'Campaign not found'}), 404

        try:
            _sync_replies_for_campaign(campaign)
        except Exception as sync_err:
            print(f"[AUTO SYNC] Recipient view sync skipped for {campaign_id}: {sync_err}")

        recipients = EmailCampaignRecipient.query.filter_by(campaign_id=campaign_id).all()
        results = [{
            'email': r.receiver_email,
            'name': r.receiver_name,
            'status': r.status,
            'replyStatus': r.reply_status,
            'replySubject': r.reply_subject,
            'replySnippet': r.reply_snippet,
            'replyBody': r.reply_body,
            'sentAt': r.sent_at.isoformat() if r.sent_at else None,
            'repliedAt': r.replied_at.isoformat() if r.replied_at else None
        } for r in recipients]
        return jsonify({'success': True, 'recipients': results}), 200
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/campaigns/<campaign_id>/rank-vendors', methods=['POST', 'OPTIONS'])
@cross_origin()
def rank_campaign_vendors(campaign_id):
    """Rank vendor email replies for a campaign using response content and the provided criteria."""
    if request.method == 'OPTIONS':
        return jsonify({}), 200

    def _compact_text(value):
        if not value:
            return ''
        return re.sub(r'\s+', ' ', str(value)).strip()

    def _heuristic_score(text, criteria_text):
        text_lower = (text or '').lower()
        criteria_lower = (criteria_text or '').lower()
        score = 0
        keyword_map = {
            'cost': ['cost', 'price', 'pricing', 'quote', 'quotation', 'discount', 'rate'],
            'quality': ['quality', 'grade', 'certified', 'certification', 'premium'],
            'delivery': ['delivery', 'lead time', 'dispatch', 'shipping', 'logistics'],
            'reliability': ['reliable', 'reliability', 'consistent', 'on time', 'on-time'],
            'payment': ['payment', 'credit', 'terms', 'advance', 'net 30', 'net 45'],
            'warranty': ['warranty', 'guarantee', 'return', 'replacement'],
            'capacity': ['capacity', 'volume', 'scalable', 'production'],
            'certification': ['iso', 'certified', 'compliance', 'license', 'registration'],
            'location': ['location', 'near', 'local', 'site', 'warehouse'],
        }
        for terms in keyword_map.values():
            for term in terms:
                if term in text_lower:
                    score += 4
        # Slight bias if the reply mentions multiple numeric quote values
        if re.search(r'\b\d+(?:\.\d+)?\b', text_lower):
            score += 4
        if 'quotation' in text_lower or 'quote' in text_lower:
            score += 8
        if 'best' in text_lower or 'lowest' in text_lower or 'competitive' in text_lower:
            score += 6
        # If the reply directly covers criteria wording, boost slightly.
        for word in re.findall(r'[a-z]{4,}', criteria_lower):
            if word in text_lower:
                score += 1
        return min(score, 100)

    try:
        payload = request.get_json(silent=True) or {}
        criteria = _compact_text(payload.get('criteria') or payload.get('question') or payload.get('rankingCriteria') or '')
        user_email = payload.get('userEmail') or payload.get('senderEmail')

        trusted_uid, uid_err = _resolve_session_user_id(payload.get('user_id') or payload.get('username'))
        if uid_err is not None:
            return uid_err[0], uid_err[1]

        _ensure_email_usage_tables()
        _ensure_campaign_reply_tracking_columns()

        campaign = EmailCampaign.query.get_or_404(campaign_id)
        if campaign.username != trusted_uid:
            return jsonify({'success': False, 'error': 'Campaign not found'}), 404
        try:
            _sync_replies_for_campaign(campaign, fallback_email=user_email, fallback_username=trusted_uid)
        except Exception as sync_err:
            print(f"[RANK VENDORS] Reply sync skipped for {campaign_id}: {sync_err}")

        recipients = EmailCampaignRecipient.query.filter_by(campaign_id=campaign_id).all()
        replied = [
            r for r in recipients
            if (r.reply_status or '').lower() == 'replied'
            and (_compact_text(r.reply_body) or _compact_text(r.reply_snippet) or _compact_text(r.reply_subject))
        ]

        if not replied:
            return jsonify({'success': False, 'error': 'No vendor replies found to rank yet.'}), 400

        vendor_payload = []
        for index, recipient in enumerate(replied):
            reply_text = _compact_text(recipient.reply_body or recipient.reply_snippet or recipient.reply_subject)
            if len(reply_text) > 3500:
                reply_text = reply_text[:3500]
            vendor_payload.append({
                'index': index,
                'vendor_name': recipient.receiver_name or recipient.receiver_email,
                'email': recipient.receiver_email,
                'reply_text': reply_text,
                'sent_at': recipient.sent_at.isoformat() if recipient.sent_at else None,
                'replied_at': recipient.replied_at.isoformat() if recipient.replied_at else None,
                'reply_subject': _compact_text(recipient.reply_subject or ''),
            })

        openai_key = os.getenv('OPENAI_API_KEY')
        if not openai_key:
            ranked = []
            for item in vendor_payload:
                score = _heuristic_score(item['reply_text'], criteria)
                ranked.append({
                    'rank': 0,
                    'score': score,
                    'vendor_name': item['vendor_name'],
                    'email': item['email'],
                    'reply_summary': item['reply_text'][:220],
                    'reason': 'Heuristic ranking used because OPENAI_API_KEY is not set.',
                    'criteria_match': [],
                    'reply_text': item['reply_text'],
                })
            ranked.sort(key=lambda x: x['score'], reverse=True)
            for idx, item in enumerate(ranked, start=1):
                item['rank'] = idx
            return jsonify({
                'success': True,
                'campaign': {'id': campaign.id, 'name': campaign.name, 'subject': campaign.subject},
                'criteria': criteria,
                'vendors': ranked,
            }), 200

        client = openai.OpenAI(api_key=openai_key)
        prompt = [
            {
                'role': 'system',
                'content': (
                    'You are a procurement analyst ranking vendor email replies like a human buyer would. '
                    'Compare every vendor reply against the criteria and the quotations mentioned in the emails. '
                    'Prioritize cost/pricing, quality, reliability, delivery performance, reputation, capacity, compliance, communication, location/logistics, technology, risk, sustainability, payment terms, lead time, after-sales support, customization, financial stability, scalability, warranty/return policies, inventory, contract flexibility, certifications, data security, ethics, client references, and supply-chain stability when relevant. '
                    'Return only valid JSON in the format: {"vendors":[{"rank":1,"vendor_name":"...","email":"...","score":92,"reply_summary":"...","reason":"...","matched_criteria":["..."],"quote_comparison":"...","strengths":["..."],"risks":["..."]}]}.'
                )
            },
            {
                'role': 'user',
                'content': json.dumps({
                    'criteria': criteria,
                    'campaign': {'id': campaign.id, 'name': campaign.name, 'subject': campaign.subject},
                    'vendors': vendor_payload,
                }, ensure_ascii=False)
            }
        ]

        response = client.chat.completions.create(
            model=os.getenv('OPENAI_MODEL', 'gpt-4o-mini'),
            messages=prompt,
            temperature=0.0,
            max_tokens=int(os.getenv('OPENAI_VENDOR_RANK_MAX_TOKENS', '1400'))
        )

        response_text = (response.choices[0].message.content or '').strip().replace('```json', '').replace('```', '').strip()
        parsed = json.loads(response_text)
        vendors = parsed.get('vendors') if isinstance(parsed, dict) else None
        if not isinstance(vendors, list):
            raise ValueError('LLM did not return a vendors array')

        cleaned = []
        for idx, item in enumerate(vendors, start=1):
            if not isinstance(item, dict):
                continue
            cleaned.append({
                'rank': int(item.get('rank') or idx),
                'score': int(item.get('score') or 0),
                'vendor_name': _compact_text(item.get('vendor_name') or ''),
                'email': _compact_text(item.get('email') or ''),
                'reply_summary': _compact_text(item.get('reply_summary') or ''),
                'reason': _compact_text(item.get('reason') or ''),
                'matched_criteria': item.get('matched_criteria') if isinstance(item.get('matched_criteria'), list) else [],
                'quote_comparison': _compact_text(item.get('quote_comparison') or ''),
                'strengths': item.get('strengths') if isinstance(item.get('strengths'), list) else [],
                'risks': item.get('risks') if isinstance(item.get('risks'), list) else [],
            })

        cleaned.sort(key=lambda x: x['rank'])
        return jsonify({
            'success': True,
            'campaign': {'id': campaign.id, 'name': campaign.name, 'subject': campaign.subject},
            'criteria': criteria,
            'vendors': cleaned,
        }), 200

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/campaigns/<campaign_id>/check-replies', methods=['POST'])
@cross_origin()
@require_auth
def check_campaign_replies(campaign_id):
    """Checks Gmail API for any replies to the campaign's sent emails."""
    try:
        _ensure_email_usage_tables()
        _ensure_campaign_reply_tracking_columns()
        campaign = EmailCampaign.query.get_or_404(campaign_id)
        if campaign.username != g.user_id:
            return jsonify({'success': False, 'error': 'Campaign not found'}), 404

        payload = request.get_json(silent=True) or {}
        sender_email = payload.get('userEmail')

        updated_count = _sync_replies_for_campaign(campaign, fallback_email=sender_email)

        return jsonify({'success': True, 'updated': updated_count}), 200

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/enrich-businesses-with-emails', methods=['POST'])
@cross_origin()
@require_auth
def enrich_businesses_with_emails():
    """
    Enrich business data with email addresses using scrap.io API
    Takes a list of businesses with website URLs and enriches them with email data
    """
    try:
        print(f"\n[EMAIL_ENRICHMENT] ========== REQUEST START ==========")
        data = request.get_json()
        businesses = data.get('businesses', [])
        # Quota is tracked per-username; must come from the verified session,
        # not the request body, or a caller could rotate fake usernames to
        # bypass their quota entirely (each "new" username starts fresh).
        username = _normalize_username(g.user_id)
        scrap_io_api_key = os.getenv('SCRAP_IO_API_KEY')
        print(f"[EMAIL_ENRICHMENT] Processing {len(businesses)} businesses for user: {username}")

        _ensure_email_usage_tables()
        quota = _get_or_create_quota(username)
        usage_before = _build_usage_summary(username, quota)

        if usage_before['remainingCount'] <= 0:
            return jsonify({
                'success': False,
                'error': 'Email extraction limit reached for this user.',
                'code': 'QUOTA_EXCEEDED',
                'usageSummary': usage_before
            }), 403
        
        if not businesses or len(businesses) == 0:
            return jsonify({
                'success': False,
                'error': 'No businesses provided'
            }), 400
        
        if not scrap_io_api_key:
            return jsonify({
                'success': False,
                'error': 'Scrap.io API key not configured',
                'code': 'API_KEY_MISSING'
            }), 401
        
        print(f"[EMAIL_ENRICHMENT] Enriching {len(businesses)} businesses with emails")
        
        def extract_emails_from_text(text):
            """Extract email addresses from text using regex"""
            if not text:
                return []
            email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
            emails = re.findall(email_pattern, str(text))
            return list(set(emails))  # Remove duplicates
        
        def get_email_from_google_places(business):
            """Try to extract email from Google Places data"""
            # Check if business data already contains formatted address, phone, website
            # that might have email info
            if business.get('formatted_address'):
                emails = extract_emails_from_text(business['formatted_address'])
                if emails:
                    return emails[0]
            
            return None
        
        def generate_common_email_patterns(domain, business_name):
            """Generate common email patterns for a domain"""
            # Remove www. if present for cleaner domain
            domain_clean = domain.replace('www.', '')
            
            # Common prefixes to try
            common_prefixes = [
                'info', 'contact', 'support', 'hello', 'business', 'sales', 
                'email', 'inquiry', 'admin', 'help', 'team'
            ]
            
            # Try with business name components
            name_parts = business_name.lower().split()[:2] if business_name else []
            
            patterns = []
            
            # Standard patterns
            for prefix in common_prefixes:
                patterns.append(f"{prefix}@{domain_clean}")
            
            # Business name patterns
            if name_parts:
                for part in name_parts:
                    part_clean = ''.join(c for c in part if c.isalnum())
                    if part_clean:
                        patterns.append(f"{part_clean}@{domain_clean}")
            
            return patterns
        
        def try_sync_website_scrape(website, business_name):
            """Try to scrape website for email addresses synchronously"""
            try:
                print(f"[EMAIL_ENRICHMENT] Starting website scrape for {website}")
                response = requests.get(
                    website if website.startswith('http') else f'https://{website}',
                    timeout=5,
                    headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
                )
                
                print(f"[EMAIL_ENRICHMENT] Website scrape status: {response.status_code}")
                if response.status_code == 200:
                    # Look for email patterns in HTML
                    emails = extract_emails_from_text(response.text)
                    if emails:
                        # Filter out common auto-generated emails
                        filtered = [e for e in emails if not any(x in e.lower() for x in ['noreply', 'no-reply', 'postmaster'])]
                        return filtered[0] if filtered else None
            except requests.exceptions.Timeout:
                print(f"[EMAIL_ENRICHMENT] Website scrape timeout for {website}")
            except Exception as e:
                print(f"[EMAIL_ENRICHMENT] Website scrape error for {website}: {str(e)}")
            
            return None
        
        def extract_email_from_business_data(business_data, domain, website):
            """Extract email from scrap.io response with multiple fallback strategies"""
            email = 'N/A'
            
            website_data = business_data.get('website_data', {})
            
            if not website_data:
                return email
            
            # Strategy 1: Check emails array (primary)
            if website_data.get('emails') and len(website_data.get('emails', [])) > 0:
                email_item = website_data['emails'][0]
                if isinstance(email_item, dict) and 'email' in email_item:
                    extracted = email_item['email']
                    if extracted and extracted != 'N/A':
                        email = extracted
                        print(f"[EMAIL_ENRICHMENT] âœ… Strategy 1 (emails array) found: {email}")
                        return email
                elif isinstance(email_item, str) and email_item.strip():
                    email = email_item
                    print(f"[EMAIL_ENRICHMENT] âœ… Strategy 1 (emails string) found: {email}")
                    return email
            
            # Strategy 2: Try contact information
            if not email or email == 'N/A':
                contact_info = website_data.get('contact_info', {})
                if contact_info and contact_info.get('email'):
                    email = contact_info['email']
                    print(f"[EMAIL_ENRICHMENT] âœ… Strategy 2 (contact_info) found: {email}")
                    return email
            
            # Strategy 3: Extract from social profiles or other fields
            if not email or email == 'N/A':
                for key in ['socials', 'social_profiles', 'contact', 'business_info']:
                    if key in website_data:
                        data = website_data[key]
                        if isinstance(data, dict):
                            for field_key, field_value in data.items():
                                if field_value and isinstance(field_value, str):
                                    found_emails = extract_emails_from_text(field_value)
                                    if found_emails:
                                        email = found_emails[0]
                                        print(f"[EMAIL_ENRICHMENT] âœ… Strategy 3 ({key}.{field_key}) found: {email}")
                                        return email
            
            # Strategy 4: Check for contact URLs and other fields
            if not email or email == 'N/A':
                for key in ['contact_page', 'about_page', 'company_info', 'all_text']:
                    if key in website_data:
                        text_data = website_data[key]
                        if text_data:
                            found_emails = extract_emails_from_text(str(text_data))
                            if found_emails:
                                email = found_emails[0]
                                print(f"[EMAIL_ENRICHMENT] âœ… Strategy 4 ({key}) found: {email}")
                                return email
            
            # Strategy 5: Fallback to phone if no email found
            if (not email or email == 'N/A') and website_data.get('phones') and len(website_data.get('phones', [])) > 0:
                phone_obj = website_data['phones'][0]
                if isinstance(phone_obj, dict) and 'phone' in phone_obj:
                    phone = phone_obj['phone']
                    email = f"Phone: {phone}"
                    print(f"[EMAIL_ENRICHMENT] â„¹ï¸  Strategy 5 (phone fallback) found: {email}")
                    return email
            
            return email
        
        def call_scrap_io_with_retry(domain, max_retries=1, delay=1):
            """Call scrap.io with minimal retries to avoid hanging"""
            import time
            
            headers = {
                'Authorization': f'Bearer {scrap_io_api_key}',
                'Content-Type': 'application/json'
            }
            
            for attempt in range(max_retries):
                try:
                    # Use domain parameter with shorter timeout
                    response = requests.get(
                        scrap_io_endpoint,
                        params={'domain': domain},
                        headers=headers,
                        timeout=8  # Reduced from 15 to 8 seconds
                    )
                    
                    if response.status_code in [200, 202]:
                        result = response.json()
                        data = result.get('data', [])
                        
                        # If we got data, return it immediately
                        if data and len(data) > 0:
                            print(f"[EMAIL_ENRICHMENT] Got data on attempt {attempt + 1}")
                            return result, True
                        
                        # If no data, don't retry - move to fallback strategy
                        print(f"[EMAIL_ENRICHMENT] No data from scrap.io, moving to fallback")
                        return result, False
                    else:
                        print(f"[EMAIL_ENRICHMENT] API Error {response.status_code}: {response.text[:100]}")
                        return {}, False
                        
                except requests.exceptions.Timeout:
                    print(f"[EMAIL_ENRICHMENT] Timeout on attempt {attempt + 1} - moving to fallback")
                    return {}, False
                except Exception as e:
                    print(f"[EMAIL_ENRICHMENT] Error on attempt {attempt + 1}: {str(e)}")
                    return {}, False
            
            return {}, False
        
        enriched_businesses = []
        scrap_io_endpoint = "https://scrap.io/api/v1/gmap/enrich"
        
        for business in businesses:
            website = business.get('website')
            
            # If no website, add business as-is with empty email
            if not website:
                business_copy = business.copy()
                business_copy['email'] = 'N/A'
                enriched_businesses.append(business_copy)
                continue
            
            # Extract domain from website URL
            domain = website.replace('https://', '').replace('http://', '').split('/')[0]
            business_name = business.get('name', '')
            
            email = 'N/A'
            
            try:
                # First, try to extract from Google Places data itself
                google_email = get_email_from_google_places(business)
                if google_email:
                    email = google_email
                    print(f"[EMAIL_ENRICHMENT] âœ… Found email from Google Places: {email}")
                else:
                    # Try calling scrap.io with retry mechanism for async processing
                    result, has_data = call_scrap_io_with_retry(domain, max_retries=2, delay=1)
                    
                    if has_data:
                        data = result.get('data', [])
                        if data and len(data) > 0:
                            business_data = data[0]
                            email = extract_email_from_business_data(business_data, domain, website)
                    
                    # If still no email, try direct website scraping
                    if email == 'N/A':
                        print(f"[EMAIL_ENRICHMENT] Trying direct website scrape for {domain}...")
                        scraped_email = try_sync_website_scrape(website, business_name)
                        if scraped_email:
                            email = scraped_email
                            print(f"[EMAIL_ENRICHMENT] âœ… Found via website scrape: {email}")
                    
                    # If still no email, generate common patterns for user reference
                    if email == 'N/A':
                        patterns = generate_common_email_patterns(domain, business_name)
                        if patterns:
                            # Use the most common pattern as suggestion
                            email = patterns[0]
                            print(f"[EMAIL_ENRICHMENT] â„¹ï¸  Using common pattern suggestion: {email}")
                
                business_copy = business.copy()
                business_copy['email'] = email
                enriched_businesses.append(business_copy)
                
            except Exception as e:
                print(f"[EMAIL_ENRICHMENT] âŒ Error processing {business_name}: {str(e)}")
                business_copy = business.copy()
                business_copy['email'] = 'Error'
                enriched_businesses.append(business_copy)
        
        processed_count = len(enriched_businesses)
        billable_count = sum(
            1 for business in enriched_businesses if _is_billable_email(business.get('email'))
        )
        charged_count = min(billable_count, usage_before['remainingCount'])
        quota.emails_used_this_month += charged_count
        db.session.add(quota)

        request_id = str(uuid4())
        cost_this_request = round(charged_count * EMAIL_EXTRACTION_UNIT_COST, 2)
        usage_after = _build_usage_summary(username, quota)
        usage_log = EmailExtractionUsageLog(
            request_id=request_id,
            username=username,
            processed_count=processed_count,
            billable_count=billable_count,
            charged_count=charged_count,
            cost_this_request=cost_this_request,
            total_cost_after=usage_after['totalCost']
        )
        db.session.add(usage_log)
        db.session.commit()

        print(f"[EMAIL_ENRICHMENT] Successfully enriched {processed_count} businesses")
        print(f"[EMAIL_ENRICHMENT] ========== REQUEST END ========== (took {request_id})")
        
        return jsonify({
            'success': True,
            'businesses': enriched_businesses,
            'enrichedCount': processed_count,
            'billableEmailCount': billable_count,
            'chargedEmailCount': charged_count,
            'costThisRequest': cost_this_request,
            'usageSummary': usage_after,
            'timestamp': datetime.now().isoformat()
        }), 200
        
    except Exception as e:
        db.session.rollback()
        print(f"[EMAIL_ENRICHMENT] Error in enrich_businesses_with_emails: {str(e)}")
        print(f"[EMAIL_ENRICHMENT] ========== REQUEST ERROR ========== ")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/enrich-businesses-with-linkedin', methods=['POST'])
@cross_origin()
@require_auth
def enrich_businesses_with_linkedin():
    try:
        data = request.get_json()
        businesses = data.get('businesses', [])
        
        # Use duckduckgo-search to intelligently grab the most relevant linkedin company URL based on business name
        
        try:
            from ddgs import DDGS
        except ImportError:
            DDGS = None
            
        enriched_businesses = []
        for b in businesses:
            raw_name = b.get('name', '')
            b['linkedin'] = "N/A"
            if raw_name and DDGS:
                try:
                    search_query = f"{raw_name} linkedin company"
                    results = DDGS().text(search_query, max_results=3)
                    for r in results:
                        if 'linkedin.com/company' in r.get('href', ''):
                            b['linkedin'] = r['href']
                            break
                except Exception as e:
                    print(f"Error searching linkedin for {raw_name}:", e)
            
            enriched_businesses.append(b)

        return jsonify({
            'success': True,
            'data': {'businesses': enriched_businesses}
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500



@app.route('/api/account', methods=['DELETE'])
@require_auth
def delete_own_account():
    """
    Self-service account deletion - removes the CALLER's own user row and
    data they own. Never accepts a target user id; the identity is always
    the verified session's g.user_id, so this can only ever delete the
    caller's own account. Used by the e2e test suite to clean up throwaway
    accounts it creates; safe to leave enabled in any environment.
    """
    user_id = g.user_id
    try:
        from core.models import Team, TeamMember, Project, UserSettingModel
        from agents.executive_assistant.models import ExecTask, ExecReminder, ExecStakeholder
        from agents.event_networking.models import ENEvent
        from agents.supply_chain.models import SCSupplier
        from agents.content_marketing.models import CMProject
        from agents.market_research.models import ResearchProject

        # Teams owned by this user cascade-delete their Projects,
        # TeamMembers, and PendingInvites (FK ondelete=CASCADE).
        for team in Team.query.filter_by(owner_id=user_id).all():
            db.session.delete(team)
        TeamMember.query.filter_by(user_id=user_id).delete()

        ExecTask.query.filter_by(user_id=user_id).delete()
        ExecReminder.query.filter_by(user_id=user_id).delete()
        ExecStakeholder.query.filter_by(user_id=user_id).delete()

        for evt in ENEvent.query.filter_by(user_id=user_id).all():
            db.session.delete(evt)  # cascades to ENAttendee

        SCSupplier.query.filter_by(user_id=user_id).delete()

        for proj in CMProject.query.filter_by(user_id=user_id).all():
            db.session.delete(proj)  # cascades to CM documents/kg/content/conversations

        ResearchProject.query.filter_by(user_id=user_id).delete()
        UserSettingModel.query.filter_by(user_id=user_id).delete()
        EmailCampaign.query.filter_by(username=user_id).delete()

        user = User.query.filter_by(email=user_id).first()
        if user:
            db.session.delete(user)

        db.session.commit()
        return jsonify({'success': True, 'message': 'Account and owned data deleted'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500


# === HEALTH CHECK ENDPOINT ===
@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint for Docker and load balancers"""
    return jsonify({
        'status': 'healthy',
        'service': 'enable-agents-api',
        'timestamp': datetime.now().isoformat()
    }), 200


@app.route('/test-connection', methods=['GET'])
@require_auth
def test_connection():
    """Frontend compatibility health endpoint."""
    return jsonify({
        'status': 'connected',
        'service': 'enable-agents-api',
        'timestamp': datetime.now().isoformat()
    }), 200


@app.route('/api/save-project', methods=['POST', 'OPTIONS'])
@cross_origin()
def save_project():
    if request.method == 'OPTIONS':
        return jsonify({}), 200
    try:
        data = request.json
        trusted_uid, uid_err = _resolve_session_user_id(data.get('username'))
        if uid_err is not None:
            return uid_err[0], uid_err[1]
        username = trusted_uid

        name = data.get('name')
        query_used = data.get('query') or data.get('query_used', '')
        leads = data.get('businesses') or data.get('leads') or []
        
        if not username or username == 'anonymous' or not name:
            return jsonify({'success': False, 'error': 'Missing authenticated user or name'}), 400
            
        # Create Project
        project = SavedProject(username=username, name=name, query_used=query_used)
        db.session.add(project)
        db.session.flush() # get project id
        
        # Add leads
        for lead_data in leads:
            lead = SavedLead(
                project_id=project.id,
                name=lead_data.get('name', ''),
                website=lead_data.get('website', ''),
                phone=lead_data.get('phone', ''),
                address=lead_data.get('address', ''),
                emails=json.dumps(lead_data.get('emails', [])),
                linkedin_links=json.dumps(lead_data.get('linkedin_urls', [])),
                social_links=json.dumps(lead_data.get('social_links', {})),
                raw_data=json.dumps(lead_data)
            )
            db.session.add(lead)
            
        db.session.commit()
        return jsonify({'success': True, 'project_id': project.id, 'message': 'Leads saved successfully'})
    except Exception as e:
        db.session.rollback()
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/score-leads', methods=['POST', 'OPTIONS'])
@cross_origin()
@require_auth
def score_leads():
    """Score leads with a hybrid pipeline: embeddings rank the full list and the LLM refines only the top matches."""
    if request.method == 'OPTIONS':
        return jsonify({}), 200

    try:
        payload = request.get_json(silent=True) or {}
        requirement = (payload.get('requirement') or payload.get('query') or '').strip()
        businesses = payload.get('businesses') or payload.get('leads') or []

        if not requirement:
            return jsonify({'success': False, 'error': 'Missing requirement text'}), 400

        def _compact_text(value):
            if not value:
                return ''
            return re.sub(r'\s+', ' ', str(value)).strip()

        def _safe_json(value):
            if isinstance(value, dict):
                return value
            if isinstance(value, str):
                try:
                    parsed = json.loads(value)
                    return parsed if isinstance(parsed, dict) else {}
                except Exception:
                    return {}
            return {}

        def _extract_lead_text(lead_obj, max_chars=1200):
            raw_data = _safe_json(lead_obj.get('raw_data'))
            lead_name = _compact_text(lead_obj.get('name') or '')
            
            # PRIORITY: Lead name first (ensures it's the focus for LLM)
            text_bits = []
            if lead_name:
                text_bits.append(lead_name)
            
            # Then add fields, EXCLUDING conflicting names from raw_data
            # This prevents raw_data descriptions (which may reference OTHER leads) from confusing the LLM
            priority_fields = [
                lead_obj.get('summary'),
                lead_obj.get('description'),
                lead_obj.get('website'),
                lead_obj.get('address'),
                lead_obj.get('phone'),
                raw_data.get('summary'),
                raw_data.get('description'),
                raw_data.get('website'),
                raw_data.get('address'),
                raw_data.get('phone'),
                raw_data.get('about'),
                raw_data.get('specialties'),
                raw_data.get('services'),
                raw_data.get('keywords'),
                raw_data.get('categories'),
                raw_data.get('category'),
                raw_data.get('industry')
            ]
            
            for field in priority_fields:
                if isinstance(field, list):
                    text_bits.extend([_compact_text(item) for item in field if _compact_text(item)])
                elif isinstance(field, dict):
                    text_bits.extend([_compact_text(v) for v in field.values() if _compact_text(v)])
                else:
                    value = _compact_text(field)
                    if value:
                        text_bits.append(value)
            
            combined = ' '.join(text_bits)
            return combined[:max_chars] if len(combined) > max_chars else combined

        def _extract_two_line_summary(lead_obj):
            text = _extract_lead_text(lead_obj)
            lead_name = _compact_text(lead_obj.get('name') or '')
            if not text:
                return lead_name or 'No summary available'
            sentences = re.split(r'[\.\!\?]\s+', text)
            summary = ' '.join(sentences[:2]).strip()
            summary = re.sub(r'\s+', ' ', summary)
            summary = summary[:220].rstrip()
            if lead_name and lead_name.lower() not in summary.lower():
                return f"{lead_name} is a relevant match for your requirement."
            return summary

        def _safe_llm_summary(lead_obj, candidate_summary):
            lead_name = _compact_text(lead_obj.get('name') or '')
            summary = _compact_text(candidate_summary or '')
            if not summary:
                return _extract_two_line_summary(lead_obj)
            if lead_name and lead_name.lower() not in summary.lower():
                return _extract_two_line_summary(lead_obj)
            return summary[:220]

        def _coerce_lead(lead, index):
            if isinstance(lead, dict):
                return lead
            if isinstance(lead, str):
                try:
                    parsed = json.loads(lead)
                    if isinstance(parsed, dict):
                        return parsed
                except Exception:
                    pass
            return {'name': str(lead), 'raw_data': {'name': str(lead)}, '_index': index}

        def _similarity_score(req_vec, lead_vec):
            if req_vec is None or lead_vec is None:
                return 0
            if not np.any(req_vec) or not np.any(lead_vec):
                return 0
            sim = 1 - cosine(req_vec, lead_vec)
            if np.isnan(sim):
                return 0
            return int(max(0.0, min(1.0, sim)) * 100)

        def _fallback_sort_key(item):
            return item.get('match_score', 0) or 0

        openai_key = os.getenv('OPENAI_API_KEY')
        if not openai_key:
            results = []
            for lead in businesses:
                lead_obj = _coerce_lead(lead, len(results))
                results.append({
                    'index': len(results),
                    'match_score': 0,
                    'short_summary': _extract_two_line_summary(lead_obj)
                })
            return jsonify({'success': True, 'results': results})

        client = openai.OpenAI()
        client.api_key = openai_key

        lead_objects = [_coerce_lead(lead, index) for index, lead in enumerate(businesses)]
        lead_texts = [_extract_lead_text(lead_obj) for lead_obj in lead_objects]

        # Embeddings provide the full-list ranking; this is the fast, deterministic layer.
        embedding_inputs = [requirement] + lead_texts
        embeddings = get_embeddings_batch(embedding_inputs)

        if not embeddings or len(embeddings) != len(embedding_inputs):
            raise ValueError('Failed to generate embeddings for scoring')

        requirement_vec = np.array(embeddings[0], dtype=np.float32)
        lead_vectors = [np.array(item, dtype=np.float32) for item in embeddings[1:]]

        base_results = []
        for index, lead_obj in enumerate(lead_objects):
            base_score = _similarity_score(requirement_vec, lead_vectors[index])
            base_results.append({
                'index': index,
                'match_score': base_score,
                'short_summary': _extract_two_line_summary(lead_obj),
                'lead_obj': lead_obj,
                'lead_text': lead_texts[index]
            })

        # LLM only touches the top matches to refine ranking and produce richer summaries.
        top_k = min(int(os.getenv('LEAD_SCORE_LLM_TOP_K', '100')), len(base_results))
        if top_k > 0:
            top_candidates = sorted(base_results, key=lambda item: item['match_score'], reverse=True)[:top_k]
            compact_candidates = [
                {
                    'index': candidate['index'],
                    'name': candidate['lead_obj'].get('name') or '',
                    'current_score': candidate['match_score'],
                    'text': candidate['lead_text'][:1200]
                }
                for candidate in top_candidates
            ]

            llm_prompt = [
                {
                    'role': 'system',
                    'content': (
                        'You rank business/organization leads for a user requirement. '
                        'CRITICAL: Use the "name" field as the authoritative lead identifier. The summary must reference THIS NAME ONLY, not any other names mentioned in the text. '
                        'Use the requirement and company text to produce a final match_score from 0 to 100 and a concise two-line summary that directly references the lead name. '
                        'If the requirement implies buying, selling, or procuring, favor companies that offer that service or product. '
                        'Return only valid JSON as an array of objects with keys: index, match_score, short_summary.'
                    )
                },
                {
                    'role': 'user',
                    'content': json.dumps({'requirement': requirement, 'companies': compact_candidates}, ensure_ascii=False)
                }
            ]

            try:
                llm_response = client.chat.completions.create(
                    model=os.getenv('OPENAI_MODEL', 'gpt-4o-mini'),
                    messages=llm_prompt,
                    temperature=0.0,
                    max_tokens=1200
                )
                response_text = (llm_response.choices[0].message.content or '').strip()
                response_text = response_text.replace('```json', '').replace('```', '').strip()
                parsed = json.loads(response_text)
                if isinstance(parsed, list):
                    parsed_map = {int(item.get('index')): item for item in parsed if isinstance(item, dict) and str(item.get('index', '')).isdigit()}
                    for result in base_results:
                        item = parsed_map.get(result['index'])
                        if not item:
                            continue
                        llm_score = int(item.get('match_score') or result['match_score'])
                        llm_score = max(0, min(100, llm_score))
                        # Blend embedding ranking with LLM refinement for better buyer/seller intent handling.
                        result['match_score'] = int(round((result['match_score'] * 0.45) + (llm_score * 0.55)))
                        result['short_summary'] = _safe_llm_summary(result['lead_obj'], item.get('short_summary') or '')
            except Exception as llm_error:
                app.logger.warning(f"[score-leads] LLM refinement skipped: {llm_error}")

        results = [
            {
                'index': int(item['index']),
                'match_score': int(item['match_score']),
                'short_summary': item['short_summary']
            }
            for item in sorted(base_results, key=_fallback_sort_key, reverse=True)
        ]

        return jsonify({'success': True, 'results': results})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


def _normalize_dedupe_value(value):
    if not value:
        return ''
    normalized = str(value).strip().lower()
    normalized = re.sub(r'\s+', ' ', normalized)
    return re.sub(r'[^a-z0-9 ]+', '', normalized)


def _lead_dedupe_key(lead_data):
    if not lead_data:
        return ''

    website = str(lead_data.get('website', '') or '').strip().lower()
    parsed_website = urlparse(website if website.startswith(('http://', 'https://')) else f'https://{website}') if website else None
    website_key = ''
    if parsed_website and parsed_website.netloc:
        website_key = parsed_website.netloc.lower().lstrip('www.')

    phone_key = re.sub(r'\D+', '', str(lead_data.get('phone', '') or ''))
    name_key = _normalize_dedupe_value(lead_data.get('name', ''))
    address_key = _normalize_dedupe_value(lead_data.get('address', ''))

    key_parts = [part for part in [name_key, website_key, phone_key] if part]
    if not key_parts:
        key_parts = [address_key] if address_key else []
    return '|'.join(key_parts)

@app.route('/api/append-project', methods=['POST', 'OPTIONS'])
@cross_origin()
def append_project():
    if request.method == 'OPTIONS':
        return jsonify({}), 200
    try:
        data = request.json
        trusted_uid, uid_err = _resolve_session_user_id(data.get('username'))
        if uid_err is not None:
            return uid_err[0], uid_err[1]
        username = trusted_uid
        project_id = data.get('projectId')
        leads = data.get('businesses') or data.get('leads') or []
        
        if not username or username == 'anonymous' or not project_id:
            return jsonify({'success': False, 'error': 'Missing authenticated user or projectId'}), 400
            
        # verify
        project = db.session.get(SavedProject, project_id)
        if not project or project.username != username:
            return jsonify({'success': False, 'error': 'Project not found.'}), 404
        
        # Check current existing leads in project to append uniquely
        existing_leads = db.session.query(SavedLead).filter_by(project_id=project_id).all()
        existing_keys = set()
        for lead in existing_leads:
            lead_payload = {}
            if lead.raw_data:
                try:
                    lead_payload = json.loads(lead.raw_data)
                except Exception:
                    lead_payload = {}
            lead_payload = {
                'name': lead_payload.get('name') or lead.name or '',
                'website': lead_payload.get('website') or lead.website or '',
                'phone': lead_payload.get('phone') or lead.phone or '',
                'address': lead_payload.get('address') or lead.address or '',
            }
            dedupe_key = _lead_dedupe_key(lead_payload)
            if dedupe_key:
                existing_keys.add(dedupe_key)
        
        added = 0
        for lead_data in leads:
            dedupe_key = _lead_dedupe_key(lead_data)
            if dedupe_key and dedupe_key in existing_keys:
                continue
                
            lead = SavedLead(
                project_id=project.id,
                name=lead_data.get('name', ''),
                website=lead_data.get('website', ''),
                phone=lead_data.get('phone', ''),
                address=lead_data.get('address', ''),
                emails=json.dumps(lead_data.get('emails', [])),
                linkedin_links=json.dumps(lead_data.get('linkedin_urls', [])),
                social_links=json.dumps(lead_data.get('social_links', {})),
                raw_data=json.dumps(lead_data)
            )
            db.session.add(lead)
            if dedupe_key:
                existing_keys.add(dedupe_key)
            added += 1
            
        db.session.commit()
        return jsonify({'success': True, 'project_id': project.id, 'message': f'Appended {added} new leads.'})
    except Exception as e:
        db.session.rollback()
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/saved-projects', methods=['GET'])
@cross_origin()
def get_saved_projects():
    try:
        username_arg = request.args.get('username')
        trusted_uid, uid_err = _resolve_session_user_id(username_arg)
        if uid_err is not None:
            return uid_err[0], uid_err[1]

        projects = db.session.query(SavedProject).filter_by(username=trusted_uid).order_by(SavedProject.created_at.desc()).all()
        result = []
        for p in projects:
            lead_count = db.session.query(SavedLead).filter_by(project_id=p.id).count()
            result.append({
                'id': p.id,
                'name': p.name,
                'query_used': p.query_used,
                'created_at': p.created_at.isoformat(),
                'lead_count': lead_count
            })
            
        return jsonify({'success': True, 'projects': result})
    except Exception as e:
         import traceback
         traceback.print_exc()
         return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/saved-projects/<int:project_id>/leads', methods=['GET'])
@cross_origin()
def get_project_leads(project_id):
    try:
        username_arg = request.args.get('username')
        trusted_uid, uid_err = _resolve_session_user_id(username_arg)
        if uid_err is not None:
            return uid_err[0], uid_err[1]

        project = db.session.get(SavedProject, project_id)

        if not project or project.username != trusted_uid:
            return jsonify({'success': False, 'error': 'Project not found'}), 404
            
        leads = db.session.query(SavedLead).filter_by(project_id=project_id).all()
        result = []
        for l in leads:
            raw_data = {}
            if l.raw_data:
                try:
                    raw_data = json.loads(l.raw_data)
                except Exception:
                    raw_data = {}

            emails = json.loads(l.emails) if l.emails else []
            linkedin_links = json.loads(l.linkedin_links) if l.linkedin_links else []
            social_links = json.loads(l.social_links) if l.social_links else {}
            result.append({
                'id': l.id,
                'name': l.name,
                'website': l.website,
                'phone': l.phone,
                'address': l.address,
                'emails': emails,
                'email': raw_data.get('email') or (emails[0] if emails else 'N/A'),
                'linkedin': raw_data.get('linkedin') or (linkedin_links[0] if linkedin_links else ''),
                'linkedin_urls': linkedin_links,
                'social_links': social_links,
                'summary': raw_data.get('summary') or raw_data.get('description') or '',
                'description': raw_data.get('description') or '',
            })
            
        return jsonify({'success': True, 'project': {'id': project.id, 'name': project.name}, 'leads': result})
    except Exception as e:
         import traceback
         traceback.print_exc()
         return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/saved-projects/<int:project_id>', methods=['DELETE', 'OPTIONS'])
@cross_origin()
def delete_saved_project(project_id):
    if request.method == 'OPTIONS':
        return jsonify({}), 200
    try:
        username_arg = request.args.get('username')
        trusted_uid, uid_err = _resolve_session_user_id(username_arg)
        if uid_err is not None:
            return uid_err[0], uid_err[1]

        project = db.session.get(SavedProject, project_id)
        if not project or project.username != trusted_uid:
            return jsonify({'success': False, 'error': 'Project not found or unauthorized'}), 404

        db.session.delete(project)
        db.session.commit()
        return jsonify({'success': True, 'message': 'Project deleted successfully'})
    except Exception as e:
         db.session.rollback()
         import traceback
         traceback.print_exc()
         return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/admin/cleanup-orphaned-data', methods=['POST'])
@cross_origin()
@require_auth
def cleanup_orphaned_data():
    """
    Delete all orphaned data that doesn't belong to a platform Project.
    This includes SavedProject/SavedLead records from Market Research
    that were created before project-scoping was implemented.
    """
    try:
        # Count before deletion
        saved_projects_count = SavedProject.query.count()
        saved_leads_count = SavedLead.query.count()

        # Delete all SavedLeads first (due to FK constraint)
        SavedLead.query.delete()
        # Delete all SavedProjects
        SavedProject.query.delete()

        db.session.commit()

        return jsonify({
            'success': True,
            'message': 'Orphaned data cleaned up successfully',
            'deleted': {
                'saved_projects': saved_projects_count,
                'saved_leads': saved_leads_count
            }
        }), 200
    except Exception as e:
        db.session.rollback()
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/sales-helper-chat', methods=['POST'])
@cross_origin()
def sales_helper_chat():
    try:
        data = request.get_json() or {}
        question = (data.get('question') or '').strip()
        leads = data.get('leads') or []
        project = data.get('project') or {}

        trusted_uid, uid_err = _resolve_session_user_id(
            data.get('user_id') or data.get('username')
        )
        if uid_err is not None:
            return uid_err[0], uid_err[1]

        if not question:
            return jsonify({'success': False, 'error': 'Missing question'}), 400

        if not leads:
            return jsonify({'success': False, 'error': 'No leads provided'}), 400

        # Limit to first 50 leads for AI model context window management
        # Keeps API tokens reasonable and ensures fast response times
        lead_batch = leads[:50]

        available_fields = set()
        for lead in lead_batch:
            available_fields.update(lead.keys())
        
        available_fields_str = ", ".join(sorted([f for f in available_fields if f != 'match_score']))

        lead_rows = []
        for index, lead in enumerate(lead_batch, start=1):
            # Build a row with all available data
            lead_items = []
            for field in sorted(lead.keys()):
                value = lead.get(field, 'N/A')
                if value and value != 'N/A' and field != 'match_score':
                    lead_items.append(f"{field.replace('_', ' ').title()}: {value}")
            
            if not lead_items:
                lead_items = [f"Name: {lead.get('name', 'Lead ' + str(index))}"]
            
            lead_rows.append(f"{index}. " + " | ".join(lead_items))

        context_text = "\n".join(lead_rows)
        project_name = project.get('name') or 'Selected saved leads list'
        project_query = project.get('query_used') or 'N/A'

        prompt = f"""You are an intelligent sales analysis assistant. Your job is to answer the user's question by analyzing the provided leads data creatively and helpfully.

Saved list name: {project_name}
Original query: {project_query}
Available data fields: {available_fields_str}

Leads data:
{context_text}

User's question:
{question}

IMPORTANT INSTRUCTIONS:
1. ALWAYS provide a helpful answer, even if some specific fields are missing.
2. Use inference and reasoning: If exact data is unavailable, derive insights from related information.
   - Example: If "Number of Employees" is missing, infer from "Company Size", "Budget", "Industry", or "Description".
3. Mention lead names when listing specific companies.
4. Be concise but comprehensive - provide actionable insights.
5. If a field is completely unavailable, note it but focus on what you CAN analyze.
6. When multiple leads match the criteria, group or rank them.
7. Do not make up details, but DO use available context to reason and infer.
"""

        # Persist incoming leads into shared context (Redis + MySQL via ContextStore)
        user_id_for_context = trusted_uid
        project_id_value = (
            str(project.get('id')) if project and project.get('id') else (project.get('name') if project else None)
        )
        session_ctx = str(project_id_value)[:36] if project_id_value else None
        ingest_entries = []
        for lead in lead_batch:
            if not isinstance(lead, dict):
                continue
            try:
                flattened = " | ".join(
                    [
                        f"{k.replace('_', ' ').title()}: {v}"
                        for k, v in lead.items()
                        if v is not None and v != ''
                    ]
                )
                entry_metadata = {'project_name': project.get('name') if project else None}
                entry_key = _shared_context_entry_key('sales_helper', 'ingest_lead', 'lead', lead)
                payload_dict = _shared_context_payload_dict(
                    'sales_helper',
                    'ingest_lead',
                    'lead',
                    project_id_value,
                    lead,
                    flattened,
                    entry_metadata,
                )
                ingest_entries.append((entry_key, payload_dict, session_ctx))
            except Exception as ing_exc:
                logging.getLogger(__name__).warning('sales_helper context ingest row skipped: %s', ing_exc)
        try:
            if ingest_entries:
                ContextStore().set_many(user_id_for_context, 'sales_helper', ingest_entries)
        except Exception as persist_exc:
            logging.getLogger(__name__).exception('sales_helper shared context ingest failed: %s', persist_exc)

        if not os.getenv('OPENAI_API_KEY'):
            # Show first 10 lead names as preview in response
            lead_names = ", ".join([lead.get('name', 'N/A') for lead in leads[:10]])
            return jsonify({
                'success': True,
                'answer': f"Selected list: {project_name}. Leads loaded: {len(leads)}. I can see these example leads: {lead_names}. Ask a more specific question to analyze them further.",
                'lead_count': len(lead_batch),
                'lead_count_total': len(leads),
                'lead_count_limited': len(leads) > len(lead_batch)
            }), 200

        client = openai.OpenAI()
        client.api_key = os.environ['OPENAI_API_KEY']

        # Before calling LLM, attempt to surface additional saved context entries for this user
        extra_context_text = ""
        try:
            user_q = trusted_uid
            if user_q:
                # simple keyword match search against stored JSON payloads/keys
                matches = ContextStore().search(user_q, question, limit=6)
                if matches:
                    extracted_lines = []
                    for m in matches:
                        try:
                            payload = json.loads(m.value) if m.value else {}
                        except Exception:
                            payload = {}
                        extracted_lines.append(f"- {payload.get('text') or m.key}")
                    extra_context_text = "\n".join(extracted_lines)
        except Exception:
            extra_context_text = ""

        if extra_context_text:
            prompt = prompt + "\n\nAdditional saved context entries for this user:\n" + extra_context_text

        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a precise sales analysis assistant."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=900,
            temperature=0.2
        )

        answer = response.choices[0].message.content.strip()
        return jsonify({
            'success': True,
            'answer': answer,
            'lead_count': len(lead_batch),
            'lead_count_total': len(leads),
            'lead_count_limited': len(leads) > len(lead_batch)
        }), 200
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


# ========== SALES HELPER PRODUCT CATALOG DOCUMENTS ==========
# Lightweight file-backed store (no DB migration needed) - mirrors the
# json-sidecar pattern already used by /file_to_json_convert et al.

_SALES_HELPER_DOCS_DIR = os.path.join(os.getcwd(), 'data', 'sales_helper_docs')
_SALES_HELPER_UPLOADS_DIR = os.path.join(os.getcwd(), 'data', 'uploads', 'sales_helper')
os.makedirs(_SALES_HELPER_DOCS_DIR, exist_ok=True)
os.makedirs(_SALES_HELPER_UPLOADS_DIR, exist_ok=True)


def _sales_helper_index_path(user_id):
    safe_user = hashlib.sha256(user_id.encode('utf-8')).hexdigest()[:24]
    return os.path.join(_SALES_HELPER_DOCS_DIR, f'{safe_user}.json')


def _load_sales_helper_docs(user_id):
    path = _sales_helper_index_path(user_id)
    if not os.path.exists(path):
        return []
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return []


def _save_sales_helper_docs(user_id, docs):
    path = _sales_helper_index_path(user_id)
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(docs, f)


def _extract_sales_helper_text(file_path, ext):
    """Extract text from an uploaded product catalog file (pdf/docx/doc/txt)."""
    if ext == 'pdf':
        text_parts = []
        with fitz.open(file_path) as pdf:
            for page in pdf:
                page_text = page.get_text()
                if page_text.strip():
                    text_parts.append(page_text)
        return "\n\n".join(text_parts)
    if ext in ('docx', 'doc'):
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    # txt and anything else falls back to plain-text read
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()


@app.route('/api/sales-helper/documents', methods=['GET'])
@cross_origin()
@require_auth
def list_sales_helper_documents():
    user_id = g.user_id
    docs = _load_sales_helper_docs(user_id)
    response_docs = [
        {k: v for k, v in d.items() if k not in ('file_path', 'extracted_text')}
        for d in docs
    ]
    return jsonify({'success': True, 'documents': response_docs}), 200


@app.route('/api/sales-helper/documents', methods=['POST'])
@cross_origin()
@require_auth
def upload_sales_helper_document():
    """Upload a product catalog document for prospect matching."""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file provided'}), 400

        file = request.files['file']
        user_id = g.user_id
        doc_type = request.form.get('type', 'product_catalog')

        filename = secure_filename(file.filename)
        ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
        if ext not in ('pdf', 'doc', 'docx', 'txt'):
            return jsonify({'success': False, 'error': 'Unsupported file type. Use PDF, Word, or TXT.'}), 400

        doc_id = str(uuid4())
        user_dir = os.path.join(_SALES_HELPER_UPLOADS_DIR, hashlib.sha256(user_id.encode('utf-8')).hexdigest()[:24])
        os.makedirs(user_dir, exist_ok=True)
        file_path = os.path.join(user_dir, f'{doc_id}_{filename}')
        file.save(file_path)
        file_size = os.path.getsize(file_path)

        try:
            extracted_text = _extract_sales_helper_text(file_path, ext)
        except Exception as extract_err:
            logging.getLogger(__name__).warning('Sales helper doc extraction failed: %s', extract_err)
            extracted_text = ''

        document = {
            'id': doc_id,
            'name': filename,
            'type': doc_type,
            'size': f'{(file_size / 1024 / 1024):.1f} MB',
            'uploadedAt': datetime.utcnow().strftime('%Y-%m-%d'),
            'status': 'processed',
            'file_path': file_path,
            'extracted_text': extracted_text[:20000],  # cap stored text
        }

        docs = _load_sales_helper_docs(user_id)
        docs.append(document)
        _save_sales_helper_docs(user_id, docs)

        # Don't echo the full extracted text back to the client
        response_doc = {k: v for k, v in document.items() if k not in ('file_path', 'extracted_text')}
        return jsonify({'success': True, 'document': response_doc}), 201
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


def _find_sales_helper_doc(doc_id, user_id):
    for d in _load_sales_helper_docs(user_id):
        if d['id'] == doc_id:
            return d
    return None


@app.route('/api/sales-helper/documents/<doc_id>/download', methods=['GET'])
@cross_origin()
@require_auth
def download_sales_helper_document(doc_id):
    from flask import send_file
    user_id = g.user_id
    doc = _find_sales_helper_doc(doc_id, user_id)
    if not doc or not os.path.exists(doc.get('file_path', '')):
        return jsonify({'error': 'Document not found'}), 404
    return send_file(doc['file_path'], as_attachment=True, download_name=doc['name'])


@app.route('/api/sales-helper/documents/<doc_id>/view', methods=['GET'])
@cross_origin()
@require_auth
def view_sales_helper_document(doc_id):
    from flask import send_file
    user_id = g.user_id
    doc = _find_sales_helper_doc(doc_id, user_id)
    if not doc or not os.path.exists(doc.get('file_path', '')):
        return jsonify({'error': 'Document not found'}), 404
    return send_file(doc['file_path'], as_attachment=False, download_name=doc['name'])


@app.route('/api/sales-helper/documents/<doc_id>', methods=['DELETE'])
@cross_origin()
@require_auth
def delete_sales_helper_document(doc_id):
    try:
        user_id = g.user_id
        docs = _load_sales_helper_docs(user_id)
        remaining = [d for d in docs if d['id'] != doc_id]
        removed = [d for d in docs if d['id'] == doc_id]
        for d in removed:
            if d.get('file_path') and os.path.exists(d['file_path']):
                os.remove(d['file_path'])
        _save_sales_helper_docs(user_id, remaining)
        return jsonify({'success': True}), 200
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/sales-helper/match-prospects', methods=['POST'])
@cross_origin()
@require_auth
def match_sales_helper_prospects():
    """Use the uploaded product catalog(s) + a leads list to generate real,
    LLM-based prospect-fit analysis (not canned demo output)."""
    try:
        data = request.get_json() or {}
        user_id = g.user_id
        leads = data.get('leads') or []
        document_ids = data.get('document_ids') or []

        if not leads:
            return jsonify({'success': False, 'error': 'No leads provided'}), 400
        if not document_ids:
            return jsonify({'success': False, 'error': 'No product catalog documents selected'}), 400

        docs = _load_sales_helper_docs(user_id)
        selected_docs = [d for d in docs if d['id'] in document_ids]
        if not selected_docs:
            return jsonify({'success': False, 'error': 'Selected documents not found'}), 404

        catalog_text = "\n\n---\n\n".join(
            f"[{d['name']}]\n{d.get('extracted_text', '')[:6000]}" for d in selected_docs
        )

        lead_rows = []
        for idx, lead in enumerate(leads[:20], start=1):
            name = lead.get('name') or f'Lead {idx}'
            summary = lead.get('summary') or lead.get('description') or ''
            lead_rows.append(f"{idx}. {name} - {summary}".strip(' -'))
        leads_text = "\n".join(lead_rows)

        if not os.getenv('OPENAI_API_KEY'):
            return jsonify({
                'success': True,
                'analysis': f"Loaded {len(selected_docs)} catalog document(s) and {len(leads)} leads. "
                            f"Add an OpenAI API key to generate real match scoring.",
            }), 200

        prompt = f"""You are a sales analyst. Given this product/service catalog and this list of prospects,
identify which prospects are the best fit and why. Be specific and reference concrete details from
the catalog when explaining fit.

PRODUCT CATALOG:
{catalog_text}

PROSPECTS:
{leads_text}

Return a concise ranked list (best fit first) with a one-line reason for each prospect."""

        client = openai.OpenAI()
        client.api_key = os.environ['OPENAI_API_KEY']
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a precise sales/product-fit analyst."},
                {"role": "user", "content": prompt},
            ],
            max_tokens=900,
            temperature=0.3,
        )

        analysis = response.choices[0].message.content.strip()
        return jsonify({'success': True, 'analysis': analysis}), 200
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


# ========== SHARED CONTEXT API ENDPOINTS ==========

@app.route('/api/context/save', methods=['POST'])
def api_context_save():
    sk_err = _require_context_service_key_or_401()
    if sk_err is not None:
        return sk_err
    try:
        payload = request.get_json() or {}
        user_id, uid_err = _effective_context_api_user_id(payload)
        if uid_err is not None:
            return uid_err[0], uid_err[1]

        project_id = payload.get('project_id')
        entries = payload.get('entries') or []
        source_agent = payload.get('source_agent') or 'unknown'
        source_action = payload.get('source_action') or 'batch_save'
        sid = str(project_id)[:36] if project_id is not None else None
        batch = []
        meta = payload.get('entry_metadata') or {}
        for e in entries:
            if not isinstance(e, dict):
                continue
            flattened = " | ".join(
                [
                    f"{k.replace('_', ' ').title()}: {v}"
                    for k, v in e.items()
                    if v is not None and v != ''
                ]
            )
            entry_type = e.get('type') or 'data'
            entry_key = _shared_context_entry_key(source_agent, source_action, entry_type, e)
            payload_dict = _shared_context_payload_dict(
                source_agent, source_action, entry_type, project_id, e, flattened, meta
            )
            batch.append((entry_key, payload_dict, sid))

        if not batch:
            return jsonify({'success': True, 'saved': 0}), 200

        ContextStore().set_many(user_id, source_agent, batch)
        return jsonify({'success': True, 'saved': len(batch)}), 200
    except Exception as ex:
        logging.getLogger(__name__).exception('api_context_save failed')
        return jsonify({'success': False, 'error': str(ex)}), 500


@app.route('/api/context/search', methods=['POST'])
def api_context_search():
    sk_err = _require_context_service_key_or_401()
    if sk_err is not None:
        return sk_err
    try:
        payload = request.get_json() or {}
        user_id, uid_err = _effective_context_api_user_id(payload)
        if uid_err is not None:
            return uid_err[0], uid_err[1]

        query = (payload.get('query') or '').strip()
        if not user_id:
            return jsonify({'success': False, 'error': 'Missing user identity'}), 400

        raw_limit = payload.get('limit')
        try:
            limit = 10 if raw_limit in (None, '') else int(raw_limit)
        except (TypeError, ValueError):
            return jsonify({'success': False, 'error': 'Invalid limit; must be an integer'}), 400
        if limit < 1:
            return jsonify({'success': False, 'error': 'Invalid limit; must be greater than 0'}), 400

        try:
            results = ContextStore().search(user_id, query, limit=limit)
        except ValueError as ve:
            return jsonify({'success': False, 'error': str(ve)}), 400
        out = []
        for r in results:
            try:
                payload_obj = json.loads(r.value) if r.value else {}
            except Exception:
                payload_obj = {}
            out.append(
                {
                    'id': r.id,
                    'text': payload_obj.get('text'),
                    'data': payload_obj.get('data'),
                    'entry_metadata': payload_obj.get('entry_metadata'),
                    'source_agent': payload_obj.get('source_agent'),
                    'source_action': payload_obj.get('source_action'),
                    'entry_type': payload_obj.get('entry_type'),
                    'created_at': r.created_at.isoformat(),
                }
            )
        return jsonify({'success': True, 'results': out}), 200
    except Exception as ex:
        logging.getLogger(__name__).exception('api_context_search failed')
        return jsonify({'success': False, 'error': str(ex)}), 500


# Register all enabled agents from backend/agents/ (must be after all models defined)
from agents.registry import register_agents
register_agents(app)

# Register connector API routes
from core.connectors.routes import bp as connectors_bp
app.register_blueprint(connectors_bp)

# Register settings API routes
from core.settings_routes import bp as settings_bp
app.register_blueprint(settings_bp)

# Register team and projects API routes
from routes.team import team_bp
from routes.projects import projects_bp
app.register_blueprint(team_bp)
app.register_blueprint(projects_bp)

# Register workflow API routes
from routes.workflows import workflows_bp, load_system_templates
app.register_blueprint(workflows_bp)

# Load system workflow templates on startup
with app.app_context():
    load_system_templates()

# Register dependencies API routes
from routes.dependencies import dependencies_bp
app.register_blueprint(dependencies_bp)

# Import workflow models for db.create_all()
from models.workflow import WorkflowTemplate, WorkflowInstance


if __name__ == '__main__':
    with app.app_context():
        db.create_all()
        load_system_templates()
    app.run(debug=True, use_reloader=False, host='0.0.0.0', port=5000)
